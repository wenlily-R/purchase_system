# -*- coding: utf-8 -*-
"""V11.164 任务3预检: uploads 全部 docx 完整性 + 字体/占位符扫描"""
import os, zipfile, re, glob

BASE = os.path.dirname(os.path.abspath(__file__))
UP = os.path.join(BASE, 'uploads')
from docx import Document

print('=== 1. docx zip 完整性 ===')
for f in sorted(glob.glob(os.path.join(UP, '**', '*.docx'), recursive=True)):
    rel = os.path.relpath(f, UP)
    size = os.path.getsize(f)
    ok_zip = zipfile.is_zipfile(f)
    # python-docx 严格打开
    try:
        doc = Document(f)
        paras = len(doc.paragraphs)
        tables = len(doc.tables)
        ok_docx = True
        err = ''
    except Exception as e:
        ok_docx = False; paras = tables = -1; err = str(e)[:80]
    # 占位符残留
    ph = ''
    if ok_docx:
        txt = '\n'.join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    txt += '\n' + c.text
        left = re.findall(r'\{[^}]*\}', txt)
        ph = '占位符残留:' + ','.join(sorted(set(left)))[:80] if left else ''
        # 乱码检查(替换符/常见乱码)
        mojibake = re.findall(r'[\ufffd]|Ã[\x80-\xbf]|â€', txt)
        if mojibake:
            ph += ' [乱码!]'
    flag = '✅' if (ok_zip and ok_docx and not ph) else '❌'
    print(f"{flag} {rel} ({size}B) zip={ok_zip} docx={ok_docx} 段落={paras} 表格={tables} {err} {ph}")

print('\n=== 2. 默认模板 tpl_maimai.docx 结构(V11.130标准: 13章节+3表格) ===')
doc = Document(os.path.join(UP, 'tpl_maimai.docx'))
texts = [p.text for p in doc.paragraphs]
nums = [t for t in texts if re.match(r'^[一二三四五六七八九十]+、', t)]
print('表格数:', len(doc.tables), '(应=3: 标题行/标的明细/落款)')
print('章节数:', len(nums), nums[:16])
print('甲方固定:', any('河曲县正成洗选煤有限责任公司' in t for t in texts))
# 字体检查(中文不指定字体在Windows可能乱码/默认字体)
from docx.oxml.ns import qn
fonts = set()
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.name: fonts.add(r.font.name)
        rPr = r._element.rPr
        if rPr is not None:
            ea = rPr.find(qn('w:eastAsia'))
            if ea is not None and ea.get(qn('w:val')): fonts.add('eastAsia=' + ea.get(qn('w:val')))
print('段落字体:', fonts)

print('\n=== 3. 已生成合同样例检查(3份contract_*.docx) ===')
for f in sorted(glob.glob(os.path.join(UP, 'contract_*.docx'))):
    rel = os.path.relpath(f, UP)
    doc = Document(f)
    txt = '\n'.join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                txt += '\n' + c.text
    left = re.findall(r'\{[^}]*\}', txt)
    moji = re.findall(r'[\ufffd]|Ã[\x80-\xbf]', txt)
    print(f"{'✅' if not left and not moji else '❌'} {rel} 段落={len(doc.paragraphs)} 表格={len(doc.tables)} 残留={sorted(set(left))[:6] if left else '无'} 乱码={'有' if moji else '无'}")
