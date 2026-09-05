#!/usr/bin/env python3
"""正成能源采购系统 — PDF完整版PRD实现"""
import os, sys, json, datetime, hashlib, uuid, functools, threading, time, base64, re, io
import urllib.request, urllib.parse
import glob, secrets
from flask import Flask, jsonify, request, render_template, redirect, url_for, session, send_from_directory, make_response
from flask_cors import CORS
import sqlite3

app = Flask(__name__, template_folder='templates', static_folder='static')

@app.before_request
def write_guard():
    """安全防护: CSRF同源校验 + 系统写锁定(V5.1)
    白名单: 登录/登出/飞书回调(飞书服务器回调, 非用户操作)
    写锁定可通过系统设置页开关(write_lock), 默认开启"""
    if request.method in ('POST', 'PUT', 'DELETE', 'PATCH'):
        p = request.path
        if p in ('/api/login', '/api/logout', '/api/feishu/callback', '/api/dingtalk/callback', '/api/dingtalk/sso') or p.startswith('/api/inquiry/vendor/'):
            return None
        # CSRF: 带Origin/Referer的跨站写请求一律拒绝(同源才放行)
        origin = request.headers.get('Origin') or request.headers.get('Referer')
        if origin:
            ohost = urllib.parse.urlparse(origin).netloc
            if ohost and ohost != request.host:
                return jsonify({'error': '请求来源校验失败, 已拒绝'}), 403
        if not can_manage_config():
            # V8.0: 全局写锁默认关闭 — 所有启用账号可进行日常业务操作
            # 管理类操作(配置/删除/撤回/编辑单据)由各接口的 can_manage_config 自行保护
            if cfg_get('write_lock', '0') == '1':
                return jsonify({'error': '系统已锁定：仅系统管理员可操作，如需操作请联系管理员'}), 403
    return None
app.secret_key = 'zhengcheng-purchase-2026-secret-key'
CORS(app, supports_credentials=True)

BASE = os.path.dirname(os.path.abspath(__file__))
DB = os.path.join(BASE, 'data', 'purchase.db')
os.makedirs(os.path.join(BASE, 'data'), exist_ok=True)

def db():
    conn = sqlite3.connect(DB, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn

# ── 安全配置: 会话Cookie加固 (V5.1) ──
# V8.0: 会话有效期优化 — 刷新页面保持登录, 8小时不操作才自动登出
app.config.update(
    SESSION_COOKIE_SAMESITE='Lax',
    SESSION_COOKIE_HTTPONLY=True,
    PERMANENT_SESSION_LIFETIME=datetime.timedelta(hours=24),  # V8.3: 24小时不操作才登出 — 刷新/关浏览器保持登录
    SESSION_COOKIE_NAME='purchase_session',
    SESSION_REFRESH_EACH_REQUEST=False,  # V8.4: 不随每个请求刷新cookie时间戳 — 提交请求后会话不失效(旧cookie始终有效)
)

def now(): return datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

def esc_html(s):
    """HTML 转义(商家免登录报价页用, 防注入)"""
    if s is None: return ''
    return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;').replace("'",'&#39;')
def today(): return datetime.date.today().strftime('%Y-%m-%d')

# ── 密码安全: PBKDF2-SHA256 (V5.1, 兼容旧MD5登录时自动迁移) ──
PBKDF2_ITER = 120000
def hash_password(pwd):
    salt = secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac('sha256', pwd.encode('utf-8'), bytes.fromhex(salt), PBKDF2_ITER).hex()
    return 'pbkdf2$' + salt + '$' + h

def verify_password(pwd, stored):
    if not stored: return False
    if stored.startswith('pbkdf2$'):
        try:
            _, salt, h = stored.split('$')
            calc = hashlib.pbkdf2_hmac('sha256', pwd.encode('utf-8'), bytes.fromhex(salt), PBKDF2_ITER).hex()
            return h == calc
        except Exception:
            return False
    return hashlib.md5(pwd.encode('utf-8')).hexdigest() == stored

def is_legacy_md5(stored):
    return bool(stored) and len(stored) == 32 and re.match(r'^[0-9a-f]{32}$', stored) is not None

# ── 登录安全: 失败锁定 + 审计 (V5.1) ──
MAX_FAILS, LOCK_MINUTES = 5, 15
def get_client_ip():
    return (request.headers.get('X-Forwarded-For') or request.remote_addr or '').split(',')[0].strip()

def login_fail_count(c, username, ip):
    r = c.execute("SELECT COUNT(*) FROM login_attempts WHERE username=? AND ip=? AND success=0 AND created_at >= datetime('now','localtime',?)",
                  (username, ip, '-%d minutes' % LOCK_MINUTES)).fetchone()[0]
    return r

# ── 自动备份: sqlite backup API, 保留14份 (V5.1) ──
# 双机兼容: Mac(生产)用原固定路径; Windows(开发)落到本机桌面同名目录, 避免写死Mac路径导致WinError5备份失败
if sys.platform.startswith('win'):
    BACKUP_DIR = os.path.join(os.path.expanduser('~'), 'Desktop', '正成能源', '04_数据库备份', '自动备份')
else:
    BACKUP_DIR = '/Users/a0/Desktop/正成能源/04_数据库备份/自动备份'  # 2026-08-17 整理: 备份统一放正成能源/04_数据库备份/自动备份
def backup_db():
    try:
        os.makedirs(BACKUP_DIR, exist_ok=True)
        name = 'purchase_' + datetime.datetime.now().strftime('%Y%m%d_%H%M%S') + '.db'
        tmp = os.path.join(BACKUP_DIR, name + '.tmp')
        src_c = sqlite3.connect(DB, timeout=30)
        dst = sqlite3.connect(tmp)
        src_c.backup(dst)
        dst.close(); src_c.close()
        os.replace(tmp, os.path.join(BACKUP_DIR, name))
        files = sorted(glob.glob(os.path.join(BACKUP_DIR, 'purchase_*.db')), reverse=True)
        for f in files[14:]:
            try: os.remove(f)
            except Exception: pass
        return name
    except Exception as e:
        return 'ERROR:' + str(e)

def last_backup_name():
    try:
        files = sorted(glob.glob(os.path.join(BACKUP_DIR, 'purchase_*.db')), reverse=True)
        if files:
            return os.path.basename(files[0]).replace('purchase_', '').replace('.db', '')
    except Exception: pass
    return ''

def _startup_backup():
    try:
        if cfg_get('last_backup_date') != today():
            name = backup_db()
            cfg_set('last_backup_date', today())
            print('  [自动备份] 启动备份完成:', name)
    except Exception as e:
        print('  [自动备份] 启动备份失败:', e)

def _daily_backup_loop():
    while True:
        time.sleep(30)
        try:
            if time.localtime().tm_hour == 3 and cfg_get('last_backup_date') != today():
                name = backup_db()
                cfg_set('last_backup_date', today())
                print('  [自动备份] 定时备份完成:', name)
        except Exception:
            pass

def gen_no(prefix, table, col, c=None):
    """通用单号: 前缀-年月-序号, 如 CG-202608-0001
    用 MAX 取当前最大序号(并发安全, 避免 COUNT 竞态导致重复单号)"""
    conn = c if c else db()
    m = datetime.date.today().strftime('%Y%m')
    r = conn.execute(f"SELECT MAX({col}) m FROM {table} WHERE {col} LIKE ?", (f"{prefix}-{m}%",)).fetchone()
    if not c: conn.close()
    cur = int(str(r['m']).split('-')[-1]) if r and r['m'] else 0
    return f"{prefix}-{m}-{cur+1:04d}"

# V11.45: 部门→单号前缀映射(申请单号前两位随部门变化, 如生产=SC 财务=CW)
DEPT_PREFIX = {
    '生产部': 'SC', '财务部': 'CW', '机电部': 'JD', '机电': 'JD', '信息部': 'XX',
    '后勤部': 'HQ', '工程部': 'GC', '维修车间': 'WX', '综合办': 'ZH', '调度': 'DD',
    '化验室': 'HY', '库房': 'KF', '绿化': 'LH', '环卫': 'HW', '生产车队': 'CD',
    '采购部': 'CG', '采购与供应链部': 'CG',
}

def dept_prefix(dept):
    """部门 → 单号前缀, 未知部门默认 SC(采购)"""
    if not dept:
        return 'SC'
    d = str(dept).strip()
    if d in DEPT_PREFIX:
        return DEPT_PREFIX[d]
    # 包含匹配(如 "机电部2" → JD)
    for k, v in DEPT_PREFIX.items():
        if k and k in d:
            return v
    return 'SC'

def gen_req_no(dept=None, c=None):
    """采购申请单号: 部门前缀 + 年月日 + 当日总序号, 如 SC2026082101(生产部当天第1张) / CW2026082102(财务部当天第2张)
    序号=当天所有部门申请单的总序号(不分部门), 前缀按部门; 用 MAX 取当天最大序号(并发安全)"""
    conn = c if c else db()
    m = datetime.date.today().strftime('%Y%m%d')
    prefix = dept_prefix(dept)
    # V11.45b: 序号按当天所有申请单(不分部门前缀), 取当天任意前缀下的最大尾号
    r = conn.execute("SELECT req_no m FROM purchase_requests WHERE req_no LIKE ?", (f"__{m}%",)).fetchall()
    cur = 0
    for row in r:
        no = str(row['m'])
        # 尾号 = 最后2位
        try:
            tail = int(no[-2:])
            if tail > cur:
                cur = tail
        except Exception:
            continue
    if not c: conn.close()
    return f"{prefix}{m}{cur+1:02d}"

def gen_contract_no(c=None):
    """合同编码规则(55.docx需求7): HQZC-SBCG-份号-年份, 如 HQZC-SBCG-019-2026
    份号=当年合同序号, 年份=签订年份; 用 MAX 取当前最大序号(并发安全)"""
    conn = c if c else db()
    year = datetime.date.today().strftime('%Y')
    r = conn.execute("SELECT MAX(contract_no) m FROM contracts WHERE contract_no LIKE ?", (f'HQZC-SBCG-%-{year}',)).fetchone()
    if not c: conn.close()
    cur = 0
    if r and r['m']:
        try: cur = int(str(r['m']).split('-')[-2])
        except Exception: cur = 0
    return f'HQZC-SBCG-{cur+1:03d}-{year}'

def log(op, action, detail, c=None):
    if c: c.execute("INSERT INTO logs(operator,action,detail,created_at) VALUES(?,?,?,?)", (op,action,detail,now()))
    else: cc=db();cc.execute("INSERT INTO logs(operator,action,detail,created_at) VALUES(?,?,?,?)",(op,action,detail,now()));cc.commit();cc.close()

def add_notif(user_ids, title, content, biz_type='', biz_id=0, conn=None):
    """V11.214: 写系统站内信(铃铛🔔) — user_ids: 用户id列表; 供审批/报修/定损等环节推送待办提醒
    独立连接调用(不传conn); 若调用方已持有写连接须传conn避免锁冲突(调用方commit)"""
    if not user_ids: return
    _c = conn or db()
    for uid in user_ids:
        try:
            _c.execute("INSERT INTO notifications(user_id,type,title,content,biz_type,biz_id) VALUES(?,?,?,?,?,?)",
                       (uid, 'approval', title, content, biz_type, biz_id))
        except Exception:
            pass
    if not conn:
        _c.commit(); _c.close()

def dict_row(r): return dict(r) if r else None

# ── Auth decorator ──
def login_required(f):
    @functools.wraps(f)
    def wrap(*a,**kw):
        if 'user_id' not in session:
            return jsonify({'error':'login required'}),401
        return f(*a,**kw)
    return wrap

def can_manage_config():
    """系统配置管理权限: 系统管理员 / 分管领导(敏感业务写操作如库存编辑/单据删除/撤回/系统设置)
    或 sys_config.config_users 里指定的用户名(逗号分隔, 仅限非敏感的管理协作账号)。
    V11.201 收紧: 采购员等业务角色不再因在 config_users 而获得管理员级写权限。"""
    if session.get('user_role') in ('系统管理员', '分管领导'):
        return True
    extra = cfg_get('config_users', '')
    return session.get('username') in [x.strip() for x in extra.split(',') if x.strip()]

def admin_required(f):
    """需求44-权限强约束: 系统管理员或指定配置管理用户可配置/改基础数据/传模板"""
    @functools.wraps(f)
    def wrap(*a,**kw):
        if not can_manage_config():
            return jsonify({'error':'仅系统管理员可操作'}),403
        return f(*a,**kw)
    return wrap

# ============================================================
# INIT DB
# ============================================================
def init_db():
    conn = db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS _schema_version (version INTEGER);
    """)
    ver = conn.execute("SELECT MAX(version) FROM _schema_version").fetchone()[0]
    if ver is None:
        # Fresh install or upgrade from v3
        # Drop old v3 tables that have incompatible schemas
        conn.executescript("""
            DROP TABLE IF EXISTS users;
            DROP TABLE IF EXISTS purchase_requests;
            DROP TABLE IF EXISTS approval_instances;
        """)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS departments (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE NOT NULL, code TEXT UNIQUE NOT NULL);
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE NOT NULL, password TEXT NOT NULL,
            name TEXT NOT NULL, phone TEXT, role TEXT DEFAULT '员工',
            dept_id INTEGER, title TEXT, is_active INTEGER DEFAULT 1,
            FOREIGN KEY(dept_id) REFERENCES departments(id)
        );
        CREATE TABLE IF NOT EXISTS budget_accounts (
            id INTEGER PRIMARY KEY AUTOINCREMENT, code TEXT UNIQUE NOT NULL, name TEXT NOT NULL,
            dept_id INTEGER, annual_budget REAL DEFAULT 0, used_budget REAL DEFAULT 0,
            fiscal_year TEXT, FOREIGN KEY(dept_id) REFERENCES departments(id)
        );
        CREATE TABLE IF NOT EXISTS categories (id INTEGER PRIMARY KEY AUTOINCREMENT, code TEXT UNIQUE NOT NULL, name TEXT NOT NULL);
        CREATE TABLE IF NOT EXISTS items (id INTEGER PRIMARY KEY AUTOINCREMENT, cat_code TEXT NOT NULL, name TEXT NOT NULL, spec TEXT, unit TEXT DEFAULT '个', price REAL DEFAULT 0, safe_stock REAL DEFAULT 0, warehouse TEXT DEFAULT '主库房', supplier TEXT DEFAULT '');
        CREATE TABLE IF NOT EXISTS suppliers (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE NOT NULL, contact TEXT, phone TEXT, category TEXT, level TEXT DEFAULT '一般供应商', bank TEXT, account TEXT, tax_id TEXT, invoice_type TEXT DEFAULT '增值税专用发票', rating REAL DEFAULT 4.0, status TEXT DEFAULT '正常');
        CREATE TABLE IF NOT EXISTS purchase_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT, req_no TEXT UNIQUE NOT NULL, dept TEXT, requester TEXT, requester_id INTEGER,
            budget_code TEXT, purpose TEXT, target_date TEXT, status TEXT DEFAULT '待审批',
            total_estimated REAL DEFAULT 0, remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), updated_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS request_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT, req_id INTEGER NOT NULL, item_name TEXT NOT NULL, spec TEXT,
            unit TEXT DEFAULT '个', quantity REAL DEFAULT 1, estimated_price REAL DEFAULT 0, total_price REAL DEFAULT 0,
            remark TEXT, FOREIGN KEY(req_id) REFERENCES purchase_requests(id)
        );
        CREATE TABLE IF NOT EXISTS purchase_orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT, order_no TEXT UNIQUE NOT NULL, req_id INTEGER,
            item_name TEXT, spec TEXT, quantity REAL DEFAULT 1, unit TEXT DEFAULT '个',
            price REAL DEFAULT 0, amount REAL DEFAULT 0, tax_rate REAL DEFAULT 13, tax_amount REAL DEFAULT 0,
            total_amount REAL DEFAULT 0, supplier TEXT, requester TEXT, category TEXT,
            owner TEXT, owner_id INTEGER, target_date TEXT, status TEXT DEFAULT '待审批',
            remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), updated_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS price_comparisons (
            id INTEGER PRIMARY KEY AUTOINCREMENT, req_id INTEGER, order_id INTEGER, supplier TEXT,
            price REAL, quantity REAL, total REAL, delivery_cycle TEXT, warranty TEXT,
            quote_file TEXT, is_selected INTEGER DEFAULT 0, remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS contracts (id INTEGER PRIMARY KEY AUTOINCREMENT, contract_no TEXT UNIQUE NOT NULL, order_id INTEGER, contract_name TEXT, supplier TEXT, amount REAL DEFAULT 0, sign_date TEXT, start_date TEXT, end_date TEXT, content TEXT, file_path TEXT, status TEXT DEFAULT '执行中', remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS deliveries (id INTEGER PRIMARY KEY AUTOINCREMENT, delivery_no TEXT UNIQUE NOT NULL, order_id INTEGER, contract_id INTEGER, supplier TEXT, item_name TEXT, spec TEXT, quantity REAL DEFAULT 0, unit TEXT DEFAULT '个', driver_name TEXT, vehicle_no TEXT, delivery_date TEXT, receiver TEXT, sign_status TEXT DEFAULT '待签收', sign_time TEXT, remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS receivings (id INTEGER PRIMARY KEY AUTOINCREMENT, receive_no TEXT UNIQUE NOT NULL, delivery_id INTEGER, order_id INTEGER, item_name TEXT, spec TEXT, quantity REAL DEFAULT 0, unit TEXT DEFAULT '个', qualified_qty REAL DEFAULT 0, defective_qty REAL DEFAULT 0, inspector TEXT, warehouse TEXT DEFAULT '主库房', status TEXT DEFAULT '待检验', received_at TEXT, remark TEXT, attachments TEXT DEFAULT '', dept TEXT DEFAULT '', batch_no TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS invoices (id INTEGER PRIMARY KEY AUTOINCREMENT, invoice_no TEXT, invoice_code TEXT, order_id INTEGER, supplier TEXT, amount REAL DEFAULT 0, tax_amount REAL DEFAULT 0, total_amount REAL DEFAULT 0, invoice_date TEXT, invoice_type TEXT DEFAULT '增值税专用发票', file_path TEXT, status TEXT DEFAULT '待验证', remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS credit_notes (id INTEGER PRIMARY KEY AUTOINCREMENT, credit_no TEXT UNIQUE NOT NULL, order_id INTEGER, category TEXT, supplier TEXT, item_name TEXT, amount REAL DEFAULT 0, invoice_no TEXT, attachments TEXT, status TEXT DEFAULT '待审批', remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), updated_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS payment_requests (id INTEGER PRIMARY KEY AUTOINCREMENT, payment_no TEXT UNIQUE NOT NULL, credit_id INTEGER, payment_type TEXT DEFAULT '正常付款', supplier TEXT, amount REAL DEFAULT 0, contract_id INTEGER, status TEXT DEFAULT '待审批', paid_at TEXT, remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS inventory (id INTEGER PRIMARY KEY AUTOINCREMENT, item_id INTEGER, item_name TEXT, spec TEXT, cat_code TEXT, unit TEXT DEFAULT '个', quantity REAL DEFAULT 0, safe_stock REAL DEFAULT 0, warehouse TEXT DEFAULT '主库房', price REAL DEFAULT 0, updated_at TEXT DEFAULT (datetime('now','localtime')), UNIQUE(item_name,spec,warehouse));
        CREATE TABLE IF NOT EXISTS requisitions (id INTEGER PRIMARY KEY AUTOINCREMENT, req_no TEXT UNIQUE NOT NULL, dept TEXT, requester TEXT, item_name TEXT, spec TEXT, quantity REAL DEFAULT 0, unit TEXT DEFAULT '个', purpose TEXT, status TEXT DEFAULT '待审批', issued_at TEXT, receiver TEXT DEFAULT '', receive_dept TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now','localtime')), updated_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS approval_flow_config (id INTEGER PRIMARY KEY AUTOINCREMENT, biz_type TEXT NOT NULL, level_no INTEGER NOT NULL, role TEXT NOT NULL, min_amount REAL DEFAULT 0, max_amount REAL DEFAULT 9999999, label TEXT DEFAULT '');
        CREATE TABLE IF NOT EXISTS approval_instances (id INTEGER PRIMARY KEY AUTOINCREMENT, biz_type TEXT NOT NULL, biz_id INTEGER NOT NULL, level_no INTEGER NOT NULL, role TEXT, approver TEXT DEFAULT '', approver_id INTEGER, status TEXT DEFAULT 'pending', comment TEXT DEFAULT '', processed_at TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS notifications (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, type TEXT, title TEXT, content TEXT, biz_type TEXT, biz_id INTEGER, is_read INTEGER DEFAULT 0, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS logs (id INTEGER PRIMARY KEY AUTOINCREMENT, operator TEXT, action TEXT, detail TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS sys_config (key TEXT PRIMARY KEY, value TEXT);
        CREATE TABLE IF NOT EXISTS feishu_instances (
            id INTEGER PRIMARY KEY AUTOINCREMENT, instance_code TEXT UNIQUE, biz_type TEXT, biz_id INTEGER,
            status TEXT DEFAULT 'pending', error TEXT, created_at TEXT DEFAULT (datetime('now','localtime')),
            updated_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS dingtalk_instances (
            id INTEGER PRIMARY KEY AUTOINCREMENT, instance_code TEXT UNIQUE, biz_type TEXT, biz_id INTEGER,
            status TEXT DEFAULT 'pending', error TEXT, created_at TEXT DEFAULT (datetime('now','localtime')),
            updated_at TEXT DEFAULT (datetime('now','localtime')), form_values TEXT);
        CREATE TABLE IF NOT EXISTS dingtalk_push_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT, biz_type TEXT, biz_id TEXT, doc_no TEXT,
            push_type TEXT DEFAULT 'auto', target_user TEXT, target_userid TEXT,
            operator TEXT DEFAULT '系统', content TEXT, created_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS reminder_log (
            rule TEXT NOT NULL, key TEXT NOT NULL, pushed_at TEXT DEFAULT (datetime('now','localtime')),
            PRIMARY KEY(rule,key));
        CREATE TABLE IF NOT EXISTS inventory_counts (
            id INTEGER PRIMARY KEY AUTOINCREMENT, count_no TEXT UNIQUE NOT NULL, status TEXT DEFAULT '盘点中',
            remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), finished_at TEXT);
        CREATE TABLE IF NOT EXISTS inventory_count_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT, count_id INTEGER NOT NULL, inventory_id INTEGER,
            item_name TEXT, book_qty REAL DEFAULT 0, actual_qty REAL DEFAULT 0);
        CREATE TABLE IF NOT EXISTS inventory_adjustments (
            id INTEGER PRIMARY KEY AUTOINCREMENT, adj_no TEXT UNIQUE NOT NULL, adj_type TEXT NOT NULL,
            inventory_id INTEGER, item_name TEXT, spec TEXT, unit TEXT DEFAULT '个', book_qty REAL DEFAULT 0,
            adj_qty REAL DEFAULT 0, reason TEXT, status TEXT DEFAULT '待审批', source TEXT DEFAULT '手动',
            created_by TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), approved_at TEXT);
        CREATE TABLE IF NOT EXISTS contract_templates (
            id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT NOT NULL, file_path TEXT,
            version TEXT DEFAULT 'V1', status TEXT DEFAULT '启用', is_default INTEGER DEFAULT 0,
            remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        -- V5.0 库存流水表: 记录每次入库/出库/作废对库存的影响(溯源用)
        CREATE TABLE IF NOT EXISTS inventory_flows (
            id INTEGER PRIMARY KEY AUTOINCREMENT, item_name TEXT, spec TEXT, unit TEXT,
            flow_type TEXT,        -- 入库 / 出库 / 作废入库 / 作废出库 / 盘点调整
            doc_type TEXT,         -- receiving / requisition / count
            doc_id INTEGER, doc_no TEXT, qty REAL DEFAULT 0, balance_after REAL DEFAULT 0,
            operator TEXT, remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        -- V5.0 出库单明细子表(一张出库单多商品)
        CREATE TABLE IF NOT EXISTS requisition_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            requisition_id INTEGER NOT NULL, item_name TEXT, spec TEXT, unit TEXT DEFAULT '个',
            quantity REAL DEFAULT 0, purpose TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        -- V11.0 三方询价
        CREATE TABLE IF NOT EXISTS inquiries (
            id INTEGER PRIMARY KEY AUTOINCREMENT, inq_no TEXT UNIQUE NOT NULL, req_id INTEGER NOT NULL,
            title TEXT, purpose TEXT, status TEXT DEFAULT '询价中', selected_supplier_id INTEGER DEFAULT 0,
            deadline TEXT DEFAULT '', created_by TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), updated_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS inquiry_suppliers (
            id INTEGER PRIMARY KEY AUTOINCREMENT, inquiry_id INTEGER NOT NULL, supplier_name TEXT NOT NULL,
            contact TEXT, phone TEXT, token TEXT UNIQUE NOT NULL,
            quote_price REAL DEFAULT 0, quote_remark TEXT, quote_time TEXT,
            quote_delivery TEXT DEFAULT '', quote_warranty TEXT DEFAULT '',
            is_selected INTEGER DEFAULT 0, created_at TEXT DEFAULT (datetime('now','localtime')));
    """)
    conn.commit()

    if conn.execute("SELECT COUNT(*) FROM departments").fetchone()[0] == 0:
        conn.execute("INSERT INTO departments VALUES(1,'生产部','SC'),(2,'维修车间','WX'),(3,'后勤部','HQ'),(4,'综合办','ZHB'),(5,'财务部','CW')")
        for _un,_pw,_nm,_ph,_role,_did,_title in [
            ('admin','admin123','系统管理员','13800000000','系统管理员',4,'系统管理员'),
            ('zhangjl','123456','张经理','13800000001','部门负责人',1,'生产部经理'),
            ('lizong','123456','李总','13800000002','分管领导',None,'副总经理'),
            ('wangcw','123456','王财务','13800000003','财务',5,'财务主管'),
            ('zongjl','123456','赵总','13800000004','总经理',None,'总经理'),
            ('yuangong','123456','刘员工','13800000005','员工',1,'操作工'),
        ]:
            conn.execute("INSERT INTO users(username,password,name,phone,role,dept_id,title,is_active) VALUES(?,?,?,?,?,?,?,1)",
                         (_un, hash_password(_pw), _nm, _ph, _role, _did, _title))
        conn.executescript(f"""
            INSERT INTO suppliers VALUES(1,'恒安劳保用品','刘经理','13811111111','后勤类','核心供应商','工行001','62220001','91410100MA3X','增值税专用发票',4.5,'正常');
            INSERT INTO suppliers VALUES(2,'晋工机械有限公司','赵工','13822222222','维修配件类','核心供应商','建行002','62220002','91410100MA5X','增值税专用发票',4.8,'正常');
            INSERT INTO suppliers VALUES(3,'河曲建材批发','王老板','13833333333','建材类','一般供应商','农行003','62220003','','增值税普通发票',3.5,'正常');
            INSERT INTO categories VALUES(1,'BB','备品备件'),(2,'SB','设备'),(3,'JC','建材'),(4,'WX','维修耗材'),(5,'BG','办公用品');
            INSERT INTO items VALUES(1,'BB','轴承6205','6205-2RS','个',58.0,100,'主库房','晋工机械');
            INSERT INTO items VALUES(2,'WX','润滑油46#','46#液压油','桶',320.0,10,'主库房','晋工机械');
            INSERT INTO items VALUES(3,'BG','一次性纸杯','230ml','箱',100.0,50,'后勤库','恒安劳保用品');
            INSERT INTO inventory VALUES(1,1,'轴承6205','6205-2RS','BB','个',200,100,'主库房',58.0,datetime('now','localtime'));
            INSERT INTO inventory VALUES(2,2,'润滑油46#','46#液压油','WX','桶',15,10,'主库房',320.0,datetime('now','localtime'));
            INSERT INTO inventory VALUES(3,3,'一次性纸杯','230ml','BG','箱',80,50,'后勤库',100.0,datetime('now','localtime'));
            INSERT INTO budget_accounts VALUES(1,'SC-BG-2026','生产部办公耗材预算',1,50000,0,'2026');
            INSERT INTO budget_accounts VALUES(2,'SC-WX-2026','生产部维修配件预算',1,200000,0,'2026');
            INSERT INTO budget_accounts VALUES(3,'HQ-BG-2026','后勤部办公预算',3,30000,0,'2026');
            -- 审批流配置：按金额
            INSERT INTO approval_flow_config VALUES(1,'purchase_request',1,'部门负责人',0,5000,'小额-部门审批');
            INSERT INTO approval_flow_config VALUES(2,'purchase_request',2,'财务',5000,20000,'中额-部门+财务');
            INSERT INTO approval_flow_config VALUES(3,'purchase_request',3,'财务',20000,9999999,'大额-部门+财务+领导');
            INSERT INTO approval_flow_config VALUES(4,'purchase_request',4,'分管领导',20000,9999999,'大额-分管领导');
            INSERT INTO approval_flow_config VALUES(5,'purchase_request',5,'总经理',50000,9999999,'超大额-总经理终审');
        """)
        conn.execute("INSERT INTO _schema_version(version) VALUES(4)")
    # ---- V4.1 飞书对接: 用户飞书ID列(幂等) ----
    try:
        conn.execute("ALTER TABLE users ADD COLUMN feishu_open_id TEXT")
    except Exception:
        pass
    # ---- V5.2 钉钉对接: 用户钉钉ID列(幂等) ----
    try:
        conn.execute("ALTER TABLE users ADD COLUMN dingtalk_userid TEXT")
    except Exception:
        pass
    # ---- V8.5 库存按 名称+规格 独立SKU(修复同名不同规格合并): 重建唯一约束 + 按流水拆分历史合并数据 ----
    try:
        if conn.execute("SELECT value FROM sys_config WHERE key='inventory_spec_migrated'").fetchone() is None:
            conn.executescript("""
                CREATE TABLE inventory_new (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, item_id INTEGER, item_name TEXT, spec TEXT,
                    cat_code TEXT, unit TEXT DEFAULT '个', quantity REAL DEFAULT 0, safe_stock REAL DEFAULT 0,
                    warehouse TEXT DEFAULT '主库房', price REAL DEFAULT 0, updated_at TEXT DEFAULT (datetime('now','localtime')),
                    tax_rate REAL DEFAULT 13, remark TEXT DEFAULT '', max_stock REAL DEFAULT 0, last_move_date TEXT,
                    expiry_date TEXT DEFAULT '', supplier TEXT DEFAULT '', UNIQUE(item_name,spec,warehouse));
            """)
            old_rows = conn.execute("SELECT * FROM inventory").fetchall()
            # 流水分组: (名称,规格) -> 净数量(入库-出库, 流水出库存负数) + 最近变动时间
            groups = {}
            for f in conn.execute("SELECT item_name, spec, unit, qty, created_at FROM inventory_flows").fetchall():
                k = (f['item_name'], f['spec'] or '')
                g = groups.setdefault(k, {'qty': 0.0, 'unit': f['unit'] or '个', 'last': f['created_at'] or ''})
                g['qty'] += float(f['qty'] or 0)
                if f['created_at'] and f['created_at'] > g['last']:
                    g['last'] = f['created_at']
            used = set()
            for r in old_rows:
                k = (r['item_name'], r['spec'] or '')
                if k in groups and k not in used:
                    used.add(k)
                    qty = groups[k]['qty']; unit = groups[k]['unit']; last = groups[k]['last']
                else:
                    # 无流水的历史行原样保留(数量不变)
                    qty = r['quantity'] or 0; unit = r['unit'] or '个'; last = r['last_move_date'] or r['updated_at'] or ''
                conn.execute(
                    "INSERT INTO inventory_new(id,item_id,item_name,spec,cat_code,unit,quantity,safe_stock,warehouse,price,updated_at,tax_rate,remark,max_stock,last_move_date,expiry_date,supplier) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (r['id'], r['item_id'], r['item_name'], r['spec'] or '', r['cat_code'] or '', unit, qty,
                     r['safe_stock'] or 0, r['warehouse'] or '主库房', r['price'] or 0, now(), r['tax_rate'] or 13,
                     r['remark'] or '', r['max_stock'] or 0, last, r['expiry_date'] or '', r['supplier'] or ''))
            # 拆分出的新规格行(历史合并数据按流水拆开为独立SKU)
            for k, g in groups.items():
                if k in used:
                    continue
                old = next((r for r in old_rows if r['item_name'] == k[0]), None)
                conn.execute(
                    "INSERT INTO inventory_new(item_name,spec,cat_code,unit,quantity,safe_stock,warehouse,price,updated_at,tax_rate,last_move_date,supplier) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                    (k[0], k[1], (old['cat_code'] or '') if old else '', g['unit'], g['qty'],
                     (old['safe_stock'] or 0) if old else 0, (old['warehouse'] or '主库房') if old else '主库房',
                     (old['price'] or 0) if old else 0, now(), (old['tax_rate'] or 13) if old else 13, g['last'],
                     (old['supplier'] or '') if old else ''))
            conn.execute("DROP TABLE inventory")
            conn.execute("ALTER TABLE inventory_new RENAME TO inventory")
            conn.execute("INSERT OR REPLACE INTO sys_config(key,value) VALUES('inventory_spec_migrated','1')")
            conn.commit()
            log('系统', '库存结构迁移', f'库存按 名称+规格 独立SKU 重建完成: 新增拆分出 {len(groups)-len(used)} 个历史规格行')
    except Exception as _e:
        try:
            log('系统', '库存结构迁移异常', str(_e)[:200])
        except Exception:
            pass
    # ---- V5 合同模板: 无模板时自动生成内置默认模板 ----
    if conn.execute("SELECT COUNT(*) FROM contract_templates").fetchone()[0] == 0:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('采购合同', 0)
            doc.add_paragraph('合同编号：{合同编号}    订单编号：{订单编号}')
            doc.add_paragraph('甲方（采购方）：{甲方名称}')
            doc.add_paragraph('地址：{甲方地址}    联系人：{甲方联系人}    电话：{甲方电话}')
            doc.add_paragraph('乙方（供应方）：{乙方名称}')
            doc.add_paragraph('地址：{乙方地址}    联系人：{乙方联系人}    电话：{乙方电话}')
            doc.add_paragraph('开户行：{乙方开户行}    账号：{乙方账号}')
            doc.add_paragraph('下单日期：{下单日期}    预计交货日期：{预计交货日期}')
            doc.add_paragraph('结算方式：{结算方式}')
            doc.add_paragraph('采购明细清单：')
            doc.add_paragraph('{明细清单}')
            doc.add_paragraph('合计金额：{合计金额}')
            doc.add_paragraph('本协议一式两份，甲乙双方各执一份，签字盖章后生效。')
            os.makedirs(os.path.join(BASE, 'uploads'), exist_ok=True)
            doc.save(os.path.join(BASE, 'uploads', 'tpl_default.docx'))
            conn.execute("INSERT INTO contract_templates(name,file_path,version,status,is_default,remark) VALUES('标准采购合同模板','tpl_default.docx','V1','启用',1,'系统内置默认模板, 可上传替换')")
        except Exception as e:
            print('内置模板生成失败:', e)
    conn.commit()
    # 各业务审批流种子(幂等): 订单/合同/挂账/付款 默认按金额5档审批
    for _biz in ('purchase_order','contract','credit','payment'):
        if conn.execute("SELECT COUNT(*) FROM approval_flow_config WHERE biz_type=?", (_biz,)).fetchone()[0] == 0:
            conn.execute(f"""INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label) VALUES
                ('{_biz}',1,'部门负责人',0,5000,'小额-部门审批'),
                ('{_biz}',2,'财务',5000,20000,'中额-部门+财务'),
                ('{_biz}',3,'财务',20000,9999999,'大额-部门+财务+领导'),
                ('{_biz}',4,'分管领导',20000,9999999,'大额-分管领导'),
                ('{_biz}',5,'总经理',50000,9999999,'超大额-总经理终审'""")
    # V5.0: 入库/出库审批流(1级部门负责人即可, 出库金额0)
    for _biz in ('receiving','requisition'):
        if conn.execute("SELECT COUNT(*) FROM approval_flow_config WHERE biz_type=?", (_biz,)).fetchone()[0] == 0:
            conn.execute(f"""INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label) VALUES
                ('{_biz}',1,'部门负责人',0,9999999,'入库/出库-部门审批')""")
    # V11.175: 驳回历史表必须先于下方补列循环创建(否则老库缺表时 ALTER 报 no such table, 新库/迁移库无法启动)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS approval_reject_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            biz_type TEXT, biz_id INTEGER,
            approver TEXT, approver_id INTEGER DEFAULT 0,
            comment TEXT, processed_at TEXT,
            source TEXT DEFAULT 'system',
            created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE INDEX IF NOT EXISTS idx_reject_logs ON approval_reject_logs(biz_type, biz_id);
    """)
    # ---- V4.4 需求文档补充: 交易模式/对账/合并开票/强制关联(幂等迁移) ----
    for _tbl, _col, _ddl in [
        # ---- V55 需求变更列(全部幂等) ----
        ('purchase_requests', 'attachments', "ALTER TABLE purchase_requests ADD COLUMN attachments TEXT DEFAULT ''"),
        ('purchase_requests', 'urgent', "ALTER TABLE purchase_requests ADD COLUMN urgent INTEGER DEFAULT 0"),
        ('purchase_orders', 'urgent', "ALTER TABLE purchase_orders ADD COLUMN urgent INTEGER DEFAULT 0"),
        ('purchase_orders', 'attachments', "ALTER TABLE purchase_orders ADD COLUMN attachments TEXT DEFAULT ''"),
        ('contracts', 'urgent', "ALTER TABLE contracts ADD COLUMN urgent INTEGER DEFAULT 0"),
        ('contracts', 'attachment', "ALTER TABLE contracts ADD COLUMN attachment TEXT DEFAULT ''"),
        ('payment_requests', 'urgent', "ALTER TABLE payment_requests ADD COLUMN urgent INTEGER DEFAULT 0"),
        ('payment_requests', 'payment_reason', "ALTER TABLE payment_requests ADD COLUMN payment_reason TEXT DEFAULT ''"),
        ('payment_requests', 'expect_pay_date', "ALTER TABLE payment_requests ADD COLUMN expect_pay_date TEXT DEFAULT ''"),
        ('payment_requests', 'invoice_type', "ALTER TABLE payment_requests ADD COLUMN invoice_type TEXT DEFAULT '专票'"),
        ('payment_requests', 'has_contract', "ALTER TABLE payment_requests ADD COLUMN has_contract TEXT DEFAULT '否'"),
        ('payment_requests', 'contract_attachment', "ALTER TABLE payment_requests ADD COLUMN contract_attachment TEXT DEFAULT ''"),
        ('payment_requests', 'payee_name', "ALTER TABLE payment_requests ADD COLUMN payee_name TEXT DEFAULT ''"),
        ('payment_requests', 'payee_account', "ALTER TABLE payment_requests ADD COLUMN payee_account TEXT DEFAULT ''"),
        ('payment_requests', 'attachments', "ALTER TABLE payment_requests ADD COLUMN attachments TEXT DEFAULT ''"),
        ('inventory', 'tax_rate', "ALTER TABLE inventory ADD COLUMN tax_rate REAL DEFAULT 13"),
        ('inventory', 'remark', "ALTER TABLE inventory ADD COLUMN remark TEXT DEFAULT ''"),
        ('inventory', 'max_stock', "ALTER TABLE inventory ADD COLUMN max_stock REAL DEFAULT 0"),
        ('inventory', 'last_move_date', "ALTER TABLE inventory ADD COLUMN last_move_date TEXT DEFAULT ''"),
        ('inventory', 'expiry_date', "ALTER TABLE inventory ADD COLUMN expiry_date TEXT DEFAULT ''"),
        ('inventory', 'supplier', "ALTER TABLE inventory ADD COLUMN supplier TEXT DEFAULT ''"),
        ('approval_flow_config', 'approver', "ALTER TABLE approval_flow_config ADD COLUMN approver TEXT DEFAULT ''"),
        ('receivings', 'remark', "ALTER TABLE receivings ADD COLUMN remark TEXT DEFAULT ''"),
        ('receivings', 'urgent', "ALTER TABLE receivings ADD COLUMN urgent INTEGER DEFAULT 0"),
        ('purchase_orders', 'trade_mode', "ALTER TABLE purchase_orders ADD COLUMN trade_mode TEXT DEFAULT '货到付款'"),
        ('suppliers', 'created_at', "ALTER TABLE suppliers ADD COLUMN created_at TEXT DEFAULT ''"),
        ('deliveries', 'contract_no', "ALTER TABLE deliveries ADD COLUMN contract_no TEXT DEFAULT ''"),
        ('receivings', 'contract_id', "ALTER TABLE receivings ADD COLUMN contract_id INTEGER"),
        ('receivings', 'contract_no', "ALTER TABLE receivings ADD COLUMN contract_no TEXT DEFAULT ''"),
        ('invoices', 'order_ids', "ALTER TABLE invoices ADD COLUMN order_ids TEXT DEFAULT ''"),
        ('invoices', 'contract_id', "ALTER TABLE invoices ADD COLUMN contract_id INTEGER"),
        ('invoices', 'contract_no', "ALTER TABLE invoices ADD COLUMN contract_no TEXT DEFAULT ''"),
        ('payment_requests', 'trade_mode', "ALTER TABLE payment_requests ADD COLUMN trade_mode TEXT DEFAULT '货到付款'"),
        ('payment_requests', 'settlement_id', "ALTER TABLE payment_requests ADD COLUMN settlement_id INTEGER"),
        ('credit_notes', 'contract_no', "ALTER TABLE credit_notes ADD COLUMN contract_no TEXT DEFAULT ''"),
        ('purchase_requests', 'rejected_reason', "ALTER TABLE purchase_requests ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('purchase_requests', 'apply_date', "ALTER TABLE purchase_requests ADD COLUMN apply_date TEXT DEFAULT ''"),
        ('contracts', 'updated_at', "ALTER TABLE contracts ADD COLUMN updated_at TEXT"),
        ('receivings', 'updated_at', "ALTER TABLE receivings ADD COLUMN updated_at TEXT"),
        ('requisitions', 'issued_at', "ALTER TABLE requisitions ADD COLUMN issued_at TEXT"),
        ('request_items', 'category', "ALTER TABLE request_items ADD COLUMN category TEXT DEFAULT ''"),
        ('request_items', 'brand_param', "ALTER TABLE request_items ADD COLUMN brand_param TEXT DEFAULT ''"),
        ('request_items', 'arrival_date', "ALTER TABLE request_items ADD COLUMN arrival_date TEXT DEFAULT ''"),
        # V11.126: 钉钉实例存审批表单值(询价定标审批需要读取领导在钉钉选的供应商)
        ('dingtalk_instances', 'form_values', "ALTER TABLE dingtalk_instances ADD COLUMN form_values TEXT"),
        # V11.180: 三方询价采购方议价 — 厂家原始报价保留在 quote_* 不动, 新增采购调整字段
        ('inquiry_suppliers', 'adj_details', "ALTER TABLE inquiry_suppliers ADD COLUMN adj_details TEXT DEFAULT ''"),
        ('inquiry_suppliers', 'adj_price', "ALTER TABLE inquiry_suppliers ADD COLUMN adj_price REAL DEFAULT 0"),
        ('inquiry_suppliers', 'adj_remark', "ALTER TABLE inquiry_suppliers ADD COLUMN adj_remark TEXT DEFAULT ''"),
        # V11.181: 驳回附件同步 — 驳回记录存附件元数据+钉钉实例号(查看详情可直达OA审批单看附件)
        ('approval_reject_logs', 'attachments', "ALTER TABLE approval_reject_logs ADD COLUMN attachments TEXT DEFAULT ''"),
        ('approval_reject_logs', 'instance_code', "ALTER TABLE approval_reject_logs ADD COLUMN instance_code TEXT DEFAULT ''"),
        # V11.184: 审批节点附件全量同步(同意/驳回上传的图片附件) — 独立表, 一条审批操作一条记录, 多次审批不覆盖
        # 列: biz_type/biz_id/action(agree|reject)/approver/comment/processed_at/attachments(JSON)/source/instance_code/created_at
        # V11.175: 驳回展示 — 各业务表补 rejected_reason(列表/详情显示最近驳回理由)
        ('requisitions', 'rejected_reason', "ALTER TABLE requisitions ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('receivings', 'rejected_reason', "ALTER TABLE receivings ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('contracts', 'rejected_reason', "ALTER TABLE contracts ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('purchase_orders', 'rejected_reason', "ALTER TABLE purchase_orders ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('credit_notes', 'rejected_reason', "ALTER TABLE credit_notes ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        ('payment_requests', 'rejected_reason', "ALTER TABLE payment_requests ADD COLUMN rejected_reason TEXT DEFAULT ''"),
        # V11.175c: 列表独立驳回列 — 各业务表补最新驳回人/时间
        ('requisitions', 'rejected_by', "ALTER TABLE requisitions ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('receivings', 'rejected_by', "ALTER TABLE receivings ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('contracts', 'rejected_by', "ALTER TABLE contracts ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('purchase_orders', 'rejected_by', "ALTER TABLE purchase_orders ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('credit_notes', 'rejected_by', "ALTER TABLE credit_notes ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('payment_requests', 'rejected_by', "ALTER TABLE payment_requests ADD COLUMN rejected_by TEXT DEFAULT ''"),
        ('requisitions', 'rejected_at', "ALTER TABLE requisitions ADD COLUMN rejected_at TEXT DEFAULT ''"),
        ('receivings', 'rejected_at', "ALTER TABLE receivings ADD COLUMN rejected_at TEXT DEFAULT ''"),
        ('contracts', 'rejected_at', "ALTER TABLE contracts ADD COLUMN rejected_at TEXT DEFAULT ''"),
        ('purchase_orders', 'rejected_at', "ALTER TABLE purchase_orders ADD COLUMN rejected_at TEXT DEFAULT ''"),
        ('credit_notes', 'rejected_at', "ALTER TABLE credit_notes ADD COLUMN rejected_at TEXT DEFAULT ''"),
        ('payment_requests', 'rejected_at', "ALTER TABLE payment_requests ADD COLUMN rejected_at TEXT DEFAULT ''"),
        # V11.185: 驳回退回重提闭环 — 各业务表补 reject_count(累计驳回次数)/resubmit_count(累计重提次数)
        #          + rejected_items(本次被驳回的明细条目id JSON; 整单驳回='__all__'; 空=无标记)
        ('purchase_requests', 'reject_count', "ALTER TABLE purchase_requests ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('purchase_requests', 'resubmit_count', "ALTER TABLE purchase_requests ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('purchase_requests', 'rejected_items', "ALTER TABLE purchase_requests ADD COLUMN rejected_items TEXT DEFAULT ''"),
        ('purchase_orders', 'reject_count', "ALTER TABLE purchase_orders ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('purchase_orders', 'resubmit_count', "ALTER TABLE purchase_orders ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('purchase_orders', 'rejected_items', "ALTER TABLE purchase_orders ADD COLUMN rejected_items TEXT DEFAULT ''"),
        ('contracts', 'reject_count', "ALTER TABLE contracts ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('contracts', 'resubmit_count', "ALTER TABLE contracts ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('contracts', 'rejected_items', "ALTER TABLE contracts ADD COLUMN rejected_items TEXT DEFAULT ''"),
        ('receivings', 'reject_count', "ALTER TABLE receivings ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('receivings', 'resubmit_count', "ALTER TABLE receivings ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('receivings', 'rejected_items', "ALTER TABLE receivings ADD COLUMN rejected_items TEXT DEFAULT ''"),
        ('requisitions', 'reject_count', "ALTER TABLE requisitions ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('requisitions', 'resubmit_count', "ALTER TABLE requisitions ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('requisitions', 'rejected_items', "ALTER TABLE requisitions ADD COLUMN rejected_items TEXT DEFAULT ''"),
        ('inquiries', 'reject_count', "ALTER TABLE inquiries ADD COLUMN reject_count INTEGER DEFAULT 0"),
        ('inquiries', 'resubmit_count', "ALTER TABLE inquiries ADD COLUMN resubmit_count INTEGER DEFAULT 0"),
        ('inquiries', 'rejected_items', "ALTER TABLE inquiries ADD COLUMN rejected_items TEXT DEFAULT ''"),
        # ---- V11.203 模块一: 合同发票条款/开票计划(1.1)+发票催收状态(1.2), 老合同默认空不受影响 ----
        ('contracts', 'invoice_clause', "ALTER TABLE contracts ADD COLUMN invoice_clause TEXT DEFAULT ''"),
        ('contracts', 'invoice_est_first', "ALTER TABLE contracts ADD COLUMN invoice_est_first TEXT DEFAULT ''"),
        ('contracts', 'invoice_est_done', "ALTER TABLE contracts ADD COLUMN invoice_est_done TEXT DEFAULT ''"),
        ('contracts', 'inv_collect_status', "ALTER TABLE contracts ADD COLUMN inv_collect_status TEXT DEFAULT ''"),
    ]:
        _cols = [r[1] for r in conn.execute(f"PRAGMA table_info({_tbl})").fetchall()]
        if _col not in _cols:
            conn.execute(_ddl)
    # ---- V11.203 模块一1.2: 合同发票登记台账 + 发票节点提醒去重表(新表, 不影响老库) ----
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS contract_invoices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contract_id INTEGER NOT NULL,
            invoice_no TEXT DEFAULT '',
            amount REAL DEFAULT 0,
            invoice_type TEXT DEFAULT '',
            received_date TEXT DEFAULT '',
            operator TEXT DEFAULT '',
            remark TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS contract_inv_reminds (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            contract_id INTEGER NOT NULL,
            remind_date TEXT NOT NULL,
            kind TEXT DEFAULT 'due',
            pushed INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now','localtime')),
            UNIQUE(contract_id, remind_date, kind)
        );
    """)
    # ---- V11.193 退库模块: 退库申请单(已领用物资退回仓库) ----
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS return_requests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            return_no TEXT UNIQUE NOT NULL,
            source_req_id INTEGER NOT NULL,
            source_req_no TEXT DEFAULT '',
            dept TEXT DEFAULT '',          -- 领用部门(自动带出)
            receiver TEXT DEFAULT '',      -- 领用人(自动带出)
            warehouse TEXT DEFAULT '主库房',
            reason TEXT DEFAULT '',        -- 退库原因(领用剩余/物料未使用/错领物料/质量问题/其他)
            reason_note TEXT DEFAULT '',   -- 自定义说明(原因=其他时)
            total_amount REAL DEFAULT 0,
            status TEXT DEFAULT '草稿',    -- 草稿/待审批/审批通过/退库已完成/已驳回/已作废
            requester TEXT DEFAULT '',
            attachments TEXT DEFAULT '[]',
            warehouse_confirm_by TEXT DEFAULT '',
            warehouse_confirm_at TEXT DEFAULT '',
            reject_count INTEGER DEFAULT 0,
            resubmit_count INTEGER DEFAULT 0,
            rejected_items TEXT DEFAULT '',
            rejected_reason TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime')),
            updated_at TEXT DEFAULT (datetime('now','localtime')),
            finished_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS return_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            return_id INTEGER NOT NULL,
            source_item_id INTEGER DEFAULT 0,  -- 源出库单明细id(requisition_items.id)
            item_name TEXT, spec TEXT, unit TEXT DEFAULT '个',
            issued_qty REAL DEFAULT 0,     -- 原领用出库数量(带出)
            returned_qty REAL DEFAULT 0,   -- 该源明细累计已退(带出,只读)
            return_qty REAL DEFAULT 0,     -- 本次退库数量(强校验: 0<qty<=可退)
            price REAL DEFAULT 0,          -- 物料单价(带出)
            amount REAL DEFAULT 0,         -- 退库金额=qty*price
            created_at TEXT DEFAULT (datetime('now','localtime'))
        );
    """)
    # requisitions 加累计已退数量字段(回写源出库单)
    _rcols = [r[1] for r in conn.execute("PRAGMA table_info(requisitions)").fetchall()]
    if 'returned_qty' not in _rcols:
        conn.execute("ALTER TABLE requisitions ADD COLUMN returned_qty REAL DEFAULT 0")
    # ---- V11.202 分批验收: receivings 加批次号(空=老整批流程单, 非空=分批验收单, 幂等) ----
    _rcvcols = [r[1] for r in conn.execute("PRAGMA table_info(receivings)").fetchall()]
    if 'batch_no' not in _rcvcols:
        conn.execute("ALTER TABLE receivings ADD COLUMN batch_no TEXT DEFAULT ''")
    # ---- V11.206 集体验收: 标记是否需集体验收 + 验收状态(空=常规, 1=需集体验收; collect_status: 空/待集体验收/已集体验收) ----
    if 'collect_accept' not in _rcvcols:
        conn.execute("ALTER TABLE receivings ADD COLUMN collect_accept INTEGER DEFAULT 0")
    if 'collect_status' not in _rcvcols:
        conn.execute("ALTER TABLE receivings ADD COLUMN collect_status TEXT DEFAULT ''")
    # 集体验收审批流配置(首次建库初始化, 幂等; 角色可后续在系统设置改)
    if conn.execute("SELECT COUNT(*) FROM approval_flow_config WHERE biz_type='collect_accept'").fetchone()[0] == 0:
        conn.execute("INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label) VALUES('collect_accept',1,'分管领导',0,1000000,'集体验收-1级')")
    # 退库审批流配置(首次建库时初始化; 幂等: 已有配置不覆盖)
    if conn.execute("SELECT COUNT(*) FROM approval_flow_config WHERE biz_type='return_request'").fetchone()[0] == 0:
        conn.execute("INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label) VALUES('return_request',1,'部门负责人',0,1000000,'退库审批-1级')")
    # ---- V11.208 维修采购独立流程(模块五): 单表+阶段字段实现 提报→定损→报价→变更二次确认→返库验收 ----
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS repair_plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plan_no TEXT UNIQUE NOT NULL,
            device_name TEXT,            -- 损坏设备名称
            fault_desc TEXT,             -- 故障描述
            dept TEXT DEFAULT '',
            requester TEXT DEFAULT '', requester_id INTEGER DEFAULT 0,
            status TEXT DEFAULT '草稿',   -- 草稿/待定损/定损通过/待报价/报价完成/变更待确认/维修中/待返库/已完成/已驳回/已作废
            stage TEXT DEFAULT '',        -- 当前环节标记(维修中=含变更次数)
            change_count INTEGER DEFAULT 0,  -- 变更二次确认次数
            quote_total REAL DEFAULT 0,   -- 厂家报价合计
            repair_company TEXT DEFAULT '', -- 维修厂家
            finish_date TEXT DEFAULT '',  -- 预计完工
            attachments TEXT DEFAULT '',
            reject_count INTEGER DEFAULT 0,
            rejected_items TEXT DEFAULT '',
            remark TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime')),
            updated_at TEXT DEFAULT ''
        );
        CREATE TABLE IF NOT EXISTS repair_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plan_id INTEGER, part_name TEXT, fault_note TEXT DEFAULT '',
            confirm_status TEXT DEFAULT '待定损',  -- 待定损/确认维修/确认不修
            price REAL DEFAULT 0, unit TEXT DEFAULT '项'
        );
        CREATE TABLE IF NOT EXISTS repair_quotes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plan_id INTEGER, company TEXT, item_name TEXT DEFAULT '',
            price REAL DEFAULT 0, duration TEXT DEFAULT '',
            status TEXT DEFAULT '报价', created_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS repair_changes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            plan_id INTEGER, add_item TEXT, add_price REAL DEFAULT 0,
            status TEXT DEFAULT '待确认',   -- 待确认/已确认/已拒绝
            created_by TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now','localtime'))
        );
    """)
    # 维修采购审批流(定损审批)默认配置: 角色可后续在系统设置改(厂长账号接入后填)
    if conn.execute("SELECT COUNT(*) FROM approval_flow_config WHERE biz_type='repair_plan'").fetchone()[0] == 0:
        conn.execute("INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label) VALUES('repair_plan',1,'分管领导',0,1000000,'维修定损-1级')")
    # ---- V11.210 设备维修完整工单(设备维修功能需求.docx): 幂等补列(老WP单默认值兼容), 不动物资采购任何表 ----
    _rp = [r[1] for r in conn.execute("PRAGMA table_info(repair_plans)").fetchall()]
    _adds = [
        ('device_no', "TEXT DEFAULT ''"),        # 设备编号
        ('fault_time', "TEXT DEFAULT ''"),       # 故障发生时间
        ('urgency', "TEXT DEFAULT '普通'"),       # 紧急等级 普通/紧急
        ('init_judge', "TEXT DEFAULT ''"),       # 初步判断 整机维修/零部件更换
        ('est_cost', "REAL DEFAULT 0"),          # 提报人预估费用
        ('repair_type', "TEXT DEFAULT ''"),      # 定损结果 内部自修/委外维修/更换新设备
        ('damage_items_json', "TEXT DEFAULT ''"),# 定损清单(故障部位/更换部件)
        ('damage_opinion', "TEXT DEFAULT ''"),   # 定损意见
        ('damage_time', "TEXT DEFAULT ''"),      # 定损时间
        ('internal_note', "TEXT DEFAULT ''"),    # 内部自修处理记录
        ('convert_req_no', "TEXT DEFAULT ''"),   # 转物资采购申请的单号
        ('vendor_selected', "TEXT DEFAULT ''"),  # 比价选定的服务商
        ('quote_files', "TEXT DEFAULT ''"),      # 报价附件
        ('entrust_no', "TEXT DEFAULT ''"),       # 维修委托单号
        ('contract_id', "INTEGER DEFAULT 0"),    # 维修服务合同id
        ('send_req_no', "TEXT DEFAULT ''"),      # 发料出库单号
        ('outer_status', "TEXT DEFAULT ''"),     # 委外状态 在厂/委外维修中/已回厂
        ('start_time', "TEXT DEFAULT ''"),       # 维修开工时间
        ('expect_finish', "TEXT DEFAULT ''"),    # 预计完工
        ('actual_finish', "TEXT DEFAULT ''"),    # 实际完工
        ('accept_result', "TEXT DEFAULT ''"),    # 验收结果 通过/不通过返修
        ('accept_opinion', "TEXT DEFAULT ''"),   # 验收意见
        ('accept_files', "TEXT DEFAULT ''"),     # 测试记录/照片附件
        ('accept_time', "TEXT DEFAULT ''"),      # 验收时间
        ('handle_type', "TEXT DEFAULT ''"),      # 验收后处理 归还部门/回收入库
        ('invoice_no', "TEXT DEFAULT ''"),       # 维修发票号
        ('invoice_amount', "REAL DEFAULT 0"),    # 维修发票金额
        ('pay_status', "TEXT DEFAULT ''"),       # 付款状态
    ]
    for col, ddl in _adds:
        if col not in _rp:
            conn.execute(f"ALTER TABLE repair_plans ADD COLUMN {col} {ddl}")
    # V11.210 维修变更四方确认: 变更单加 四方确认人/时间/结果(提报人/采购专员/机电厂长/机修车间主任), 永久留痕不可删
    _rc = [r[1] for r in conn.execute("PRAGMA table_info(repair_changes)").fetchall()]
    if 'confirm1_by' not in _rc:
        conn.executescript("""
            ALTER TABLE repair_changes ADD COLUMN confirm1_by TEXT DEFAULT '';  -- 报修提报人
            ALTER TABLE repair_changes ADD COLUMN confirm1_at TEXT DEFAULT '';
            ALTER TABLE repair_changes ADD COLUMN confirm2_by TEXT DEFAULT '';  -- 采购专员
            ALTER TABLE repair_changes ADD COLUMN confirm2_at TEXT DEFAULT '';
            ALTER TABLE repair_changes ADD COLUMN confirm3_by TEXT DEFAULT '';  -- 机电厂长
            ALTER TABLE repair_changes ADD COLUMN confirm3_at TEXT DEFAULT '';
            ALTER TABLE repair_changes ADD COLUMN confirm4_by TEXT DEFAULT '';  -- 机修车间主任
            ALTER TABLE repair_changes ADD COLUMN confirm4_at TEXT DEFAULT '';
            ALTER TABLE repair_changes ADD COLUMN change_reason TEXT DEFAULT ''; -- 变更原因
            ALTER TABLE repair_changes ADD COLUMN add_part TEXT DEFAULT '';      -- 新增配件
            ALTER TABLE repair_changes ADD COLUMN add_labor REAL DEFAULT 0;      -- 新增工时费
        """)
    # V11.216 报价表补 配件费/工时费/质保/预计完工 列(幂等)
    _rq = [r[1] for r in conn.execute("PRAGMA table_info(repair_quotes)").fetchall()]
    for _col, _ddl in (('part_cost', 'REAL DEFAULT 0'), ('labor_cost', 'REAL DEFAULT 0'),
                       ('warranty', "TEXT DEFAULT ''"), ('finish_date', "TEXT DEFAULT ''")):
        if _col not in _rq:
            conn.execute(f"ALTER TABLE repair_quotes ADD COLUMN {_col} {_ddl}")
    # ---- V11.196 工作台公告 ----
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS notices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            content TEXT DEFAULT '',        -- 富文本HTML(含内嵌图片引用 /uploads/xxx)
            scope TEXT DEFAULT 'all',       -- 'all'=全员 | 'roles:分管领导' 按角色 | 'users:温丽,穆娇' 按用户
            pinned INTEGER DEFAULT 0,       -- 置顶
            status TEXT DEFAULT '草稿',     -- 草稿/已发布/已撤销
            publisher TEXT DEFAULT '',      -- 发布人
            publish_at TEXT DEFAULT '',     -- 发布时间
            effective_at TEXT DEFAULT '',   -- 生效时间(空=发布即生效)
            expire_at TEXT DEFAULT '',      -- 失效时间(空=永不过期, 到期自动隐藏)
            created_at TEXT DEFAULT (datetime('now','localtime')),
            updated_at TEXT DEFAULT (datetime('now','localtime'))
        );
        CREATE TABLE IF NOT EXISTS notice_action_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            notice_id INTEGER, title TEXT, action TEXT, operator TEXT,
            detail TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now','localtime'))
        );
    """)
    # ---- V5.1 安全加固: 登录审计/系统元数据(幂等) ----
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS login_attempts (
            id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, ip TEXT,
            success INTEGER DEFAULT 0, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS system_meta (
            key TEXT PRIMARY KEY, value TEXT);
    """)
    # V11.175: 驳回历史表 — 每次驳回追加一条(人/时间/理由/来源), 多次驳回不覆盖; 列表/详情/打印导出共用
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS approval_reject_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            biz_type TEXT, biz_id INTEGER,
            approver TEXT, approver_id INTEGER DEFAULT 0,
            comment TEXT, processed_at TEXT,
            source TEXT DEFAULT 'system',
            created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE INDEX IF NOT EXISTS idx_reject_logs ON approval_reject_logs(biz_type, biz_id);
    """)
    # V11.184: 审批流转操作记录表(同意/驳回统一) — 每条审批操作(含上传附件元数据)独立一行, 多次审批不覆盖
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS approval_action_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            biz_type TEXT, biz_id INTEGER,
            action TEXT DEFAULT 'agree',          -- agree=同意 / reject=驳回
            approver TEXT, approver_id INTEGER DEFAULT 0,
            comment TEXT DEFAULT '',              -- 审批意见文字
            processed_at TEXT,
            attachments TEXT DEFAULT '[]',        -- JSON: [{fileName,fileId,spaceId,fileSize,fileType}]
            source TEXT DEFAULT 'system',         -- dingtalk / system
            instance_code TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE INDEX IF NOT EXISTS idx_action_logs ON approval_action_logs(biz_type, biz_id);
    """)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS order_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            order_id INTEGER NOT NULL,
            item_name TEXT, spec TEXT, unit TEXT DEFAULT '个',
            quantity REAL DEFAULT 0, price REAL DEFAULT 0, amount REAL DEFAULT 0,
            tax_rate REAL DEFAULT 13, tax_amount REAL DEFAULT 0, total_amount REAL DEFAULT 0,
            remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS expenses (
            id INTEGER PRIMARY KEY AUTOINCREMENT, expense_no TEXT UNIQUE NOT NULL,
            expense_date TEXT, category TEXT, amount REAL DEFAULT 0, reason TEXT,
            supplier TEXT, invoice_no TEXT, remark TEXT, created_by TEXT,
            created_at TEXT DEFAULT (datetime('now','localtime')));
        CREATE TABLE IF NOT EXISTS settlements (
            id INTEGER PRIMARY KEY AUTOINCREMENT, settlement_no TEXT UNIQUE NOT NULL,
            period TEXT, supplier TEXT, contract_ids TEXT, order_ids TEXT,
            total_amount REAL DEFAULT 0, status TEXT DEFAULT '待确认',
            remark TEXT, created_at TEXT DEFAULT (datetime('now','localtime')), confirmed_at TEXT);
        -- V55 预警中心: 预警记录(首页集中展示/标记已处理/日志留存)
        CREATE TABLE IF NOT EXISTS alert_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT, alert_type TEXT, level TEXT DEFAULT 'orange',
            title TEXT, content TEXT, biz_type TEXT, biz_id INTEGER,
            status TEXT DEFAULT 'pending', created_at TEXT DEFAULT (datetime('now','localtime')),
            processed_at TEXT, processed_by TEXT);
    """)
    conn.commit(); conn.close()

# ============================================================
# ── AUTH API ──
# ============================================================

# V11.161: 品牌AI分析缓存(同品牌24小时只调一次AI, 避免询价详情/列表反复卡10秒)
_BRAND_AI_CACHE = {}
_BRAND_AI_CACHE_TTL = 86400  # 24小时

def ai_analyze_brand(brand_name):
    """V11.115: 调用AI分析品牌优缺点 — V11.161: 加24h内存缓存+超时降至3s(修复详情加载卡死)"""
    _key = (brand_name or '').strip()
    if not _key:
        return None
    _now = time.time()
    if _key in _BRAND_AI_CACHE:
        _cached_at, _cached_val = _BRAND_AI_CACHE[_key]
        if _now - _cached_at < _BRAND_AI_CACHE_TTL:
            return _cached_val
    import json as _json
    import urllib.request as _req
    try:
        # 使用Agnes AI API
        api_url = "https://apihub.agnes-ai.com/v1/chat/completions"
        api_key = "«redacted:sk-…»"  # 从环境变量或配置读取
        
        # 构建提示词
        prompt = f"请简要分析'{brand_name}'品牌的优缺点，各用一句话描述。格式：优点：xxx\n缺点：xxx"
        
        payload = _json.dumps({
            "model": "agnes-2.0-flash",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 200
        }).encode('utf-8')
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        req = _req.Request(api_url, data=payload, headers=headers, method='POST')
        with _req.urlopen(req, timeout=3) as response:
            result = _json.loads(response.read().decode('utf-8'))
        text = result['choices'][0]['message']['content']
        
        # 解析返回结果
        优点 = ""
        缺点 = ""
        for line in text.split('\n'):
            if '优点' in line:
                优点 = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
            elif '缺点' in line:
                缺点 = line.split('：')[-1].strip() if '：' in line else line.split(':')[-1].strip()
        
        _out = {'优点': 优点, '缺点': 缺点}
        _BRAND_AI_CACHE[_key] = (_now, _out)
        return _out
    except Exception as e:
        print(f'[品牌分析AI调用失败] {e}')
        _BRAND_AI_CACHE[_key] = (_now, {'优点': '', '缺点': ''})  # 失败也缓存, 避免反复调
        return None

def search_brand_info(supplier_name, category):
    """V11.91: 简单品牌优缺点分析（本地数据+AI补充）"""
    # 行业常见品牌知识
    brand_knowledge = {
        '长城': {'优点': '国产老牌，性价比高', '缺点': '精度一般'},
        '恒力': {'优点': '民营石化龙头，品质稳定', '缺点': '价格略高'},
        '中石化': {'优点': '央企，质量可靠', '缺点': '交货周期长'},
        '宝钢': {'优点': '国产钢材龙头', '缺点': '价格偏高'},
        '鞍钢': {'优点': '北方钢厂，性价比高', '缺点': '运输距离远'},
        '柳工': {'优点': '国产工程机械龙头', '缺点': '二手保值率一般'},
        '徐工': {'优点': '规模大，服务网点多', '缺点': '价格中等'},
        '三一': {'优点': '创新强', '缺点': '售后需预约'},
        '施耐德': {'优点': '国际品牌，质量稳定', '缺点': '价格高'},
        '西门子': {'优点': '德系品质，可靠性高', '缺点': '价格昂贵'},
        'ABB': {'优点': '电力领域领先', '缺点': '交期长'},
        '海尔': {'优点': '国产家电龙头', '缺点': '工业品线弱'},
        '美的': {'优点': '性价比高', '缺点': '高端线弱'},
        '格力': {'优点': '空调技术强', '缺点': '品类单一'},
        '华为': {'优点': '技术领先', '缺点': '价格高，供货紧'},
        '中兴': {'优点': '性价比高', '缺点': '品牌力弱'},
        '联想': {'优点': '国内PC龙头', '缺点': '高端线弱'},
        '戴尔': {'优点': '商务稳定', '缺点': '价格高'},
        '惠普': {'优点': '外设强', '缺点': 'PC线一般'},
        '金立': {'优点': '备用机', '缺点': '已退市'},
        '得力': {'优点': '办公用品龙头，性价比高', '缺点': '高端线弱'},
        '晨光': {'优点': '文具品牌知名度高', '缺点': '工业品线弱'},
        '坚朗': {'优点': '建筑五金龙头，质量稳定，品类齐全', '缺点': '价格偏高'},
        '悍高': {'优点': '家居五金知名品牌，功能设计好', '缺点': '工业领域应用少'},
        'A': {'优点': '知名品牌', '缺点': '需验厂'},
        'B': {'优点': '知名品牌', '缺点': '需验厂'},
        '测试': {'优点': '测试用', '缺点': '仅用于测试'},
        'a': {'优点': '测试品牌', '缺点': '待验证'},
        'b': {'优点': '测试品牌', '缺点': '待验证'},
        'c': {'优点': '测试品牌', '缺点': '待验证'},
        'aa': {'优点': '测试品牌', '缺点': '待验证'},
        '测试a': {'优点': '测试用', '缺点': '仅用于测试'},
    }
    # 匹配供应商名称（优先匹配更长的品牌名）
    # 按品牌名长度降序排序，避免单字母'A'、'B'误匹配
    for brand, info in sorted(brand_knowledge.items(), key=lambda x: len(x[0]), reverse=True):
        if brand in supplier_name or supplier_name in brand:
            return info
    # 按行业类别推荐
    if category:
        if '钢材' in category or '建材' in category:
            return {'优点': '本地供应', '缺点': '需验厂'}
        if '仪表' in category or '阀门' in category:
            return {'优点': '专业厂家', '缺点': '交期1-2周'}
    # 本地知识库没有，尝试AI分析
    ai_result = ai_analyze_brand(supplier_name)
    if ai_result and (ai_result.get('优点') or ai_result.get('缺点')):
        return ai_result
    return {'优点': '', '缺点': ''}

@app.route('/api/login', methods=['POST'])
def api_login():
    d = request.json or {}
    username = (d.get('username') or '').strip()
    ip = get_client_ip()
    conn = db()
    # V11.138: 取消失败锁定(用户要求) — 不限制连续失败次数, 仅记录审计
    u = conn.execute("SELECT u.*,d.name as dept_name FROM users u LEFT JOIN departments d ON u.dept_id=d.id WHERE u.username=? AND u.is_active=1", (username,)).fetchone()
    if not u or not verify_password(d.get('password',''), u['password']):
        conn.execute("INSERT INTO login_attempts(username,ip,success) VALUES(?,?,0)", (username, ip))
        conn.commit()
        conn.close()
        log('系统', '登录失败', '账号 %s (IP %s) 密码错误' % (username, ip))
        return jsonify({'error': '用户名或密码错误'}), 401
    # 旧MD5密码自动升级为PBKDF2(一次登录即完成迁移)
    if is_legacy_md5(u['password']):
        conn.execute("UPDATE users SET password=? WHERE id=?", (hash_password(d.get('password','')), u['id']))
    conn.execute("DELETE FROM login_attempts WHERE username=? AND ip=?", (username, ip))
    conn.execute("INSERT INTO login_attempts(username,ip,success) VALUES(?,?,1)", (username, ip))
    conn.commit(); conn.close()
    # 会话固定防护: 登录成功后重建会话
    session.clear()
    session['user_id'] = u['id']; session['username'] = u['username']; session['user_name'] = u['name']; session['user_role'] = u['role']
    session['dept_id'] = u['dept_id']; session['dept_name'] = u['dept_name'] or ''
    session['login_time'] = now()
    session.permanent = True  # V8.0: 持久会话 — 刷新保持登录, 8小时不操作才登出
    log('系统', '登录成功', '%s (%s) 登录系统' % (u['name'], u['role']))
    return jsonify({'success':True, 'user':{'id':u['id'],'username':u['username'],'name':u['name'],'role':u['role'],'dept_name':u['dept_name'] or '','can_config': can_manage_config()}})

@app.route('/api/logout', methods=['POST'])
def api_logout():
    log('系统', '退出登录', '%s 退出系统' % session.get('user_name',''))
    session.clear(); return jsonify({'success':True})

@app.route('/api/me')
@login_required
def api_me():
    return jsonify({'id':session['user_id'],'name':session['user_name'],'role':session['user_role'],'dept_id':session['dept_id'],'dept_name':session['dept_name'],'can_config': can_manage_config()})

@app.route('/api/users')
@login_required
def api_users():
    # V11.64: 用户列表仅 管理员/领导/财务 可见(防账号信息泄露)
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理', '财务'):
        return jsonify([])
    conn = db(); rows = conn.execute("SELECT u.*,d.name as dept_name FROM users u LEFT JOIN departments d ON u.dept_id=d.id ORDER BY u.id").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/users', methods=['POST'])
@admin_required
def api_users_create():
    d = request.json or {}
    username = (d.get('username') or '').strip()
    name = (d.get('name') or '').strip()
    role = d.get('role') or '员工'
    if not username or not name:
        return jsonify({'error': '用户名和姓名必填'}), 400
    pw = d.get('password') or '123456'
    if len(pw) < 6:
        return jsonify({'error': '初始密码长度至少6位'}), 400
    conn = db()
    if conn.execute("SELECT 1 FROM users WHERE username=?", (username,)).fetchone():
        conn.close(); return jsonify({'error': '用户名已存在'}), 400
    conn.execute("INSERT INTO users(username,password,name,phone,role,dept_id,title,is_active) VALUES(?,?,?,?,?,?,?,1)",
                 (username, hash_password(pw), name, d.get('phone',''), role, d.get('dept_id'), d.get('title','')))
    conn.commit(); conn.close()
    log(session.get('user_name',''), '新增用户', '创建账号 %s (%s, %s)' % (username, name, role))
    return jsonify({'success': True, 'message': '用户已创建, 初始密码: %s' % pw})

@app.route('/api/users/<int:uid>/reset-password', methods=['POST'])
@admin_required
def api_reset_password(uid):
    d = request.json or {}
    new_pw = d.get('new_password') or ''
    if len(new_pw) < 6:
        return jsonify({'error': '新密码长度至少6位'}), 400
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not u:
        conn.close(); return jsonify({'error': '用户不存在'}), 404
    conn.execute("UPDATE users SET password=? WHERE id=?", (hash_password(new_pw), uid))
    conn.commit(); conn.close()
    log(session.get('user_name',''), '重置密码', '管理员重置用户 %s 的密码' % u['name'])
    return jsonify({'success': True, 'message': '密码已重置'})

@app.route('/api/users/<int:uid>', methods=['PUT'])
@admin_required
def api_user_update(uid):
    """V11.157: 修改用户信息 — 账号/姓名/角色/部门/电话/职务/钉钉绑定/状态 可改"""
    d = request.json or {}
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not u:
        conn.close(); return jsonify({'error': '用户不存在'}), 404
    username = (d.get('username') or u['username'] or '').strip()
    if not username:
        conn.close(); return jsonify({'error': '账号不能为空'}), 400
    # 账号唯一性检查(排除自己)
    _dup = conn.execute("SELECT id FROM users WHERE username=? AND id<>?", (username, uid)).fetchone()
    if _dup:
        conn.close(); return jsonify({'error': '该账号已被其他用户使用'}), 400
    name = (d.get('name') or u['name'] or '').strip()
    if not name:
        conn.close(); return jsonify({'error': '姓名不能为空'}), 400
    role = d.get('role') or u['role'] or '员工'
    # V11.157d: dept_id 只在显式传入时更新(未传=保留原值; 传空串=未分配)
    if 'dept_id' in d:
        if d.get('dept_id') in ('', None):
            dept_id = None
        else:
            try: dept_id = int(d['dept_id'])
            except Exception: dept_id = u['dept_id']
    else:
        dept_id = u['dept_id']
    phone = (d.get('phone') or u['phone'] or '')
    title = (d.get('title') or u['title'] or '')
    # 钉钉绑定: 传空字符串=解绑, 不传=保留原值
    if 'dingtalk_userid' in d:
        dingtalk_userid = (d.get('dingtalk_userid') or '').strip()
    else:
        dingtalk_userid = u['dingtalk_userid'] or ''
    is_active = d.get('is_active')
    if is_active is not None:
        is_active = 1 if (is_active is True or str(is_active) in ('1', 'true')) else 0
    else:
        is_active = u['is_active']
    conn.execute("UPDATE users SET username=?, name=?, role=?, dept_id=?, phone=?, title=?, dingtalk_userid=?, is_active=? WHERE id=?",
                 (username, name, role, dept_id, phone, title, dingtalk_userid, is_active, uid))
    conn.commit(); conn.close()
    log(session.get('user_name',''), '修改用户', '更新用户 %s (%s, 部门%d, 钉钉%s, 状态%d)' % (username, role, dept_id or 0, dingtalk_userid or '无', is_active))
    return jsonify({'success': True, 'message': '用户信息已更新'})

@app.route('/api/users/<int:uid>', methods=['DELETE'])
@admin_required
def api_user_delete(uid):
    """V11.166: 删除用户 — 物理删除users记录
    保护规则: 不能删除自己; 不能删除系统管理员角色(防止权限失控);
    业务单据只存姓名不存user_id, 删除不影响历史单据显示"""
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not u:
        conn.close(); return jsonify({'error': '用户不存在'}), 404
    if uid == session.get('user_id'):
        conn.close(); return jsonify({'error': '不能删除当前登录账号'}), 400
    if u['role'] == '系统管理员':
        conn.close(); return jsonify({'error': '系统管理员账号不能删除(可改为停用)'}), 400
    conn.execute("DELETE FROM users WHERE id=?", (uid,))
    conn.commit(); conn.close()
    log(session.get('user_name',''), '删除用户', '删除账号 %s (%s, %s)' % (u['username'], u['name'], u['role']))
    return jsonify({'success': True, 'message': '用户 %s 已删除' % u['name']})

@app.route('/api/change-password', methods=['POST'])
@login_required
def api_change_password():
    d = request.json or {}
    new_pw = d.get('new_password') or ''
    if len(new_pw) < 6:
        return jsonify({'error': '新密码长度至少6位'}), 400
    if new_pw == d.get('old_password'):
        return jsonify({'error': '新密码不能与旧密码相同'}), 400
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
    if not u or not verify_password(d.get('old_password',''), u['password']):
        conn.close(); return jsonify({'error': '旧密码不正确'}), 400
    conn.execute("UPDATE users SET password=? WHERE id=?", (hash_password(new_pw), u['id']))
    conn.commit(); conn.close()
    log(session.get('user_name',''), '修改密码', '用户 %s 修改了登录密码' % u['name'])
    return jsonify({'success': True, 'message': '密码已更新'})

@app.route('/api/departments')
@login_required
def api_departments():
    conn = db(); rows = conn.execute("SELECT * FROM departments").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

# ============================================================
# ── APPROVAL FLOW HELPER ──
# ============================================================
def get_approval_config(biz_type, amount):
    conn = db()
    rows = conn.execute("SELECT * FROM approval_flow_config WHERE biz_type=? AND ?>=min_amount AND ?<=max_amount ORDER BY level_no", (biz_type, amount, amount)).fetchall()
    conn.close()
    return rows

def create_approvals(biz_type, biz_id, amount, submitter=''):
    """V5.0: 按审批流配置生成审批实例
    - 节点配置了具体审批人(approver=用户名) → 绑定该用户
    - 未配置(留空) → 按角色在 users 表找有效用户
    - V11.183: submitter参数预留(提交人) — 钉钉模板固定审批人模式下不校验发起人=审批人,
      且穆娇自提交的单由模板固定审批人正常处理(实测发起成功), 因此不做自动换人, 保持系统与钉钉审批人一致
    支持每个环节独立配置/更换审批负责人, 页面可视化维护, 无需改代码"""
    configs = get_approval_config(biz_type, amount)
    conn = db()
    for cfg in configs:
        approver_name = ''
        approver_id = None
        if cfg['approver'] and str(cfg['approver']).strip():
            u = conn.execute("SELECT * FROM users WHERE username=? AND is_active=1", (str(cfg['approver']).strip(),)).fetchone()
            if u:
                approver_name = u['name'] or u['username']
                approver_id = u['id']
        if approver_id is None:
            # 按角色自动找有效用户
            u2 = conn.execute("SELECT * FROM users WHERE role=? AND is_active=1 ORDER BY id LIMIT 1", (cfg['role'],)).fetchone()
            if u2:
                approver_name = u2['name'] or u2['username']
                approver_id = u2['id']
        conn.execute("INSERT INTO approval_instances(biz_type,biz_id,level_no,role,approver,approver_id) VALUES(?,?,?,?,?,?)",
                     (biz_type, biz_id, cfg['level_no'], cfg['role'], approver_name, approver_id))
    conn.commit(); conn.close()

def check_all_approved(biz_type, biz_id):
    conn = db()
    r = conn.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending'", (biz_type, biz_id)).fetchone()[0]
    conn.close()
    return r == 0

def do_approve(biz_type, biz_id, approver, approver_id, action='approved', comment='', signature=''):
    """V7.0: 审批操作权限锁死 — 仅当前节点指定审批人/系统管理员可同意驳回
    - 当前节点审批人 = 审批实例 approver(节点配置指定的人) 或 角色对应有效用户
    - 其余用户(含其他角色领导)一律拒绝; 钉钉端权限与系统完全同步
    - V11.2: signature 电子签名(dataURL图片)随审批记录保存"""
    conn = db()
    cur = conn.execute("SELECT * FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no LIMIT 1", (biz_type, biz_id)).fetchone()
    if not cur:
        conn.close(); return {'success':False,'error':'无待审批节点'}
    # ── 权限校验(7.0优化2): 仅当前节点指定审批人/系统管理员 ──
    me = conn.execute("SELECT * FROM users WHERE id=?", (approver_id,)).fetchone() if approver_id else None
    is_admin = me and me['role'] == '系统管理员'
    node_approver_id = cur['approver_id'] if cur['approver_id'] else None
    if not node_approver_id:
        # 审批实例未绑定具体人(按角色) → 解析角色对应有效用户
        u = conn.execute("SELECT * FROM users WHERE role=? AND is_active=1 ORDER BY id LIMIT 1", (cur['role'],)).fetchone()
        node_approver_id = u['id'] if u else None
    if not is_admin and (not me or me['id'] != node_approver_id):
        conn.close()
        return {'success':False,'error':'仅当前审批节点指定审批人可操作'}
    # 签名限制: 图片dataURL太长(可达几十KB), 截断保护
    sig = (signature or '').strip()
    if sig and not sig.startswith('data:image/'):
        sig = ''
    if len(sig) > 200000:
        sig = sig[:200000]
    conn.execute("UPDATE approval_instances SET status=?, approver=?, approver_id=?, comment=?, processed_at=?, signature=? WHERE id=?",
                 (action, approver, approver_id, comment, now(), sig, cur['id']))
    conn.commit(); conn.close()
    return {'success':True}

# ============================================================
# V4.1 ── 飞书对接: 审批同步 + 机器人预警 (零依赖, urllib)
# ============================================================
FS_API = 'https://open.feishu.cn/open-apis'
FS_BIZ = {  # biz_type -> (审批定义名称, 表单控件前缀)
    'purchase_request': ('采购申请审批', 'SQ'),
    'purchase_order':   ('采购订单审批', 'CG'),
    'contract':         ('合同审批', 'HT'),
    'credit':           ('挂账审批', 'GZ'),
    'payment':          ('付款审批', 'FK'),
    'receiving':        ('入库审批', 'RK'),
    'requisition':      ('出库审批', 'CK'),
}
FS_PRE = {'purchase_request': 'SQ', 'purchase_order': 'CG', 'contract': 'HT', 'credit': 'GZ', 'payment': 'FK', 'receiving': 'RK', 'requisition': 'CK'}
FS_NODE_MAX = 3  # 审批定义中的审批节点数(与系统链路最大级数一致)

def cfg_get(key, default=''):
    c = db()
    try: r = c.execute("SELECT value FROM sys_config WHERE key=?", (key,)).fetchone()
    except Exception: r = None
    c.close()
    return r['value'] if r else default

def cfg_set(key, value):
    c = db()
    c.execute("INSERT INTO sys_config(key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, str(value)))
    c.commit(); c.close()

def feishu_enabled():
    return cfg_get('feishu_enabled') == '1'

_TOKEN = {'t': None, 'exp': 0}
def fs_token():
    if _TOKEN['t'] and time.time() < _TOKEN['exp'] - 120: return _TOKEN['t']
    app_id = cfg_get('feishu_app_id'); app_secret = cfg_get('feishu_app_secret')
    if not app_id or not app_secret: raise RuntimeError('飞书应用未配置: 请先在"飞书设置"填写 app_id / app_secret')
    code, resp = fs_post('/auth/v3/tenant_access_token/internal', {'app_id': app_id, 'app_secret': app_secret}, auth=False)
    if code != 0: raise RuntimeError(f'获取tenant_access_token失败(code={code}): {resp.get("msg","")}')
    _TOKEN['t'] = resp.get('tenant_access_token'); _TOKEN['exp'] = time.time() + int(resp.get('expire', 7200))
    return _TOKEN['t']

def fs_post(path, payload, auth=True, timeout=12):
    body = json.dumps(payload).encode('utf-8')
    headers = {'Content-Type': 'application/json; charset=utf-8'}
    if auth:
        try: headers['Authorization'] = 'Bearer ' + fs_token()
        except Exception as e: return -1, {'msg': str(e)}
    req = urllib.request.Request(FS_API + path, data=body, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            d = json.loads(r.read().decode('utf-8'))
        return d.get('code', -1), d
    except Exception as e:
        return -1, {'msg': f'网络错误: {e}'}

def fs_get(path, timeout=12):
    headers = {}
    try: headers['Authorization'] = 'Bearer ' + fs_token()
    except Exception as e: return -1, {'msg': str(e)}
    req = urllib.request.Request(FS_API + path, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            d = json.loads(r.read().decode('utf-8'))
        return d.get('code', 0), d
    except Exception as e:
        return -1, {'msg': f'网络错误: {e}'}

def find_user_by_role(role):
    """审批角色 → 具体用户(严格校验): 必须 is_active=1 有效用户
    V5.0+: 不再静默回退 admin 顶替 — 角色无有效用户时返回 None, 由调用方决定发起失败"""
    c = db()
    u = c.execute("SELECT * FROM users WHERE role=? AND is_active=1 ORDER BY id LIMIT 1", (role,)).fetchone()
    c.close()
    return u

def find_approver_for_role(role, channel='dingtalk'):
    """钉钉/飞书审批发起时的审批人解析(严格): 有效用户 + 已绑定对应通道ID
    返回 None 表示该角色无满足条件的审批人(需在通讯录/系统配置中补齐)"""
    c = db()
    if channel == 'dingtalk':
        u = c.execute("SELECT * FROM users WHERE role=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!='' ORDER BY id LIMIT 1", (role,)).fetchone()
    else:
        u = c.execute("SELECT * FROM users WHERE role=? AND is_active=1 AND feishu_open_id IS NOT NULL AND feishu_open_id!='' ORDER BY id LIMIT 1", (role,)).fetchone()
    c.close()
    return u

def find_user_by_name(name):
    c = db()
    u = c.execute("SELECT * FROM users WHERE name=? AND is_active=1 ORDER BY id LIMIT 1", (name,)).fetchone()
    c.close()
    return u

def conn_check_user(username):
    """审批流配置时校验: 指定审批人用户名存在且为有效用户; 返回 users Row 或 None"""
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE username=? AND is_active=1", (username,)).fetchone()
    conn.close()
    return u

def admin_open_id():
    c = db()
    u = c.execute("SELECT * FROM users WHERE feishu_open_id IS NOT NULL AND feishu_open_id!='' ORDER BY id LIMIT 1").fetchone()
    c.close()
    return u['feishu_open_id'] if u else ''

def fs_approval_codes():
    try: return json.loads(cfg_get('feishu_approval_codes', '{}'))
    except Exception: return {}

def fs_approval_code(biz_type):
    return fs_approval_codes().get(biz_type, '')

# ---- 审批定义: 一键创建(幂等) ----
def fs_form_controls(biz):
    pre = FS_PRE[biz]
    return [
        {'id': pre+'_no', 'label': {'text': '单据编号'}, 'type': 'input', 'required': True, 'placeholder': '如 SQ-202608-0001'},
        {'id': pre+'_name', 'label': {'text': '内容摘要'}, 'type': 'input', 'required': True},
        {'id': pre+'_amount', 'label': {'text': '金额(元)'}, 'type': 'number', 'required': True},
        {'id': pre+'_applicant', 'label': {'text': '申请人'}, 'type': 'input', 'required': True},
        {'id': pre+'_date', 'label': {'text': '提交时间'}, 'type': 'input', 'required': False},
    ]

def fs_init_definitions():
    """调用飞书API创建5个审批定义(幂等); 若某定义已存在则沿用其代码"""
    ao = admin_open_id()
    if not ao: return {'success': False, 'error': '请先在②用户飞书绑定中绑定至少一个用户的 open_id'}
    out = {}
    for biz, (title, pre) in FS_BIZ.items():
        ac = 'zhengcheng_' + biz
        code_r, resp = fs_post('/approval/v4/approvals/create', {
            'approval_name': title,
            'approval_code': ac,
            'description': '正成能源智慧采购系统 - ' + title,
            'form': {'form_controls': fs_form_controls(biz)},
            'node_list': [
                {'id': 'node_1', 'name': '审批节点1', 'type': 'APPROVER', 'approver': {'type': 'FIXED', 'ids': [ao]}},
                {'id': 'node_2', 'name': '审批节点2', 'type': 'APPROVER', 'approver': {'type': 'FIXED', 'ids': [ao]}},
                {'id': 'node_3', 'name': '审批节点3', 'type': 'APPROVER', 'approver': {'type': 'FIXED', 'ids': [ao]}},
            ],
        })
        if code_r == 0:
            got = resp.get('data', {}).get('approval', {}).get('approval_code') or ac
            out[biz] = got
        elif '重复' in str(resp.get('msg', '')) or code_r in (1060012, 1060013, 1060014):
            out[biz] = ac  # 已存在, 沿用
        else:
            out[biz] = f'ERROR:{resp.get("msg","")}'
    codes = {k: v for k, v in out.items() if not str(v).startswith('ERROR')}
    if codes: cfg_set('feishu_approval_codes', json.dumps(codes, ensure_ascii=False))
    return {'success': True, 'results': out}

# ---- 单据信息 -> 飞书审批表单 ----
def fs_biz_info(biz_type, biz_id):
    c = db()
    r = None
    if biz_type == 'purchase_request':
        r = c.execute("SELECT * FROM purchase_requests WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['req_no'], r['purpose'] or r['req_no'], r['total_estimated'], r['requester'] or '系统', r['created_at'])
    if biz_type == 'purchase_order':
        r = c.execute("SELECT * FROM purchase_orders WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['order_no'], f"{r['item_name']} {r['spec'] or ''}".strip(), r['total_amount'], r['owner'] or '系统', r['created_at'])
    if biz_type == 'contract':
        r = c.execute("SELECT * FROM contracts WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['contract_no'], r['contract_name'] or r['contract_no'], r['amount'], '系统', r['created_at'])
    if biz_type == 'credit':
        r = c.execute("SELECT * FROM credit_notes WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['credit_no'], r['item_name'] or r['credit_no'], r['amount'], '系统', r['created_at'])
    if biz_type == 'payment':
        r = c.execute("SELECT * FROM payment_requests WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['payment_no'], f"{r['payment_type']} {r['supplier'] or ''}".strip(), r['amount'], '系统', r['created_at'])
    if biz_type == 'inquiry_approval':
        r = c.execute("SELECT * FROM inquiries WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['inq_no'], r['title'] or r['inq_no'], 0, r['created_by'] or '', '')
    # V11.172: 修复缺少 if 条件的 receiving 分支 — 原本无条件查 receivings 表,
    # 导致 requisition 等类型走到这里查不到记录返回 None, 钉钉审批永远发不出去
    if biz_type == 'receiving':
        r = c.execute("SELECT * FROM receivings WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['receive_no'], r['item_name'] or r['receive_no'], r['quantity'] or 0, r['inspector'] or '系统', r['created_at'])
    if biz_type == 'requisition':
        r = c.execute("SELECT * FROM requisitions WHERE id=?", (biz_id,)).fetchone()
        c.close()
        if not r: return None
        return (r['req_no'], r['item_name'] or r['req_no'], r['quantity'] or 0, r['requester'] or '系统', r['created_at'])
    c.close(); return None

def fs_start_instance(biz_type, biz_id):
    """单据进入待审批后, 同步发起飞书审批实例; 成功返回instance_code, 失败返回None"""
    try:
        if not feishu_enabled(): return None
        ac = fs_approval_code(biz_type)
        if not ac: return None
        c = db()
        if c.execute("SELECT COUNT(*) FROM feishu_instances WHERE biz_type=? AND biz_id=? AND status NOT IN ('error','cancelled')", (biz_type, biz_id)).fetchone()[0] > 0:
            c.close(); return None
        levels = c.execute("SELECT DISTINCT role FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no", (biz_type, biz_id)).fetchall()
        c.close()
        if not levels: return None
        approvers, missing = [], []
        for lv in levels:
            # V5.0+: 严格校验 — 审批人必须是 有效用户(is_active=1) 且已绑定飞书通讯录成员ID
            u = find_approver_for_role(lv['role'], 'feishu')
            if u and u['feishu_open_id']: approvers.append(u['feishu_open_id'])
            else: missing.append(lv['role'])
        if missing:
            log('系统', '飞书审批未发起', f"{biz_type}#{biz_id} 以下角色无有效审批人(未启用或未绑定飞书通讯录): {','.join(missing)}")
            return None
        info = fs_biz_info(biz_type, biz_id)
        if not info: return None
        pre = FS_PRE[biz_type]
        form = [
            {'id': pre+'_no', 'value': str(info[0])},
            {'id': pre+'_name', 'value': str(info[1])},
            {'id': pre+'_amount', 'value': str(info[2])},
            {'id': pre+'_applicant', 'value': str(info[3])},
            {'id': pre+'_date', 'value': str(info[4] or '')[:16]},
        ]
        initiator = admin_open_id()
        ua = find_user_by_name(str(info[3]))
        if ua and ua['feishu_open_id']: initiator = ua['feishu_open_id']
        if not initiator: initiator = approvers[0]
        node_list = [[o] for o in approvers]
        payload = {
            'approval_code': ac, 'open_id': initiator, 'form': form,
            'node_approver_openid_list': node_list,
            'uuid': f'zc-{biz_type}-{biz_id}',
        }
        code_r, resp = fs_post('/approval/v4/instances', payload)
        if code_r != 0 and len(node_list) < FS_NODE_MAX:
            # 节点数与定义不一致时, 补齐到定义节点数(多余节点归最后审批人)重试
            node_list = node_list + [node_list[-1]] * (FS_NODE_MAX - len(node_list))
            payload['node_approver_openid_list'] = node_list
            code_r, resp = fs_post('/approval/v4/instances', payload)
        if code_r == 0:
            inst = resp.get('data', {}).get('instance_code', '')
            c = db()
            c.execute("INSERT INTO feishu_instances(instance_code,biz_type,biz_id,status) VALUES(?,?,?,'pending')", (inst, biz_type, biz_id))
            c.commit(); c.close()
            log('系统', '发起飞书审批', f"{info[0]} 审批人{len(approvers)}级")
            fs_send(approvers[0], f"📋 新的{FS_BIZ[biz_type][0]}待处理\n{info[0]} {info[1]}\n金额 **¥{float(info[2] or 0):,.0f}**\n请前往飞书【审批】应用处理", 'blue')
            return inst
        c = db()
        c.execute("INSERT INTO feishu_instances(instance_code,biz_type,biz_id,status,error) VALUES(?,?,?,'error',?)",
                  (f'ERR-{biz_type}-{biz_id}-{int(time.time())}', biz_type, biz_id, json.dumps(resp, ensure_ascii=False)[:500]))
        c.commit(); c.close()
        log('系统', '飞书审批发起失败', f"{biz_type}#{biz_id}: {json.dumps(resp, ensure_ascii=False)[:200]}")
        return None
    except Exception as e:
        log('系统', '飞书审批发起异常', f"{biz_type}#{biz_id}: {e}")
        return None

# ---- 机器人消息 ----
def fs_send(target, text, color='blue', id_type='open_id'):
    try:
        card = {
            'config': {'wide_screen_mode': True},
            'header': {'template': color, 'title': {'tag': 'plain_text', 'content': '采购系统提醒'}},
            'elements': [
                {'tag': 'div', 'text': {'tag': 'lark_md', 'content': text}},
                {'tag': 'hr'},
                {'tag': 'note', 'elements': [{'tag': 'plain_text', 'content': '正成能源智慧采购系统'}]},
            ],
        }
        code_r, resp = fs_post(f'/im/v1/messages?receive_id_type={id_type}', {
            'receive_id': target, 'msg_type': 'interactive', 'content': json.dumps(card, ensure_ascii=False),
        })
        return code_r == 0
    except Exception:
        return False

# ---- 审批结果同步(幂等) ----
# V11.213: 集中式单据号/名称/金额 CASE 片段(待办/审批中心/工作台/超时预警共用)
# 覆盖全部审批类型, 新增业务类型必须在此登记, 防止待办列表 biz_no/biz_name 为 NULL 显示空
def _ap_case(field):
    """field: 'no' | 'name' | 'amount' — 返回按 biz_type 映射到各业务表的 CASE WHEN 片段(别名 ai)"""
    no = ("CASE WHEN ai.biz_type='purchase_request' THEN (SELECT pr.req_no FROM purchase_requests pr WHERE pr.id=ai.biz_id)"
          " WHEN ai.biz_type='purchase_order' THEN (SELECT po.order_no FROM purchase_orders po WHERE po.id=ai.biz_id)"
          " WHEN ai.biz_type='contract' THEN (SELECT ct.contract_no FROM contracts ct WHERE ct.id=ai.biz_id)"
          " WHEN ai.biz_type='credit' THEN (SELECT cn.credit_no FROM credit_notes cn WHERE cn.id=ai.biz_id)"
          " WHEN ai.biz_type='payment' THEN (SELECT pp.payment_no FROM payment_requests pp WHERE pp.id=ai.biz_id)"
          " WHEN ai.biz_type='receiving' THEN (SELECT rv.receive_no FROM receivings rv WHERE rv.id=ai.biz_id)"
          " WHEN ai.biz_type='requisition' THEN (SELECT rq.req_no FROM requisitions rq WHERE rq.id=ai.biz_id)"
          " WHEN ai.biz_type='return_request' THEN (SELECT rt.return_no FROM return_requests rt WHERE rt.id=ai.biz_id)"
          " WHEN ai.biz_type='repair_plan' THEN (SELECT rp.plan_no FROM repair_plans rp WHERE rp.id=ai.biz_id)"
          " WHEN ai.biz_type='collect_accept' THEN (SELECT rv.receive_no FROM receivings rv WHERE rv.id=ai.biz_id)"
          " WHEN ai.biz_type='inquiry_approval' THEN (SELECT iq.inq_no FROM inquiries iq WHERE iq.id=ai.biz_id)"
          " ELSE '' END")
    name = ("CASE WHEN ai.biz_type='purchase_request' THEN (SELECT pr.purpose FROM purchase_requests pr WHERE pr.id=ai.biz_id)"
            " WHEN ai.biz_type='purchase_order' THEN (SELECT po.item_name FROM purchase_orders po WHERE po.id=ai.biz_id)"
            " WHEN ai.biz_type='contract' THEN (SELECT ct.contract_name FROM contracts ct WHERE ct.id=ai.biz_id)"
            " WHEN ai.biz_type='credit' THEN (SELECT cn.item_name FROM credit_notes cn WHERE cn.id=ai.biz_id)"
            " WHEN ai.biz_type='payment' THEN (SELECT pp.payment_reason FROM payment_requests pp WHERE pp.id=ai.biz_id)"
            " WHEN ai.biz_type='receiving' THEN (SELECT rv.item_name FROM receivings rv WHERE rv.id=ai.biz_id)"
            " WHEN ai.biz_type='requisition' THEN (SELECT rq.item_name FROM requisitions rq WHERE rq.id=ai.biz_id)"
            " WHEN ai.biz_type='return_request' THEN (SELECT rt.reason FROM return_requests rt WHERE rt.id=ai.biz_id)"
            " WHEN ai.biz_type='repair_plan' THEN (SELECT rp.device_name FROM repair_plans rp WHERE rp.id=ai.biz_id)"
            " WHEN ai.biz_type='collect_accept' THEN (SELECT rv.item_name FROM receivings rv WHERE rv.id=ai.biz_id)"
            " WHEN ai.biz_type='inquiry_approval' THEN (SELECT iq.purpose FROM inquiries iq WHERE iq.id=ai.biz_id)"
            " ELSE '' END")
    amount = ("CASE WHEN ai.biz_type='purchase_request' THEN (SELECT pr.total_estimated FROM purchase_requests pr WHERE pr.id=ai.biz_id)"
              " WHEN ai.biz_type='contract' THEN (SELECT ct.amount FROM contracts ct WHERE ct.id=ai.biz_id)"
              " WHEN ai.biz_type='receiving' THEN (SELECT rv.quantity FROM receivings rv WHERE rv.id=ai.biz_id)"
              " WHEN ai.biz_type='requisition' THEN (SELECT rq.quantity FROM requisitions rq WHERE rq.id=ai.biz_id)"
              " WHEN ai.biz_type='repair_plan' THEN (SELECT rp.est_cost FROM repair_plans rp WHERE rp.id=ai.biz_id)"
              " ELSE 0 END")
    return {'no': no, 'name': name, 'amount': amount}[field]
def biz_parent_status(biz_type, result):
    m = {
        'purchase_request': ('已通过', '已驳回'), 'purchase_order': ('审批通过', '已驳回'),
        'contract': ('执行中', '已驳回'), 'credit': ('已通过', '已驳回'), 'payment': ('已通过', '已驳回'),
        'receiving': ('已入库', '已驳回'), 'requisition': ('已出库', '已驳回'),
        'return_request': ('审批通过', '已驳回'),  # V11.193 退库: 审批通过=待仓库清点入库(库存不立即加)
        'repair_plan': ('审批通过', '审批驳回'),  # V11.210 维修金额分级审批: 通过=审批通过(可录报价), 驳回=审批驳回(退回定损)
    }
    ok, no = m.get(biz_type, ('已通过', '已驳回'))
    return ok if result == 'ok' else no

def biz_table(biz_type):
    return {'purchase_request': 'purchase_requests', 'purchase_order': 'purchase_orders',
            'contract': 'contracts', 'credit': 'credit_notes', 'payment': 'payment_requests',
            'receiving': 'receivings', 'requisition': 'requisitions',
            'return_request': 'return_requests',  # V11.193 退库
            'collect_accept': 'receivings',  # V11.206 集体验收: 父单据=入库单
            'repair_plan': 'repair_plans',  # V11.208 维修采购
            'inquiry_approval': 'inquiries'}[biz_type]  # V11.133: biz_id=询价单id

# ============================================================
# V11.64: 数据级权限 — 按角色过滤列表数据(前端隐藏菜单+后端过滤数据, 双保险)
# 规则: 管理员/领导全看; 采购员看采购相关; 库管员看库房相关; 财务看单据+财务
# ============================================================
def can_see_all():
    """管理员/领导 可看全部数据"""
    return session.get('user_role') in ('系统管理员', '分管领导', '总经理')

def filter_scope(role):
    """返回角色可看的业务域: full=全部 / buy=采购 / stock=库房 / fin=财务 / own=仅自己的"""
    if role in ('系统管理员', '分管领导', '总经理'):
        return 'full'
    if role == '采购员':
        return 'buy'
    if role == '库管员':
        return 'stock'
    if role == '财务':
        return 'fin'
    return 'own'  # 员工/部门负责人

def can_see_price():
    """价格可见: 管理员/领导/财务/采购员(谈价需要), 库管员/员工/部门负责人不可见"""
    return session.get('user_role') in ('系统管理员', '分管领导', '总经理', '财务', '采购员')

def mask_price(d):
    """价格脱敏: 无权限的角色, 金额字段置 None(前端显示 ***)"""
    for k in ('total_estimated', 'total_amount', 'price', 'amount', 'est_amount', 'tax_amount', 'quote_price', 'estimated_price'):
        if k in d and d[k] is not None:
            d[k] = None
    return d

def _scope_where(alias):
    """按角色返回 SQL WHERE 片段(过滤数据范围):
    own=只自己的(requester_id=当前用户) / 其余返回空(不过滤, 由调用方决定)"""
    role = session.get('user_role')
    if role in ('员工',):
        return f"{alias}requester_id={session.get('user_id', 0)}"
    if role == '部门负责人':
        return f"{alias}requester_id={session.get('user_id', 0)}"  # 部门负责人仅看自己(暂简化, 后续按部门)
    return ''  # 管理员/领导/采购/库管/财务 由各接口决定(业务域不同)

def log_approval_action(biz_type, biz_id, action, approver, approver_id=0, comment='', processed_at='', attachments=None, source='system', instance_code='', conn=None):
    """V11.184: 审批流转操作记录(同意/驳回统一) — 每条审批操作独立一行, 附件元数据随行存, 多次审批不覆盖。
    action: agree/reject; attachments: [{fileName,fileId,spaceId,fileSize,fileType}]
    钉钉审批(来源dingtalk, 含实例号)与系统审批(来源system)都走这里; 系统驳回同时写 approval_reject_logs(兼容旧逻辑)"""
    try:
        _att = json.dumps(attachments or [], ensure_ascii=False) if attachments else '[]'
        _pt = processed_at or now()
        if conn is not None:
            conn.execute("INSERT INTO approval_action_logs(biz_type,biz_id,action,approver,approver_id,comment,processed_at,attachments,source,instance_code) VALUES(?,?,?,?,?,?,?,?,?,?)",
                         (biz_type, biz_id, action, str(approver or '')[:50], int(approver_id or 0),
                          str(comment or '')[:500], _pt, _att, source, str(instance_code or '')[:100]))
        else:
            c = db()
            c.execute("INSERT INTO approval_action_logs(biz_type,biz_id,action,approver,approver_id,comment,processed_at,attachments,source,instance_code) VALUES(?,?,?,?,?,?,?,?,?,?)",
                      (biz_type, biz_id, action, str(approver or '')[:50], int(approver_id or 0),
                       str(comment or '')[:500], _pt, _att, source, str(instance_code or '')[:100]))
            c.commit(); c.close()
    except Exception:
        pass


def get_approval_action_logs(biz_type, biz_id):
    """V11.184: 读取审批流转操作记录(时间正序), 附件JSON已解析"""
    try:
        c = db()
        rows = c.execute("SELECT * FROM approval_action_logs WHERE biz_type=? AND biz_id=? ORDER BY id ASC", (biz_type, biz_id)).fetchall()
        c.close()
        out = []
        for r in rows:
            d = dict_row(r)
            try:
                d['attachments'] = json.loads(d.get('attachments') or '[]') if d.get('attachments') else []
            except Exception:
                d['attachments'] = []
            out.append(d)
        return out
    except Exception:
        return []


def _src_of(approver, approver_id):
    """V11.175: 判定驳回来源 — 钉钉轮询回调(approver=钉钉或approver_id=0)标dingtalk; 系统审批标system"""
    if str(approver) == '钉钉' or not int(approver_id or 0):
        return 'dingtalk'
    return 'system'

def log_reject(biz_type, biz_id, approver, approver_id, comment, source='system', conn=None, attachments=None, instance_code=''):
    """V11.175: 追加驳回历史(每次驳回一条, 不覆盖) + 同步父单据rejected_reason(取最新)
    V11.181: attachments(钉钉驳回附件元数据JSON)/instance_code(钉钉实例号) 一并存储
    conn: 复用调用方连接(避免SQLite写锁冲突); 无则自开连接"""
    try:
        _att_json = json.dumps(attachments or [], ensure_ascii=False) if attachments else ''
        if conn is not None:
            conn.execute("INSERT INTO approval_reject_logs(biz_type,biz_id,approver,approver_id,comment,processed_at,source,attachments,instance_code) VALUES(?,?,?,?,?,?,?,?,?)",
                         (biz_type, biz_id, str(approver or '')[:50], int(approver_id or 0), str(comment or '')[:500], now(), source, _att_json, str(instance_code or '')[:100]))
            try:
                # V11.175c: 父单据同步最新驳回人/时间/理由(列表独立驳回列展示)
                conn.execute(f"UPDATE {biz_table(biz_type)} SET rejected_reason=?, rejected_by=?, rejected_at=? WHERE id=?",
                             (str(comment or '')[:200], str(approver or '')[:50], now(), biz_id))
            except Exception:
                try:
                    conn.execute(f"UPDATE {biz_table(biz_type)} SET rejected_reason=? WHERE id=?", (str(comment or '')[:200], biz_id))
                except Exception:
                    pass
            return
        c = db()
        c.execute("INSERT INTO approval_reject_logs(biz_type,biz_id,approver,approver_id,comment,processed_at,source,attachments,instance_code) VALUES(?,?,?,?,?,?,?,?,?)",
                  (biz_type, biz_id, str(approver or '')[:50], int(approver_id or 0), str(comment or '')[:500], now(), source, _att_json, str(instance_code or '')[:100]))
        # 父单据 rejected_reason 同步为最新驳回理由(展示用; 完整历史查 reject_logs)
        try:
            c.execute(f"UPDATE {biz_table(biz_type)} SET rejected_reason=?, rejected_by=?, rejected_at=? WHERE id=?",
                      (str(comment or '')[:200], str(approver or '')[:50], now(), biz_id))
        except Exception:
            try:
                c.execute(f"UPDATE {biz_table(biz_type)} SET rejected_reason=? WHERE id=?", (str(comment or '')[:200], biz_id))
            except Exception:
                pass
        c.commit(); c.close()
    except Exception:
        pass

def get_reject_logs(biz_type, biz_id):
    """V11.175: 读取驳回历史(按时间正序); V11.181: 含附件元数据+钉钉实例号"""
    try:
        c = db()
        rows = c.execute("SELECT approver, approver_id, comment, processed_at, source, attachments, instance_code FROM approval_reject_logs WHERE biz_type=? AND biz_id=? ORDER BY id ASC", (biz_type, biz_id)).fetchall()
        c.close()
        out = []
        for r in rows:
            d = dict_row(r)
            try:
                d['attachments'] = json.loads(d.get('attachments') or '[]') if d.get('attachments') else []
            except Exception:
                d['attachments'] = []
            out.append(d)
        return out
    except Exception:
        return []

def finish_approvals(biz_type, biz_id, result='ok', approver='飞书', approver_id=0, comment='飞书审批同步', attachments=None, instance_code='', rejected_items=None):
    """result: 'ok'=通过, 'reject'=驳回; 同步更新审批节点/父单据/飞书实例状态 (幂等)
    V11.181: attachments(驳回附件元数据)/instance_code(钉钉实例号) 随驳回记录存储
    场景覆盖: ①系统内逐级审批(每过一级pending减1, 全过才置父状态)
    ②飞书回调(回调时节点全pending, 一次性全部通过) ③驳回(全部pending置rejected)"""
    c = db()
    pending = c.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending'", (biz_type, biz_id)).fetchone()[0]
    approved = c.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='approved'", (biz_type, biz_id)).fetchone()[0]
    if result == 'ok':
        if pending > 0:
            if approved == 0:
                # 飞书回调场景: 节点全pending但审批已全过 → 一次性全部置approved
                c.execute("UPDATE approval_instances SET status='approved', approver=?, approver_id=?, comment=?, processed_at=? WHERE biz_type=? AND biz_id=? AND status='pending'",
                          (approver, approver_id, comment, now(), biz_type, biz_id))
                # V11.184: 同意操作记录(含钉钉上传附件)
                log_approval_action(biz_type, biz_id, 'agree', approver, approver_id, comment or '', now(),
                                    attachments or [], _src_of(approver, approver_id), instance_code, conn=c)
            else:
                c.close(); return False  # 系统内逐级审批中, 还有待审节点, 不改父状态
        else:
            # V11.184: 末节点同意(系统逐级审批最后一关) — 记录操作
            log_approval_action(biz_type, biz_id, 'agree', approver, approver_id, comment or '', now(),
                                attachments or [], _src_of(approver, approver_id), instance_code, conn=c)
        st = biz_parent_status(biz_type, 'ok')
    else:
        c.execute("UPDATE approval_instances SET status='rejected', approver=?, approver_id=?, comment=?, processed_at=? WHERE biz_type=? AND biz_id=? AND status='pending'",
                  (approver, approver_id, comment, now(), biz_type, biz_id))
        # V11.175: 追加驳回历史(人/时间/理由/来源), 多次驳回不覆盖; 复用本连接避免写锁冲突
        # V11.181: 附件元数据+钉钉实例号随驳回记录存储
        log_reject(biz_type, biz_id, approver, approver_id, comment, source=_src_of(approver, approver_id), conn=c,
                   attachments=attachments or [], instance_code=instance_code)
        # V11.184: 驳回操作记录(统一日志, 含附件)
        log_approval_action(biz_type, biz_id, 'reject', approver, approver_id, comment or '', now(),
                            attachments or [], _src_of(approver, approver_id), instance_code, conn=c)
        # V11.186: 驳回后单据回到"草稿"(可编辑后重新提交) — 非终态'已驳回'
        st = '草稿'
    # V11.206: 集体验收审批 — 独立处理: 父单据=receivings, 通过只置 collect_status(不动 status 状态机, 由常规入库审批继续流转)
    if biz_type == 'collect_accept':
        if result == 'ok':
            c.execute("UPDATE receivings SET collect_status='已集体验收', updated_at=? WHERE id=?", (now(), biz_id))
        else:
            c.execute("UPDATE receivings SET collect_accept=0, collect_status='', updated_at=? WHERE id=?", (now(), biz_id))
        # 集体验收通过/驳回均通知提交人(采购员) — 通过后可继续走入库审批
        try:
            _rn = c.execute("SELECT * FROM receivings WHERE id=?", (biz_id,)).fetchone()
            _no = _rn['receive_no'] if _rn else ''
            log_approval_action('collect_accept', biz_id, 'agree' if result == 'ok' else 'reject', approver, approver_id, comment or '', now(), attachments or [], _src_of(approver, approver_id), instance_code, conn=c)
            if result == 'ok':
                log(_op_name() or approver, '集体验收通过', f'{_no} 集体验收确认完成, 可提交入库审批')
        except Exception:
            pass
        c.commit(); c.close()
        return True
    # V11.126: 询价定标审批无通用 biz_table(父单据=inquiries), 单独处理; 其他走通用表
    if biz_type == 'inquiry_approval':
        # V11.185/186: 询价驳回 → 回"询价中"(采购可改报价/重新提交审批), 累计驳回次数留痕
        if result != 'ok':
            _rc = c.execute("SELECT COALESCE(reject_count,0) FROM inquiries WHERE id=?", (biz_id,)).fetchone()
            _rc_n = (_rc[0] if _rc else 0) + 1
            c.execute("UPDATE inquiries SET status='询价中', updated_at=?, reject_count=?, rejected_items=? WHERE id=?",
                      (now(), _rc_n, '__all__', biz_id))
        else:
            c.execute("UPDATE inquiries SET status=?, updated_at=? WHERE id=?", (st, now(), biz_id))
    else:
        _tbl = biz_table(biz_type)
        # V11.185: 驳回退回闭环 — 累计驳回次数+标记条目(通用表); 通过/正常只改状态
        if result != 'ok' and _tbl:
            _rc = c.execute(f"SELECT COALESCE(reject_count,0) FROM {_tbl} WHERE id=?", (biz_id,)).fetchone()
            _rc_n = (_rc[0] if _rc else 0) + 1
            c.execute(f"UPDATE {_tbl} SET status=?, updated_at=?, reject_count=?, rejected_items=? WHERE id=?",
                      (st, now(), _rc_n, rejected_items if rejected_items is not None else '__all__', biz_id))
        else:
            c.execute(f"UPDATE {_tbl} SET status=?, updated_at=? WHERE id=?", (st, now(), biz_id))
    # V11.67: 询价定标审批通过 → 同步询价单状态(定标审批中→已生成订单); 驳回→恢复询价中
    if biz_type == 'purchase_order':
        if result == 'ok' and st in ('已通过', '审批通过'):
            c.execute("UPDATE inquiries SET status='已生成订单', updated_at=? WHERE id IN (SELECT i.id FROM inquiries i JOIN inquiry_suppliers s ON s.inquiry_id=i.id WHERE s.id=(SELECT selected_supplier_id FROM purchase_orders WHERE id=?))", (now(), biz_id))
    # V11.126: 询价定标审批(手动提交或三家自动提交) — 通过→更新审批记录+根据选定供应商(优先钉钉表单, 兜底最低价)生成订单生效; 驳回→恢复询价中
    elif biz_type == 'inquiry_approval':
        if result == 'ok':
            _sel_id = None
            try:
                _inst = c.execute("SELECT * FROM dingtalk_instances WHERE biz_type=? AND biz_id=? ORDER BY id DESC LIMIT 1", (biz_type, biz_id)).fetchone()
                if _inst and 'form_values' in _inst.keys() and _inst['form_values']:
                    _fv = json.loads(_inst['form_values'])
                    for _it in (_fv if isinstance(_fv, list) else []):
                        if isinstance(_it, dict) and (_it.get('name') == '选定供应商' or _it.get('name') == '选定供应商 '):
                            _vals = _it.get('value', '')
                            # V11.137: 钉钉单选返回可能为纯文本"厂家A"或数组["厂家A"]或JSON串 — 统一归一
                            _sel_txt = ''
                            if isinstance(_vals, str):
                                _s = _vals.strip()
                                if _s.startswith('[') and _s.endswith(']'):
                                    try:
                                        _lst = json.loads(_s)
                                        if isinstance(_lst, list) and _lst:
                                            _sel_txt = str(_lst[0]).strip()
                                    except Exception:
                                        _sel_txt = _s
                                else:
                                    _sel_txt = _s
                            elif isinstance(_vals, list) and _vals:
                                _sel_txt = str(_vals[0]).strip()
                            # ① 选项是"厂家A/B/C" → 按添加顺序id升序排名映射(与Excel比价单从左到右一致, 厂家A=最左)
                            _rank_map = {'厂家A': 0, '厂家B': 1, '厂家C': 2}
                            # V11.155: 领导选"按各物资最低价择优采购" → 不指定厂家, 审批通过后采购员分项定标(每项默认最低价)
                            if '最低价择优' in _sel_txt:
                                _sel_id = 'ALL_LOWEST'
                            elif _sel_txt in _rank_map:
                                _ranked = c.execute(
                                    "SELECT id FROM inquiry_suppliers WHERE inquiry_id=? ORDER BY id ASC",
                                    (biz_id,)).fetchall()
                                _ri = _rank_map[_sel_txt]
                                if _ri < len(_ranked):
                                    _sel_id = _ranked[_ri]['id']
                            # ② 选项是供应商名称(如"厂家2 (¥1750)") → 按名称匹配
                            elif _sel_txt:
                                _m = c.execute("SELECT id FROM inquiry_suppliers WHERE inquiry_id=? AND supplier_name=? ORDER BY id LIMIT 1",
                                               (biz_id, _sel_txt.split(' (')[0])).fetchone()
                                if _m:
                                    _sel_id = _m['id']
                                elif _sel_txt.strip().isdigit():
                                    _sel_id = int(_sel_txt)
                            break
            except Exception:
                _sel_id = None
            if not _sel_id:
                _r2 = c.execute("SELECT id FROM inquiry_suppliers WHERE inquiry_id=? AND quote_price>0 ORDER BY quote_price ASC LIMIT 1", (biz_id,)).fetchone()
                _sel_id = _r2['id'] if _r2 else None
            # V11.155: 领导选"按各物资最低价择优采购" → 不自动生成订单, 状态=待定标, 采购员分项定标(每项默认最低价)
            if _sel_id == 'ALL_LOWEST':
                c.execute("UPDATE inquiries SET status='待定标', updated_at=? WHERE id=?", (now(), biz_id))
                c.execute("UPDATE inquiry_approvals SET selected_supplier_id=NULL, status='已完成' WHERE inquiry_id=?", (biz_id,))
                try:
                    c.execute("ALTER TABLE inquiries ADD COLUMN approve_note TEXT DEFAULT ''")
                except Exception:
                    pass
                c.execute("UPDATE inquiries SET approve_note='领导已同意按各物资最低价择优采购, 请分项定标' WHERE id=?", (biz_id,))
                # 清理提交审批时生成的草稿订单(不指定厂家, 改由分项定标生成, 防残留)
                try:
                    _drafts = c.execute("SELECT id FROM purchase_orders WHERE inquiry_id=? AND status='草稿'", (biz_id,)).fetchall()
                    for _dd in _drafts:
                        c.execute("DELETE FROM order_items WHERE order_id=?", (_dd['id'],))
                        c.execute("DELETE FROM purchase_orders WHERE id=?", (_dd['id'],))
                except Exception:
                    pass
            elif _sel_id:
                try:
                    _sup = c.execute("SELECT * FROM inquiry_suppliers WHERE id=?", (_sel_id,)).fetchone()
                    _iq = c.execute("SELECT * FROM inquiries WHERE id=?", (biz_id,)).fetchone()
                    if _sup and _iq:
                        _pr = c.execute("SELECT * FROM purchase_requests WHERE id=?", (_iq['req_id'],)).fetchone()
                        _items = c.execute("SELECT * FROM request_items WHERE req_id=? ORDER BY id", (_iq['req_id'],)).fetchall()
                        # V11.180: 采购已议价时用调整后报价生成订单(原始报价留痕)
                        _sup_d = dict(_sup)
                        _total = float(inquiry_eff_price(_sup_d, 'quote_price'))
                        _remark = '询价单:%s 供应商:%s 报价¥%.0f' % (_iq['inq_no'], _sup['supplier_name'], _total)
                        _adj_rm = (_sup_d.get('adj_remark') or '').strip()
                        if _adj_rm:
                            _remark += '; 采购议价备注:%s' % _adj_rm
                        # 优先完善提交时生成的草稿订单, 没有才新建(防重复订单)
                        _draft = c.execute("SELECT id FROM purchase_orders WHERE inquiry_id=? AND status='草稿' ORDER BY id LIMIT 1", (biz_id,)).fetchone()
                        if _draft:
                            _oid = _draft['id']
                            c.execute("UPDATE purchase_orders SET supplier=?, price=?, amount=?, total_amount=?, remark=?, status='已通过', updated_at=? WHERE id=?",
                                      (_sup['supplier_name'], _total, _total, _total, _remark, now(), _oid))
                            c.execute("DELETE FROM order_items WHERE order_id=?", (_oid,))
                        else:
                            _no = gen_no('CG', 'purchase_orders', 'order_no', c)
                            c.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
                                supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments,status,inquiry_id) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                                      (_no, _iq['req_id'], _iq['title'][:50], '', 1, '个', _total, _total, 0, 0, _total,
                                       _sup['supplier_name'], _iq['created_by'], '后勤类', _iq['created_by'], 1, (_iq['deadline'] or '')[:10], '货到付款',
                                       _remark, 0, json.dumps([], ensure_ascii=False), '已通过', _iq['id']))
                            _oid = c.execute("SELECT id FROM purchase_orders WHERE order_no=?", (_no,)).fetchone()[0]
                        # V11.170: 明细价格优先用商家报价(quote_details按全量物资顺序存unit_price),
                        # 不再按申请参考金额分摊(参考金额常为0, 导致首项吃全额/后项变0, 合同丢明细)
                        # V11.180: 采购已议价时用调整后明细
                        _qd = {}
                        try:
                            _qd_list = json.loads(inquiry_eff_price(_sup_d, 'quote_details'))
                            for _qi, _q in enumerate(_qd_list):
                                _qd[_qi] = _q
                        except Exception:
                            _qd = {}
                        _grand = 0.0
                        for _idx, _it in enumerate(_items):
                            _qty = float(_it['quantity'] or 1)
                            _q = _qd.get(_idx, {})
                            _price = float(_q.get('unit_price') or 0) or 0
                            _amt = round(_price * _qty, 2)
                            _grand += _amt
                            c.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                                      (_oid, _it['item_name'], _it['spec'] or '', _it['unit'] or '个', _qty, _price, _amt, 0, 0, _amt, ''))
                        # 兜底: 商家未填单价(旧数据) → 回退按申请参考金额比例分摊报价总额
                        if _grand <= 0:
                            c.execute("DELETE FROM order_items WHERE order_id=?", (_oid,))
                            _base = sum(float(x['total_price'] or 0) for x in _items) if _items else 0
                            _grand2 = 0.0
                            for _idx, _it in enumerate(_items):
                                _qty = float(_it['quantity'] or 1)
                                if _base > 0 and _idx < len(_items) - 1:
                                    _amt = _total * (float(_it['total_price'] or 0) / _base)
                                else:
                                    _amt = _total - _grand2
                                _amt = round(_amt, 2)
                                _price = round(_amt / _qty, 2) if _qty else 0
                                _grand2 += _amt
                                c.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                                          (_oid, _it['item_name'], _it['spec'] or '', _it['unit'] or '个', _qty, _price, _amt, 0, 0, _amt, ''))
                        c.execute("UPDATE inquiries SET status='已生成订单', selected_supplier_id=?, updated_at=? WHERE id=?", (_sel_id, now(), biz_id))
                        c.execute("UPDATE inquiry_approvals SET selected_supplier_id=?, status='已完成' WHERE inquiry_id=?", (_sel_id, biz_id))
                        c.execute("UPDATE inquiry_suppliers SET is_selected=1 WHERE id=?", (_sel_id,))
                        # 注意: 此处不写日志 — 事务未提交时开新连接会 database is locked, 末尾已有提交后日志
                except Exception as e:
                    pass
            else:
                c.execute("UPDATE inquiry_approvals SET status='已批准' WHERE inquiry_id=?", (biz_id,))
        else:
            c.execute("UPDATE inquiry_approvals SET status='已驳回' WHERE inquiry_id=?", (biz_id,))
            c.execute("UPDATE inquiries SET status='询价中', updated_at=? WHERE id=?", (now(), biz_id))
    if result == 'ok':
        if biz_type == 'receiving' and st == '已入库':
            do_receiving_stock(c, biz_id)
        elif biz_type == 'requisition' and st == '已出库':
            do_requisition_stock(c, biz_id)
    # V7.0: 全节点自动推送 — 单据流转到下一级时, 向下一级审批人钉钉推送(待办+通知)
    try:
        if result == 'ok':
            nxt = c.execute("SELECT * FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no LIMIT 1", (biz_type, biz_id)).fetchone()
            if nxt:
                u = None
                if nxt['approver_id']:
                    u = c.execute("SELECT * FROM users WHERE id=?", (nxt['approver_id'],)).fetchone()
                if not u:
                    u = c.execute("SELECT * FROM users WHERE role=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!='' ORDER BY id LIMIT 1", (nxt['role'],)).fetchone()
                if u and u['dingtalk_userid']:
                    no = c.execute(f"SELECT * FROM {biz_table(biz_type)} WHERE id=?", (biz_id,)).fetchone()
                    doc_no = no['req_no'] if 'req_no' in no.keys() else (no['order_no'] if 'order_no' in no.keys() else (no['contract_no'] if 'contract_no' in no.keys() else (no['receive_no'] if 'receive_no' in no.keys() else (no['payment_no'] if 'payment_no' in no.keys() else ''))))
                    c.execute("INSERT INTO dingtalk_push_log(biz_type,biz_id,doc_no,push_type,target_user,target_userid,operator,content) VALUES(?,?,?,?,?,?,?,?)",
                              (biz_type, str(biz_id), doc_no, 'auto', u['name'] or u['username'], u['dingtalk_userid'],
                               approver, f"下一节点审批: {nxt['role']} 待审批"))
                    c.commit()
                    # 异步推送避免阻塞审批
                    def _push(uid, no_, biz_type_, biz_id_):
                        try:
                            dt_send_todo([uid], '📋 新的审批待办', f"{no_} 等待您的审批",
                                         f"节点: {nxt['role']}", biz_type_, biz_id_, push_type='auto', operator=approver)
                        except Exception:
                            pass
                    threading.Thread(target=_push, args=(u['dingtalk_userid'], doc_no, biz_type, biz_id), daemon=True).start()
    except Exception:
        pass
    # V55需求1-2: 先款后货合同生效(执行中)后, 自动生成待入库记录展示在入库板块
    # V11.11: 合同审批通过(执行中)后, 无论交易模式(货到付款/先款后货/自定义)均自动生成待入库记录
    if biz_type == 'contract' and st == '执行中':
        _ct = c.execute("SELECT * FROM contracts WHERE id=?", (biz_id,)).fetchone()
        if _ct and _ct['order_id']:
            _po = c.execute("SELECT * FROM purchase_orders WHERE id=?", (_ct['order_id'],)).fetchone()
            if _po:
                _exist = c.execute("SELECT 1 FROM receivings WHERE order_id=? AND status!='已入库'", (_po['id'],)).fetchone()
                if not _exist:
                    _oi = c.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (_po['id'],)).fetchall()
                    _name = _oi[0]['item_name'] if _oi else _po['item_name']
                    _qty = sum(float(x['quantity']) for x in _oi) if _oi else _po['quantity']
                    # V11.31: 自动带出部门(申请单→订单→入库单链)
                    _dept = _po['requester'] and '' or ''
                    try:
                        if _po['req_id']:
                            _pr = c.execute("SELECT dept FROM purchase_requests WHERE id=?", (_po['req_id'],)).fetchone()
                            if _pr and _pr['dept']:
                                _dept = _pr['dept']
                    except Exception:
                        pass
                    _rno = gen_no('RK', 'receivings', 'receive_no', c)
                    # V11.152: 多物资订单自动生成入库时, 明细完整存items_json(入库验收列表/详情可显示全部物资名)
                    _items_json = json.dumps(
                        [{'item_name': x['item_name'], 'spec': x['spec'] or '', 'quantity': x['quantity'],
                          'unit': x['unit'] or '个', 'price': x['price'] or 0} for x in _oi],
                        ensure_ascii=False) if _oi else ''
                    _name = (_oi[0]['item_name'] + ' 等%d项' % len(_oi)) if len(_oi) > 1 else (_oi[0]['item_name'] if _oi else _po['item_name'])
                    # V11.198: 自动生成的入库单默认暂估(is_est=1, 货到票未到先暂估入账, 收到发票后发票核对红冲转正式; 需正式的可在提交时/手动入库选正式)
                    c.execute("INSERT INTO receivings(receive_no,delivery_id,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,dept,items_json,is_est) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,0)",
                              (_rno, None, _po['id'], _name, '', _qty, '个', 0, '待入库', now(), '合同生效后自动进入入库板块(整批%d项)' % (len(_oi) if _oi else 1), _dept, _items_json))
    c.execute("UPDATE feishu_instances SET status='synced', updated_at=? WHERE biz_type=? AND biz_id=? AND status='pending'", (now(), biz_type, biz_id))
    c.execute("UPDATE dingtalk_instances SET status='synced', updated_at=? WHERE biz_type=? AND biz_id=? AND status='pending'", (now(), biz_type, biz_id))
    c.commit(); c.close()
    log(approver, f'{biz_type}审批{"通过" if result=="ok" else "驳回"}', f'{biz_type}#{biz_id} → {st}')

    # V11.178: 驳回后钉钉通知提交人(含单据号/驳回人/理由) — 事务已提交, 线程异步发避免阻塞/锁冲突
    if result != 'ok':
        _src = _src_of(approver, approver_id)
        try:
            def _notify_sub(_bt, _bid, _ap, _cm, _sc):
                try:
                    notify_submitter_rejected(_bt, _bid, _ap, _cm, source=_sc)
                except Exception:
                    pass
            threading.Thread(target=_notify_sub, args=(biz_type, biz_id, approver, comment, _src), daemon=True).start()
        except Exception:
            pass

    # V11.70: 审批办结通知发起人(站内信+钉钉工作通知) — V11.190: 用find_doc_submitter(合同/挂账/付款经order_id关联订单取提交人, 修复合同通过后不通知)
    # 注意: 主连接c已在上面close, 这里必须新开连接(原代码用已关闭连接查单号→异常被吞→通知从未发出)
    if result == 'ok' and st in ('已通过', '审批通过', '已入库', '已出库', '已签合同', '执行中'):
        try:
            _doc_no = ''
            _dc = db()
            try:
                _d = _dc.execute(f"SELECT * FROM {biz_table(biz_type)} WHERE id=?", (biz_id,)).fetchone()
            except Exception:
                _d = None
            _dc.close()
            if _d:
                for _k in ('req_no', 'order_no', 'contract_no', 'receive_no', 'payment_no', 'credit_no'):
                    if _k in _d.keys() and _d[_k]:
                        _doc_no = str(_d[_k]); break
            if not _doc_no:
                _doc_no = f'{biz_type}#{biz_id}'
            _u = find_doc_submitter(biz_type, biz_id)
            if _u and _u.get('dingtalk_userid'):
                import threading as _th
                def _notify():
                    try:
                        from app import dt_send_todo
                        dt_send_todo(
                            [_u['dingtalk_userid']],
                            f'✅ 审批通过 · {_doc_no}',
                            f'您提交的{_doc_no}已审批通过，可继续后续操作',
                            f'单据: {_doc_no}',
                            biz_type, str(biz_id),
                            push_type='done',
                            operator=approver
                        )
                    except Exception:
                        pass
                _th.Thread(target=_notify, args=(), daemon=True).start()
        except Exception:
            pass
    return True

def fs_sync_result(instance_code, result):
    c = db()
    r = c.execute("SELECT * FROM feishu_instances WHERE instance_code=?", (instance_code,)).fetchone()
    c.close()
    if not r or r['status'] in ('synced', 'error'): return
    finish_approvals(r['biz_type'], r['biz_id'], 'ok' if result == 'approved' else 'reject', '飞书', 0, f'飞书审批{result}')

def fs_mark_cancelled(instance_code):
    c = db()
    c.execute("UPDATE feishu_instances SET status='cancelled', updated_at=? WHERE instance_code=?", (now(), instance_code))
    c.commit(); c.close()

# ---- 回调解密(encrypt_key模式) ----
def fs_decrypt(body):
    key_b64 = cfg_get('feishu_encrypt_key')
    if not key_b64: return None
    key = base64.b64decode(key_b64)
    iv = key[:16]
    raw = base64.b64decode(body.get('encrypt', ''))
    plain = None
    try:
        from Crypto.Cipher import AES
        plain = AES.new(key, AES.MODE_CBC, iv).decrypt(raw)
    except ImportError:
        try:
            from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
            dec = Cipher(algorithms.AES(key), modes.CBC(iv)).decryptor()
            plain = dec.update(raw) + dec.finalize()
        except ImportError:
            raise RuntimeError('回调使用加密模式但缺少解密库: pip install pycryptodome, 或在飞书后台把加密策略改为"不加密"')
    if plain and 0 < plain[-1] <= 16: plain = plain[:-plain[-1]]
    return json.loads(plain.decode('utf-8'))

# ---- 预警提醒(全部走飞书机器人, 系统内不开发提醒) ----
def fs_send_reminders(force=False):
    if not feishu_enabled(): return
    approve_hours = int(cfg_get('feishu_approve_hours', '24') or 24)
    order_days = int(cfg_get('feishu_order_days', '3') or 3)
    report_chat = cfg_get('feishu_report_chat')
    today_s = today()
    c = db()
    def pushed(rule, key):
        if force: return False
        k = f"{today_s}:{key}"
        if c.execute("SELECT 1 FROM reminder_log WHERE rule=? AND key=?", (rule, k)).fetchone(): return True
        c.execute("INSERT INTO reminder_log(rule,key) VALUES(?,?)", (rule, k))
        return False
    def admins_open():
        return [r['feishu_open_id'] for r in c.execute("SELECT feishu_open_id FROM users WHERE role='系统管理员' AND feishu_open_id IS NOT NULL AND feishu_open_id!=''").fetchall()]
    # ① 超时未审批 → 推给对应审批角色
    rows = c.execute("""SELECT ai.* FROM approval_instances ai
        WHERE ai.status='pending' AND ai.created_at <= datetime('now','localtime',?)
        ORDER BY ai.created_at""", (f'-{approve_hours} hours',)).fetchall()
    by_role = {}
    for r in rows: by_role.setdefault(r['role'], []).append(r)
    for role, lst in by_role.items():
        u = find_user_by_role(role)
        if not u or not u['feishu_open_id']: continue
        if not pushed('approve_timeout', role):
            fs_send(u['feishu_open_id'], f"⏰ **{role}** 有 **{len(lst)}** 条审批超过 {approve_hours} 小时未处理:\n" +
                    '\n'.join(f"- {x['biz_type']}#{x['biz_id']} (提交于 {x['created_at']})" for x in lst[:10]), 'orange')
        for x in lst:  # 超过48小时升级给管理员
            try: age_h = int((datetime.datetime.now() - datetime.datetime.strptime(x['created_at'][:19], '%Y-%m-%d %H:%M:%S')).total_seconds() // 3600)
            except Exception: age_h = approve_hours
            if age_h >= 48 and not pushed('approve_escalate', x['id']):
                for ao in admins_open():
                    fs_send(ao, f"🚨 审批升级: {x['biz_type']}#{x['biz_id']} 已超 {age_h} 小时, 审批角色【{role}】, 请介入处理", 'red')
    # ② 待采未下单 → 推给申请人
    rows2 = c.execute("""SELECT pr.* FROM purchase_requests pr
        WHERE pr.status='已通过' AND pr.created_at <= datetime('now','localtime',?)
        AND NOT EXISTS (SELECT 1 FROM purchase_orders po WHERE po.req_id=pr.id)""", (f'-{order_days} days',)).fetchall()
    for r in rows2:
        u = find_user_by_name(r['requester'])
        if not u or not u['feishu_open_id']: continue
        if not pushed('req_no_order', r['id']):
            fs_send(u['feishu_open_id'], f"📝 采购申请 **{r['req_no']}** 已通过审批 {order_days} 天仍未转采购订单\n金额 ¥{r['total_estimated'] or 0:,.0f} · {r['purpose'] or ''}\n请及时下单", 'orange')
    # ③ 逾期未回货(送货单未签收) → 推给订单负责人
    rows3 = c.execute("""SELECT d.*, po.owner FROM deliveries d
        LEFT JOIN purchase_orders po ON d.order_id=po.id
        WHERE d.sign_status='待签收' AND d.delivery_date < date('now')""").fetchall()
    for r in rows3:
        u = find_user_by_name(r['owner'])
        if not u or not u['feishu_open_id']: continue
        if not pushed('delivery_overdue', r['id']):
            fs_send(u['feishu_open_id'], f"🚚 送货单 **{r['delivery_no']}** 计划 {r['delivery_date']} 到货, 至今未签收\n物资: {r['item_name']} x{r['quantity']}{r['unit']}\n请立即联系供应商催货", 'red')
    # ④ 订单超目标日 → 推给负责人
    rows4 = c.execute("""SELECT * FROM purchase_orders WHERE target_date < date('now') AND status NOT IN ('已完成','已关闭','已挂账','已入库')""").fetchall()
    for r in rows4:
        u = find_user_by_name(r['owner'])
        if not u or not u['feishu_open_id']: continue
        if not pushed('order_overdue', r['id']):
            fs_send(u['feishu_open_id'], f"📋 采购订单 **{r['order_no']}** 目标日 {r['target_date']} 已过, 当前状态【{r['status']}】\n物资: {r['item_name']} x{r['quantity']}{r['unit']} · 金额 ¥{r['total_amount'] or 0:,.0f}", 'orange')
    # ⑤ 每周一 09:00 汇总周报 → 报表群 + 管理员
    if datetime.date.today().weekday() == 0 and datetime.datetime.now().strftime('%H') == '09' and not pushed('weekly_report', 'wk'):
        stats = {
            '总额': c.execute("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE created_at >= datetime('now','localtime','-7 days')").fetchone()[0],
            '待审批': c.execute("SELECT COUNT(*) FROM approval_instances WHERE status='pending'").fetchone()[0],
            '进行中订单': c.execute("SELECT COUNT(*) FROM purchase_orders WHERE status NOT IN ('已完成','已关闭','已挂账')").fetchone()[0],
            '库存预警': c.execute("SELECT COUNT(*) FROM inventory WHERE quantity<=safe_stock AND safe_stock>0").fetchone()[0],
            '超时未签收': c.execute("SELECT COUNT(*) FROM deliveries WHERE sign_status='待签收' AND delivery_date < date('now')").fetchone()[0],
        }
        text = (f"📊 **本周采购周报**\n- 近7天采购总额: ¥{stats['总额']:,.0f}\n- 当前待审批: {stats['待审批']} 条\n"
                f"- 进行中订单: {stats['进行中订单']} 条\n- 库存预警: {stats['库存预警']} 项\n- 超时未签收: {stats['超时未签收']} 单")
        if report_chat: fs_send(report_chat, text, 'blue', id_type='chat_id')
        for ao in admins_open(): fs_send(ao, text, 'blue')
    c.commit(); c.close()

def dt_poll_loop():
    """钉钉审批结果即时同步线程:
    V11.28: 15s→60s 固定轮询(API消耗降75%)
    V11.168: 动态轮询 — 有pending审批时15秒快轮询(钉钉审批通过后系统侧快速同步),
             无pending时60秒慢轮询(平时不耗API); 兼顾同步速度与API配额"""
    while True:
        try:
            if dingtalk_enabled():
                _pend = dt_pending_count()
                dt_poll_results()
                dt_retry_failed_instances()
                dt_terminate_stale()
                # V11.203 模块一1.2: 发票节点到期自动提醒采购专员(内部节流15分钟+每合同每天去重)
                try:
                    check_invoice_node_reminders()
                except Exception:
                    pass
                # 有待审批 → 15秒后再查(审批结果同步快); 无待审批 → 60秒慢轮询(省API)
                time.sleep(15 if _pend > 0 else 60)
            else:
                time.sleep(60)
        except Exception:
            time.sleep(60)


def dt_pending_count():
    """当前待同步的钉钉审批实例数(仅本地DB查询, 不耗钉钉API)"""
    try:
        c = db()
        n = c.execute("""SELECT COUNT(*) FROM dingtalk_instances WHERE status='pending'
            AND updated_at >= datetime('now','localtime','-7 days')""").fetchone()[0]
        c.close()
        return n or 0
    except Exception:
        return 0


def dt_terminate_stale():
    """终态终止兜底: 系统侧已审批完(无pending节点)但钉钉实例仍在RUNNING → 终止。
    解决 terminate 偶发 internalError(审批流刚创建未稳定): 15秒轮询持续重试, 直到成功。"""
    try:
        if not dingtalk_enabled(): return
        c = db()
        rows = c.execute("SELECT * FROM dingtalk_instances WHERE status IN ('pending','synced')").fetchall()
        for r in rows:
            try:
                n_pending = c.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending'", (r['biz_type'], r['biz_id'])).fetchone()[0]
            except Exception:
                continue
            if n_pending > 0:
                continue
            inst = dt_query_instance(r['instance_code'])
            if not inst:
                continue
            if str(inst.get('status', '')) in ('RUNNING', 'NEW'):
                u = find_user_by_role('系统管理员')
                uid = u['dingtalk_userid'] if u and u['dingtalk_userid'] else dt_first_bound_userid()
                dt_terminate_instance(r['instance_code'], uid)
        c.close()
    except Exception:
        pass


def scheduler_loop():
    """预警调度: 每60秒扫描一次; 虾(xia_enabled=1)启用后预警全部交给虾独立服务, 系统本体不再推送"""
    while True:
        try:
            time.sleep(60)
            # V55: 系统内预警中心(首页集中展示)始终重建, 不受虾/飞书开关影响
            try:
                if cfg_get('warn_enabled', '1') == '1': build_alerts()
            except Exception:
                pass
            if cfg_get('xia_enabled') == '1': continue  # 虾接管提醒推送
            if feishu_enabled(): fs_send_reminders()
            if dingtalk_enabled(): dt_send_reminders()
        except Exception:
            pass

# ============================================================
# V5.2 ── 钉钉对接: 审批同步 + 工作通知 + 快捷审批 (urllib; 旧式回调需 pycryptodome)
# 审批模板需在钉钉OA审批后台手工创建(见《钉钉对接配置手册.md》), 模板process_code填入配置
# 流程: 单据进入待审批 → 发起钉钉审批实例(快捷审批) → 审批人钉钉处理 → 事件回调同步回系统
# ============================================================
DT_API = 'https://oapi.dingtalk.com'
DT_NEW_API = 'https://api.dingtalk.com'
DT_BIZ = {  # biz_type -> 审批模板名称
    'purchase_request': '采购申请审批',
    'purchase_order':   '采购订单审批',
    'contract':         '合同审批',
    'credit':           '挂账审批',
    'payment':          '付款审批',
    'receiving':        '入库审批',
    'requisition':      '出库审批',
    'return_request':   '退库审批',  # V11.193
    'collect_accept':   '集体验收审批',  # V11.206
    'repair_plan':      '维修采购审批',  # V11.208
    'inquiry_approval': '采购比价单审批',  # V11.142: 补全, 否则dt_send_todo抛KeyError
}
DT_FORM = [('单据编号', 'text'), ('内容摘要', 'text'), ('金额(元)', 'text'), ('申请人', 'text'), ('提交时间', 'text')]

_DT_TOKEN = {'t': None, 'exp': 0}
_DT_TOKEN_NEW = {'t': None, 'exp': 0}
_DT_TICKET = {'t': None, 'exp': 0}

def dingtalk_enabled():
    return cfg_get('dingtalk_enabled') == '1'

def dt_token():
    if _DT_TOKEN['t'] and time.time() < _DT_TOKEN['exp'] - 120: return _DT_TOKEN['t']
    ak = cfg_get('dingtalk_app_key'); sk = cfg_get('dingtalk_app_secret')
    if not ak or not sk: raise RuntimeError('钉钉应用未配置: 请先在"钉钉设置"填写 AppKey / AppSecret')
    url = f'{DT_API}/gettoken?appkey={urllib.parse.quote(ak)}&appsecret={urllib.parse.quote(sk)}'
    try:
        with urllib.request.urlopen(url, timeout=12) as r:
            d = json.loads(r.read().decode('utf-8'))
    except Exception as e:
        raise RuntimeError(f'钉钉网络错误: {e}')
    if d.get('errcode') != 0: raise RuntimeError(f'获取钉钉access_token失败: {d.get("errmsg", d)}')
    _DT_TOKEN['t'] = d.get('access_token'); _DT_TOKEN['exp'] = time.time() + int(d.get('expires_in', 7200))
    return _DT_TOKEN['t']

def dt_new_token():
    """新版接口 token (api.dingtalk.com/v1.0)"""
    if _DT_TOKEN_NEW['t'] and time.time() < _DT_TOKEN_NEW['exp'] - 120: return _DT_TOKEN_NEW['t']
    ak = cfg_get('dingtalk_app_key'); sk = cfg_get('dingtalk_app_secret')
    if not ak or not sk: raise RuntimeError('钉钉应用未配置: 请先在"钉钉设置"填写 AppKey / AppSecret')
    body = json.dumps({'appKey': ak, 'appSecret': sk}).encode('utf-8')
    req = urllib.request.Request(f'{DT_NEW_API}/v1.0/oauth2/accessToken', data=body,
                                 headers={'Content-Type': 'application/json'})
    try:
        with urllib.request.urlopen(req, timeout=12) as r:
            d = json.loads(r.read().decode('utf-8'))
    except Exception as e:
        raise RuntimeError(f'钉钉新版接口网络错误: {e}')
    t = d.get('accessToken')
    if not t: raise RuntimeError(f'获取钉钉新版accessToken失败: {d}')
    _DT_TOKEN_NEW['t'] = t; _DT_TOKEN_NEW['exp'] = time.time() + int(d.get('expireIn', 7200))
    return _DT_TOKEN_NEW['t']

def dt_new_post(path, payload, timeout=12):
    """新版接口 POST, 返回 (err_code, resp)"""
    url = DT_NEW_API + path
    try:
        tok = dt_new_token()
    except Exception as e:
        return -1, {'msg': str(e)}
    body = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(url, data=body, headers={
        'Content-Type': 'application/json', 'x-acs-dingtalk-access-token': tok})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            d = json.loads(r.read().decode('utf-8'))
        return 0, d
    except urllib.error.HTTPError as e:
        try:
            d = json.loads(e.read().decode('utf-8'))
            return 1, d
        except Exception:
            return 1, {'msg': f'HTTP {e.code}'}
    except Exception as e:
        return -1, {'msg': f'网络错误: {e}'}

def dt_post(path, payload, timeout=12):
    """旧版接口 POST (oapi.dingtalk.com), 返回 (errcode, resp); 工作通知/手机号查询/SSO 用"""
    try:
        tok = dt_token()
    except Exception as e:
        return -1, {'msg': str(e)}
    sep = '&' if '?' in path else '?'
    url = DT_API + path + sep + 'access_token=' + urllib.parse.quote(tok)
    body = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(url, data=body, headers={
        'Content-Type': 'application/json'})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            d = json.loads(r.read().decode('utf-8'))
        return d.get('errcode', -1), d
    except urllib.error.HTTPError as e:
        try:
            d = json.loads(e.read().decode('utf-8'))
            return d.get('errcode', 1), d
        except Exception:
            return 1, {'msg': f'HTTP {e.code}'}
    except Exception as e:
        return -1, {'msg': f'网络错误: {e}'}

def dt_agent_id():
    try: return int(cfg_get('dingtalk_agent_id', '0') or 0)
    except Exception: return 0

def dt_approval_codes():
    try: return json.loads(cfg_get('dingtalk_approval_codes', '{}'))
    except Exception: return {}

def dt_approval_code(biz_type):
    return dt_approval_codes().get(biz_type, '')

def dt_public_url():
    f = os.path.join(BASE, 'data', 'public_url.txt')
    try:
        u = open(f).read().strip()
        return u if u.startswith('http') else ''
    except Exception:
        return ''

def dt_userid_by_mobile(mobile):
    code, resp = dt_post('/topapi/v2/user/getbymobile', {'mobile': mobile})
    if code != 0: return None, resp.get('msg', '查询失败')
    r = resp.get('result', {}) or {}
    return r.get('userid', ''), r.get('name', '')


def dt_first_bound_userid():
    """返回系统内第一个已绑定钉钉ID的用户(作发起人兜底)"""
    try:
        c = db()
        r = c.execute("SELECT dingtalk_userid FROM users WHERE dingtalk_userid IS NOT NULL AND dingtalk_userid != '' ORDER BY id LIMIT 1").fetchone()
        c.close()
        return r['dingtalk_userid'] if r else ''
    except Exception:
        return ''

def dt_union_id():
    """V11.8: 获取钉钉 unionId(上传附件必填, 缓存到 sys_config.dingtalk_union_id)
    取第一个已绑定 userid → /topapi/v2/user/get 拿 unionid; 失败逐用户尝试"""
    try:
        cached = cfg_get('dingtalk_union_id', '')
        if cached:
            return cached
        c = db()
        users = c.execute("SELECT name, dingtalk_userid FROM users WHERE dingtalk_userid IS NOT NULL AND dingtalk_userid != '' ORDER BY id").fetchall()
        c.close()
        for u in users:
            try:
                code_r, resp = dt_post('/topapi/v2/user/get', {'userid': u['dingtalk_userid']})
                if code_r == 0 and isinstance(resp, dict):
                    uid = (resp.get('result') or {}).get('unionid', '')
                    if uid:
                        cfg_set('dingtalk_union_id', uid)
                        return uid
            except Exception:
                continue
        return ''
    except Exception:
        return ''

def dt_biz_info(biz_type, biz_id):
    return fs_biz_info(biz_type, biz_id)  # 复用单据信息提取

def start_instances(biz_type, biz_id):
    """单据进入待审批后, 向已配置的外部通道(飞书/钉钉)同步发起审批实例"""
    fs_start_instance(biz_type, biz_id)
    dt_start_instance(biz_type, biz_id)

def dt_actioner_key(biz_type):
    """自选审批人节点的 actionerKey (模板相关, 从配置读, 缺失返回 '')"""
    try:
        return json.loads(cfg_get('dingtalk_actioner_keys', '{}')).get(biz_type, '')
    except Exception:
        return ''

def dt_cat_option(budget_name):
    """预算科目名 → 采购申请模板'采购类别'下拉选项(固定资产/行政办公/酒店日耗/工程采购)"""
    n = str(budget_name or '')
    if '固定' in n or '设备' in n or '维修' in n or '配件' in n: return '固定资产'
    if '酒店' in n or '耗材' in n or '客房' in n: return '酒店日耗'
    if '工程' in n or '施工' in n or '建设' in n: return '工程采购'
    return '行政办公'  # 兜底


def dt_date_ts(date_str):
    """'yyyy-MM-dd' → 毫秒时间戳字符串 (DDDateField 要求)"""
    try:
        d = datetime.datetime.strptime(str(date_str).strip()[:10], '%Y-%m-%d')
        return str(int(d.timestamp() * 1000))
    except Exception:
        return str(int(time.time() * 1000))


def dt_build_form(biz_type, biz_id, info):
    """V6.0: 按钉钉审批模板字段组装表单值（单据完整详情结构化展示）。

    采购申请审批模板字段(2026-08-06 schema 实测):
      部门=DepartmentField(必填,部门ID数组如[1]) / 采购类别=DDSelectField(必填,固定选项:
      固定资产/行政办公/酒店日耗/工程采购) / 采购事由=TextField(必填) /
      交付日期=DDDateField(必填,yyyy-MM-dd) / 附件=DDAttachment(可选) / 备注=TextareaField

    V6.0 增强(6.0.docx 需求1):
      - 单据基础信息全展示: 编号/申请人/部门/时间/类型/总金额/数量/紧急等级
      - 各单据类型全部表单字段结构化展示在"备注"(审批人钉钉端无需跳回系统)
    """
    today = datetime.date.today().strftime('%Y-%m-%d')
    # V11.133: inquiry_approval 必须加入分支判断, 否则询价审批永远走回退表单(钉钉不显示选商家)
    # V11.202: 加入 return_request(退库审批), 否则退库审批钉钉表单永远走回退字段(模板控件名对不上必失败)
    if biz_type in ('purchase_request', 'contract', 'purchase_order', 'receiving', 'requisition', 'payment', 'inquiry_approval', 'return_request', 'collect_accept', 'repair_plan'):
        c = db()
        r = c.execute(f"SELECT * FROM {biz_table(biz_type)} WHERE id=?", (biz_id,)).fetchone()
        if r:
            detail = dt_build_detail(biz_type, r, c)
            if biz_type == 'purchase_request':
                # V11.15: 申请模板控件=部门/采购类别/采购事由/交付日期/备注/附件
                cat = dt_cat_option(r['budget_code'] or r['dept'] or '')
                purpose = str(r['purpose'] or '').strip()[:200]
                # V11.199: 钉钉必填控件(采购事由TextField)空值会820001 required error — 空事由用占位保证必填
                if not purpose:
                    purpose = f"(未填采购事由, 详见备注) {r['req_no']}"
                target = str(r['target_date'] or today)[:10]
                if not target or target == 'None':
                    target = today
                form = [
                    {'name': '部门', 'value': '[1]'},
                    {'name': '采购类别', 'value': cat},
                    {'name': '采购事由', 'value': purpose},
                    {'name': '交付日期', 'value': target},
                    {'name': '备注', 'value': detail[:1900]},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'contract':
                # V11.15: 合同模板控件=合同编号/对方单位名称/合同总额（元）/图片/备注/附件(全角括号实测)
                _amt = float(r['amount'] or 0)
                form = [
                    {'name': '合同编号', 'value': r['contract_no'] or ''},
                    {'name': '对方单位名称', 'value': r['contract_name'] or ''},
                    {'name': '合同总额（元）', 'value': '%.2f' % _amt},
                    {'name': '图片', 'value': '[]'},
                    {'name': '备注', 'value': detail[:1900]},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'receiving':
                # V11.14: 入库模板控件=入库日期/备注/附件(表格控件"入库明细"格式无法兼容, 明细走备注文本)
                rdetail = dt_build_detail('receiving', r, c)
                form = [
                    {'name': '入库日期', 'value': str(r['received_at'] or today)[:10]},
                    {'name': '备注', 'value': rdetail[:1900]},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'repair_plan':
                # V11.208: 维修采购审批模板(备注控件展示完整详情, 附件控件带故障照片)
                rdetail = dt_build_detail('repair_plan', r, c)
                form = [
                    {'name': '维修说明', 'value': str(r['device_name'] or '')[:100]},
                    {'name': '备注', 'value': rdetail[:1900]},
                ]
                attach = dt_build_attachment('repair_plan', r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'requisition':
                # V11.14: 出库模板控件=出库日期/备注/附件(表格控件"采购明细"格式无法兼容, 明细走备注文本)
                rdetail = dt_build_detail('requisition', r, c)
                form = [
                    {'name': '出库日期', 'value': str(r['issued_at'] or today)[:10]},
                    {'name': '备注', 'value': rdetail[:1900]},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'collect_accept':
                # V11.206: 集体验收审批表单(与入库同控件: 入库日期/备注/附件; 明细走备注文本)
                rdetail = dt_build_detail('receiving', r, c)
                form = [
                    {'name': '入库日期', 'value': str(r['received_at'] or today)[:10]},
                    {'name': '备注', 'value': '👥 集体验收申请\n\n' + rdetail[:1850]},
                ]
                attach = dt_build_attachment('receiving', r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'return_request':
                # V11.202 退库审批(模板字段与出库一致: 退库日期/备注/附件 — 用户按此建模板并填码)
                rdetail = dt_build_detail('return_request', r, c)
                form = [
                    {'name': '退库日期', 'value': str(r['created_at'] or today)[:10]},
                    {'name': '备注', 'value': (rdetail or ('退库单 %s 待审批' % r['return_no']))[:1900]},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'payment':
                # V11.15: 付款模板控件=单据编号/内容摘要/金额(元)/申请人
                form = [
                    {'name': '单据编号', 'value': r['payment_no'] or ''},
                    {'name': '内容摘要', 'value': detail[:1900]},
                    {'name': '金额(元)', 'value': '%.2f' % float(r['amount'] or 0)},
                    {'name': '申请人', 'value': r['payee_name'] or r['requester'] or '系统'},
                ]
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
            elif biz_type == 'inquiry_approval':
                # V11.76/V11.133: 询价审批表单=询价详情+供应商报价+选定供应商(单选)
                # ⚠️ 控件名必须与钉钉模板完全一致(实测带尾随空格"选定供应商 ", 详情控件叫"三方报价详情")
                c2 = db()
                iq = c2.execute("SELECT * FROM inquiries WHERE id=?", (biz_id,)).fetchone()
                if not iq:
                    c2.close(); return []
                # 查询供应商(按添加顺序id升序=Excel比价单从左到右顺序, 厂家A=最左/厂家C=最右, 与Excel一一对应不混淆)
                sups = c2.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=? ORDER BY id ASC", (biz_id,)).fetchall()
                _abc = ('厂家A', '厂家B', '厂家C')
                supplier_opts = []
                supplier_details = []
                for _i, si in enumerate(sups):
                    _tag = _abc[_i] if _i < 3 else ('厂家' + chr(65 + _i))
                    # V11.180: 采购已议价时钉钉审批展示调整后最终报价(原始报价留痕可查)
                    _eff_price = inquiry_eff_price(dict(si), 'quote_price')
                    _is_adj = inquiry_is_adjusted(dict(si))
                    supplier_opts.append({'value': str(si['id']), 'text': '%s (¥%.0f)' % (si['supplier_name'], _eff_price or 0)})
                    detail = '%s[%s] 总价（含税含运）¥%.0f' % (_tag, si['supplier_name'], _eff_price or 0)
                    if _is_adj:
                        detail += '【采购已议价】'
                        _ori = float(si['quote_price'] or 0)
                        if _ori:
                            detail += '(原始报价¥%.0f)' % _ori
                    if si['quote_brand']:
                        detail += ' 品牌:%s' % si['quote_brand']
                    if si['quote_remark']:
                        detail += ' 厂家备注:%s' % si['quote_remark']
                    # V11.180: 采购内部备注(领导看谈判情况)
                    _adj_rm = (si['adj_remark'] or '').strip() if 'adj_remark' in si.keys() else ''
                    if _adj_rm:
                        detail += ' 【采购议价备注】%s' % _adj_rm
                    # V11.162: 含税单价/含税总价 明细同步进钉钉审批(调整后明细优先)
                    try:
                        _qd = json.loads(inquiry_eff_price(dict(si), 'quote_details'))
                        if _qd:
                            _qd_txt = []
                            for _q in _qd:
                                if _q and float(_q.get('unit_price') or 0) > 0:
                                    _qp = float(_q['unit_price'])
                                    _qq = float(_q.get('qty') or 1)
                                    _qd_txt.append('%s含税¥%.2f×%s=¥%.2f' % (_q.get('item_name') or '物料', _qp, _qq, _qp * _qq))
                            if _qd_txt:
                                detail += ' | 明细:' + '；'.join(_qd_txt)
                    except Exception:
                        pass
                    supplier_details.append(detail)
                c2.close()
                # V11.137/V11.150: 模板"选定供应商"控件选项=厂家A/厂家B/厂家C(截图确认)
                # V11.150: 厂家A/B/C按添加顺序映射(与Excel比价单从左到右一致, A=最左C=最右), 领导选A即Excel第一家, 不混淆
                # 之前按报价排序导致领导分不清对应Excel哪家(报价顺序≠添加顺序)
                # V11.155: 默认值=按各物资最低价择优采购(领导不指定厂家时直接同意即可, 醒目); 如需指定厂家再改选A/B/C
                # V11.155d: 三方报价详情顶部加醒目提示(领导知道选定供应商可点选, 避免直接点同意错意)
                _abc = ('厂家A', '厂家B', '厂家C')
                _default_opt = '按各物资最低价择优采购'
                _pick_hint = '【请选择采购方式】下方"选定供应商"为可点选下拉框，默认=按各物资最低价择优采购（推荐）；如需指定厂家请点开下拉改选（厂家A=报价表最左/厂家B=中间/厂家C=最右）'
                form = [
                    {'name': '询价单号', 'value': iq['inq_no'] or ''},
                    {'name': '物资名称', 'value': (iq['title'] or '')[:50]},
                    {'name': '三方报价详情', 'value': '⚠️' + _pick_hint + '\n\n' + ('\n'.join(supplier_details) if supplier_details else '暂无报价')},
                    {'name': '选定供应商 ', 'value': _default_opt if _default_opt else ''},  # ⚠️ 纯文本非数组
                    {'name': '备注', 'value': '默认=按各物资最低价择优采购(每项选报价最低的厂家, 采购员分项定标生成订单); 如需指定厂家, 请改选: 厂家A=Excel最左/厂家B=中间/厂家C=最右'},
                ]
                # V11.135: 比价单Excel作为钉钉审批附件(领导可直接查看完整比价表)
                try:
                    _xlsx = gen_inquiry_xlsx_file(biz_id)
                    if _xlsx:
                        form.append({'name': '附件', 'value': json.dumps(
                            [{'path': os.path.join(BASE, 'uploads', _xlsx), 'name': _xlsx, 'cat': '比价单'}],
                            ensure_ascii=False)})
                except Exception as _e:
                    log('系统', '比价单附件生成失败', f'inquiry#{biz_id}: {str(_e)[:100]}')
                return form
            else:  # purchase_order
                cat = dt_cat_option(r['category'] or r['item_name'] or '')
                purpose = f"订单 {r['order_no']} {r['item_name'] or ''}"[:200]
                target = str(r['target_date'] or today)[:10]
                form = [
                    {'name': '部门', 'value': '["选项一"]'},
                    {'name': '销售方式', 'value': '["选项一"]'},
                    {'name': '订单图片', 'value': '[]'},
                    {'name': '备注', 'value': detail[:1900]},
                ]
                # 附件: 订单Excel凭证 → 钉钉附件
                attach = dt_build_attachment(biz_type, r, c)
                if attach:
                    form.append({'name': '附件', 'value': json.dumps(attach, ensure_ascii=False)})
                c.close()
                return form
    # 其他业务类型: 回退到原有字段组装(按各自模板配置)
    return [
        {'name': '采购名称', 'value': str(info[1])[:60]},
        {'name': '采购数量', 'value': '1'},
        {'name': '金额（元）', 'value': f"{float(info[2] or 0):.2f}"},
        {'name': '用途说明', 'value': f"{info[0]} {info[1]}"[:200]},
        {'name': '日期', 'value': today},
    ]


def dt_build_detail(biz_type, r, c):
    """V6.0: 组装单据完整详情结构化文本(钉钉备注字段, 审批人无需跳回系统)"""
    from collections import OrderedDict
    lines = OrderedDict()
    f = lambda v: str(v) if v is not None else ''
    # V11.206: 集体验收审批详情 = 入库单详情 + 验收标记说明
    if biz_type == 'collect_accept':
        biz_type = 'receiving'
    # V11.208: 维修采购定损审批详情
    if biz_type == 'repair_plan':
        lines['单据编号'] = r['plan_no']; lines['单据类型'] = '维修采购'
        lines['损坏设备'] = r['device_name']; lines['故障描述'] = str(r['fault_desc'] or '')[:200]
        lines['所属部门'] = r['dept'] or '-'; lines['发起人'] = r['requester'] or '-'
        its = c.execute("SELECT * FROM repair_items WHERE plan_id=? ORDER BY id", (r['id'],)).fetchall()
        if its:
            lines['更换部件'] = '、'.join(f"{x['part_name']}({x['fault_note'] or '待定损'})" for x in its)[:300]
        lines['维修厂家'] = r['repair_company'] or '待报价'
        lines['报价合计'] = f"¥{float(r['quote_total'] or 0):,.2f}"
        return '\n'.join(f"{k}: {v}" for k, v in lines.items())
    if biz_type == 'purchase_request':
        lines['单据编号'] = r['req_no']; lines['单据类型'] = '采购申请'
        lines['申请人'] = r['requester']; lines['申请部门'] = r['dept']
        lines['申请时间'] = str(r['created_at'] or '')[:16]
        lines['预算归属'] = r['budget_code'] or '-'
        lines['采购用途'] = r['purpose']
        lines['需求到货'] = str(r['target_date'] or '')[:10]
        lines['预估总金额'] = f"¥{float(r['total_estimated'] or 0):,.2f}"
        lines['紧急等级'] = '🚨加急' if r['urgent'] else '普通'
        # 明细子表
        its = c.execute("SELECT * FROM request_items WHERE req_id=? ORDER BY id", (r['id'],)).fetchall()
        if its:
            lines['商品明细'] = ''
            for i, it in enumerate(its, 1):
                extra = ' '.join(x for x in [it['category'] if 'category' in it.keys() else '', it['brand_param'] if 'brand_param' in it.keys() else ''] if x)
                arr = it['arrival_date'] if 'arrival_date' in it.keys() and it['arrival_date'] else ''
                lines['商品明细'] += f"{i}. {it['item_name']} {it['spec'] or ''} x{it['quantity']}{it['unit'] or ''} 单价¥{float(it['estimated_price'] or 0):.2f} 小计¥{float(it['total_price'] or 0):.2f}" + (f" ({extra})" if extra else '') + (f" 到货:{arr}" if arr else '') + "\n"
        if r['remark']: lines['备注'] = r['remark']
    elif biz_type == 'contract':
        lines['合同编号'] = r['contract_no']; lines['单据类型'] = '采购合同'
        lines['合同名称'] = r['contract_name']; lines['合作供应商'] = r['supplier']
        lines['合同总金额'] = f"¥{float(r['amount'] or 0):,.2f}"
        lines['签订日期'] = str(r['sign_date'] or '')[:10]
        lines['履约周期'] = f"{str(r['start_date'] or '')[:10]} ~ {str(r['end_date'] or '')[:10]}"
        lines['紧急等级'] = '🚨加急' if r['urgent'] else '普通'
        # V11.203 模块一1.1/1.2: 发票条款/开票计划/发票回收状态 — 系统详情与钉钉审批备注同步可见
        if 'invoice_clause' in r.keys() and r['invoice_clause']:
            lines['发票条款'] = r['invoice_clause']
        if 'invoice_est_first' in r.keys() and r['invoice_est_first']:
            lines['预计首次开票'] = r['invoice_est_first']
        if 'invoice_est_done' in r.keys() and r['invoice_est_done']:
            lines['预计全部开票完成'] = r['invoice_est_done']
        if 'inv_collect_status' in r.keys() and r['inv_collect_status']:
            lines['发票回收'] = r['inv_collect_status']
        # 关联订单/入库: 履约与质保信息
        if r['order_id']:
            po = c.execute("SELECT * FROM purchase_orders WHERE id=?", (r['order_id'],)).fetchone()
            if po:
                lines['来源订单'] = po['order_no']
                lines['付款方式'] = po['trade_mode'] or '-'
        if r['content']:
            txt = str(r['content'])[:400]
            lines['合同内容摘要'] = txt
        if r['remark']: lines['备注'] = r['remark']
    elif biz_type == 'purchase_order':
        lines['订单编号'] = r['order_no']; lines['单据类型'] = '采购订单'
        lines['供应商'] = r['supplier']; lines['申请人'] = r['requester']
        lines['下单时间'] = str(r['created_at'] or '')[:16]
        lines['交易模式'] = r['trade_mode'] or '货到付款'
        lines['紧急等级'] = '🚨加急' if r['urgent'] else '普通'
        lines['订单金额'] = f"¥{float(r['total_amount'] or 0):,.2f}"
        lines['物资名称'] = r['item_name']; lines['规格型号'] = r['spec'] or ''
        lines['数量'] = f"{r['quantity']}{r['unit'] or ''}"
        lines['单价'] = f"¥{float(r['price'] or 0):.2f}"
        lines['税率'] = f"{r['tax_rate'] or 13}%"
        lines['需求到货'] = str(r['target_date'] or '')[:10]
        if r['remark']: lines['备注'] = r['remark']
    elif biz_type == 'receiving':
        lines['入库单号'] = r['receive_no']; lines['单据类型'] = '入库单'
        # V11.29: 归属部门(钉钉审批也可见)
        lines['归属部门'] = r['dept'] if 'dept' in r.keys() and r['dept'] else '-'
        lines['关联订单'] = r['order_no'] if 'order_no' in r.keys() and r['order_no'] else ('#' + str(r['order_id']) if r['order_id'] else '-')
        lines['供应商'] = r['supplier'] if 'supplier' in r.keys() and r['supplier'] else '-'
        lines['仓库'] = r['warehouse'] or '主库房'
        lines['验收人'] = r['inspector'] or '-'
        lines['提交时间'] = str(r['created_at'] or '')[:16]
        lines['合格数量'] = f"{r['qualified_qty'] or r['quantity'] or 0}{r['unit'] or ''}"
        # V11.202 分批验收: 钉钉审批详情同步展示 批次/本批入库数量/订单待验收待定余量(仅关联订单或分批单)
        if 'batch_no' in r.keys() and r['batch_no']:
            lines['批次'] = "%s · %s" % (r['batch_no'], '暂估入库' if r['is_est'] else '正式入库')
        if r['order_id']:
            try:
                _ost = _order_rcv_stats(c, r['order_id'])
                lines['本批入库数量'] = "%g%s" % ((_rcv_doc_qty(dict_row(r)) or 0), r['unit'] or '')
                lines['订单验收进度'] = "订单总%g｜已验收入库%g｜待验收待定%g" % (_ost['order_total'], _ost['accepted'], _ost['pending'])
            except Exception:
                pass
        # 明细(手动入库单 items_json / 关联订单 order_items)
        its = []
        if r['items_json']:
            try: its = json.loads(r['items_json'])
            except Exception: its = []
        if not its and r['order_id']:
            its = [dict(x) for x in c.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (r['order_id'],)).fetchall()]
        if its:
            lines['商品明细'] = ''
            for i, it in enumerate(its, 1):
                lines['商品明细'] += f"{i}. {it.get('item_name','')} {it.get('spec','') or ''} x{it.get('quantity',0)}{it.get('unit','') or ''}" + (f" 单价¥{float(it.get('price',0) or 0):.2f}" if it.get('price') is not None else '') + "\n"
        if r['remark']: lines['备注'] = r['remark']
    elif biz_type == 'requisition':
        lines['出库单号'] = r['req_no']; lines['单据类型'] = '出库单'
        lines['领用部门'] = r['dept'] or '-'; lines['领用人'] = r['requester'] or '-'
        # V11.25: 领取人/领取部门(出库追溯)
        lines['领取人'] = r['receiver'] or r['requester'] or '-'
        lines['领取部门'] = r['receive_dept'] or r['dept'] or '-'
        lines['提交时间'] = str(r['created_at'] or '')[:16]
        lines['数量'] = f"{r['quantity'] or 0}{r['unit'] or ''}"
        lines['用途'] = r['purpose'] or '-'
        its = c.execute("SELECT * FROM requisition_items WHERE requisition_id=? ORDER BY id", (r['id'],)).fetchall()
        if its:
            lines['商品明细'] = ''
            for i, it in enumerate(its, 1):
                lines['商品明细'] += f"{i}. {it['item_name']} {it['spec'] or ''} x{it['quantity']}{it['unit'] or ''}\n"
    elif biz_type == 'payment':
        lines['付款单号'] = r['payment_no']; lines['单据类型'] = '付款申请'
        lines['付款事由'] = r['payment_reason'] or '-'
        lines['供应商'] = r['supplier'] or '-'; lines['付款金额'] = f"¥{float(r['amount'] or 0):,.2f}"
        lines['期望付款日期'] = str(r['expect_pay_date'] or '')[:10]
        lines['发票类型'] = r['invoice_type'] or '-'
        lines['是否签合同'] = '是' if r['has_contract'] == '是' else '否'
        lines['收款人'] = f"{r['payee_name'] or '-'} {r['payee_account'] or ''}"
        lines['交易模式'] = r['trade_mode'] or '-'
        lines['紧急等级'] = '🚨加急' if r['urgent'] else '普通'
        if r['remark']: lines['备注'] = r['remark']
    elif biz_type == 'return_request':
        # V11.202 退库审批钉钉详情(退库审批模板的备注字段内容)
        lines['退库单号'] = r['return_no']; lines['单据类型'] = '退库单'
        lines['来源出库单'] = r['source_req_no'] or ('#' + str(r['source_req_id']) if r['source_req_id'] else '-')
        lines['退库部门'] = r['dept'] or '-'
        lines['退库人'] = r['receiver'] or r['requester'] or '-'
        lines['仓库'] = r['warehouse'] or '-'
        lines['退库原因'] = ((r['reason'] or '') + (('：' + str(r['reason_note'])) if r['reason_note'] else '')) or '-'
        lines['退库金额'] = f"¥{float(r['total_amount'] or 0):,.2f}"
        lines['提交时间'] = str(r['created_at'] or '')[:16]
        its = c.execute("SELECT * FROM return_items WHERE return_id=? ORDER BY id", (r['id'],)).fetchall()
        if its:
            lines['商品明细'] = ''
            for i, it in enumerate(its, 1):
                lines['商品明细'] += f"{i}. {it['item_name']} {it['spec'] or ''} x{it['return_qty']}{it['unit'] or ''}\n"
    # 组装多行文本: "字段: 值"
    out = []
    for k, v in lines.items():
        if k == '备注' and v:
            out.append(f"【{k}】\n{v}")
        elif v or k in ('商品明细', '入库明细'):
            out.append(f"【{k}】\n{v}".rstrip())
    return "\n".join(out)


def gen_doc_voucher(biz_type, biz_id, kind, title):
    """V11.9b: 单据附件凭证 — 复用系统现有下载接口生成的文件(与页面下载完全一致)
    申请/订单/入库/出库 → 调对应 /download 接口拿xlsx存uploads; 已存在则复用"""
    try:
        no = None
        c = db()
        r = c.execute(f"SELECT * FROM {biz_table(biz_type)} WHERE id=?", (biz_id,)).fetchone()
        if r:
            no = r['order_no'] if kind == 'order' else (r['req_no'] if kind == 'prequest' else (r['receive_no'] if kind == 'receiving' else r['req_no']))
        c.close()
        if not no:
            return None
        fname = f"voucher_{kind}_{no}.xlsx"
        fpath = os.path.join(BASE, 'uploads', fname)
        if os.path.exists(fpath):
            return fname
        # 内部调用现有下载接口(与页面"下载"按钮完全同一份生成逻辑)
        url = {'prequest': f'/api/prequests/{biz_id}/download',
               'order': f'/api/orders/{biz_id}/download',
               'receiving': f'/api/receivings/{biz_id}/download',
               'requisition': f'/api/requisitions/{biz_id}/download'}.get(kind)
        if not url:
            return None
        client = app.test_client()
        client.post('/api/login', json={'username': 'admin', 'password': 'admin123'})
        resp = client.get(url)
        if resp.status_code != 200 or not resp.data:
            log('系统', '单据凭证生成失败', f'{biz_type}#{biz_id}: 下载接口HTTP {resp.status_code}')
            return None
        os.makedirs(os.path.join(BASE, 'uploads'), exist_ok=True)
        with open(fpath, 'wb') as f:
            f.write(resp.data)
        return fname
    except Exception as e:
        log('系统', '单据凭证生成失败', f'{biz_type}#{biz_id}: {str(e)[:120]}')
        return None


def dt_build_attachment(biz_type, r, c):
    """V11.9: 组装钉钉附件字段值 — 每类单据自动生成标准Excel凭证后上传
    申请/订单/入库/出库 → 自动生成凭证xlsx存uploads; 合同 → 已有docx
    返回 [{path,name,cat}] 本地路径标记, 由 dt_resolve_attachments 上传钉钉"""
    try:
        files = []
        # 合同: 自动生成的合同文件
        if biz_type == 'contract' and r['file_path']:
            files.append({'path': os.path.join(BASE, 'uploads', r['file_path']), 'name': r['file_path'],
                          'cat': '合同类'})
        # V11.9: 申请/订单/入库/出库 → 自动生成Excel凭证(领导审批可见)
        gen_map = {'purchase_request': ('prequest', '采购申请单'),
                   'purchase_order': ('order', '采购订单'),
                   'receiving': ('receiving', '入库验收单'),
                   'requisition': ('requisition', '出库单')}
        if biz_type in gen_map and r['id']:
            try:
                fn = gen_doc_voucher(biz_type, r['id'], gen_map[biz_type][0], gen_map[biz_type][1])
                if fn:
                    files.append({'path': os.path.join(BASE, 'uploads', fn), 'name': fn, 'cat': gen_map[biz_type][1]})
            except Exception as e:
                log('系统', '单据凭证生成失败', f'{biz_type}#{r["id"]}: {str(e)[:120]}')
        # 申请: attachments JSON 数组(纯文件名)
        for key in ('attachments', 'attachment'):
            if key in r.keys() and r[key]:
                try:
                    lst = json.loads(r[key]) if isinstance(r[key], str) else (r[key] or [])
                except Exception:
                    lst = [r[key]]
                for a in lst:
                    if isinstance(a, str) and a.strip():
                        files.append({'path': os.path.join(BASE, 'uploads', a.strip()), 'name': a.strip(), 'cat': '附件'})
        # 只保留真实存在的文件
        exist = [f for f in files if os.path.exists(f['path'])]
        return exist
    except Exception:
        return []


def dt_storage_space_id():
    """V8.5: 获取钉钉企业存储空间(ORG全员空间) — 审批附件存放于此, 组织内成员均可预览下载
    优先级: ①sys_config缓存(有效数字ID) ②审批钉盘空间(workflow域, 无需Storage权限) ③创建ORG空间(需Storage.Space.Write)
    返回 '' 表示不可用"""
    sid = cfg_get('dingtalk_storage_space_id', '')
    if sid and str(sid).isdigit():
        return sid
    if sid:  # 无效占位符(如 SPACE123) → 清除
        cfg_set('dingtalk_storage_space_id', '')
    # ① 审批钉盘空间(无需额外权限, 稳定可用)
    try:
        c, r = dt_new_post('/v1.0/workflow/processInstances/spaces/infos/query',
                           {'userId': dt_first_bound_userid() or '', 'agentId': dt_agent_id()})
        if c == 0 and r.get('result', {}).get('spaceId'):
            sid = str(r['result']['spaceId'])
            cfg_set('dingtalk_storage_space_id', sid)
            return sid
    except Exception as e:
        log('系统', '钉钉审批空间查询异常', str(e)[:120])
    # ② 创建ORG空间(需 Storage.Space.Write 权限)
    try:
        c2, r2 = dt_new_post('/v1.0/storage/spaces', {
            'option': {'name': '采购系统审批附件', 'ownerType': 'ORG', 'quota': 0}})
        if c2 == 0 and r2.get('space', {}).get('id'):
            sid = str(r2['space']['id'])
            cfg_set('dingtalk_storage_space_id', sid)
            return sid
        log('系统', '钉钉存储空间创建失败', json.dumps(r2, ensure_ascii=False)[:200])
    except Exception as e:
        log('系统', '钉钉存储空间创建异常', str(e)[:150])
    return ''


def dt_upload_file_to_dingtalk(path, filename):
    """V8.5: 本地文件上传钉钉企业存储(v1.0 storage: GetFileUploadInfo→签名URL直传→CommitFile)
    返回 {spaceId, fileName, fileSize, fileType, fileId} — OA审批附件组件必需结构。失败返回 None"""
    try:
        sid = dt_storage_space_id()
        if not sid:
            return None
        fname = os.path.basename(filename)
        ext = os.path.splitext(fname)[1].lower().lstrip('.')
        ftype = 'image' if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp') else 'file'
        with open(path, 'rb') as f:
            data = f.read()
        fsize = len(data)
        try:
            import hashlib
            md5 = hashlib.md5(data).hexdigest()
        except Exception:
            md5 = ''
        # 1. 获取上传信息(签名URL + uploadKey) — V11.8: URL必须带 ?unionId=
        _uid = dt_union_id()
        if not _uid:
            log('系统', '钉钉附件上传失败', f'{fname}: 未获取到unionId')
            return None
        c, r = dt_new_post(f'/v1.0/storage/spaces/{sid}/files/uploadInfos/query?unionId={_uid}', {
            'multipart': False,
            'protocol': 'HEADER_SIGNATURE',
            'option': {'preCheckParam': {'name': fname, 'size': fsize, 'md5': md5}}})
        if c != 0:
            # V10.4 自动重试: 空间授权过期(no privilege/permissionDenied) → 清缓存重新获取空间+unionId → 重试一次
            _msg = json.dumps(r, ensure_ascii=False)[:300] if r else ''
            if 'privilege' in _msg or 'permissionDenied' in _msg:
                log('系统', '钉钉空间授权过期', f'{fname}: {_msg[:80]} → 清缓存重取空间重试')
                cfg_set('dingtalk_storage_space_id', '')
                cfg_set('dingtalk_union_id', '')
                sid2 = dt_storage_space_id()
                _uid2 = dt_union_id()
                if sid2 and _uid2:
                    c, r = dt_new_post(f'/v1.0/storage/spaces/{sid2}/files/uploadInfos/query?unionId={_uid2}', {
                        'multipart': False,
                        'protocol': 'HEADER_SIGNATURE',
                        'option': {'preCheckParam': {'name': fname, 'size': fsize, 'md5': md5}}})
            if c != 0:
                log('系统', '钉钉附件上传失败', f'{fname}: 获取上传信息失败 {json.dumps(r, ensure_ascii=False)[:200]}')
                return None
        upload_key = r.get('uploadKey', '')
        hsi = r.get('headerSignatureInfo') or {}
        urls = hsi.get('resourceUrls') or []
        headers = hsi.get('headers') or {}
        if not upload_key or not urls:
            log('系统', '钉钉附件上传失败', f'{fname}: 无签名URL/uploadKey')
            return None
        # 2. 上传文件二进制到签名URL(直传 OSS) — V11.8: 必须用 http.client 手动PUT
        #    (urllib 带 data 会自动加 Content-Type, 破坏 OSS 签名 → 403 SignatureDoesNotMatch)
        try:
            import http.client
            from urllib.parse import urlparse
            _u = urlparse(urls[0])
            _conn = http.client.HTTPSConnection(_u.hostname, timeout=60)
            _conn.putrequest('PUT', _u.path + ('?' + _u.query if _u.query else ''))
            for _k, _v in (headers or {}).items():
                _conn.putheader(_k, _v)
            _conn.putheader('Content-Length', str(len(data)))
            _conn.endheaders()
            _conn.send(data)
            _resp = _conn.getresponse()
            _body = _resp.read()
            _conn.close()
            if _resp.status != 200:
                log('系统', '钉钉附件上传失败', f'{fname}: OSS直传HTTP {_resp.status} {_body[:120]}')
                return None
        except Exception as e:
            log('系统', '钉钉附件上传异常', f'{fname}: {str(e)[:120]}')
            return None
        # 3. 提交文件(dentry) → 获得 fileId — V11.8: URL带 ?unionId=
        c2, r2 = dt_new_post(f'/v1.0/storage/spaces/{sid}/files/commit?unionId={_uid}', {
            'name': fname, 'parentId': '0', 'uploadKey': upload_key})
        if c2 != 0:
            log('系统', '钉钉附件上传失败', f'{fname}: 提交失败 {json.dumps(r2, ensure_ascii=False)[:200]}')
            return None
        dentry = r2.get('dentry') or {}
        fid = dentry.get('id') or dentry.get('uuid') or ''
        if not fid:
            log('系统', '钉钉附件上传失败', f'{fname}: 提交成功但无fileId')
            return None
        return {'spaceId': sid, 'fileName': fname, 'fileSize': fsize, 'fileType': ftype, 'fileId': fid}
    except Exception as e:
        log('系统', '钉钉附件上传异常', str(e)[:150])
        return None


def dt_resolve_attachments(form, biz_type, biz_id):
    """V6.0: 表单中附件字段(本地路径标记) → 上传钉钉存储 → 替换为钉钉附件结构
    返回上传成功的附件数"""
    n = 0
    for fld in form:
        if fld.get('name') != '附件':
            continue
        try:
            local = json.loads(fld['value']) if isinstance(fld['value'], str) else fld['value']
        except Exception:
            continue
        if not isinstance(local, list):
            continue
        up = []
        for it in local:
            p = it.get('path', '')
            if not p or not os.path.exists(p):
                continue
            d = dt_upload_file_to_dingtalk(p, it.get('name', os.path.basename(p)))
            if d:
                up.append(d); n += 1
        if up:
            fld['value'] = json.dumps(up, ensure_ascii=False)
        else:
            # 无附件可上传 → 移除附件字段(避免空附件报错)
            form.remove(fld)
        break
    return n


def dt_start_instance(biz_type, biz_id):
    """单据进入待审批后, 同步发起钉钉审批实例(新版接口); 成功返回instance_id, 失败返回None"""
    try:
        if not dingtalk_enabled(): return None
        code = dt_approval_code(biz_type)
        if not code:
            # V11.202: 缺模板码时留痕, 不再静默(用户查"为何没推到钉钉"有据可依)
            log('系统', '钉钉审批未发起', f'{biz_type}#{biz_id}: 未配置钉钉审批模板码(请在系统设置-钉钉设置填写 {DT_BIZ.get(biz_type, biz_type)} 的PROC模板码)')
            return None
        ak = dt_actioner_key(biz_type)
        c = db()
        if c.execute("SELECT COUNT(*) FROM dingtalk_instances WHERE biz_type=? AND biz_id=? AND status NOT IN ('error','cancelled')", (biz_type, biz_id)).fetchone()[0] > 0:
            c.close(); return None
        levels = c.execute("SELECT DISTINCT role, approver FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no", (biz_type, biz_id)).fetchall()
        c.close()
        if not levels: return None
        approvers, missing = [], []
        for lv in levels:
            # V5.0+: 优先用审批节点配置的具体审批人(其钉钉ID); 未配置则按角色严格解析
            u = None
            if lv['approver']:
                cc = db()
                u = cc.execute("SELECT * FROM users WHERE name=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!=''", (lv['approver'],)).fetchone()
                if u is None:
                    # 配置的名字可能是用户名
                    u = cc.execute("SELECT * FROM users WHERE username=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!=''", (lv['approver'],)).fetchone()
                cc.close()
            if u is None:
                u = find_approver_for_role(lv['role'], 'dingtalk')
            if u and u['dingtalk_userid']: approvers.append(u['dingtalk_userid'])
            else: missing.append(f"{lv['role']}{(':'+lv['approver']) if lv['approver'] else ''}")
        if missing:
            # 需求(2026-08-06): 模板里选的审批人必须是钉钉通讯录成员, 系统对应角色必须是有效用户
            # 任一角色缺失 → 发起失败并明确记录, 不再用 admin 静默顶替
            log('系统', '钉钉审批未发起', f"{biz_type}#{biz_id} 以下节点无有效审批人(未启用或未绑定钉钉通讯录): {','.join(missing)}")
            return None
        info = dt_biz_info(biz_type, biz_id)
        if not info: return None
        originator = ''
        ua = find_user_by_name(str(info[3]))
        if ua and ua['dingtalk_userid']: originator = ua['dingtalk_userid']
        if not originator:
            # 兜底: 第一个已绑定钉钉ID的审批人, 再不行用系统内任意已绑定用户
            originator = approvers[0] if approvers else dt_first_bound_userid()
        if not originator:
            log('系统', '钉钉审批未发起', f"{biz_type}#{biz_id} 无已绑定钉钉ID的用户")
            return None
        form = dt_build_form(biz_type, biz_id, info)
        # V6.0: 系统附件 → 上传钉钉存储 → 挂载审批附件(钉钉端在线预览/下载)
        try:
            n_attach = dt_resolve_attachments(form, biz_type, biz_id)
            if n_attach:
                log('系统', '钉钉审批附件', f"{info[0]} 挂载{n_attach}个附件")
        except Exception:
            pass
        payload = {
            'originatorUserId': originator, 'processCode': code, 'deptId': 1,
            'formComponentValues': form,
        }
        if ak:
            # 模板含"自选审批人"节点时指定审批人; 否则走模板默认审批流
            payload['targetSelectActioners'] = [{'actionerKey': ak, 'actionerUserIds': approvers}]
        # 2026-08-06: 新API(v1.0/workflow/processInstances)对该模板报
        # processInstanceStartFailed, 旧API(/topapi/processinstance/create)实测成功 → 创建走旧API
        try:
            _create_payload = {
                'process_code': code,
                'originator_user_id': originator,
                'dept_id': 1,
                'form_component_values': form,
                'agent_id': dt_agent_id(),
            }
            # V8.0: 模板审批流为"发起人自选"时, 旧API用 approvers 指定审批人
            # (模板固定审批人时传了也不生效, 需在钉钉后台把审批节点改为自选/按角色)
            # V8.2: 只要有审批人就传 approvers(自选节点即生效, 无需 actionerKey)
            if approvers:
                _create_payload['approvers'] = approvers
            code_r, resp = dt_post('/topapi/processinstance/create', _create_payload)
            if code_r == 0:
                iid = resp.get('process_instance_id', '') or resp.get('process_instance', {}).get('instance_id', '')
                c = db()
                c.execute("INSERT INTO dingtalk_instances(instance_code,biz_type,biz_id,status) VALUES(?,?,?,'pending')", (iid, biz_type, biz_id))
                c.commit(); c.close()
                log('系统', '发起钉钉审批', f"{info[0]} 审批人{len(approvers)}级")
                dt_send_todo(approvers, f"新的{DT_BIZ[biz_type]}待处理", f"{info[0]} {info[1]}", f"金额 ¥{float(info[2] or 0):,.0f}", biz_type, biz_id)
                return iid
        except Exception as e:
            log('系统', '钉钉审批发起异常', f"{biz_type}#{biz_id}: {e}")
            return None
        c = db()
        # V11.182: 防死循环 — 同一单据已存在 error 记录则只更新错误信息, 不再无限插新记录
        # (dt_retry_failed_instances 每轮重试失败曾导致单据#22 堆积2194条ERR, 数据库膨胀)
        _ex = c.execute("SELECT id FROM dingtalk_instances WHERE biz_type=? AND biz_id=? AND status='error' ORDER BY id DESC LIMIT 1", (biz_type, biz_id)).fetchone()
        if _ex:
            c.execute("UPDATE dingtalk_instances SET error=?, updated_at=? WHERE id=?",
                      (json.dumps(resp, ensure_ascii=False)[:500], now(), _ex['id']))
        else:
            c.execute("INSERT INTO dingtalk_instances(instance_code,biz_type,biz_id,status,error) VALUES(?,?,?,'error',?)",
                      (f'ERR-{biz_type}-{biz_id}-{int(time.time())}', biz_type, biz_id, json.dumps(resp, ensure_ascii=False)[:500]))
        c.commit(); c.close()
        log('系统', '钉钉审批发起失败', f"{biz_type}#{biz_id}: {json.dumps(resp, ensure_ascii=False)[:200]}")
        return None
    except Exception as e:
        log('系统', '钉钉审批发起异常', f"{biz_type}#{biz_id}: {e}")
        return None

# ---- 工作通知(action_card, "去处理"直达系统审批页; 审批实例本身钉钉会另行通知) ----
def dt_send_todo(userids, title, text, extra='', biz_type='', biz_id=0, push_type='auto', operator='系统'):
    """V7.0: 钉钉工作通知推送(双触达: 工作台待办+工作通知) + 推送记录留存
    push_type: auto=自动节点推送 / urgent=手动加急提醒 / overdue=超期提醒 / alert=业务提醒 / test=测试
    V8.3: 审批类提醒(auto/urgent)改用钉钉OA审批原生通知(实例创建后钉钉自动通知审批人),
    不再经工作通知机器人(jiao助手)推送; 开关 sys_config.dingtalk_oa_notify_only=1 时生效"""
    try:
        # V11.172: 手动加急(urgent)强制推送 — 用户点加急=明确要催办, 不被 oa_notify_only 拦截;
        # 仅自动节点推送(auto)走OA审批原生通知(钉钉会自动提醒审批人, 避免重复打扰)
        if push_type == 'auto' and cfg_get('dingtalk_oa_notify_only', '0') == '1':
            return False
        agent = dt_agent_id()
        if not agent: return False
        userids = [u for u in userids if u]
        if not userids: return False
        # V11.174: "去处理"优先跳钉钉OA审批实例页(审批人点开直接在钉钉里审批, 不再绕采购系统网页)
        # V11.192: 改为跳系统免登直达页 — 原n.dingtalk.com OA H5 对申请人/非当前审批人打开白屏/跳登录
        # (钉钉OA H5 需审批登录态+可信域名); 系统 /sso/goto 自动钉钉免登+跳对应单据详情/审批弹窗, 复制链接兜底
        url = ''
        _inst_code = ''
        if biz_type and biz_id:
            try:
                c2 = db()
                _row = c2.execute("SELECT instance_code FROM dingtalk_instances WHERE biz_type=? AND biz_id=? AND status NOT IN ('error','cancelled') ORDER BY id DESC LIMIT 1", (biz_type, biz_id)).fetchone()
                c2.close()
                if _row and _row['instance_code'] and not str(_row['instance_code']).startswith('ERR-'):
                    _inst_code = str(_row['instance_code'])
            except Exception:
                pass
        # V11.192: 跳系统免登直达页(带单据定位), 全部通知/加急统一; 审批提醒用act=approve, 结果通知用act=detail
        _act = 'approve' if push_type in ('auto', 'urgent', 'overdue') else 'detail'
        _pu = dt_public_url()
        if _pu:
            url = _pu.rstrip('/') + f'/sso/goto?biz={biz_type}&id={biz_id}&act={_act}' if biz_type and biz_id else _pu.rstrip('/') + '/#approvals'
        _copy_url = url
        msg = {'msgtype': 'action_card', 'action_card': {
            'title': title,
            'markdown': text + ('\n' + extra if extra else '') + (f'\n\n📎 打不开请点上方按钮，或复制链接到浏览器：\n{_copy_url}' if _copy_url else ''),
            'btn_orientation': '1',
            'btn_json_list': [{'title': '去处理' if push_type in ('auto', 'urgent', 'overdue') else '查看详情', 'action_url': url}] if url else [],
        }}
        code_r, resp = dt_post('/topapi/message/corpconversation/asyncsend_v2', {
            'agent_id': agent, 'userid_list': ','.join(userids), 'msg': msg,
        })
        if code_r == 0:
            # 推送记录留存(7.0优化1: 每次推送留痕可查)
            try:
                c = db()
                names = []
                for uid in userids:
                    u = c.execute("SELECT name FROM users WHERE dingtalk_userid=?", (uid,)).fetchone()
                    names.append(u['name'] if u else uid)
                c.execute("INSERT INTO dingtalk_push_log(biz_type,biz_id,doc_no,push_type,target_user,target_userid,operator,content) VALUES(?,?,?,?,?,?,?,?)",
                          (biz_type or '', str(biz_id), '', push_type, '、'.join(names), ','.join(userids),
                           operator, f"{title} | {text}"[:200]))
                c.commit(); c.close()
            except Exception:
                pass
        return code_r == 0
    except Exception:
        return False


def dt_urgent_remind(biz_type, biz_id, operator):
    """V7.0优化1-2: 手动加急提醒 — 向当前待审批节点负责人钉钉二次推送
    限频: 同一单据 10 分钟内不可重复(可后台配置 urgent_limit_min)
    返回 (success, message)"""
    try:
        if not dingtalk_enabled():
            return False, '钉钉未启用'
        # 限频检查(默认1分钟 — V8.0: 原10分钟太慢, 加急提醒应可快速连续触发)
        limit_min = int(cfg_get('urgent_limit_min', '1') or 1)
        c = db()
        last = c.execute("SELECT MAX(created_at) FROM dingtalk_push_log WHERE biz_type=? AND biz_id=? AND push_type='urgent'", (biz_type, str(biz_id))).fetchone()[0]
        if last:
            try:
                from datetime import datetime as _dt
                last_dt = _dt.strptime(last, '%Y-%m-%d %H:%M:%S')
                if (_dt.now() - last_dt).total_seconds() < limit_min * 60:
                    c.close()
                    return False, f'{limit_min}分钟内已加急提醒过，请稍后再试'
            except Exception:
                pass
        # 找当前待审批节点负责人
        cur = c.execute("SELECT * FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no LIMIT 1", (biz_type, biz_id)).fetchone()
        if not cur:
            c.close(); return False, '该单据无待审批节点'
        u = None
        if cur['approver_id']:
            u = c.execute("SELECT * FROM users WHERE id=?", (cur['approver_id'],)).fetchone()
        # V11.176: 优先按审批实例配置的审批人(用户名/姓名)精确匹配 — 原逻辑直接按角色取第一个
        # 用户, 导致审批人配置是xingguo(邢果)却推给了同角色的赵培姝
        if not u and cur['approver']:
            u = c.execute("SELECT * FROM users WHERE username=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!='' ORDER BY id LIMIT 1", (cur['approver'],)).fetchone()
            if u is None:
                u = c.execute("SELECT * FROM users WHERE name=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!='' ORDER BY id LIMIT 1", (cur['approver'],)).fetchone()
        if not u:
            u = c.execute("SELECT * FROM users WHERE role=? AND is_active=1 AND dingtalk_userid IS NOT NULL AND dingtalk_userid!='' ORDER BY id LIMIT 1", (cur['role'],)).fetchone()
        c.close()
        if not u or not u['dingtalk_userid']:
            return False, '当前审批人未绑定钉钉'
        # V8.3: 审批提醒统一走OA审批原生通知时, 手动加急不再经工作通知机器人推送
        # V11.172: 修复"加急按钮是死的" — 用户手动点加急=明确要催办, 必须实际推送,
        # 不再被 oa_notify_only 拦截(该配置只管自动节点推送, 不影响手动加急)
        doc_no = dt_biz_info(biz_type, biz_id)
        doc_no = doc_no[0] if doc_no else f'{biz_type}#{biz_id}'
        title = '⏰ 人工加急提醒，请尽快审批'
        text = f"{doc_no} 审批等待中，发起人已加急催办，请尽快处理"
        ok = dt_send_todo([u['dingtalk_userid']], title, text,
                          f"单据: {doc_no}", biz_type, biz_id, push_type='urgent', operator=operator)
        return ok, ('加急提醒已发送给 ' + (u['name'] or u['username'])) if ok else '加急提醒发送失败'
    except Exception as e:
        return False, f'加急提醒异常: {str(e)[:80]}'

# ---- 审批结果同步(幂等) ----
def dt_sync_result(instance_id, result, comment='', approver='', approver_id=0, processed_at='', attachments=None):
    """V6.0: 钉钉审批结果回写系统(幂等)
    - result: agree/reject
    - comment: 钉钉审批意见(驳回原因等) → 回写单据留痕
    - V11.175: approver/approver_id/processed_at = 钉钉实际驳回人/时间(从实例详情提取), 实时同步
    - V11.181: attachments = 驳回附件元数据列表[{fileName,fileId,spaceId,...}], 存驳回记录供查看
    审批状态双向同步: 钉钉同意→系统单据通过; 钉钉驳回→系统单据驳回+原因留痕"""
    c = db()
    r = c.execute("SELECT * FROM dingtalk_instances WHERE instance_code=?", (instance_id,)).fetchone()
    c.close()
    if not r or r['status'] in ('synced', 'error'): return
    _approver = approver or '钉钉'
    _approver_id = int(approver_id or 0)
    _proc_at = processed_at or now()
    if comment:
        # 审批意见留痕: 写入审批实例 comment + 单据 rejected_reason
        c2 = db()
        c2.execute("UPDATE approval_instances SET comment=?, approver=?, approver_id=?, processed_at=? WHERE biz_type=? AND biz_id=? AND status='pending'",
                   (str(comment)[:200], _approver, _approver_id, _proc_at, r['biz_type'], r['biz_id']))
        try:
            c2.execute(f"UPDATE {biz_table(r['biz_type'])} SET rejected_reason=? WHERE id=? AND rejected_reason IS NOT NULL",
                       (str(comment)[:200], r['biz_id']))
        except Exception:
            pass
        c2.commit(); c2.close()
        log('钉钉', '审批意见回写', f"{r['biz_type']}#{r['biz_id']} 驳回人:{_approver} 理由:{str(comment)[:100]}")
    finish_approvals(r['biz_type'], r['biz_id'], 'ok' if result == 'agree' else 'reject', _approver, _approver_id, comment or f'钉钉审批{result}',
                     attachments=attachments or [], instance_code=instance_id)
    # V11.28: 审批结果即时通知申请人(钉钉工作通知, 一次审批一次调用, 消耗可忽略)
    # V11.178: 仅"通过"在此通知(驳回由 finish_approvals 统一通知, 含理由, 避免重复)
    if result == 'agree':
        try:
            if dingtalk_enabled():
                _b = biz_table(r['biz_type'])
                _c = db()
                try:
                    _row = _c.execute(f"SELECT * FROM {_b} WHERE id=?", (r['biz_id'],)).fetchone()
                except Exception:
                    _row = None
                _c.close()
                if _row is not None:
                    _doc_no = ''
                    for _k in ('req_no', 'order_no', 'contract_no', 'receive_no', 'payment_no'):
                        if _k in _row.keys() and _row[_k]:
                            _doc_no = str(_row[_k]); break
                    _requester = None
                    for _k in ('requester', 'created_by', 'apply_by'):
                        if _k in _row.keys() and _row[_k]:
                            _requester = str(_row[_k]); break
                    if _requester:
                        _u = db()
                        try:
                            _usr = _u.execute("SELECT * FROM users WHERE name=? LIMIT 1", (_requester,)).fetchone()
                        except Exception:
                            _usr = None
                        _u.close()
                        if _usr and _usr['dingtalk_userid']:
                            dt_send_todo([_usr['dingtalk_userid']], f'审批结果通知：{_doc_no}',
                                         f"您提交的{_doc_no} 经钉钉审批 **已通过 ✅**",
                                         biz_type=r['biz_type'], biz_id=r['biz_id'], push_type='result',  # V11.145: 用result类型绕过oa_notify_only限制
                                         operator='系统')
        except Exception:
            pass


# ---- 审批结果轮询兜底(不依赖事件回调; corpId 缺失时也能同步结果) ----
def dt_query_instance(instance_id):
    """查询钉钉审批实例状态; 返回 dict 或 None"""
    try:
        code_r, resp = dt_post('/topapi/processinstance/get', {'process_instance_id': instance_id})
        if code_r != 0:
            return None
        return resp.get('process_instance') or resp.get('result') or resp
    except Exception:
        return None


def dt_extract_reject_op(inst):
    """V11.175: 从钉钉实例详情提取驳回操作(人/意见/时间)。
    V11.181: 同时提取附件(operation_records.files 及 表单附件控件), 供系统同步展示。
    钉钉 operation_records 里每个节点一条, 找出 result=REJECT 的那条。"""
    try:
        ops = inst.get('operation_records') or inst.get('operationRecords') or []
        if not isinstance(ops, list):
            return None
        for op in ops:
            if not isinstance(op, dict):
                continue
            if str(op.get('operation_result') or op.get('result') or '').upper() in ('REJECT', 'REFUSE'):
                _userid = op.get('userid') or op.get('user_id') or ''
                _name = op.get('user_name') or ''
                if not _name:
                    _u = find_user_by_dingtalk_id(_userid)
                    _name = _u['name'] if _u else ''
                if not _name:
                    _name = '钉钉'
                _comment = op.get('remark') or op.get('comment') or ''
                _ts = op.get('date') or op.get('operate_time') or op.get('create_time') or ''
                if _ts:
                    try:
                        import datetime as _dtm
                        # 钉钉返回毫秒时间戳
                        _ts = _dtm.datetime.fromtimestamp(int(_ts) / 1000).strftime('%Y-%m-%d %H:%M:%S')
                    except Exception:
                        _ts = str(_ts)[:19]
                # V11.181: 提取附件(操作记录files 或 表单附件控件)
                _att = []
                try:
                    _files = op.get('files') or op.get('attachments') or []
                    if isinstance(_files, list):
                        for _fi in _files:
                            if isinstance(_fi, dict) and (_fi.get('fileName') or _fi.get('file_name') or _fi.get('fileId') or _fi.get('file_id')):
                                _att.append({'fileName': _fi.get('fileName') or _fi.get('file_name') or '附件',
                                             'fileId': _fi.get('fileId') or _fi.get('file_id') or '',
                                             'spaceId': _fi.get('spaceId') or _fi.get('space_id') or '',
                                             'fileSize': _fi.get('fileSize') or _fi.get('file_size') or 0,
                                             'fileType': _fi.get('fileType') or _fi.get('file_type') or ''})
                except Exception:
                    pass
                if not _att:
                    # 表单附件控件(如 单据凭证xlsx/图片) 一并带上
                    try:
                        for _fc in (inst.get('form_component_values') or []):
                            if not isinstance(_fc, dict) or '附件' not in str(_fc.get('name') or ''):
                                continue
                            _v = _fc.get('value') or ''
                            _fl = json.loads(_v) if isinstance(_v, str) and _v.strip().startswith('[') else (_v if isinstance(_v, list) else [])
                            for _fi in (_fl if isinstance(_fl, list) else []):
                                if isinstance(_fi, dict) and _fi.get('fileId'):
                                    _att.append({'fileName': _fi.get('fileName') or '附件',
                                                 'fileId': str(_fi.get('fileId') or ''),
                                                 'spaceId': str(_fi.get('spaceId') or ''),
                                                 'fileSize': _fi.get('fileSize') or 0,
                                                 'fileType': _fi.get('fileType') or ''})
                    except Exception:
                        pass
                return {'approver': _name, 'approver_id': 0, 'comment': _comment, 'processed_at': _ts,
                        'attachments': _att}
    except Exception:
        pass
    return None


def find_user_by_dingtalk_id(dtid):
    """按钉钉userid查系统用户(供驳回人回填)"""
    if not dtid:
        return None
    try:
        c = db()
        r = c.execute("SELECT * FROM users WHERE dingtalk_userid=?", (dtid,)).fetchone()
        c.close()
        return r
    except Exception:
        return None


def dt_extract_agree_op(inst):
    """V11.184: 从钉钉实例详情提取同意操作(人/意见/时间/附件)。旧版operation_records无附件,
    附件通道: ①新版API operationRecords[].files(需开通, 未开通自动跳过) ②表单附件控件(发起时挂载)"""
    try:
        ops = inst.get('operation_records') or inst.get('operationRecords') or []
        if not isinstance(ops, list):
            return None
        for op in ops:
            if not isinstance(op, dict):
                continue
            if str(op.get('operation_result') or op.get('result') or '').upper() in ('AGREE', 'APPROVE'):
                _userid = op.get('userid') or op.get('user_id') or ''
                _name = op.get('user_name') or ''
                if not _name:
                    _u = find_user_by_dingtalk_id(_userid)
                    _name = _u['name'] if _u else ''
                if not _name:
                    _name = '钉钉'
                _comment = op.get('remark') or op.get('comment') or ''
                _ts = op.get('date') or op.get('operate_time') or op.get('create_time') or ''
                if _ts:
                    try:
                        import datetime as _dtm
                        _ts = _dtm.datetime.fromtimestamp(int(_ts) / 1000).strftime('%Y-%m-%d %H:%M:%S')
                    except Exception:
                        _ts = str(_ts)[:19]
                # 附件: 操作记录files + 表单附件控件
                _att = []
                try:
                    _files = op.get('files') or op.get('attachments') or []
                    if isinstance(_files, list):
                        for _fi in _files:
                            if isinstance(_fi, dict) and (_fi.get('fileName') or _fi.get('file_name') or _fi.get('fileId') or _fi.get('file_id')):
                                _att.append({'fileName': _fi.get('fileName') or _fi.get('file_name') or '附件',
                                             'fileId': _fi.get('fileId') or _fi.get('file_id') or '',
                                             'spaceId': _fi.get('spaceId') or _fi.get('space_id') or '',
                                             'fileSize': _fi.get('fileSize') or _fi.get('file_size') or 0,
                                             'fileType': _fi.get('fileType') or _fi.get('file_type') or ''})
                except Exception:
                    pass
                if not _att:
                    try:
                        for _fc in (inst.get('form_component_values') or []):
                            if not isinstance(_fc, dict) or '附件' not in str(_fc.get('name') or ''):
                                continue
                            _v = _fc.get('value') or ''
                            _fl = json.loads(_v) if isinstance(_v, str) and _v.strip().startswith('[') else (_v if isinstance(_v, list) else [])
                            for _fi in (_fl if isinstance(_fl, list) else []):
                                if isinstance(_fi, dict) and _fi.get('fileId'):
                                    _att.append({'fileName': _fi.get('fileName') or '附件',
                                                 'fileId': str(_fi.get('fileId') or ''),
                                                 'spaceId': str(_fi.get('spaceId') or ''),
                                                 'fileSize': _fi.get('fileSize') or 0,
                                                 'fileType': _fi.get('fileType') or ''})
                    except Exception:
                        pass
                return {'approver': _name, 'approver_id': 0, 'comment': _comment, 'processed_at': _ts,
                        'attachments': _att}
    except Exception:
        pass
    return None


def dt_poll_results():
    """遍历 dingtalk_instances 中 pending 状态实例, 查询钉钉侧结果并回写系统
    V11.28: 只查最近7天内活跃的pending实例(老实例不查, API消耗大降)"""
    try:
        c = db()
        rows = c.execute("""SELECT * FROM dingtalk_instances WHERE status='pending'
            AND updated_at >= datetime('now','localtime','-7 days')""").fetchall()
        c.close()
        n = 0
        for r in rows:
            try:
                inst = dt_query_instance(r['instance_code'])
                if not inst:
                    continue
                # V11.126: 保存审批表单值(询价定标审批需读取领导选的供应商) — 顺带用本次查询结果, 不额外耗API
                try:
                    _fv = inst.get('form_component_values') or inst.get('formComponentValues')
                    if _fv:
                        _c2 = db()
                        _c2.execute("UPDATE dingtalk_instances SET form_values=?, updated_at=? WHERE id=?", (json.dumps(_fv, ensure_ascii=False), now(), r['id']))
                        _c2.commit(); _c2.close()
                except Exception:
                    pass
                st = str(inst.get('status', ''))
                if st in ('APPROVED', 'COMPLETED'):
                    # COMPLETED(旧API) 需结合 result 判断 agree/reject
                    _res = str(inst.get('result', 'agree') or 'agree').lower()
                    if _res in ('refuse', 'reject'):
                        # V11.177: COMPLETED+refuse 也是驳回 — 提取驳回人/理由/时间(原只传reject丢理由)
                        _op = dt_extract_reject_op(inst)
                        if _op:
                            dt_sync_result(r['instance_code'], 'reject', _op['comment'], _op['approver'], _op['approver_id'], _op['processed_at'], _op.get('attachments') or [])
                        else:
                            dt_sync_result(r['instance_code'], 'reject')
                    else:
                        # V11.184: 同意也提取审批人/意见/附件(仅末节点有操作记录时)
                        _op = dt_extract_agree_op(inst)
                        if _op and _op.get('processed_at'):
                            dt_sync_result(r['instance_code'], 'agree' if _res in ('agree', 'agree_ok') else 'reject',
                                           _op['comment'], _op['approver'], _op['approver_id'], _op['processed_at'], _op.get('attachments') or [])
                        else:
                            dt_sync_result(r['instance_code'], 'agree' if _res in ('agree', 'agree_ok') else 'reject')
                    n += 1
                elif st in ('REJECTED',):
                    # V11.175: 提取钉钉实际驳回人+驳回意见 → 实时同步(人/时间/理由)
                    _op = dt_extract_reject_op(inst)
                    if _op:
                        dt_sync_result(r['instance_code'], 'reject', _op['comment'], _op['approver'], _op['approver_id'], _op['processed_at'], _op.get('attachments') or [])
                    else:
                        dt_sync_result(r['instance_code'], 'reject')
                    n += 1
                elif st in ('TERMINATED', 'CANCELED'):
                    dt_sync_result(r['instance_code'], 'refuse'); n += 1
            except Exception:
                continue
        return n
    except Exception:
        return 0


def dt_retry_failed_instances():
    """error 状态的钉钉实例自动重试(表单错误等修复后无需人工操作); 已过3分钟的才重试
    V11.182: 单单据error记录超5条(重试超5轮)则放弃 — 审批人无效等配置问题不再无限重试烧API"""
    try:
        c = db()
        # 只挑 error 次数≤5 的单据重试(同单已重试过多=配置问题, 放弃等人工修复)
        rows = c.execute("""SELECT d.id, d.biz_type, d.biz_id FROM dingtalk_instances d
            WHERE d.status='error' AND d.created_at <= datetime('now','localtime','-3 minutes')
            AND (SELECT COUNT(*) FROM dingtalk_instances e WHERE e.biz_type=d.biz_type AND e.biz_id=d.biz_id AND e.status='error') <= 5
            ORDER BY d.id LIMIT 5""").fetchall()
        c.close()
        for r in rows:
            try:
                dt_start_instance(r['biz_type'], r['biz_id'])
            except Exception:
                continue
    except Exception:
        pass


def dt_terminate_instance(instance_code, userid):
    """终止钉钉审批实例(系统侧已终态时调用, 避免钉钉侧继续挂起); 偶发 internalError 自动重试"""
    for attempt in range(3):
        try:
            code_r, resp = dt_new_post('/v1.0/workflow/processInstances/terminate', {
                'processInstanceId': instance_code,
                'operatingUserId': userid,
            })
            if code_r == 0:
                c = db()
                c.execute("UPDATE dingtalk_instances SET status='terminated', updated_at=? WHERE instance_code=? AND status IN ('pending','synced')", (now(), instance_code))
                c.commit(); c.close()
                log('系统', '终止钉钉审批', instance_code)
                return True
            if attempt < 2:
                time.sleep(2)
        except Exception as e:
            if attempt < 2:
                time.sleep(2)
            else:
                log('系统', '终止钉钉审批异常', f"{instance_code}: {e}")
    log('系统', '终止钉钉审批失败', f"{instance_code}: {json.dumps(resp, ensure_ascii=False)[:200]}")
    return False


def dt_sync_now(biz_type, biz_id):
    """系统内审批动作后立即同步钉钉: ①查询该单据钉钉实例最新状态回写系统;
    ②若系统侧已终态(无pending节点)且钉钉实例仍在运行 → 终止钉钉实例"""
    try:
        if not dingtalk_enabled(): return
        c = db()
        # finish_approvals 已把本地实例置 synced, 但钉钉侧实例可能仍在 RUNNING →
        # 必须查 pending+synced 两态, 才能真正终止挂起的钉钉实例
        rows = c.execute("SELECT * FROM dingtalk_instances WHERE biz_type=? AND biz_id=? AND status IN ('pending','synced')", (biz_type, biz_id)).fetchall()
        n_pending = c.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending'", (biz_type, biz_id)).fetchone()[0]
        c.close()
        for r in rows:
            inst = dt_query_instance(r['instance_code'])
            if not inst:
                continue
            st = str(inst.get('status', ''))
            if st == 'APPROVED':
                dt_sync_result(r['instance_code'], 'agree')
            elif st in ('TERMINATED', 'CANCELED'):
                dt_sync_result(r['instance_code'], 'refuse')
            elif n_pending == 0 and st in ('RUNNING', 'NEW'):
                # 系统侧已审批完但钉钉实例还在跑 → 终止, 避免双通道不一致
                u = find_user_by_role('系统管理员')
                uid = u['dingtalk_userid'] if u and u['dingtalk_userid'] else dt_first_bound_userid()
                dt_terminate_instance(r['instance_code'], uid)
    except Exception:
        pass

# ---- 旧式回调加解密(AES-CBC; 需 pycryptodome; 新式事件订阅为明文JSON无需解密) ----
def _dt_aes_key():
    k = cfg_get('dingtalk_aes_key', '').strip()
    if not k: return None
    try:
        return base64.b64decode(k + '=' * ((4 - len(k) % 4) % 4))
    except Exception:
        return None

def dt_decrypt_old(encrypt):
    """旧式回调解密: 返回 dict(JSON) 或 str(纯文本如 success)"""
    from Crypto.Cipher import AES  # pycryptodome
    key = _dt_aes_key()
    if not key or not cfg_get('dingtalk_callback_token'):
        raise RuntimeError('钉钉回调加解密未配置(dingtalk_callback_token / dingtalk_aes_key)')
    raw = base64.b64decode(encrypt)
    cipher = AES.new(key, AES.MODE_CBC, key[:16])
    plain = cipher.decrypt(raw)
    if plain and 0 < plain[-1] <= 32: plain = plain[:-plain[-1]]
    msg_len = int.from_bytes(plain[16:20], 'big')
    text = plain[20:20 + msg_len].decode('utf-8')
    try:
        return json.loads(text)
    except Exception:
        return text

def dt_encrypt_old(msg):
    from Crypto.Cipher import AES  # pycryptodome
    key = _dt_aes_key()
    corp = cfg_get('dingtalk_corp_id', '')
    if not key or not corp: raise RuntimeError('钉钉回调加解密未配置')
    data = msg.encode('utf-8')
    # 钉钉旧式加密尾巴必须是 corp_id(不是 token), 否则钉钉解密校验失败
    payload = os.urandom(16) + len(data).to_bytes(4, 'big') + data + corp.encode('utf-8')
    pad = 32 - len(payload) % 32
    payload += bytes([pad]) * pad
    cipher = AES.new(key, AES.MODE_CBC, key[:16])
    return base64.b64encode(cipher.encrypt(payload)).decode('utf-8')

def dt_sign(timestamp, nonce, encrypt):
    token = cfg_get('dingtalk_callback_token', '')
    return hashlib.sha1(''.join(sorted([token, timestamp, nonce, encrypt])).encode('utf-8')).hexdigest()

# ---- 预警提醒(钉钉工作通知; 去重窗口8小时; 审批超时阈值默认8小时) ----
def dt_send_reminders(force=False):
    if not dingtalk_enabled(): return
    approve_hours = int(cfg_get('dingtalk_approve_hours', '8') or 8)
    order_days = int(cfg_get('dingtalk_order_days', '3') or 3)
    c = db()
    c.execute("DELETE FROM reminder_log WHERE pushed_at <= datetime('now','localtime','-8 hours')")
    today_s = today()
    def pushed(rule, key):
        if force: return False
        k = f"{today_s}:{key}"
        if c.execute("SELECT 1 FROM reminder_log WHERE rule=? AND key=?", (rule, k)).fetchone(): return True
        c.execute("INSERT INTO reminder_log(rule,key) VALUES(?,?)", (rule, k))
        return False
    def admins_dt():
        return [r['dingtalk_userid'] for r in c.execute("SELECT dingtalk_userid FROM users WHERE role='系统管理员' AND dingtalk_userid IS NOT NULL AND dingtalk_userid!=''").fetchall()]
    def role_userid(role):
        u = find_user_by_role(role)
        return u['dingtalk_userid'] if u and u['dingtalk_userid'] else ''
    # ① 审批超时未处理(默认>8小时) → 对应审批角色; >48小时升级管理员
    rows = c.execute("""SELECT ai.* FROM approval_instances ai
        WHERE ai.status='pending' AND ai.created_at <= datetime('now','localtime',?)
        ORDER BY ai.created_at""", (f'-{approve_hours} hours',)).fetchall()
    by_role = {}
    for r in rows: by_role.setdefault(r['role'], []).append(r)
    for role, lst in by_role.items():
        uid = role_userid(role)
        if not uid: continue
        if not pushed('dt_approve_timeout', role):
            dt_send_todo([uid], f"⏰ {role} 有 {len(lst)} 条审批超过 {approve_hours} 小时未处理",
                         '\n'.join(f"- {x['biz_type']}#{x['biz_id']} (提交于 {x['created_at']})" for x in lst[:10]))
        for x in lst:
            try: age_h = int((datetime.datetime.now() - datetime.datetime.strptime(x['created_at'][:19], '%Y-%m-%d %H:%M:%S')).total_seconds() // 3600)
            except Exception: age_h = approve_hours
            if age_h >= 48 and not pushed('dt_approve_escalate', x['id']):
                for ao in admins_dt():
                    dt_send_todo([ao], f"🚨 审批升级: {x['biz_type']}#{x['biz_id']} 已超 {age_h} 小时", f"审批角色【{role}】, 请介入处理")
    # ② 待采未下单(已通过N天未转订单) → 申请人 (按人合并消息, 减少API调用)
    rows2 = c.execute("""SELECT pr.* FROM purchase_requests pr
        WHERE pr.status='已通过' AND pr.created_at <= datetime('now','localtime',?)
        AND NOT EXISTS (SELECT 1 FROM purchase_orders po WHERE po.req_id=pr.id)""", (f'-{order_days} days',)).fetchall()
    by_req = {}
    for r in rows2:
        u = find_user_by_name(r['requester'])
        if not u or not u['dingtalk_userid']: continue
        by_req.setdefault(u['dingtalk_userid'], []).append(r)
    for uid, lst in by_req.items():
        if not pushed('dt_req_no_order', uid):
            lines = [f"- {r['req_no']} ¥{r['total_estimated'] or 0:,.0f} · {r['purpose'] or ''}" for r in lst[:10]]
            dt_send_todo([uid], f"📝 {len(lst)} 条采购申请已通过 {order_days} 天仍未转订单", '\n'.join(lines), push_type='alert')
    # ③ 逾期未回货(送货单未签收) → 订单负责人 (按人合并)
    rows3 = c.execute("""SELECT d.*, po.owner FROM deliveries d
        LEFT JOIN purchase_orders po ON d.order_id=po.id
        WHERE d.sign_status='待签收' AND d.delivery_date < date('now')""").fetchall()
    by_owner = {}
    for r in rows3:
        u = find_user_by_name(r['owner'])
        if not u or not u['dingtalk_userid']: continue
        by_owner.setdefault(u['dingtalk_userid'], []).append(r)
    for uid, lst in by_owner.items():
        if not pushed('dt_delivery_overdue', uid):
            lines = [f"- {r['delivery_no']} {r['item_name']} x{r['quantity']}{r['unit']} (计划 {r['delivery_date']})" for r in lst[:10]]
            dt_send_todo([uid], f"🚚 {len(lst)} 张送货单逾期未签收", '\n'.join(lines), push_type='alert')
    # ④ 订单超目标日 → 负责人 (按人合并)
    rows4 = c.execute("""SELECT * FROM purchase_orders WHERE target_date < date('now') AND status NOT IN ('已完成','已关闭','已挂账','已入库')""").fetchall()
    by_owner4 = {}
    for r in rows4:
        u = find_user_by_name(r['owner'])
        if not u or not u['dingtalk_userid']: continue
        by_owner4.setdefault(u['dingtalk_userid'], []).append(r)
    for uid, lst in by_owner4.items():
        if not pushed('dt_order_overdue', uid):
            lines = [f"- {r['order_no']} 目标日 {r['target_date']} 状态【{r['status']}】 ¥{r['total_amount'] or 0:,.0f}" for r in lst[:10]]
            dt_send_todo([uid], f"📋 {len(lst)} 条采购订单超目标日未完成", '\n'.join(lines), push_type='alert')
    # ⑤ 库存预警(≤安全库存) → 系统管理员 (合并成一条)
    rows5 = c.execute("""SELECT i.* FROM inventory i WHERE i.quantity<=i.safe_stock AND i.safe_stock>0 ORDER BY i.item_name""").fetchall()
    if rows5 and not pushed('dt_inventory_alert', 'all'):
        lines = [f"- {r['item_name']} 当前 {r['quantity']:g}{r['unit']}, 安全库存 {r['safe_stock']:g}{r['unit']} (仓库: {r['warehouse']})" for r in rows5[:15]]
        for ao in admins_dt():
            dt_send_todo([ao], f"⚠️ {len(rows5)} 项库存低于安全库存", '\n'.join(lines), push_type='alert')
    # ⑥ 每周一 09:00 周报 → 管理员
    if datetime.date.today().weekday() == 0 and datetime.datetime.now().strftime('%H') == '09' and not pushed('dt_weekly_report', 'wk'):
        stats = {
            '采购总额': c.execute("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE created_at >= datetime('now','localtime','-7 days')").fetchone()[0],
            '待审批': c.execute("SELECT COUNT(*) FROM approval_instances WHERE status='pending'").fetchone()[0],
            '进行中订单': c.execute("SELECT COUNT(*) FROM purchase_orders WHERE status NOT IN ('已完成','已关闭','已挂账')").fetchone()[0],
            '库存预警': c.execute("SELECT COUNT(*) FROM inventory WHERE quantity<=safe_stock AND safe_stock>0").fetchone()[0],
            '超时未签收': c.execute("SELECT COUNT(*) FROM deliveries WHERE sign_status='待签收' AND delivery_date < date('now')").fetchone()[0],
        }
        text = '📊 本周采购周报\n' + '\n'.join(f'- {k}: {v:,.0f}' for k, v in stats.items())
        for ao in admins_dt():
            dt_send_todo([ao], '📊 本周采购周报', text, push_type='alert')
    c.commit(); c.close()

@app.route('/api/dingtalk/callback', methods=['GET', 'POST'])
def api_dingtalk_callback():
    """钉钉事件订阅回调: 新式明文JSON / 旧式AES加密; 审批结果同步回系统"""
    if request.method == 'GET':
        enc_param = request.args.get('encrypt', '')
        if not enc_param:
            # 新式事件订阅注册校验: 钉钉 GET 回调地址期望纯文本 success
            return 'success'
        ts = request.args.get('timestamp', ''); nonce = request.args.get('nonce', '')
        try:
            body = dt_decrypt_old(enc_param)
            inner = body.get('Random', '') if isinstance(body, dict) else ''
            enc = dt_encrypt_old(inner or 'success')
            return jsonify({'msg_signature': dt_sign(ts, nonce, enc), 'timeStamp': ts, 'nonce': nonce, 'encrypt': enc})
        except Exception as e:
            log('系统', '钉钉回调验证失败', str(e))
            return jsonify({'msg_signature': '', 'timeStamp': ts, 'nonce': nonce, 'encrypt': ''})
    try:
        body = request.get_json(force=True, silent=True) or {}
    except Exception:
        body = {}
    try:
        if body.get('encrypt'):
            body = dt_decrypt_old(body['encrypt'])
    except Exception as e:
        log('系统', '钉钉回调解密失败', str(e))
        return 'success'  # 钉钉要求纯文本success; 返回它避免重试轰炸
    if not isinstance(body, dict):
        return 'success'
    etype = body.get('eventType', '') or body.get('EventType', '')
    data = body.get('data', body) or {}
    if str(etype) in ('check_url', 'check_url_encrypt') or 'check_url' in str(etype):
        # 旧式加密回调校验: 钉钉 POST 加密 check_url, 必须返回加密的 success JSON(尾巴=corp_id)
        ts = request.args.get('timestamp', ''); nonce = request.args.get('nonce', '')
        try:
            enc = dt_encrypt_old('success')
            return jsonify({'msg_signature': dt_sign(ts, nonce, enc), 'timeStamp': ts, 'nonce': nonce, 'encrypt': enc})
        except Exception:
            return 'success'
    if 'bpms' in str(etype):
        iid = data.get('processInstanceId', '')
        result = data.get('result', '')
        # V6.0: 审批意见/驳回原因 → 回写系统留痕
        comment = str(data.get('remark', '') or '').strip()
        if not comment:
            records = data.get('approvalRecords') or []
            for rec in records:
                if rec.get('status') == 'REFUSE' and rec.get('remark'):
                    comment = rec['remark']
                    break
        if iid and result in ('agree', 'refuse'):
            # V11.126: 审批结束时拉取一次实例详情, 保存表单值(询价定标审批要读领导选的供应商)
            try:
                _inst = dt_query_instance(iid)
                if _inst:
                    _fv = _inst.get('form_component_values') or _inst.get('formComponentValues')
                    if _fv:
                        _c2 = db()
                        _c2.execute("UPDATE dingtalk_instances SET form_values=?, updated_at=? WHERE instance_code=?",
                                    (json.dumps(_fv, ensure_ascii=False), now(), iid))
                        _c2.commit(); _c2.close()
            except Exception:
                pass
            dt_sync_result(iid, result, comment)
    return 'success'  # 钉钉事件回调统一要求纯文本success

@app.route('/api/dingtalk/config', methods=['GET', 'POST'])
@login_required
def api_dingtalk_config():
    if request.method == 'POST':
        if not can_manage_config(): return jsonify({'error': '仅系统管理员可配置'}), 403
        d = request.json or {}
        for k in ('dingtalk_app_key','dingtalk_app_secret','dingtalk_agent_id','dingtalk_corp_id',
                  'dingtalk_callback_token','dingtalk_aes_key','dingtalk_approve_hours','dingtalk_order_days'):
            if k in d: cfg_set(k, d[k])
        if 'dingtalk_enabled' in d: cfg_set('dingtalk_enabled', '1' if d['dingtalk_enabled'] else '0')
        codes = d.get('dingtalk_approval_codes') or {}
        if isinstance(codes, dict) and codes:
            cfg_set('dingtalk_approval_codes', json.dumps(codes, ensure_ascii=False))
        changed = [k for k in ('dingtalk_app_key','dingtalk_app_secret','dingtalk_agent_id','dingtalk_corp_id',
                  'dingtalk_callback_token','dingtalk_aes_key','dingtalk_approve_hours','dingtalk_order_days') if k in d]
        if 'dingtalk_enabled' in d: changed.append('dingtalk_enabled')
        if isinstance(codes, dict) and codes: changed.append('dingtalk_approval_codes')
        log(session.get('user_name',''), '修改钉钉配置', '变更项: %s' % (','.join(changed) or '无'))
        return jsonify({'success': True})
    c = db()
    rows = c.execute("SELECT key,value FROM sys_config WHERE key LIKE 'dingtalk_%'").fetchall()
    c.close()
    cfg = {r['key']: r['value'] for r in rows}
    cfg['dingtalk_approval_codes'] = dt_approval_codes()
    cfg['public_url'] = dt_public_url()
    return jsonify(cfg)

@app.route('/api/dingtalk/test', methods=['POST'])
@login_required
def api_dingtalk_test():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    try:
        tk = dt_token()
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})
    target = (request.json or {}).get('userid', '')
    if target:
        ok = dt_send_todo([target], '✅ 采购系统钉钉对接测试成功', '这是一条测试工作通知', push_type='test')
        return jsonify({'success': ok, 'token': tk[:12] + '...', 'msg': '测试通知已发送' if ok else '发送失败, 请检查AgentId/权限'})
    return jsonify({'success': True, 'token': tk[:12] + '...'})

@app.route('/api/dingtalk/register-callback', methods=['POST'])
@login_required
def api_dingtalk_register_callback():
    """一键注册钉钉事件订阅(旧式call_back接口, 兼容稳定; 需应用开通事件订阅能力)"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    url = dt_public_url()
    if not url: return jsonify({'success': False, 'error': '未获取到公网地址(检查隧道与 data/public_url.txt)'})
    cb = url.rstrip('/') + '/api/dingtalk/callback'
    # 旧式回调: aes_key 必须恰好43位(字母数字), token 随机串
    if not cfg_get('dingtalk_aes_key'):
        cfg_set('dingtalk_aes_key', ''.join(secrets.choice('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789') for _ in range(43)))
    if not cfg_get('dingtalk_callback_token'):
        cfg_set('dingtalk_callback_token', secrets.token_hex(8))
    payload = {'call_back_tag': ['bpms_instance_change', 'bpms_task_change'],
               'token': cfg_get('dingtalk_callback_token'), 'aes_key': cfg_get('dingtalk_aes_key'), 'url': cb}
    body = json.dumps(payload).encode('utf-8')
    try:
        req = urllib.request.Request(DT_API + '/call_back/register_call_back?access_token=' + dt_token(), data=body,
            headers={'Content-Type': 'application/json; charset=utf-8'})
        with urllib.request.urlopen(req, timeout=15) as r:
            resp = json.loads(r.read().decode('utf-8'))
        if resp.get('errcode') == 0:
            return jsonify({'success': True, 'callbackId': resp.get('callbackId', ''), 'url': cb})
        # 71009=验证请求未通过(公网隧道/解密问题); 其余错误原样返回
        return jsonify({'success': False, 'error': resp.get('errmsg', '注册失败'), 'url': cb})
    except Exception as e:
        err = str(e)
        rd = getattr(e, 'read', None)
        if rd:
            try: err = rd().decode('utf-8', 'ignore')[:300]
            except Exception: pass
        return jsonify({'success': False, 'error': err, 'url': cb})

@app.route('/api/dingtalk/bind', methods=['POST'])
@login_required
def api_dingtalk_bind():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    d = request.json or {}
    c = db()
    c.execute("UPDATE users SET dingtalk_userid=? WHERE id=?", (str(d.get('userid', '')).strip(), int(d.get('user_id', 0))))
    c.commit(); c.close()
    log(session.get('user_name',''), '绑定钉钉', '用户#%s 绑定 userid=%s' % (d.get('user_id'), d.get('userid')))
    return jsonify({'success': True})

@app.route('/api/config-users')
@login_required
def api_config_users():
    """配置管理授权名单: 系统管理员可查看/添加/移除被授权用户(用自己账号登录即可配置系统/钉钉/飞书)"""
    if not can_manage_config():
        return jsonify({'error': '无权限'}), 403
    extra = [x.strip() for x in cfg_get('config_users', '').split(',') if x.strip()]
    conn = db()
    us = conn.execute("SELECT username, name, role, is_active FROM users ORDER BY id").fetchall()
    conn.close()
    return jsonify({'authorized': extra, 'users': [dict_row(u) for u in us],
                    'can_edit': session.get('user_role') == '系统管理员'})

@app.route('/api/config-users', methods=['POST'])
@login_required
def api_config_users_save():
    """保存授权名单: 仅系统管理员可改(防止被授权人自己扩权)"""
    if session.get('user_role') != '系统管理员':
        return jsonify({'error': '仅系统管理员可调整授权名单'}), 403
    d = request.json or {}
    names = d.get('usernames') or []
    names = [str(x).strip() for x in names if str(x).strip()]
    conn = db()
    # 校验用户名都存在
    for n in names:
        if not conn.execute("SELECT 1 FROM users WHERE username=?", (n,)).fetchone():
            conn.close(); return jsonify({'error': '账号不存在: %s' % n}), 400
    conn.execute("INSERT INTO sys_config(key,value) VALUES('config_users',?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                 (','.join(names),))
    conn.commit(); conn.close()
    log(session['user_name'], '调整配置授权', '授权名单: %s' % (','.join(names) or '空'))
    return jsonify({'success': True, 'authorized': names})

@app.route('/api/dingtalk/unbind', methods=['POST'])
@login_required
def api_dingtalk_unbind():
    """55.docx需求5: 钉钉账号解除绑定"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    d = request.json or {}
    c = db()
    c.execute("UPDATE users SET dingtalk_userid='' WHERE id=?", (int(d.get('user_id', 0)),))
    c.commit(); c.close()
    log(session['user_name'], '解绑钉钉', '用户#%s' % d.get('user_id'))
    return jsonify({'success': True})

@app.route('/api/dingtalk/lookup', methods=['POST'])
@login_required
def api_dingtalk_lookup():
    """按手机号查钉钉userid(需通讯录权限: 手机号获取userId)"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    mobile = str((request.json or {}).get('mobile', '')).strip()
    if not mobile: return jsonify({'success': False, 'error': '请输入手机号'})
    userid, name = dt_userid_by_mobile(mobile)
    if userid: return jsonify({'success': True, 'userid': userid, 'name': name})
    return jsonify({'success': False, 'error': '未找到该手机号对应的钉钉用户(需钉钉通讯录权限)'})

# ---- V7.0: 手动加急提醒(发起人/管理员可点, 限频, 留痕) ----
@app.route('/api/dingtalk/urgent-remind', methods=['POST'])
@login_required
def api_urgent_remind():
    """手动向当前待审批负责人钉钉加急提醒; 权限: 单据发起人/系统管理员; 10分钟内限一次"""
    d = request.json or {}
    biz_type = str(d.get('biz_type', '')); biz_id = int(d.get('biz_id', 0) or 0)
    if not biz_type or not biz_id:
        return jsonify({'error': '参数缺失'}), 400
    # 权限: 系统管理员 或 单据发起人
    me = db().execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
    if not (me and me['role'] == '系统管理员'):
        conn = db()
        try:
            r = conn.execute(f"SELECT requester, requester_id FROM {biz_table(biz_type)} WHERE id=?", (biz_id,)).fetchone()
        except Exception:
            r = None
        conn.close()
        is_owner = r and (r['requester'] == session['user_name'] or r['requester_id'] == session['user_id'])
        if not is_owner:
            return jsonify({'error': '仅单据发起人或系统管理员可加急提醒'}), 403
    ok, msg = dt_urgent_remind(biz_type, biz_id, session['user_name'])
    if ok:
        log(session['user_name'], '钉钉加急提醒', f"{biz_type}#{biz_id} {msg}")
        return jsonify({'success': True, 'message': msg})
    return jsonify({'error': msg}), 400

# ---- V7.0: 超期未审批单据清单(首页预警板块, 支持一键加急) ----
@app.route('/api/approvals/overdue')
@login_required
def api_overdue_approvals():
    """超期未审批: 待审批单据停留超过可配置小时数(默认24h), 返回当前节点+审批人+停留时长"""
    hours = float(cfg_get('approve_hours', '24') or 24)
    conn = db()
    rows = conn.execute("""
        SELECT ai.*,
            %s as doc_no
        FROM approval_instances ai
        WHERE ai.status='pending' AND ai.created_at <= datetime('now','localtime', ?)
        ORDER BY ai.created_at ASC""" % _ap_case('no'), (f'-{int(hours)} hours',)).fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        # 停留时长
        try:
            from datetime import datetime as _dt
            created = _dt.strptime(str(d['created_at'])[:19], '%Y-%m-%d %H:%M:%S')
            d['wait_hours'] = round((_dt.now() - created).total_seconds() / 3600, 1)
        except Exception:
            d['wait_hours'] = None
        out.append(d)
    conn.close()
    return jsonify(out)

# ---- V7.0: 钉钉推送记录查询(留痕) ----
@app.route('/api/dingtalk/push-logs')
@login_required
def api_dingtalk_push_logs():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    conn = db()
    rows = conn.execute("SELECT * FROM dingtalk_push_log ORDER BY id DESC LIMIT 200").fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/dingtalk/status', methods=['GET'])
@login_required
def api_dingtalk_status():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    c = db()
    bound = c.execute("SELECT COUNT(*) FROM users WHERE dingtalk_userid IS NOT NULL AND dingtalk_userid!=''").fetchone()[0]
    insts = c.execute("SELECT COUNT(*) FROM dingtalk_instances").fetchone()[0]
    pend = c.execute("SELECT COUNT(*) FROM dingtalk_instances WHERE status='pending'").fetchone()[0]
    c.close()
    return jsonify({
        'enabled': dingtalk_enabled(), 'config_ok': bool(cfg_get('dingtalk_app_key') and cfg_get('dingtalk_app_secret')),
        'token_ok': bool(_DT_TOKEN['t']), 'bound_users': bound, 'instances': insts, 'pending_sync': pend,
        'codes': dt_approval_codes(), 'agent_ok': dt_agent_id() > 0, 'public_url': dt_public_url(),
    })

@app.route('/api/dingtalk/remind-now', methods=['POST'])
@login_required
def api_dingtalk_remind_now():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    dt_send_reminders(force=True)
    return jsonify({'success': True})

@app.route('/api/dingtalk/attachment-check', methods=['GET'])
@login_required
def api_dingtalk_attachment_check():
    """V9.1: 钉钉附件链路自检 — 逐段诊断: token/空间/上传权限/文件写入
    用于权限开通后的快速验证"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    out = {'steps': []}
    # 1. token
    try:
        tok = dt_token()
        out['steps'].append({'name': 'access_token', 'ok': bool(tok), 'detail': f"len={len(tok)}" if tok else '获取失败'})
    except Exception as e:
        out['steps'].append({'name': 'access_token', 'ok': False, 'detail': str(e)[:100]})
    # 2. 存储空间ID
    sid = dt_storage_space_id()
    out['steps'].append({'name': '存储空间', 'ok': bool(sid), 'detail': f"spaceId={sid}" if sid else '不可用(需开通Storage权限)'})
    # 3. 上传权限(获取上传信息) — V11.8: 带 unionId
    if sid:
        import hashlib
        _uid = dt_union_id()
        test_data = b'dingtalk attachment check'
        md5 = hashlib.md5(test_data).hexdigest()
        c, r = dt_new_post(f'/v1.0/storage/spaces/{sid}/files/uploadInfos/query?unionId={_uid}', {
            'multipart': False, 'protocol': 'HEADER_SIGNATURE',
            'option': {'preCheckParam': {'name': '_check.txt', 'size': len(test_data), 'md5': md5}}})
        if c == 0:
            out['steps'].append({'name': '上传权限', 'ok': True, 'detail': 'Storage.UploadInfo.Read 已开通 ✅'})
            # 4. 实际上传+提交(完整闭环)
            try:
                d = dt_upload_file_to_dingtalk_simple(test_data, '_check.txt')
                out['steps'].append({'name': '文件上传', 'ok': bool(d), 'detail': json.dumps(d, ensure_ascii=False)[:200] if d else '失败'})
            except Exception as e:
                out['steps'].append({'name': '文件上传', 'ok': False, 'detail': str(e)[:120]})
        else:
            msg = (r or {}).get('message', '') if isinstance(r, dict) else ''
            out['steps'].append({'name': '上传权限', 'ok': False, 'detail': msg[:180] or json.dumps(r, ensure_ascii=False)[:180]})
    out['ok'] = all(s['ok'] for s in out['steps'])
    return jsonify(out)

def dt_upload_file_to_dingtalk_simple(data, filename):
    """简化上传: bytes → 钉盘, 返回 {spaceId,fileName,fileSize,fileType,fileId} — V11.8: 带unionId"""
    import hashlib, mimetypes
    sid = dt_storage_space_id()
    if not sid: return None
    _uid = dt_union_id()
    if not _uid: return None
    fname = os.path.basename(filename)
    ext = os.path.splitext(fname)[1].lower().lstrip('.')
    ftype = 'image' if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp') else 'file'
    md5 = hashlib.md5(data).hexdigest()
    c, r = dt_new_post(f'/v1.0/storage/spaces/{sid}/files/uploadInfos/query?unionId={_uid}', {
        'multipart': False, 'protocol': 'HEADER_SIGNATURE',
        'option': {'preCheckParam': {'name': fname, 'size': len(data), 'md5': md5}}})
    if c != 0:
        # V10.4 自动重试: 空间授权过期 → 清缓存重取 → 重试一次
        _msg = json.dumps(r, ensure_ascii=False)[:300] if r else ''
        if 'privilege' in _msg or 'permissionDenied' in _msg:
            cfg_set('dingtalk_storage_space_id', '')
            cfg_set('dingtalk_union_id', '')
            sid2 = dt_storage_space_id()
            _uid2 = dt_union_id()
            if sid2 and _uid2:
                c, r = dt_new_post(f'/v1.0/storage/spaces/{sid2}/files/uploadInfos/query?unionId={_uid2}', {
                    'multipart': False, 'protocol': 'HEADER_SIGNATURE',
                    'option': {'preCheckParam': {'name': fname, 'size': len(data), 'md5': md5}}})
        if c != 0: return None
    upload_key = r.get('uploadKey', '')
    hsi = r.get('headerSignatureInfo') or {}
    urls = hsi.get('resourceUrls') or []
    headers = hsi.get('headers') or {}
    if not upload_key or not urls: return None
    try:
        import http.client
        from urllib.parse import urlparse
        _u = urlparse(urls[0])
        _conn = http.client.HTTPSConnection(_u.hostname, timeout=60)
        _conn.putrequest('PUT', _u.path + ('?' + _u.query if _u.query else ''))
        for _k, _v in (headers or {}).items():
            _conn.putheader(_k, _v)
        _conn.putheader('Content-Length', str(len(data)))
        _conn.endheaders()
        _conn.send(data)
        _resp = _conn.getresponse()
        _body = _resp.read()
        _conn.close()
        if _resp.status != 200:
            return None
    except Exception:
        return None
    c2, r2 = dt_new_post(f'/v1.0/storage/spaces/{sid}/files/commit?unionId={_uid}', {
        'name': fname, 'parentId': '0', 'uploadKey': upload_key})
    if c2 != 0: return None
    dentry = r2.get('dentry') or {}
    fid = dentry.get('id') or dentry.get('uuid') or ''
    if not fid: return None
    return {'spaceId': sid, 'fileName': fname, 'fileSize': len(data), 'fileType': ftype, 'fileId': fid}

@app.route('/api/dingtalk/sync-instances', methods=['POST'])
@login_required
def api_dingtalk_sync_instances():
    """把系统内所有待审批但未在钉钉发起的单据, 补发到钉钉"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    c = db()
    rows = c.execute("SELECT DISTINCT biz_type, biz_id FROM approval_instances WHERE status='pending'").fetchall()
    c.close()
    n = 0
    for r in rows:
        if dt_start_instance(r['biz_type'], r['biz_id']): n += 1
    return jsonify({'success': True, 'pushed': n, 'total': len(rows)})

@app.route('/api/dingtalk/sync-results', methods=['POST'])
@login_required
def api_dingtalk_sync_results():
    """手动触发: 查询钉钉审批结果并回写系统(轮询兜底, 不依赖事件回调/corpId)"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    n = dt_poll_results()
    return jsonify({'success': True, 'synced': n})

@app.route('/api/dingtalk/instances', methods=['GET'])
@login_required
def api_dingtalk_instances():
    c = db()
    rows = c.execute("SELECT * FROM dingtalk_instances ORDER BY id DESC LIMIT 30").fetchall()
    c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/dingtalk/sso', methods=['POST'])
def api_dingtalk_sso():
    """钉钉H5免登: 前端dd.config后取authCode, 换userid并匹配系统用户登录(需已在②中绑定)"""
    d = request.json or {}
    code = str(d.get('authCode', '')).strip()
    if not code: return jsonify({'error': '缺少authCode'}), 400
    code_r, resp = dt_post('/topapi/v2/user/getuserinfo', {'code': code})
    if code_r != 0: return jsonify({'error': resp.get('msg', '免登失败')}), 400
    userid = (resp.get('result', {}) or {}).get('userid', '')
    if not userid: return jsonify({'error': '未获取到钉钉userid'}), 400
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE dingtalk_userid=? AND is_active=1", (userid,)).fetchone()
    conn.close()
    if not u:
        return jsonify({'error': '该钉钉账号未绑定系统用户, 请联系管理员在「钉钉设置 → 用户钉钉绑定」中绑定'}), 403
    session['user_id'] = u['id']; session['username'] = u['username']; session['user_name'] = u['name']; session['user_role'] = u['role']
    session['dept_id'] = u['dept_id']; session['dept_name'] = ''
    log(u['name'], '钉钉免登登录', f'userid={userid}')
    return jsonify({'success': True, 'name': u['name'], 'role': u['role']})

@app.route('/api/dingtalk/jsapi-ticket')
def api_dingtalk_jsapi_ticket():
    """钉钉JSAPI签名参数(dd.config), 钉钉内打开页面时使用"""
    if not (_DT_TICKET['t'] and time.time() < _DT_TICKET['exp'] - 120):
        try:
            url = f"{DT_API}/get_jsapi_ticket?access_token={dt_token()}"
            with urllib.request.urlopen(url, timeout=12) as r:
                d = json.loads(r.read().decode('utf-8'))
            if d.get('errcode') != 0: return jsonify({'error': d.get('errmsg', '获取ticket失败')}), 400
            _DT_TICKET['t'] = d.get('ticket'); _DT_TICKET['exp'] = time.time() + int(d.get('expires_in', 7200))
        except Exception as e:
            return jsonify({'error': str(e)}), 400
    nonce = secrets.token_hex(8); ts = str(int(time.time()))
    jsurl = request.args.get('url', request.url_root.rstrip('/') + '/')
    sig = hashlib.sha1(f"jsapi_ticket={_DT_TICKET['t']}&noncestr={nonce}&timestamp={ts}&url={jsurl}".encode('utf-8')).hexdigest()
    return jsonify({'agentId': cfg_get('dingtalk_agent_id', ''), 'corpId': cfg_get('dingtalk_corp_id', ''),
                    'timeStamp': ts, 'nonceStr': nonce, 'signature': sig, 'ticket_ok': True})


# V11.192: 钉钉通知/加急"去处理"免登直达页 — 申请人/审批人点开即自动免登进系统对应单据
# 根因修复: 原action_card跳 n.dingtalk.com OA H5 需钉钉审批登录态+可信域名, 通知对象(申请人/非审批人)
# 打开白屏/跳登录 → 改为跳系统单据直达(免登鉴权由 /api/dingtalk/sso 完成, 系统域名无需在钉钉后台配置可信域名)
@app.route('/sso/goto')
def api_sso_goto():
    bt = str(request.args.get('biz', '')).strip()
    bid = str(request.args.get('id', '')).strip()
    act = str(request.args.get('act', 'detail')).strip()
    if act not in ('detail', 'approve'):
        act = 'detail'
    # 单据号/名称预取(免登前展示, 用户知道点的是什么)
    _doc_no, _doc_tip = '', ''
    _tip = '单据详情' if act != 'approve' else '单据审批'
    try:
        if bt and bid.isdigit():
            _t = biz_table(bt)
            if _t:
                _c = db()
                _row = _c.execute(f"SELECT * FROM {_t} WHERE id=?", (int(bid),)).fetchone()
                _c.close()
                if _row is not None:
                    for _k in ('req_no', 'order_no', 'contract_no', 'receive_no', 'payment_no', 'credit_no', 'inq_no'):
                        if _k in _row.keys() and _row[_k]:
                            _doc_no = str(_row[_k]); break
                    if not _doc_no:
                        _doc_no = f'{bt}#{bid}'
                    _tip = '单据审批' if act == 'approve' else '单据详情'
    except Exception:
        pass
    _DT_BIZ_CN = DT_BIZ.get(bt, '单据')
    return f'''<!doctype html><html lang="zh"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1,user-scalable=no">
<title>正成能源采购 · {_DT_BIZ_CN}</title>
<style>body{{margin:0;font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:linear-gradient(160deg,#0d2c54,#1a4a7a);min-height:100vh;display:flex;align-items:center;justify-content:center}}
.card{{background:#fff;border-radius:14px;padding:28px 24px;width:86%;max-width:340px;box-shadow:0 10px 30px rgba(0,0,0,.25);text-align:center}}
.logo{{font-size:28px}}h2{{font-size:16px;color:#16325c;margin:8px 0 4px}}p{{font-size:12px;color:#666;margin:4px 0;line-height:1.7;word-break:break-all}}
.btn{{display:block;width:100%;padding:11px 0;border:none;border-radius:8px;font-size:14px;font-weight:600;margin-top:12px;cursor:pointer}}
.btn-p{{background:#1a6bff;color:#fff}}.btn-o{{background:#f0f4fa;color:#16325c}}
.tip{{font-size:11px;color:#999;margin-top:10px}}</style></head><body>
<div class="card"><div class="logo">📋</div><h2>{_DT_BIZ_CN} · {_tip}</h2>
<p>{_doc_no or '单据加载中…'}</p>
<div id="st" style="font-size:12px;color:#1a6bff;margin-top:8px">正在验证钉钉身份，请稍候…</div>
<a class="btn btn-p" id="bOpen" style="display:none;text-decoration:none" href="/?dtopen={bt}:{bid}:{act}">✅ 已登录，打开单据</a>
<a class="btn btn-o" id="bCopy" style="display:none" href="javascript:void(0)">📋 复制链接(手机浏览器打开)</a>
<p class="tip" id="tipTxt" style="display:none">非钉钉环境/免登失败：请复制下方链接，用<b>手机浏览器或电脑浏览器</b>打开后按系统账号登录即可查看单据。</p>
<script src="https://g.alicdn.com/dingding/dingtalk-jsapi/2.10.4/dingtalk.open.js"></script>
<script>
(function(){{
  const UA=navigator.userAgent||'', IS_DT=UA.includes('DingTalk');
  const link=location.origin+'/?dtopen={bt}:{bid}:{act}';
  async function go(){{
    try{{
      const cfg=await fetch('/api/dingtalk/jsapi-ticket?url='+encodeURIComponent(location.href.split('#')[0]),{{credentials:'include'}}).then(r=>r.json());
      if(!cfg.agentId||!cfg.corpId||cfg.error)throw new Error(cfg.error||'钉钉未配置');
      dd.config({{agentId:cfg.agentId,corpId:cfg.corpId,timeStamp:cfg.timeStamp,nonceStr:cfg.nonceStr,signature:cfg.signature,jsApiList:['runtime.permission.requestAuthCode']}});
      dd.ready(function(){{
        dd.runtime.permission.requestAuthCode({{corpId:cfg.corpId,onSuccess:async function(r){{
          const res=await fetch('/api/dingtalk/sso',{{method:'POST',headers:{{'Content-Type':'application/json'}},credentials:'include',body:JSON.stringify({{authCode:r.code}})}}).then(x=>x.json());
          if(res.success){{location.href=link}}
          else{{document.getElementById('st').textContent='免登未绑定系统账号：'+res.error;showCopy()}}
        }},onFail:function(e){{document.getElementById('st').textContent='钉钉授权未完成';showCopy()}}}});
      }});
      dd.error(function(e){{document.getElementById('st').textContent='签名校验失败，请用浏览器打开';showCopy()}});
    }}catch(e){{document.getElementById('st').textContent='免登不可用：'+e.message;showCopy()}}
  }}
  function showCopy(){{
    const b=document.getElementById('bCopy');b.style.display='block';b.onclick=function(){{
      if(navigator.clipboard&&navigator.clipboard.writeText){{navigator.clipboard.writeText(link).then(()=>{{document.getElementById('tipTxt').textContent='✅ 链接已复制：'+link}})}}else{{
        const ta=document.createElement('textarea');ta.value=link;document.body.appendChild(ta);ta.select();document.execCommand('copy');document.body.removeChild(ta);
        document.getElementById('tipTxt').textContent='✅ 链接已复制：'+link}}
    }};
    const tip=document.getElementById('tipTxt');tip.style.display='block';tip.textContent='非钉钉环境/免登失败：请复制链接，用浏览器打开后按系统账号登录查看。\\n'+link;
    document.getElementById('bOpen').style.display='block';
  }}
  if(IS_DT){{go()}}else{{document.getElementById('st').textContent='请在钉钉客户端打开，或复制链接到浏览器查看';showCopy()}}
}})();
</script></body></html>'''

@app.route('/api/approval-flow')
@login_required
def api_approval_flow():
    conn = db(); rows = conn.execute("SELECT * FROM approval_flow_config ORDER BY biz_type,level_no").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/approval-flow', methods=['POST'])
@login_required
def api_save_approval_flow():
    """V5.0: 审批流设置 — 管理员按 单据类型×金额区间 配置审批环节(角色)
    保存后新单据按新配置生成审批链; 审批人=角色对应 users 表有效用户, 钉钉/飞书自动同步"""
    if not can_manage_config():
        return jsonify({'error': '仅系统管理员可设置审批流'}), 403
    d = request.json or {}
    biz_type = str(d.get('biz_type', '')).strip()
    levels = d.get('levels') or []
    if biz_type not in ('purchase_request', 'purchase_order', 'contract', 'credit', 'payment', 'receiving', 'requisition', 'return_request', 'inquiry_approval', 'collect_accept', 'repair_plan'):
        return jsonify({'error': '未知单据类型'}), 400
    valid_roles = ('部门负责人', '库管员', '采购员', '财务', '分管领导', '总经理')
    parsed = []
    for lv in levels:
        role = str(lv.get('role', '')).strip()
        approver = str(lv.get('approver', '') or '').strip()  # 具体审批人用户名, 留空=按角色
        try:
            mn = float(lv.get('min_amount', 0) or 0)
            mx = float(lv.get('max_amount', 9999999) or 9999999)
        except Exception:
            return jsonify({'error': '金额区间必须为数字'}), 400
        if role not in valid_roles:
            return jsonify({'error': f'审批角色必须是: {"、".join(valid_roles)}'}), 400
        if mn < 0 or mx < mn:
            return jsonify({'error': '金额区间无效(最小不能小于0, 且最小≤最大)'}), 400
        if approver:
            # 校验具体审批人存在且有效
            u = conn_check_user(approver)
            if not u:
                return jsonify({'error': f'审批人「{approver}」不存在或未启用'}), 400
        parsed.append({'role': role, 'approver': approver, 'min_amount': mn, 'max_amount': mx})
    if not parsed:
        return jsonify({'error': '至少配置一级审批'}), 400
    conn = db()
    # 区间连续性校验: 排序后 每级 min 必须 ≤ 下一级 max(允许重叠取并集, 简化为: 第1级从0开始)
    sorted_lv = sorted(parsed, key=lambda x: x['min_amount'])
    if sorted_lv[0]['min_amount'] != 0:
        return jsonify({'error': '第一级审批金额下限必须为 0'}), 400
    conn.execute("DELETE FROM approval_flow_config WHERE biz_type=?", (biz_type,))
    for i, lv in enumerate(sorted_lv, 1):
        label = f"{i}级-{lv['role']}" + (f"-{lv['approver']}" if lv['approver'] else "")
        conn.execute("INSERT INTO approval_flow_config(biz_type,level_no,role,min_amount,max_amount,label,approver) VALUES(?,?,?,?,?,?,?)",
                     (biz_type, i, lv['role'], lv['min_amount'], lv['max_amount'], label, lv['approver']))
    conn.commit(); conn.close()
    log(session['user_name'], '修改审批流', f'{biz_type} → {len(sorted_lv)}级: ' + ' > '.join(f"{l['role']}({l['approver'] or '按角色'})[{l['min_amount']:.0f}-{l['max_amount']:.0f}]" for l in sorted_lv))
    return jsonify({'success': True, 'message': f'审批流已更新：{biz_type} {len(sorted_lv)} 级'})

# ---- V5.0: 审批人配置检查(角色→有效用户→钉钉/飞书通讯录绑定状态) ----
@app.route('/api/approval-config/check')
@login_required
def api_approval_config_check():
    """逐单据类型+金额分级展示审批链, 每级角色对应审批人是否有效/是否绑定通讯录
    让管理员一眼看出哪个角色缺人/未绑定, 满足'模板选的审批人必须是钉钉通讯录成员'的配置要求"""
    conn = db()
    flows = conn.execute("SELECT * FROM approval_flow_config ORDER BY biz_type,level_no").fetchall()
    users = conn.execute("SELECT id,username,name,role,is_active,dingtalk_userid,feishu_open_id FROM users ORDER BY role,id").fetchall()
    conn.close()
    user_map = {}
    for u in users:
        d = dict_row(u)
        user_map.setdefault(d['role'], []).append(d)
    # 各级角色去重后的完整清单(用于检查缺人)
    roles_used = sorted(set(r['role'] for r in flows))
    out_flows = []
    for f in flows:
        d = dict_row(f)
        role_users = user_map.get(d['role'], [])
        active = [u for u in role_users if u['is_active'] == 1]
        bound_dt = [u for u in active if u.get('dingtalk_userid')]
        bound_fs = [u for u in active if u.get('feishu_open_id')]
        # 配置的具体审批人: 解析用户名 → 显示姓名+绑定状态
        cfg_approver = ''
        cfg_approver_ok_dt = False
        if d.get('approver'):
            cfg_approver = d['approver']
            for u in role_users:
                if (u['username'] == d['approver'] or (u['name'] or '') == d['approver']):
                    cfg_approver = u['name'] or u['username']
                    cfg_approver_ok_dt = bool(u['is_active'] == 1 and u.get('dingtalk_userid'))
                    break
        d['approver'] = cfg_approver
        d['approver_bound_dt'] = cfg_approver_ok_dt
        d['role_users'] = role_users
        d['active_users'] = [u['name'] for u in active]
        d['dingtalk_bound'] = [u['name'] for u in bound_dt]
        d['feishu_bound'] = [u['name'] for u in bound_fs]
        # 状态: 配置了具体审批人→看该人; 否则看角色下是否有绑钉钉的有效用户
        d['config_ok_dt'] = cfg_approver_ok_dt if cfg_approver else bool(bound_dt)
        d['config_ok_fs'] = bool(bound_fs)
        out_flows.append(d)
    roles_status = []
    for role in roles_used:
        ru = user_map.get(role, [])
        roles_status.append({
            'role': role,
            'total': len(ru),
            'active': sum(1 for u in ru if u['is_active'] == 1),
            'dingtalk_bound': sum(1 for u in ru if u['is_active'] == 1 and u.get('dingtalk_userid')),
            'feishu_bound': sum(1 for u in ru if u['is_active'] == 1 and u.get('feishu_open_id')),
            'users': ru,
        })
    return jsonify({'flows': out_flows, 'roles': roles_status})

@app.route('/api/approvals/pending')
@login_required
def api_approvals_pending():
    conn = db()
    role = session['user_role']
    rows = conn.execute("""
        SELECT ai.*, %s as biz_no, %s as biz_name, %s as biz_amount
        FROM approval_instances ai WHERE ai.status='pending'
        AND NOT EXISTS (SELECT 1 FROM approval_instances y WHERE y.biz_type=ai.biz_type AND y.biz_id=ai.biz_id AND y.status='pending' AND y.level_no < ai.level_no)
        AND (ai.role=? OR ai.role='部门负责人' AND ? IN ('部门负责人','系统管理员') OR ai.approver_id=?)
        ORDER BY ai.id DESC LIMIT 50
    """ % (_ap_case('no'), _ap_case('name'), _ap_case('amount')), (role, role, session.get('user_id', 0))).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/approvals/rejected')
@login_required
def api_approvals_rejected():
    """55.docx需求4: 审批未通过数据独立板块(按业务类别分组汇总)"""
    conn = db()
    rows = conn.execute("""SELECT ai.biz_type,
        %s as biz_no,
             MAX(ai.comment) last_comment, COUNT(*) cnt, MAX(ai.processed_at) processed_at
        FROM approval_instances ai WHERE ai.status='rejected'
        GROUP BY ai.biz_type, ai.biz_id ORDER BY ai.biz_type, cnt DESC""" % _ap_case('no')).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/approvals/all-pending')
@login_required
def api_all_pending():
    conn = db()
    rows = conn.execute("""
        SELECT ai.*, %s as biz_no, %s as biz_name, %s as biz_amount
        FROM approval_instances ai WHERE ai.status='pending'
        AND NOT EXISTS (SELECT 1 FROM approval_instances y WHERE y.biz_type=ai.biz_type AND y.biz_id=ai.biz_id AND y.status='pending' AND y.level_no < ai.level_no)
        ORDER BY ai.id DESC LIMIT 50
    """ % (_ap_case('no'), _ap_case('name'), _ap_case('amount'))).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/approvals/<biz_type>/<int:biz_id>/list')
@login_required
def api_approval_list(biz_type, biz_id):
    conn = db(); rows = conn.execute("SELECT * FROM approval_instances WHERE biz_type=? AND biz_id=? ORDER BY level_no", (biz_type,biz_id)).fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/approvals/<biz_type>/<int:biz_id>/reject-logs')
@login_required
def api_reject_logs(biz_type, biz_id):
    """V11.175: 驳回历史(人/时间/理由/来源), 各模块详情弹窗展示用; V11.181: 含附件+钉钉实例号"""
    return jsonify(get_reject_logs(biz_type, biz_id))


@app.route('/api/approvals/<biz_type>/<int:biz_id>/action-logs')
@login_required
def api_action_logs(biz_type, biz_id):
    """V11.184: 审批流转操作日志(同意/驳回统一, 每条含审批人/时间/意见/附件元数据)"""
    return jsonify(get_approval_action_logs(biz_type, biz_id))


@app.route('/api/dingtalk/attachment-download', methods=['POST'])
@login_required
def api_dt_attachment_download():
    """V11.184: 下载钉钉审批附件到本地(多通道) → 返回可访问URL。
    通道①: 本地已缓存(同fileId已下载过) 通道②: 旧版 /topapi/processinstance/file/download
    通道③: 新版 storage downloadInfos(签名URL) 全部失败 → OA审批单链接兜底+友好提示"""
    d = request.json or {}
    instance_id = d.get('instance_id') or ''
    file_id = d.get('file_id') or ''
    space_id = d.get('space_id') or ''
    fname = d.get('file_name') or '附件'
    if not file_id:
        return jsonify({'error': '缺少参数'}), 400
    _dir = os.path.join(BASE, 'uploads', 'dingtalk_att')
    os.makedirs(_dir, exist_ok=True)
    _safe = re.sub(r'[^\w.\-\u4e00-\u9fff]', '_', fname)[:60]
    # 通道①: 本地已缓存
    try:
        for _f in os.listdir(_dir):
            if _f.startswith(file_id + '_') or (file_id + '_') in _f:
                return jsonify({'success': True, 'url': '/uploads/dingtalk_att/' + _f, 'name': fname, 'cached': True})
    except Exception:
        pass
    _fn = '%s_%s' % (file_id, _safe)
    _path = os.path.join(_dir, _fn)
    oa_url = 'https://n.dingtalk.com/dingtalk/web/process/%s' % instance_id if instance_id else ''
    err_chain = []
    try:
        # 通道②: 旧版审批附件下载(需 qyapi_aflow_att_auth_code)
        if instance_id:
            code, resp = dt_post('/topapi/processinstance/file/download', {
                'agent_id': dt_agent_id(), 'process_instance_id': instance_id,
                'file_id': file_id, 'space_id': space_id or ''})
            resp = resp if isinstance(resp, dict) else {}
            if code == 0:
                dl_url = resp.get('download_url') or resp.get('url') or resp.get('fileUrl') or ''
                data_b = resp.get('data') if isinstance(resp, dict) else None
                if dl_url:
                    import urllib.request as _ur
                    with open(_path, 'wb') as _f:
                        _f.write(_ur.urlopen(dl_url, timeout=30).read())
                    return jsonify({'success': True, 'url': '/uploads/dingtalk_att/' + _fn, 'name': fname})
                if isinstance(data_b, str) and data_b:
                    import base64 as _b64
                    try:
                        raw = _b64.b64decode(data_b)
                        with open(_path, 'wb') as _f:
                            _f.write(raw)
                        return jsonify({'success': True, 'url': '/uploads/dingtalk_att/' + _fn, 'name': fname})
                    except Exception:
                        pass
            err_chain.append(str(resp.get('sub_msg') or resp.get('errmsg') or f'code={code}')[:100])
        # 通道③: 新版 storage downloadInfos(签名URL, 与上传同接口族)
        try:
            _uid = dt_union_id()
            _sid = space_id or dt_storage_space_id()
            if _uid and _sid:
                c2, r2 = dt_new_post(f'/v1.0/storage/spaces/{_sid}/files/downloadInfos/query?unionId={_uid}',
                                     {'fileId': file_id})
                if c2 == 0:
                    dl2 = (r2.get('downloadInfos') or [{}])[0].get('resourceUrl') or r2.get('resourceUrl') or ''
                    if dl2:
                        import urllib.request as _ur2
                        with open(_path, 'wb') as _f:
                            _f.write(_ur2.urlopen(dl2, timeout=30).read())
                        return jsonify({'success': True, 'url': '/uploads/dingtalk_att/' + _fn, 'name': fname})
                else:
                    err_chain.append(str(r2.get('message') or '')[:80])
        except Exception as e3:
            err_chain.append(str(e3)[:80])
        # 全部失败: 友好提示 + OA链接兜底(用户可前往钉钉查看原件)
        _hint = '该审批附件获取失败，请前往钉钉审批单查看'
        if oa_url:
            return jsonify({'error': _hint, 'oa_url': oa_url, 'detail': '; '.join(err_chain)[:150]}), 200
        return jsonify({'error': _hint, 'detail': '; '.join(err_chain)[:150]}), 200
    except Exception as e:
        return jsonify({'error': '该审批附件获取失败，请前往钉钉审批单查看', 'oa_url': oa_url, 'detail': str(e)[:100]}), 200

def find_doc_submitter(biz_type, biz_id):
    """V11.190: 查单据提交人(系统用户对象) — 通过/驳回通知共用。
    各表字段不同: 申请requester/订单owner/入库inspector/出库requester; 合同/挂账/付款表无提交人字段
    → 经 order_id 关联订单取 owner/requester, 再兜底 credit 关联。返回用户Row或None"""
    try:
        _b = biz_table(biz_type)
        _c = db()
        try:
            _row = _c.execute(f"SELECT * FROM {_b} WHERE id=?", (biz_id,)).fetchone()
        except Exception:
            _row = None
        _c.close()
        if _row is None:
            return None
        _submitter = None
        for _k in ('requester', 'owner', 'inspector', 'created_by', 'apply_by'):
            if _k in _row.keys() and _row[_k]:
                _submitter = str(_row[_k]); break
        if not _submitter:
            try:
                _oid = _row['order_id'] if 'order_id' in _row.keys() else None
                if _oid:
                    _po = db()
                    _po2 = _po.execute("SELECT owner, requester FROM purchase_orders WHERE id=?", (_oid,)).fetchone()
                    _po.close()
                    if _po2:
                        _submitter = str(_po2['owner'] or _po2['requester'] or '')
                if not _submitter and biz_type == 'payment' and 'credit_id' in _row.keys() and _row['credit_id']:
                    _cc = db()
                    _cn = _cc.execute("SELECT order_id FROM credit_notes WHERE id=?", (_row['credit_id'],)).fetchone()
                    _cc.close()
                    if _cn and _cn['order_id']:
                        _po = db()
                        _po2 = _po.execute("SELECT owner, requester FROM purchase_orders WHERE id=?", (_cn['order_id'],)).fetchone()
                        _po.close()
                        if _po2:
                            _submitter = str(_po2['owner'] or _po2['requester'] or '')
            except Exception:
                pass
        if not _submitter:
            return None
        _u = db()
        try:
            _usr = _u.execute("SELECT * FROM users WHERE name=? AND is_active=1 LIMIT 1", (_submitter,)).fetchone()
            if _usr is None:
                _usr = _u.execute("SELECT * FROM users WHERE username=? AND is_active=1 LIMIT 1", (_submitter,)).fetchone()
        except Exception:
            _usr = None
        _u.close()
        # V11.190: 转dict — 调用方用 .get() 访问(sqlite3.Row无get方法会抛异常, 导致审批通过通知被吞)
        if _usr is not None:
            try:
                return dict(_usr)
            except Exception:
                return _usr
        return None
    except Exception:
        return None


def notify_submitter_rejected(biz_type, biz_id, approver, comment, source='system'):
    """V11.178: 单据被驳回后 → 系统站内信(铃铛) + 钉钉工作通知提交人(含单据号/驳回人/理由/处理建议)
    V11.216: 站内信不依赖钉钉启用/绑定(原逻辑dingtalk未启用或用户未绑钉钉直接return→用户收不到任何通知)"""
    try:
        _b = biz_table(biz_type)
        _c = db()
        try:
            _row = _c.execute(f"SELECT * FROM {_b} WHERE id=?", (biz_id,)).fetchone()
        except Exception:
            _row = None
        _c.close()
        if _row is None:
            return
        # 单据号
        _doc_no = ''
        for _k in ('req_no', 'order_no', 'contract_no', 'receive_no', 'payment_no', 'plan_no'):
            if _k in _row.keys() and _row[_k]:
                _doc_no = str(_row[_k]); break
        if not _doc_no:
            _doc_no = f'{biz_type}#{biz_id}'
        # 提交人(各表字段不同: 申请requester/订单owner/入库inspector/出库requester)
        # V11.179: 合同/挂账/付款表无提交人字段 → 通过 order_id 关联订单取 owner, 再兜底credit关联
        _submitter = None
        for _k in ('requester', 'owner', 'inspector', 'created_by', 'apply_by'):
            if _k in _row.keys() and _row[_k]:
                _submitter = str(_row[_k]); break
        if not _submitter:
            try:
                if biz_type == 'contract' and _row['order_id']:
                    _po = _c2 = db()
                    _po2 = _po.execute("SELECT owner, requester FROM purchase_orders WHERE id=?", (_row['order_id'],)).fetchone()
                    _po.close()
                    if _po2:
                        _submitter = str(_po2['owner'] or _po2['requester'] or '')
                elif biz_type == 'payment' and _row['credit_id']:
                    _cc = db()
                    _cn = _cc.execute("SELECT order_id FROM credit_notes WHERE id=?", (_row['credit_id'],)).fetchone()
                    _cc.close()
                    if _cn and _cn['order_id']:
                        _po = db()
                        _po2 = _po.execute("SELECT owner, requester FROM purchase_orders WHERE id=?", (_cn['order_id'],)).fetchone()
                        _po.close()
                        if _po2:
                            _submitter = str(_po2['owner'] or _po2['requester'] or '')
                elif biz_type == 'credit' and _row['order_id']:
                    _po = db()
                    _po2 = _po.execute("SELECT owner, requester FROM purchase_orders WHERE id=?", (_row['order_id'],)).fetchone()
                    _po.close()
                    if _po2:
                        _submitter = str(_po2['owner'] or _po2['requester'] or '')
            except Exception:
                pass
        if not _submitter:
            return
        _u = db()
        try:
            _usr = _u.execute("SELECT * FROM users WHERE name=? AND is_active=1 LIMIT 1", (_submitter,)).fetchone()
            if _usr is None:
                _usr = _u.execute("SELECT * FROM users WHERE username=? AND is_active=1 LIMIT 1", (_submitter,)).fetchone()
        except Exception:
            _usr = None
        _u.close()
        if not _usr:
            return
        _approver = approver or '审批人'
        _reason = (comment or '').strip() or '（未填写理由）'
        _src_txt = '【系统审批】' if source == 'system' else '【钉钉审批】'
        title = f'❌ 您的{_doc_no} 被驳回'
        text = f"您提交的 **{_doc_no}** 经{_src_txt}被驳回，请查看原因并及时修改后重新提交。\n\n驳回人：{_approver}\n驳回理由：{_reason}"
        # V11.216: 站内信(铃铛) 不依赖钉钉 — 提交人登录必能看到
        try:
            add_notif([_usr['id']], title, f"您提交的{_doc_no}被驳回\n驳回人:{_approver}\n理由:{_reason}", biz_type, biz_id)
        except Exception:
            pass
        # 钉钉工作通知(绑了钉钉才发)
        try:
            if _usr['dingtalk_userid'] and dingtalk_enabled():
                dt_send_todo([_usr['dingtalk_userid']], title, text,
                             f"单据: {_doc_no}", biz_type, biz_id, push_type='result', operator='系统')
        except Exception:
            pass
        log('系统', '驳回通知提交人', f"{biz_type}#{biz_id} → {_usr['name']} 理由:{_reason[:60]}")
    except Exception:
        pass


@app.route('/api/approvals/<biz_type>/<int:biz_id>/approve', methods=['POST'])
@login_required
def api_approve_action(biz_type, biz_id):
    d = request.json
    # V7.0: 审批操作权限锁死 — 仅当前节点指定审批人/系统管理员可操作
    conn = db()
    cur = conn.execute("SELECT * FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='pending' ORDER BY level_no LIMIT 1", (biz_type, biz_id)).fetchone()
    if cur:
        me = conn.execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
        is_admin = me and me['role'] == '系统管理员'
        node_approver_id = cur['approver_id'] if cur['approver_id'] else None
        if not node_approver_id:
            u = conn.execute("SELECT * FROM users WHERE role=? AND is_active=1 ORDER BY id LIMIT 1", (cur['role'],)).fetchone()
            node_approver_id = u['id'] if u else None
        if not is_admin and (not me or me['id'] != node_approver_id):
            conn.close()
            return jsonify({'error': '仅当前审批节点指定审批人可操作'}), 403
    conn.close()
    if d.get('action') == 'rejected':
        # 驳回: 全部待审节点驳回 + 父单据置为已驳回 (同步飞书实例)
        # 注: 提交人通知由 finish_approvals 内统一触发(避免重复)
        # V11.185: 条目级驳回 — 前端可勾选问题条目id列表(rejected_items), 整单驳回不传=全部标记
        _rej_items = d.get('rejected_items')
        if isinstance(_rej_items, list):
            _rej_items = json.dumps(_rej_items, ensure_ascii=False) if _rej_items else None
        finish_approvals(biz_type, biz_id, 'reject', session['user_name'], session['user_id'], d.get('comment',''),
                         rejected_items=_rej_items)
        dt_sync_now(biz_type, biz_id)  # 立即同步钉钉: 终止挂起的审批实例
        return jsonify({'success':True})
    sig = d.get('signature', '')
    r = do_approve(biz_type, biz_id, session['user_name'], session['user_id'], 'approved', d.get('comment',''), signature=sig)
    if not r['success']: return jsonify(r), 400
    finish_approvals(biz_type, biz_id, 'ok', session['user_name'], session['user_id'], d.get('comment',''))
    dt_sync_now(biz_type, biz_id)  # 立即同步钉钉: 查询最新状态/终态时终止实例
    return jsonify({'success':True})


@app.route('/api/approvals/<biz_type>/<int:biz_id>/resubmit', methods=['POST'])
@login_required
def api_generic_resubmit(biz_type, biz_id):
    """V11.185: 通用"再次提交审批" — 驳回后的单据(不新建)在原单上重新进入审批流。
    适用: purchase_request/purchase_order/contract/receiving/requisition/inquiry_approval
    流程: 校验已驳回 → (可选带items修改明细) → 清驳回标记 → resubmit_count+1 → 重建审批实例 → 重新发起钉钉
    日志: 写 approval_action_logs(action=resubmit) 完整留痕"""
    try:
        d = request.json or {}
        conn = db()
        tbl = biz_table(biz_type)
        if not tbl:
            conn.close(); return jsonify({'error': '该单据类型不支持重新提交'}), 400
        row = conn.execute(f"SELECT * FROM {tbl} WHERE id=?", (biz_id,)).fetchone()
        if not row:
            conn.close(); return jsonify({'error': '单据不存在'}), 404
        # 仅"被驳回退回的草稿"可重提: 状态=草稿 且 reject_count>0(曾驳回); 纯手动草稿走原提交接口
        if row['status'] != '草稿' or not (int(row['reject_count'] or 0) > 0):
            conn.close(); return jsonify({'error': '仅被驳回退回草稿的单据可重新提交（手动草稿请在编辑后直接提交审批）'}), 400
        # 提交人本人/管理员 才可重提(取各表提交人字段)
        who = row['requester'] if 'requester' in row.keys() else (row['created_by'] if 'created_by' in row.keys() else '')
        me = conn.execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
        is_admin = me and me['role'] == '系统管理员'
        if not is_admin and who and me and me['name'] != who:
            conn.close(); return jsonify({'error': '仅提交人本人可重新提交'}), 403
        # 金额(各表字段不同)
        amount = 0
        for k in ('total_estimated', 'total_amount', 'amount', 'est_amount'):
            if k in row.keys() and row[k]:
                try: amount = float(row[k] or 0); break
                except Exception: pass
        # V11.185: 可选带修改的明细(编辑后重提) — 申请单重建request_items
        items = d.get('items')
        if isinstance(items, list) and items and biz_type == 'purchase_request':
            conn.execute("DELETE FROM request_items WHERE req_id=?", (biz_id,))
            _tot = 0
            for it in items:
                _tp = float(it.get('quantity', 1) or 1) * float(it.get('estimated_price', 0) or 0)
                _tot += _tp
                conn.execute("INSERT INTO request_items(req_id,item_name,spec,unit,quantity,estimated_price,total_price,remark,category,brand_param,arrival_date,attach) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                             (biz_id, it.get('item_name',''), it.get('spec',''), it.get('unit','个'), float(it.get('quantity',1)),
                              float(it.get('estimated_price',0)), _tp, it.get('remark',''),
                              it.get('category',''), it.get('brand_param',''), it.get('arrival_date',''),
                              it.get('attach','') or ''))
            amount = _tot
            conn.execute(f"UPDATE {tbl} SET total_estimated=?, updated_at=? WHERE id=?", (_tot, now(), biz_id))
        # 重提: 状态回待审批 + 清驳回标记 + 次数+1
        _rs_n = int(row['resubmit_count'] or 0) + 1
        conn.execute(f"UPDATE {tbl} SET status='待审批', rejected_reason='', rejected_items='', resubmit_count=?, updated_at=? WHERE id=?",
                     (_rs_n, now(), biz_id))
        # 重建审批实例 + 清旧钉钉实例
        conn.execute("DELETE FROM approval_instances WHERE biz_type=? AND biz_id=?", (biz_type, biz_id))
        conn.execute("DELETE FROM dingtalk_instances WHERE biz_type=? AND biz_id=?", (biz_type, biz_id))
        conn.commit(); conn.close()
        # 记录重提操作日志(重新开连接, 事务已提交)
        log_approval_action(biz_type, biz_id, 'resubmit', session['user_name'], session['user_id'],
                            f'第{_rs_n}次重新提交审批' + ('(含明细修改)' if items else ''), now(), [], 'system', '')
        create_approvals(biz_type, biz_id, amount, submitter=session['user_name'])
        try:
            start_instances(biz_type, biz_id)
        except Exception:
            pass
        log(session['user_name'], '重新提交审批', f'{biz_type}#{biz_id} 第{_rs_n}次重提')
        return jsonify({'success': True, 'resubmit_count': _rs_n})
    except Exception as e:
        return jsonify({'error': str(e)[:120]}), 500

# ============================================================
# V11.27 ── 预录签名(签名版): 领导手写一次, 审批通过的单据自动盖章
# ============================================================
@app.route('/api/signature/save', methods=['POST'])
@login_required
def api_signature_save():
    """保存预录签名(dataURL PNG) → uploads/signatures/ → sys_config 记文件名"""
    d = request.json
    data_url = (d.get('dataUrl') or '')
    name = (d.get('name') or '').strip()[:20] or session['user_name']
    if not data_url.startswith('data:image/png'):
        return jsonify({'error': '签名格式错误'}), 400
    import base64
    try:
        raw = base64.b64decode(data_url.split(',')[1])
    except Exception:
        return jsonify({'error': '签名数据解码失败'}), 400
    os.makedirs(os.path.join(BASE, 'uploads', 'signatures'), exist_ok=True)
    fn = f'sig_{name}_{datetime.datetime.now().strftime("%Y%m%d%H%M%S")}.png'
    with open(os.path.join(BASE, 'uploads', 'signatures', fn), 'wb') as f:
        f.write(raw)
    conn = db()
    conn.execute("INSERT OR REPLACE INTO sys_config(key,value) VALUES('leader_signature',?)", (fn,))
    conn.execute("INSERT OR REPLACE INTO sys_config(key,value) VALUES('leader_signature_name',?)", (name,))
    conn.commit(); conn.close()
    log(session['user_name'], '保存预录签名', name)
    return jsonify({'success': True, 'file': fn})

@app.route('/api/signature/info')
@login_required
def api_signature_info():
    """返回预录签名状态(有无/名字/图片URL)"""
    conn = db()
    fn = conn.execute("SELECT value FROM sys_config WHERE key='leader_signature'").fetchone()
    nm = conn.execute("SELECT value FROM sys_config WHERE key='leader_signature_name'").fetchone()
    conn.close()
    f = fn[0] if fn else ''
    return jsonify({'has': bool(f), 'name': (nm[0] if nm else '') or '',
                    'url': '/uploads/signatures/' + f if f else ''})

@app.route('/api/signature/delete', methods=['POST'])
@login_required
def api_signature_delete():
    """删除预录签名"""
    conn = db()
    fn = conn.execute("SELECT value FROM sys_config WHERE key='leader_signature'").fetchone()
    if fn and fn[0]:
        try:
            p = os.path.join(BASE, 'uploads', 'signatures', fn[0])
            if os.path.exists(p): os.remove(p)
        except Exception: pass
    conn.execute("DELETE FROM sys_config WHERE key IN ('leader_signature','leader_signature_name')")
    conn.commit(); conn.close()
    log(session['user_name'], '删除预录签名', '')
    return jsonify({'success': True})

def get_leader_sign():
    """返回 (签名文件名, 签名人名字); 无则 ('','')"""
    conn = db()
    fn = conn.execute("SELECT value FROM sys_config WHERE key='leader_signature'").fetchone()
    nm = conn.execute("SELECT value FROM sys_config WHERE key='leader_signature_name'").fetchone()
    conn.close()
    return ((fn[0] if fn else ''), (nm[0] if nm else '') or '')

def stamp_leader_sign(ws, sign_row, biz_type='', biz_id=0):
    """V11.27: 单据Excel盖章 — 若单据已审批通过且配置了预录签名, 在签字区盖签名图+说明
    ws: openpyxl worksheet; sign_row: 签字区行号; 签名图插入到签字文本右侧空白区"""
    try:
        sign_fn, sign_name = get_leader_sign()
        if not sign_fn:
            return
        # 判断单据是否已审批通过(有 approved 节点)
        conn = db()
        ok = conn.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type=? AND biz_id=? AND status='approved'",
                          (biz_type, biz_id)).fetchone()[0] > 0
        conn.close()
        if not ok:
            return
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Font as XLFont
        img_path = os.path.join(BASE, 'uploads', 'signatures', sign_fn)
        if not os.path.exists(img_path):
            return
        img = XLImage(img_path)
        # V11.71: 签名尺寸(150x50), 紧凑放在H7(部门负责人右侧)
        try:
            from PIL import Image as PILImage
            pil_img = PILImage.open(img_path)
            orig_w, orig_h = pil_img.size
            max_w, max_h = 150, 50
            scale = min(max_w / orig_w, max_h / orig_h, 1.0)
            img.width = int(orig_w * scale)
            img.height = int(orig_h * scale)
        except Exception:
            img.width = 150; img.height = 50
        # 放在H7(部门负责人文字右侧)
        anchor = ws.cell(row=7, column=8).coordinate
        img.anchor = anchor
        ws.add_image(img)
        # 签名说明小字 (自动找签字区下方第一个非合并单元格)
        try:
            from openpyxl.cell.cell import MergedCell
            nr = sign_row + 1
            note = None
            while nr < sign_row + 20:
                cc = ws.cell(row=nr, column=6)
                if not isinstance(cc, MergedCell):
                    note = cc; break
                nr += 1
            if note is not None:
                note.value = '本单据经系统电子审批，签名由系统自动生成'
                note.font = XLFont(name='宋体', size=7, color='999999')
        except Exception:
            pass
    except Exception as e:
        print('stamp_leader_sign err:', e)

# ============================================================
# ── PURCHASE REQUESTS (多行申请) ──
# ============================================================
@app.route('/api/prequests')
@login_required
def api_prequests():
    # V11.64: 数据权限 — 员工/部门负责人只看自己的; 采购员/财务/领导/库管员看各自域
    role = session.get('user_role')
    scope = filter_scope(role)
    conn = db()
    if scope == 'own':
        rows = conn.execute("SELECT * FROM purchase_requests WHERE requester_id=? ORDER BY id DESC LIMIT 100", (session.get('user_id', 0),)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM purchase_requests ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        # V11.64: 库管员/员工/部门负责人 价格脱敏
        if scope in ('stock', 'own') and not can_see_price():
            d = mask_price(d)
        d['req_type'] = r['req_type'] if 'req_type' in r.keys() and r['req_type'] else '物资采购'
        # V11.49: 采购进度状态(红=未联系厂家/黄=已下单在途/绿=已到货)
        d['progress'] = 'none'
        if d.get('status') == '已通过':
            po = conn.execute("SELECT status FROM purchase_orders WHERE req_id=? ORDER BY id DESC LIMIT 1", (r['id'],)).fetchone()
            if po and po['status'] == '已入库':
                d['progress'] = 'arrived'      # 绿: 已入库
            elif po:
                d['progress'] = 'shipped'      # 黄: 已下单/在途
            else:
                d['progress'] = 'contact'      # 红: 未联系厂家
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/prequests/next_no')
@login_required
def api_prequest_next_no():
    # V11.45: 单号前缀随部门
    dept = request.args.get('dept') or ''
    return jsonify({'req_no': gen_req_no(dept)})

@app.route('/api/prequests/<int:rid>')
@login_required
def api_prequest(rid):
    conn = db()
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (rid,)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (rid,)).fetchall()
    approvals = conn.execute("SELECT * FROM approval_instances WHERE biz_type='purchase_request' AND biz_id=? ORDER BY level_no", (rid,)).fetchall()
    conn.close()
    return jsonify({'request':dict_row(pr),'items':[dict_row(i) for i in items],'approvals':[dict_row(a) for a in approvals]})

@app.route('/api/prequests', methods=['POST'])
@login_required
def api_create_prequest():
    d = request.json
    conn = db()
    items = d.get('items', [])
    total = sum(float(i.get('quantity',1)) * float(i.get('estimated_price',0)) for i in items)
    apply_date = d.get('apply_date') or datetime.date.today().strftime('%Y-%m-%d')
    # V11.199: 提交审批(非草稿)必填采购事由/部门 — 空事由钉钉必填控件发起报820001, 源头拦截
    if not d.get('draft'):
        if not str(d.get('purpose') or '').strip():
            conn.close(); return jsonify({'error': '请填写采购事由（这批物资买来做什么）'}), 400
        if not str(d.get('dept') or '').strip():
            conn.close(); return jsonify({'error': '请选择申请部门'}), 400
    # 并发安全: 单号冲突(UNIQUE)时重新生成重试(最多5次)
    no = ''
    for _try in range(5):
        no = gen_req_no(d.get('dept', ''), conn)
        try:
            # V11.154: draft=true → 存草稿不提交审批(采购员检查后再手动提交)
            _status = '草稿' if d.get('draft') else '待审批'
            conn.execute("""INSERT INTO purchase_requests(req_no,dept,requester,requester_id,budget_code,purpose,target_date,total_estimated,remark,attachments,urgent,apply_date,req_type,status)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (no, d.get('dept',''), session['user_name'], session['user_id'], d.get('budget_code',''),
                 d.get('purpose',''), d.get('target_date'), total, d.get('remark',''),
                 json.dumps(d.get('attachments') or [], ensure_ascii=False), 1 if d.get('urgent') else 0, apply_date,
                 d.get('req_type') or '物资采购', _status))
            break
        except sqlite3.IntegrityError:
            continue
    else:
        conn.close(); return jsonify({'error': '单号生成冲突，请重试'}), 500
    prid = conn.execute("SELECT id FROM purchase_requests WHERE req_no=?", (no,)).fetchone()[0]
    for it in items:
        tp = float(it.get('quantity',1)) * float(it.get('estimated_price',0))
        conn.execute("INSERT INTO request_items(req_id,item_name,spec,unit,quantity,estimated_price,total_price,remark,category,brand_param,arrival_date,attach) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                     (prid, it.get('item_name',''), it.get('spec',''), it.get('unit','个'), float(it.get('quantity',1)),
                      float(it.get('estimated_price',0)), tp, it.get('remark',''),
                      it.get('category',''), it.get('brand_param',''), it.get('arrival_date',''),
                      it.get('attach','') or ''))
    conn.commit()
    # V11.154: 草稿不创建审批实例(采购员检查后手动提交)
    if not d.get('draft'):
        create_approvals('purchase_request', prid, total, submitter=session['user_name'])
        start_instances('purchase_request', prid)   # 飞书/钉钉同步发起审批(未配置则跳过)
    conn.close()
    log(session['user_name'], '创建采购申请', f'{no} 共{len(items)}项 ¥{total:.0f}{" (草稿)" if d.get("draft") else ""}')
    return jsonify({'success':True, 'req_no':no, 'id':prid})

@app.route('/api/inventory/<int:iid>/replenish', methods=['POST'])
@login_required
def api_inventory_replenish(iid):
    """V11.33: 库存预警一键补货 — 低于安全库存的物资自动生成采购申请(补到安全线)"""
    conn = db()
    inv = conn.execute("SELECT * FROM inventory WHERE id=?", (iid,)).fetchone()
    if not inv:
        conn.close(); return jsonify({'error': '库存物资不存在'}), 404
    qty = float(inv['quantity'] or 0); safe = float(inv['safe_stock'] or 0)
    if safe <= 0 or qty >= safe:
        conn.close(); return jsonify({'error': '该物资未设置安全库存或未低于安全线'}), 400
    buy_qty = safe - qty  # 补到安全线
    price = float(inv['price'] or 0)
    est = buy_qty * price
    no = ''
    for _try in range(5):
        # V11.45: 补货申请归属采购部(CG前缀)
        no = gen_req_no('采购与供应链部', conn)
        try:
            conn.execute("""INSERT INTO purchase_requests(req_no,dept,requester,requester_id,budget_code,purpose,target_date,total_estimated,remark,urgent,apply_date)
                VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                (no, '采购与供应链部', session['user_name'], session['user_id'], '',
                 f'库存补货: {inv["item_name"]}', datetime.date.today().strftime('%Y-%m-%d'), est,
                 f'自动补货(库存不足): 当前{inv["quantity"]:g}{inv["unit"]}, 安全线{safe:g}{inv["unit"]}, 补货{buy_qty:g}{inv["unit"]}', 0,
                 datetime.date.today().strftime('%Y-%m-%d')))
            break
        except sqlite3.IntegrityError:
            continue
    else:
        conn.close(); return jsonify({'error': '单号生成冲突，请重试'}), 500
    prid = conn.execute("SELECT id FROM purchase_requests WHERE req_no=?", (no,)).fetchone()[0]
    conn.execute("INSERT INTO request_items(req_id,item_name,spec,unit,quantity,estimated_price,total_price,remark,category,brand_param,arrival_date) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                 (prid, inv['item_name'], inv['spec'] or '', inv['unit'] or '个', buy_qty, price, est,
                  '库存自动补货', inv['cat_code'] or '', '', ''))
    conn.commit()
    create_approvals('purchase_request', prid, est, submitter=session['user_name'])
    try: start_instances('purchase_request', prid)
    except Exception: pass
    conn.close()
    log(session['user_name'], '库存一键补货', f'{no} {inv["item_name"]} x{buy_qty:g} 自动补货')
    return jsonify({'success': True, 'req_no': no, 'id': prid, 'qty': buy_qty})

@app.route('/api/prequests/<int:rid>/reject', methods=['POST'])
@login_required
def api_reject_prequest(rid):
    d = request.json or {}
    conn = db()
    conn.execute("UPDATE purchase_requests SET status='已驳回', rejected_reason=?, updated_at=? WHERE id=?", (d.get('reason',''), now(), rid))
    conn.execute("UPDATE approval_instances SET status='rejected' WHERE biz_type='purchase_request' AND biz_id=? AND status='pending'", (rid,))
    conn.commit(); conn.close()
    log(session['user_name'], '驳回采购申请', f'申请#{rid}: {d.get("reason","")}')
    return jsonify({'success':True})

@app.route('/api/prequests/<int:rid>/resubmit', methods=['POST'])
@login_required
def api_resubmit_prequest(rid):
    """V8.0: 采购申请编辑/重提 — 任意状态可修改保存
    - 已驳回/草稿: 保存后重新进入审批流(待审批)
    - 待审批/已通过: 保存后回到待审批重新走流程(撤回原审批)
    - 已通过且已被订单引用: 禁止编辑"""
    d = request.json
    conn = db()
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (rid,)).fetchone()
    if not pr:
        conn.close(); return jsonify({'error': '申请不存在'}), 404
    if pr['status'] == '已通过':
        cnt = conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (rid,)).fetchone()[0]
        if cnt > 0:
            conn.close(); return jsonify({'error': '该申请已被订单引用，不可修改（可先删除订单）'}), 400
    if pr['status'] == '已作废':
        conn.close(); return jsonify({'error': '已作废申请不可修改'}), 400
    items = d.get('items') or []
    if not items: items = [dict(i) for i in conn.execute("SELECT * FROM request_items WHERE req_id=?", (rid,)).fetchall()]
    total = sum(float(i.get('quantity',1)) * float(i.get('estimated_price',0)) for i in items)
    # V11.187: 存草稿(draft=true) — 保存内容但状态保持草稿(不进审批/不清驳回标记); 提交则清标记+回待审批
    if d.get('draft'):
        conn.execute("UPDATE purchase_requests SET purpose=?, dept=?, budget_code=?, target_date=?, remark=?, req_type=?, urgent=?, attachments=?, total_estimated=?, status='草稿', updated_at=? WHERE id=?",
                     (d.get('purpose', pr['purpose']), d.get('dept', pr['dept']), d.get('budget_code', pr['budget_code']),
                      d.get('target_date', pr['target_date']), d.get('remark', pr['remark']),
                      d.get('req_type', pr['req_type'] if 'req_type' in pr.keys() else '物资采购'),
                      1 if d.get('urgent') else (pr['urgent'] if 'urgent' in pr.keys() else 0),
                      json.dumps(d.get('attachments') or [], ensure_ascii=False) if d.get('attachments') is not None else (pr['attachments'] if 'attachments' in pr.keys() else '[]'),
                      total, now(), rid))
    else:
        conn.execute("UPDATE purchase_requests SET purpose=?, dept=?, budget_code=?, target_date=?, remark=?, req_type=?, urgent=?, attachments=?, total_estimated=?, status='待审批', rejected_reason='', rejected_items='', resubmit_count=resubmit_count+1, updated_at=? WHERE id=?",
                     (d.get('purpose', pr['purpose']), d.get('dept', pr['dept']), d.get('budget_code', pr['budget_code']),
                      d.get('target_date', pr['target_date']), d.get('remark', pr['remark']),
                      d.get('req_type', pr['req_type'] if 'req_type' in pr.keys() else '物资采购'),
                      1 if d.get('urgent') else (pr['urgent'] if 'urgent' in pr.keys() else 0),
                      json.dumps(d.get('attachments') or [], ensure_ascii=False) if d.get('attachments') is not None else (pr['attachments'] if 'attachments' in pr.keys() else '[]'),
                      total, now(), rid))
    if items:
        # V11.154: 传了明细才重建(编辑时); 不传则保留现有明细(草稿提交审批场景)
        conn.execute("DELETE FROM request_items WHERE req_id=?", (rid,))
        for it in items:
            tp = float(it.get('quantity',1)) * float(it.get('estimated_price',0))
            conn.execute("INSERT INTO request_items(req_id,item_name,spec,unit,quantity,estimated_price,total_price,remark,category,brand_param,arrival_date,attach) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                         (rid, it.get('item_name',''), it.get('spec',''), it.get('unit','个'), float(it.get('quantity',1)),
                          float(it.get('estimated_price',0)), tp, it.get('remark',''),
                          it.get('category',''), it.get('brand_param',''), it.get('arrival_date',''),
                          it.get('attach','') or ''))
    if d.get('draft'):
        # V11.187: 存草稿 — 保留现有审批实例记录(驳回历史仍在), 不进审批流不发钉钉
        conn.commit(); conn.close()
        log(session['user_name'], '修改采购申请', f'申请#{rid} 保存为草稿(未提交)')
        return jsonify({'success':True, 'draft':True})
    conn.execute("DELETE FROM approval_instances WHERE biz_type='purchase_request' AND biz_id=?", (rid,))
    conn.execute("DELETE FROM dingtalk_instances WHERE biz_type='purchase_request' AND biz_id=?", (rid,))
    conn.commit()
    create_approvals('purchase_request', rid, total, submitter=session['user_name'])
    start_instances('purchase_request', rid)
    conn.close()
    log(session['user_name'], '修改采购申请', f'申请#{rid} 重新进入审批')
    return jsonify({'success':True})

# ============================================================
# ── ORDERS ──
# ============================================================
@app.route('/api/orders')
@login_required
def api_orders():
    # V11.159: 订单列表 — 员工仅看自己发起的(采购/库管/财务/领导/管理员全看)
    if session.get('user_role') == '员工':
        conn = db()
        rows = conn.execute("SELECT * FROM purchase_orders WHERE requester_id=? ORDER BY id DESC LIMIT 100", (session.get('user_id', 0),)).fetchall()
        out = []
        for r in rows:
            d = dict_row(r)
            cnt = conn.execute("SELECT COUNT(*), COALESCE(SUM(quantity),0) FROM order_items WHERE order_id=?", (r['id'],)).fetchone()
            d['item_count'] = cnt[0] or 1
            d['total_qty'] = cnt[1] or r['quantity']
            d['progress'] = 'none'
            _ost = d.get('status')
            if _ost in ('草稿', '待审批', '已驳回'):
                d['progress'] = 'none'
            else:
                _rc = conn.execute("SELECT 1 FROM receivings WHERE order_id=? AND status IN ('已入库','待检验','待入库','已挂账','已核销') LIMIT 1", (r['id'],)).fetchone()
                d['progress'] = 'done' if _rc else 'warn'
            out.append(d)
        conn.close()
        return jsonify(out)
    conn = db(); rows = conn.execute("SELECT * FROM purchase_orders ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        cnt = conn.execute("SELECT COUNT(*), COALESCE(SUM(quantity),0) FROM order_items WHERE order_id=?", (r['id'],)).fetchone()
        d['item_count'] = cnt[0] or 1
        d['total_qty'] = cnt[1] or r['quantity']
        # V11.126: 订单采购进度(与申请列表V11.49同语义: 红=未联系厂家/黄=已下单在途/绿=已到货)
        d['progress'] = 'none'
        _ost = d.get('status')
        if _ost in ('草稿', '待审批', '已驳回'):
            d['progress'] = 'none'
        elif _ost == '已入库' or conn.execute("SELECT 1 FROM receivings WHERE order_id=? AND status='已入库' LIMIT 1", (r['id'],)).fetchone():
            d['progress'] = 'arrived'
        else:
            _its = conn.execute("SELECT status FROM order_items WHERE order_id=?", (r['id'],)).fetchall()
            _sts = set((x['status'] or '未联系') for x in _its) if _its else set()
            if _sts and _sts <= {'已到货'}:
                d['progress'] = 'arrived'
            elif _sts & {'已发货', '已联系'}:
                d['progress'] = 'shipped'
            else:
                d['progress'] = 'contact'
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/orders/<int:oid>')
@login_required
def api_order(oid):
    conn = db()
    o = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
    approvals = conn.execute("SELECT * FROM approval_instances WHERE biz_type='purchase_order' AND biz_id=? ORDER BY level_no", (oid,)).fetchall()
    pcs = conn.execute("SELECT * FROM price_comparisons WHERE order_id=?", (oid,)).fetchall()
    items = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()
    # V11.151: 补 has_contract — 订单是否已有有效合同(前端据此隐藏"生成采购合同"按钮, 防重复生成)
    _has_ct = conn.execute(
        "SELECT 1 FROM contracts WHERE order_id=? AND status NOT IN ('已作废','已撤回','撤回') LIMIT 1",
        (oid,)).fetchone() is not None
    # V11.202 分批验收: 订单验收统计(须在 conn 关闭前计算)
    _st = None
    try:
        _st = _order_rcv_stats(conn, oid)
    except Exception:
        _st = None
    conn.close()
    _od = dict_row(o)
    _od['has_contract'] = _has_ct
    return jsonify({'order': _od, 'items': [dict_row(i) for i in items],
                    'approvals': [dict_row(a) for a in approvals],
                    'comparisons': [dict_row(p) for p in pcs],
                    'rcv_stats': _st})


@app.route('/api/orders/<int:oid>/receiving-batch', methods=['POST'])
@login_required
def api_order_receiving_batch(oid):
    """V11.202 分批验收: 订单新增一批验收入库单(独立走审批, 审批通过+仓库确认后才加库存)。
    校验: 本批数量>0 / 每明细行与订单累计都不超量 / 老整批流程单自动作废(防重复入库); 原暂估/正式逻辑完全复用。"""
    if session.get('user_role') not in ('库管员', '部门负责人', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：分批验收入库仅限库管员/领导使用'}), 403
    d = request.json or {}
    is_est = 1 if int(d.get('is_est', 0) or 0) == 1 else 0
    conn = db()
    po = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
    if not po:
        conn.close(); return jsonify({'error': '订单不存在'}), 404
    if po['status'] in ('草稿', '待审批', '已驳回', '已作废', '已取消', '已入库', '已核销', '全部已验收'):
        conn.close(); return jsonify({'error': f'订单当前状态({po["status"]})不可新增验收批次'}), 400
    st = _order_rcv_stats(conn, oid)
    if st['pending'] <= 0.001:
        conn.close(); return jsonify({'error': '该订单已全部验收完成，无需再新增批次'}), 400
    # 老整批单正在审批中 → 必须先撤回(防止双流程重复入库)
    _pend_full = conn.execute("SELECT id FROM receivings WHERE order_id=? AND (batch_no IS NULL OR batch_no='') AND status='待审批' LIMIT 1", (oid,)).fetchone()
    if _pend_full:
        conn.close(); return jsonify({'error': '该订单原有整批入库单正在审批中，请先在入库验收中撤回该单后再分批验收'}), 400
    oi = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()
    qty_items = d.get('items')
    # 解析本批每行数量: list按订单明细行序 / dict按物资名
    parsed = []  # (order_item, qty)
    if qty_items and isinstance(qty_items, list):
        for idx, it in enumerate(oi):
            q = 0.0
            if idx < len(qty_items):
                try:
                    q = float((qty_items[idx].get('quantity') if isinstance(qty_items[idx], dict) else qty_items[idx]) or 0)
                except Exception:
                    q = 0.0
            parsed.append((it, q))
    elif qty_items and isinstance(qty_items, dict):
        for it in oi:
            try:
                q = float(qty_items.get(it['item_name'], 0) or 0)
            except Exception:
                q = 0.0
            parsed.append((it, q))
    else:
        conn.close(); return jsonify({'error': '请填写本批验收数量'}), 400
    total_batch = 0.0
    errs = []
    for it, q in parsed:
        if q < 0:
            errs.append(f"「{it['item_name']}」数量不能为负")
            continue
        if q > 0:
            total_batch += q
            # 单行余量(订单行数量 - 已入库行累计; 规格一端为空的旧数据宽容匹配)
            line_qty = float(it['quantity'] or 0)
            line_acc = next((pl['accepted'] for pl in st['per_line'] if pl['item_name'] == it['item_name'] and ((pl['spec'] or '') == (it['spec'] or '') or not (pl['spec'] or '') or not (it['spec'] or ''))), 0.0)
            if q > line_qty - line_acc + 1e-9:
                errs.append(f"「{it['item_name']}」本批{q}超过可验余量{max(0, line_qty - line_acc):g}")
    if errs:
        conn.close(); return jsonify({'error': '；'.join(errs)}), 400
    if total_batch <= 0:
        conn.close(); return jsonify({'error': '本批验收数量必须大于0'}), 400
    if st['accepted'] + total_batch > st['order_total'] + 1e-9:
        conn.close(); return jsonify({'error': f'累计验收超量: 订单总数{st["order_total"]:g}，已验收{st["accepted"]:g}，本批{total_batch:g}'}), 400
    # 批次号 = 已有批次序号+1
    _n = conn.execute("SELECT COUNT(*) FROM receivings WHERE order_id=? AND batch_no IS NOT NULL AND batch_no<>''", (oid,)).fetchone()[0]
    batch_no = '第%d批' % (_n + 1)
    # 自动作废未入库存的老整批流程单(待入库/入库中/待检验/草稿/已驳回), 防双流程重复入库
    _old_docs = conn.execute("SELECT id, receive_no, status FROM receivings WHERE order_id=? AND (batch_no IS NULL OR batch_no='') AND status IN ('待入库','入库中','待检验','草稿','已驳回')", (oid,)).fetchall()
    for _od2 in _old_docs:
        conn.execute("UPDATE receivings SET status='已作废', remark=COALESCE(remark,'')||' 订单启用分批验收,整批单自动作废' WHERE id=?", (_od2['id'],))
        conn.execute("UPDATE approval_instances SET status='rejected', comment='订单启用分批验收,整批单作废' WHERE biz_type='receiving' AND biz_id=? AND status='pending'", (_od2['id'],))
    # 部门带出(申请链)
    _dept = ''
    try:
        if po['req_id']:
            _pr = conn.execute("SELECT dept FROM purchase_requests WHERE id=?", (po['req_id'],)).fetchone()
            if _pr and _pr['dept']:
                _dept = _pr['dept']
    except Exception:
        pass
    # 明细JSON(本批数量+订单单价)
    _lines = []
    for it, q in parsed:
        if q > 0:
            _lines.append({'item_name': it['item_name'], 'spec': it['spec'] or '', 'quantity': q,
                           'unit': it['unit'] or '个', 'price': float(it['price'] or 0),
                           'tax_rate': float(it['tax_rate'] or 13)})
    _typ_txt = '暂估入库' if is_est else '正式入库'
    _items_json = json.dumps(_lines, ensure_ascii=False)
    _first = _lines[0]
    _name = (_first['item_name'] + ' 等%d项' % len(_lines)) if len(_lines) > 1 else _first['item_name']
    rno = gen_no('RK', 'receivings', 'receive_no', conn)
    conn.execute("""INSERT INTO receivings(receive_no,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,dept,items_json,is_est,batch_no,inspector,warehouse)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                 (rno, oid, _name, '', total_batch, _first['unit'] or '个', total_batch, '待审批', now(),
                  f'分批验收{batch_no}·{_typ_txt}', _dept, _items_json, is_est, batch_no,
                  session.get('user_name', ''), d.get('warehouse', '主库房')))
    rid = conn.execute("SELECT id FROM receivings WHERE receive_no=?", (rno,)).fetchone()[0]
    conn.commit()
    create_approvals('receiving', rid, 0, submitter=session['user_name'])
    conn.close()
    try:
        start_instances('receiving', rid)
    except Exception as e:
        print('receiving-batch start_instances err:', e)
    log(session['user_name'], '分批验收入库', f'订单{po["order_no"]} {batch_no} {_typ_txt} {total_batch:g}件 入库单{rno} 待审批')
    return jsonify({'success': True, 'message': f'{batch_no}({_typ_txt}) {total_batch:g}件已提交审批，审批通过后自动增加库存', 'receive_no': rno, 'id': rid, 'batch_no': batch_no})

@app.route('/api/order_items/<int:iid>/status', methods=['POST'])
@login_required
def api_update_order_item_status(iid):
    """V11.78: 更新订单物品状态"""
    d = request.json
    new_status = d.get('status', '未联系')
    conn = db()
    item = conn.execute("SELECT * FROM order_items WHERE id=?", (iid,)).fetchone()
    if not item:
        conn.close()
        return jsonify({'error': '物品不存在'}), 404
    conn.execute("UPDATE order_items SET status=?, updated_at=? WHERE id=?", (new_status, now(), iid))
    conn.commit()
    conn.close()
    log(session['user_name'], '订单物品状态', '订单物品#%d: %s → %s' % (iid, item['status'] or '未联系', new_status))
    return jsonify({'success': True})

@app.route('/api/orders', methods=['POST'])
@login_required
def api_create_order():
    """同一批下单商品为一张订单: 订单头(汇总) + order_items 明细行; 一次审批"""
    d = request.json
    conn = db()
    tm = (d.get('trade_mode') or '货到付款').strip() or '货到付款'
    # V11.3: 交易模式支持自定义, 不局限于 货到付款/先款后货
    items = d.get('items') or []
    if not items and d.get('item_name'):
        items = [{'item_name': d.get('item_name'), 'spec': d.get('spec'), 'unit': d.get('unit','个'),
                  'quantity': d.get('quantity',1), 'price': d.get('price',0), 'tax_rate': d.get('tax_rate',13)}]
    if not items:
        conn.close(); return jsonify({'error': '请至少添加一个商品'}), 400
    # 防重复下单: 从申请导入时(req_id已传), 校验该申请存在且未生成过订单
    req_id = d.get('req_id')
    if req_id:
        pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (req_id,)).fetchone()
        if not pr:
            conn.close(); return jsonify({'error': '来源申请不存在'}), 400
        if pr['status'] != '已通过':
            conn.close(); return jsonify({'error': f'来源申请当前状态({pr["status"]})不可下单'}), 400
        if conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (req_id,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该申请已下单，请勿重复下单'}), 400
    no = gen_no('CG', 'purchase_orders', 'order_no', conn)
    rows = []
    total_qty = 0.0; grand_amt = 0.0; grand_tax = 0.0; grand_total = 0.0
    for it in items:
        qty = float(it.get('quantity',1) or 1)
        price = float(it.get('price',0) or 0)
        tr = float(it.get('tax_rate', d.get('tax_rate',13)) or 13)
        amt = qty*price; tax = amt*tr/100
        total_qty += qty; grand_amt += amt; grand_tax += tax; grand_total += amt+tax
        rows.append((it.get('item_name',''), it.get('spec','') or '', it.get('unit','个') or '个', qty, price, amt, tr, tax, amt+tax))
    first = rows[0]
    conn.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
        supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (no, d.get('req_id'), first[0], first[1], total_qty, first[2], first[4], grand_amt, first[6], grand_tax, grand_total,
         d.get('supplier',''), d.get('requester',''), d.get('category','后勤类'), session['user_name'], session['user_id'],
         d.get('target_date'), tm, d.get('remark',''), 1 if d.get('urgent') else 0,
         json.dumps(d.get('attachments') or [], ensure_ascii=False)))
    oid = conn.execute("SELECT id FROM purchase_orders WHERE order_no=?", (no,)).fetchone()[0]
    for r in rows:
        conn.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
            (oid, r[0], r[1], r[2], r[3], r[4], r[5], r[6], r[7], r[8], ''))
    rno = None
    if tm == '货到付款':
        # 整批商品一张入库单, 待【确认验收入库】批量转入正式库存
        # V11.152e: 防重复 — 该订单已有任意未入库状态的入库单(待入库/待检验/入库中/待审批)则不再生成
        _dup = conn.execute("SELECT id FROM receivings WHERE order_id=? AND status NOT IN ('已入库','已作废') LIMIT 1", (oid,)).fetchone()
        if not _dup:
            # V11.31: 自动带出部门(申请单链)
            _dept = ''
            try:
                if d.get('req_id'):
                    _pr = conn.execute("SELECT dept FROM purchase_requests WHERE id=?", (d.get('req_id'),)).fetchone()
                    if _pr and _pr['dept']:
                        _dept = _pr['dept']
            except Exception:
                pass
            rno = gen_no('RK', 'receivings', 'receive_no', conn)
            # V11.152: 多物资订单自动生成入库时, 明细完整存items_json(入库验收显示全部物资名)
            _items_json = json.dumps(
                [{'item_name': r[0], 'spec': r[1] or '', 'quantity': r[3], 'unit': r[2] or '个', 'price': r[4] or 0} for r in rows],
                ensure_ascii=False)
            _name = (first[0] + ' 等%d项' % len(rows)) if len(rows) > 1 else first[0]
            # V11.198: 自动生成的入库单不预设类型(is_est=0待定, 提交审批时手动选择暂估/正式)
            conn.execute("INSERT INTO receivings(receive_no,delivery_id,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,dept,items_json,is_est) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,0)",
                (rno, None, oid, _name, '', total_qty, first[2], 0, '待入库', now(), '货到付款: 下单后自动进入入库板块(整批%d项)' % len(rows), _dept, _items_json))
    conn.commit()
    create_approvals('purchase_order', oid, grand_total, submitter=session['user_name'])   # 一张订单一次审批
    start_instances('purchase_order', oid)
    conn.close()
    log(session['user_name'], '创建采购订单', '%s 共%d项商品 ¥%.0f' % (no, len(rows), grand_total))
    return jsonify({'success':True, 'order_no': no, 'id': oid, 'receive_no': rno,
                    'total_qty': total_qty, 'total_amount': grand_total, 'item_count': len(rows)})

@app.route('/api/orders/<int:oid>/submit', methods=['POST'])
@login_required
def api_order_submit(oid):
    """V11.3: 草稿订单提交审批 — 确认商家信息无误后, 创建审批实例并向上审批"""
    conn = db()
    o = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
    if not o:
        conn.close(); return jsonify({'error': '订单不存在'}), 404
    if o['status'] != '草稿':
        conn.close(); return jsonify({'error': '仅草稿状态的订单可提交审批'}), 400
    if conn.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type='purchase_order' AND biz_id=? AND status='pending'", (oid,)).fetchone()[0] > 0:
        conn.close(); return jsonify({'error': '该订单已在审批中'}), 400
    conn.execute("UPDATE purchase_orders SET status='待审批', updated_at=? WHERE id=?", (now(), oid))
    conn.commit()
    amount = float(o['total_amount'] or 0)
    create_approvals('purchase_order', oid, amount, submitter=session['user_name'])
    try:
        start_instances('purchase_order', oid)
    except Exception as e:
        print('order submit start_instances err:', e)
    conn.close()
    log(session['user_name'], '订单提交审批', '%s 由草稿提交审批 ¥%.0f' % (o['order_no'], amount))
    return jsonify({'success': True, 'order_no': o['order_no'], 'status': '待审批',
                    'message': '订单已提交审批，审批通过后自动进入后续流程'})

# ============================================================
# ── NOTIFICATIONS ──
# ============================================================
@app.route('/api/notifications')
@login_required
def api_notifications():
    conn = db()
    rows = conn.execute("SELECT * FROM notifications WHERE user_id=? AND is_read=0 ORDER BY id DESC LIMIT 20", (session['user_id'],)).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/notifications/read', methods=['POST'])
@login_required
def api_read_notifications():
    conn = db()
    conn.execute("UPDATE notifications SET is_read=1 WHERE user_id=?", (session['user_id'],))
    conn.commit(); conn.close()
    return jsonify({'success':True})

# ============================================================
# ── STATS ──
# ============================================================
@app.route('/api/reset-db', methods=['POST'])
def api_reset_db():
    conn = db()
    conn.executescript("""
        PRAGMA writable_schema=ON;
        DELETE FROM sqlite_master WHERE type='table';
        PRAGMA writable_schema=OFF;
        VACUUM;
    """)
    conn.close()
    init_db()
    return jsonify({'success':True,'message':'数据库已重置'})

@app.route('/api/stats')
@login_required
def api_stats():
    conn = db()
    stats = {
        'pending_approvals': conn.execute("SELECT COUNT(*) FROM approval_instances WHERE status='pending'").fetchone()[0],
        'active_requests': conn.execute("SELECT COUNT(*) FROM purchase_requests WHERE status='待审批'").fetchone()[0],
        'active_orders': conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE status NOT IN ('已完成','已关闭','已挂账')").fetchone()[0],
        'stock_warnings': conn.execute("SELECT COUNT(*) FROM inventory WHERE quantity<=safe_stock AND safe_stock>0").fetchone()[0],
        'total_purchase': conn.execute("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders").fetchone()[0],
        'overdue_items': conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE target_date<date('now') AND status NOT IN ('已完成','已关闭')").fetchone()[0],
    }
    conn.close()
    return jsonify(stats)

# ============================================================
# V55 ── 首页工作台聚合接口: 快捷导航/审批专区/预警中心/数据看板
# ============================================================
@app.route('/api/dashboard')
@login_required
def api_dashboard():
    role = session['user_role']; uid = session['user_id']
    can_price = role in ('系统管理员', '财务', '分管领导', '总经理')
    def pv(v):  # 财务数据脱敏: 无权限返回None
        return round(float(v or 0), 2) if can_price else None
    c = db()
    m = datetime.date.today().strftime('%Y-%m')
    # ── ① 指标卡 ──
    month_purchase = c.execute("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE created_at LIKE ?", (m + '%',)).fetchone()[0]
    pending_in = c.execute("SELECT COUNT(*) FROM receivings WHERE status='待入库'").fetchone()[0]
    pending_appr = c.execute("SELECT COUNT(*) FROM approval_instances WHERE status='pending'").fetchone()[0]
    urgent_cnt = 0
    for t in ('purchase_requests', 'purchase_orders', 'contracts', 'payment_requests'):
        urgent_cnt += c.execute("SELECT COUNT(*) FROM %s WHERE urgent=1 AND status IN ('待审批','审批通过','已通过','执行中')" % t).fetchone()[0]
    overdue_pay_amt = c.execute("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status IN ('待审批','已通过') AND expect_pay_date!='' AND expect_pay_date<date('now')").fetchone()[0]
    stock_low_cnt = c.execute("SELECT COUNT(*) FROM inventory WHERE safe_stock>0 AND quantity<=safe_stock").fetchone()[0]
    # 55.docx看板补全: 待付款总额/订单完成率/库存总量/超储数量/重点供应商占比
    pending_pay_amt = c.execute("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status IN ('待审批','已通过')").fetchone()[0]
    done_cnt = c.execute("SELECT COUNT(*) FROM purchase_orders WHERE status IN ('已完成','已入库','已核销')").fetchone()[0]
    total_cnt = c.execute("SELECT COUNT(*) FROM purchase_orders").fetchone()[0]
    stock_total_qty = c.execute("SELECT COALESCE(SUM(quantity),0) FROM inventory").fetchone()[0]
    stock_over_cnt = c.execute("SELECT COUNT(*) FROM inventory WHERE max_stock>0 AND quantity>=max_stock").fetchone()[0]
    top_sups = c.execute("SELECT supplier, SUM(total_amount) amt FROM purchase_orders WHERE supplier!='' GROUP BY supplier ORDER BY amt DESC LIMIT 5").fetchall()
    kpi = {
        'month_purchase': pv(month_purchase), 'pending_in': pending_in,
        'pending_appr': pending_appr, 'urgent_cnt': urgent_cnt,
        'overdue_pay': pv(overdue_pay_amt), 'stock_low': stock_low_cnt,
        'pending_pay': pv(pending_pay_amt),
        'order_complete_rate': round(done_cnt*100.0/total_cnt, 1) if total_cnt else 0,
        'stock_total': stock_total_qty, 'stock_over': stock_over_cnt,
        'top_suppliers': [{'name': r['supplier'], 'amount': pv(r['amt'])} for r in top_sups],
        'can_price': can_price,
    }
    # ── ② 审批消息专区 ──
    def biz_no_expr():
        return _ap_case('no')  # V11.213 集中CASE(含维修/退库/集体验收/询价审批)
    def biz_name_expr():
        return _ap_case('name')  # V11.213 集中CASE
    def urgent_expr():
        return "(CASE WHEN ai.biz_type='purchase_request' THEN (SELECT pr.urgent FROM purchase_requests pr WHERE pr.id=ai.biz_id) WHEN ai.biz_type='purchase_order' THEN (SELECT po.urgent FROM purchase_orders po WHERE po.id=ai.biz_id) WHEN ai.biz_type='contract' THEN (SELECT ct.urgent FROM contracts ct WHERE ct.id=ai.biz_id) WHEN ai.biz_type='payment' THEN (SELECT pp.urgent FROM payment_requests pp WHERE pp.id=ai.biz_id) ELSE 0 END)"
    my_pending = c.execute("""SELECT ai.*, %s as biz_no, %s as biz_name, %s as urgent
        FROM approval_instances ai WHERE ai.status='pending'
        AND NOT EXISTS (SELECT 1 FROM approval_instances y WHERE y.biz_type=ai.biz_type AND y.biz_id=ai.biz_id AND y.status='pending' AND y.level_no < ai.level_no)
        AND (ai.role=? OR (ai.role='部门负责人' AND ? IN ('部门负责人','系统管理员')) OR ai.approver_id=?)
        ORDER BY %s DESC, ai.id DESC LIMIT 30""" % (biz_no_expr(), biz_name_expr(), urgent_expr(), urgent_expr()),
        (role, role, session.get('user_id', 0))).fetchall()
    my_urgent = [dict_row(x) for x in my_pending if x['urgent']]
    i_started = c.execute("""SELECT * FROM (
            SELECT '申请' bt, req_no no, purpose name, status, created_at FROM purchase_requests WHERE requester_id=?
            UNION ALL SELECT '订单', order_no, item_name, status, created_at FROM purchase_orders WHERE owner_id=?
            UNION ALL SELECT '合同', contract_no, contract_name, status, created_at FROM contracts WHERE remark LIKE ?
            UNION ALL SELECT '付款', payment_no, payment_reason, status, created_at FROM payment_requests WHERE supplier!=''
        ) ORDER BY created_at DESC LIMIT 8""", (uid, uid, '%%%s%%' % session['user_name'])).fetchall()
    i_done = c.execute("""SELECT biz_type, biz_id, role, status, comment, processed_at FROM approval_instances
        WHERE approver_id=? AND status IN ('approved','rejected') ORDER BY processed_at DESC LIMIT 8""", (uid,)).fetchall()
    rejected = c.execute("""SELECT biz_type, biz_id, COUNT(*) cnt, MAX(comment) last_comment FROM approval_instances
        WHERE status='rejected' GROUP BY biz_type, biz_id ORDER BY cnt DESC LIMIT 8""").fetchall()
    # ── ③ 预警消息中心 ──
    alerts = c.execute("SELECT * FROM alert_items WHERE status='pending' ORDER BY CASE level WHEN 'red' THEN 0 WHEN 'orange' THEN 1 ELSE 2 END, id DESC LIMIT 30").fetchall()
    # V11.60: 询价超3天未完成 → 动态加入预警(避免询价卡死没人管)
    try:
        inq_stale = c.execute("""SELECT id, inq_no, title, created_at FROM inquiries
            WHERE status IN ('询价中') AND created_at < datetime('now','localtime','-3 days')
            ORDER BY created_at""").fetchall()
        for q in inq_stale:
            days = max(1, int((datetime.datetime.now() - datetime.datetime.strptime(q['created_at'][:19], '%Y-%m-%d %H:%M:%S')).total_seconds() / 86400))
            alerts.insert(0, {
                'id': 0, 'alert_type': '询价超时', 'level': 'orange',
                'title': f"询价单 {q['inq_no']}",
                'content': f"已发起{days}天未完成, 请及时跟进商家报价/选中",
                'biz_type': 'inquiry', 'biz_id': q['id'], 'created_at': q['created_at'],
                'status': 'pending', 'link': f"sw('inquiries')",
            })
    except Exception:
        pass
    # V11.203 模块一1.2: 发票节点到期/超期 → 系统内动态预警(钉钉推送由 check_invoice_node_reminders 负责)
    try:
        _today_s = datetime.date.today().strftime('%Y-%m-%d')
        _ic = c.execute("SELECT * FROM contracts WHERE status='执行中' AND (invoice_est_first!='' OR invoice_est_done!='')").fetchall()
        for _ct in _ic:
            _st2 = _contract_inv_stats(c, _ct['id'])
            _pend2 = float(_ct['amount'] or 0) - _st2['received_amount']
            _lv = ''
            _ac = ''
            if _ct['invoice_est_first'] and _today_s >= _ct['invoice_est_first'][:10] and _st2['received_count'] == 0:
                _lv = 'orange'
                _ac = '已到预计首次开票日(%s)，尚未收到任何发票' % _ct['invoice_est_first'][:10]
            elif _ct['invoice_est_done'] and _today_s > _ct['invoice_est_done'][:10] and _pend2 > 0.01:
                _lv = 'red'
                _ac = '已超过约定开票完成日(%s)，仍未收票 ¥%.2f' % (_ct['invoice_est_done'][:10], _pend2)
            if _lv:
                alerts.insert(0, {'id': 0, 'alert_type': '发票催收', 'level': _lv,
                                  'title': '合同 %s' % _ct['contract_no'],
                                  'content': _ac + '（供应商:%s），请采购专员及时跟进取票并登记' % (_ct['supplier'] or '-'),
                                  'biz_type': 'contract', 'biz_id': _ct['id'],
                                  'created_at': _ct['updated_at'] or '', 'status': 'pending',
                                  'link': "sw('contracts')"})
    except Exception:
        pass
    # ── ④ 数据看板 ──
    trend = []
    for i in range(5, -1, -1):
        d0 = datetime.date.today().replace(day=1) - datetime.timedelta(days=31 * i)
        mm = d0.strftime('%Y-%m')
        s = c.execute("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE created_at LIKE ?", (mm + '%',)).fetchone()[0]
        trend.append({'month': mm, 'amount': round(s or 0, 2)})
    cat_ratio = c.execute("""SELECT po.category, COALESCE(SUM(po.total_amount),0) amt FROM purchase_orders po
        WHERE po.total_amount>0 GROUP BY po.category ORDER BY amt DESC LIMIT 8""").fetchall()
    in_total = c.execute("SELECT COUNT(*) FROM receivings WHERE status='已入库'").fetchone()[0]
    in_ontime = 0
    if in_total:
        try: c.execute("ALTER TABLE receivings ADD COLUMN completed_at TEXT")
        except Exception: pass
        in_ontime = c.execute("""SELECT COUNT(*) FROM receivings r JOIN purchase_orders po ON r.order_id=po.id
            WHERE r.status='已入库' AND po.target_date!='' AND (r.completed_at IS NULL OR r.completed_at<=po.target_date)""").fetchone()[0]
    ontime_rate = round(in_ontime * 100.0 / in_total, 1) if in_total else None
    avg_h = c.execute("""SELECT AVG((julianday(processed_at)-julianday(created_at))*24) FROM approval_instances
        WHERE status IN ('approved','rejected') AND processed_at IS NOT NULL""").fetchone()[0]
    c.close()
    return jsonify({
        'kpi': kpi,
        'my_pending': [dict_row(x) for x in my_pending],
        'my_urgent': my_urgent,
        'i_started': [dict_row(x) for x in i_started],
        'i_done': [dict_row(x) for x in i_done],
        'rejected': [dict_row(x) for x in rejected],
        'alerts': [dict_row(x) for x in alerts],
        'trend': trend,
        'cat_ratio': [dict_row(x) for x in cat_ratio],
        'ontime_rate': ontime_rate, 'approve_avg_hours': round(float(avg_h), 1) if avg_h else None,
    })

# ============================================================
# ── COMMON APIs (reused from v3) ──
# ============================================================
@app.route('/api/categories')
@login_required
def api_categories():
    conn = db(); rows = conn.execute("SELECT * FROM categories").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/items')
@login_required
def api_items():
    conn = db(); rows = conn.execute("SELECT i.*,c.name as cat_name FROM items i LEFT JOIN categories c ON i.cat_code=c.code").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/suppliers')
@login_required
def api_suppliers():
    # V11.159: 供应商档案(含联系方式) — 员工不显示(采购/库管/财务/领导/管理员可见)
    if session.get('user_role') == '员工':
        return jsonify([])
    conn = db(); rows = conn.execute("SELECT * FROM suppliers").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/suppliers', methods=['POST'])
@admin_required
def api_create_supplier():
    """需求44-基础档案: 供应商完整档案(含开户行/账号, 供合同自动填充乙方信息)"""
    d = request.json
    c = db()
    if c.execute("SELECT 1 FROM suppliers WHERE name=?", (d.get('name', ''),)).fetchone():
        c.close(); return jsonify({'error': '供应商已存在'}), 400
    c.execute("""INSERT INTO suppliers(name,contact,phone,category,level,bank,account,tax_id,invoice_type,rating,status)
        VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
        (d.get('name', ''), d.get('contact', ''), d.get('phone', ''), d.get('category', ''), d.get('level', '一般供应商'),
         d.get('bank', ''), d.get('account', ''), d.get('tax_id', ''), d.get('invoice_type', '增值税专用发票'),
         d.get('rating', 4.0), d.get('status', '正常')))
    c.commit(); c.close()
    log(session['user_name'], '新增供应商', d.get('name', ''))
    return jsonify({'success': True})

@app.route('/api/suppliers/<int:sid>', methods=['POST'])
@admin_required
def api_update_supplier(sid):
    d = request.json
    c = db()
    c.execute("""UPDATE suppliers SET name=?,contact=?,phone=?,category=?,level=?,bank=?,account=?,tax_id=?,invoice_type=?,rating=?,status=? WHERE id=?""",
        (d.get('name', ''), d.get('contact', ''), d.get('phone', ''), d.get('category', ''), d.get('level', '一般供应商'),
         d.get('bank', ''), d.get('account', ''), d.get('tax_id', ''), d.get('invoice_type', '增值税专用发票'),
         d.get('rating', 4.0), d.get('status', '正常'), sid))
    c.commit(); c.close()
    log(session['user_name'], '编辑供应商', f'#{sid}')
    return jsonify({'success': True})

# ============================================================
# ── 三方询价 (V11.0) ──
# 业务流: 采购申请审批通过 → 发起三方询价(最多3家) → 商家免登录报价(/inq/<token>)
#         → 比价选中一家 → 自动生成采购订单(进订单审批)
# ============================================================
@app.route('/api/inquiries')
@login_required
def api_inquiries():
    """询价单列表(含家数/已报价数/最低报价/审批状态)
    V11.217: 三家未全部报价前 min_price 返回 None(列表不提前泄露最低价); 全部报价后才显示"""
    conn = db()
    rows = conn.execute("""
        SELECT i.*, pr.req_no, pr.purpose, pr.dept,
            (SELECT COUNT(*) FROM inquiry_suppliers s WHERE s.inquiry_id=i.id) AS sup_count,
            (SELECT COUNT(*) FROM inquiry_suppliers s WHERE s.inquiry_id=i.id AND s.quote_price>0) AS quoted_count,
            (SELECT status FROM inquiry_approvals WHERE inquiry_id=i.id ORDER BY id DESC LIMIT 1) AS approval_status
        FROM inquiries i LEFT JOIN purchase_requests pr ON i.req_id=pr.id
        ORDER BY i.id DESC LIMIT 100
    """).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict_row(r)
        # V11.217: 全部受邀供应商报完价 → 才带最低价(否则 None 前端显示'待开标')
        _sc = d['sup_count'] or 0
        _qc = d['quoted_count'] or 0
        if _sc >= 2 and _qc >= _sc:
            conn2 = db()
            _m = conn2.execute("SELECT MIN(quote_price) m FROM inquiry_suppliers WHERE inquiry_id=? AND quote_price>0", (d['id'],)).fetchone()[0]
            conn2.close()
            d['min_price'] = _m
            d['all_quoted'] = True
        else:
            d['min_price'] = None
            d['all_quoted'] = False
        d['_quoted_label'] = f"{_qc}/{_sc}"
        out.append(d)
    return jsonify(out)

@app.route('/api/inquiries/next_no')
@login_required
def api_inquiries_next_no():
    conn = db()
    no = gen_no('XJ', 'inquiries', 'inq_no', conn)
    conn.close()
    return jsonify({'inq_no': no})

@app.route('/api/inquiries/eligible')
@login_required
def api_inquiries_eligible():
    """可询价申请列表: 已通过 且 未询价过 且 未下单"""
    conn = db()
    rows = conn.execute("""
        SELECT pr.*,
            (SELECT COUNT(*) FROM request_items ri WHERE ri.req_id=pr.id) AS item_count,
            (SELECT SUM(COALESCE(ri.total_price,0)) FROM request_items ri WHERE ri.req_id=pr.id) AS item_total
        FROM purchase_requests pr
        WHERE pr.status='已通过'
          AND NOT EXISTS (SELECT 1 FROM inquiries i WHERE i.req_id=pr.id AND i.status!='已取消')
          AND NOT EXISTS (SELECT 1 FROM purchase_orders po WHERE po.req_id=pr.id)
        ORDER BY pr.id DESC LIMIT 50
    """).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/inquiries', methods=['POST'])
@login_required
def api_create_inquiry():
    """发起三方询价: 选中已通过申请 + 最多3家供应商(每家生成免登录报价链接)"""
    d = request.json
    req_id = d.get('req_id')
    suppliers = d.get('suppliers') or []
    if not req_id:
        return jsonify({'error': '请选择要询价的申请'}), 400
    if not suppliers or len(suppliers) > 3:
        return jsonify({'error': '请添加2-3家询价供应商'}), 400
    # V11.77: 允许2-3家询价，少于2家不允许
    valid_names = [x for x in suppliers if (x.get('name') or '').strip()]
    if len(valid_names) < 2:
        return jsonify({'error': '询价供应商需选择2家及以上，方可发起询价'}), 400
    conn = db()
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (req_id,)).fetchone()
    if not pr:
        conn.close(); return jsonify({'error': '申请不存在'}), 400
    if pr['status'] != '已通过':
        conn.close(); return jsonify({'error': '申请当前状态(%s)不可询价' % pr['status']}), 400
    if conn.execute("SELECT COUNT(*) FROM inquiries WHERE req_id=? AND status!='已取消'", (req_id,)).fetchone()[0] > 0:
        conn.close(); return jsonify({'error': '该申请已发起过询价'}), 400
    if conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (req_id,)).fetchone()[0] > 0:
        conn.close(); return jsonify({'error': '该申请已下单，无需询价'}), 400
    no = gen_no('XJ', 'inquiries', 'inq_no', conn)
    title = (pr['purpose'] or '')[:80]
    # V11.24: 报价截止时间 — V11.205 精确到分钟(统一开标): 传 deadline(YYYY-MM-DD 或 YYYY-MM-DD HH:MM), 不传默认7天后
    import datetime as _dt
    try:
        dl = (d.get('deadline') or '').strip().replace('T', ' ')
        if dl:
            if len(dl) <= 10:
                _dl = _dt.datetime.strptime(dl[:10], '%Y-%m-%d').replace(hour=23, minute=59)
            else:
                _dl = _dt.datetime.strptime(dl[:16], '%Y-%m-%d %H:%M')
            if _dl < _dt.datetime.now():
                _dl = _dt.datetime.now() + _dt.timedelta(minutes=30)  # 防填过去时间
        else:
            _dl = _dt.datetime.now() + _dt.timedelta(days=7)
        deadline = _dl.strftime('%Y-%m-%d %H:%M')
    except Exception:
        deadline = (_dt.datetime.now() + _dt.timedelta(days=7)).strftime('%Y-%m-%d %H:%M')
    conn.execute("INSERT INTO inquiries(inq_no,req_id,title,purpose,status,deadline,created_by) VALUES(?,?,?,?,?,?,?)",
                 (no, req_id, title, pr['purpose'], '询价中', deadline, session['user_name']))
    iid = conn.execute("SELECT id FROM inquiries WHERE inq_no=?", (no,)).fetchone()[0]
    for s in suppliers[:3]:
        name = (s.get('name') or '').strip()
        if not name:
            continue
        token = uuid.uuid4().hex
        conn.execute("INSERT INTO inquiry_suppliers(inquiry_id,supplier_name,contact,phone,token) VALUES(?,?,?,?,?)",
                     (iid, name, s.get('contact', ''), s.get('phone', ''), token))
    conn.commit(); conn.close()
    log(session['user_name'], '发起三方询价', '%s 申请#%s %d家' % (no, req_id, len([x for x in suppliers if (x.get('name') or '').strip()])))
    return jsonify({'success': True, 'inq_no': no, 'id': iid})

def inquiry_eff_price(s, key='quote_price'):
    """V11.180: 取供应商"有效报价" — 采购已议价(adj_price>0)时用调整价, 否则用厂家原始报价
    key: 'quote_price'(含运总价) 或 'quote_details'(行明细)"""
    try:
        if key == 'quote_price':
            adj = float(s.get('adj_price') or 0)
            if adj > 0:
                return adj
            return float(s.get('quote_price') or 0)
        if key == 'quote_details':
            adj = (s.get('adj_details') or '').strip()
            if adj:
                try:
                    _l = json.loads(adj)
                    if isinstance(_l, list) and _l:
                        return _l
                except Exception:
                    pass
            return s.get('quote_details') or '[]'
    except Exception:
        pass
    return s.get(key) or ('[]' if key == 'quote_details' else 0)


def inquiry_is_adjusted(s):
    """V11.180: 该供应商报价是否已被采购调整过(adj_price>0 或 adj_details非空)"""
    try:
        if float(s.get('adj_price') or 0) > 0:
            return True
        if (s.get('adj_details') or '').strip():
            return True
    except Exception:
        pass
    return False


@app.route('/api/inquiries/<int:iid>/adjust', methods=['POST'])
@login_required
def api_inquiry_adjust(iid):
    """V11.180: 采购方议价 — 修改供应商含税单价/含税总价/含运总价 + 采购内部备注
    仅采购员/系统管理员可操作; 厂家原始报价(quote_*)永不覆盖, 留痕可溯"""
    me = db().execute("SELECT * FROM users WHERE id=?", (session['user_id'],)).fetchone()
    if not (me and me['role'] in ('系统管理员', '采购员', '分管领导', '总经理')):
        return jsonify({'error': '仅采购岗位可修改报价'}), 403
    d = request.json or {}
    sid = int(d.get('supplier_id') or 0)
    conn = db()
    s = conn.execute("SELECT * FROM inquiry_suppliers WHERE id=? AND inquiry_id=?", (sid, iid)).fetchone()
    if not s:
        conn.close(); return jsonify({'error': '供应商不存在'}), 404
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    if i['status'] not in ('询价中', '待定标'):
        conn.close(); return jsonify({'error': '当前状态不可议价（仅询价中/待定标可操作）'}), 400
    # V11.205: 开标前禁止议价(报价不可见阶段不允许内部改价)
    if _inq_locked(i['deadline']):
        conn.close(); return jsonify({'error': '报价未开标（截止 %s），开标前不可议价' % (i['deadline'] or '')}), 400
    # 行明细调整(可选): [{unit_price, qty}] 按申请物资顺序
    adj_details = d.get('adj_details')
    adj_price = float(d.get('adj_price') or 0)
    adj_remark = (d.get('adj_remark') or '').strip()[:500]
    if adj_details is not None:
        if not isinstance(adj_details, list):
            conn.close(); return jsonify({'error': '调整明细格式错误'}), 400
        conn.execute("UPDATE inquiry_suppliers SET adj_details=?, adj_price=?, adj_remark=?, is_selected=0 WHERE id=?",
                     (json.dumps(adj_details, ensure_ascii=False), adj_price, adj_remark, sid))
    else:
        conn.execute("UPDATE inquiry_suppliers SET adj_price=?, adj_remark=? WHERE id=?", (adj_price, adj_remark, sid))
    conn.commit(); conn.close()
    log(session['user_name'], '采购议价', f'询价#{iid} {s["supplier_name"]} 调整后含运总价¥{adj_price:.0f} 备注:{adj_remark[:40]}')
    return jsonify({'success': True, 'message': '议价已保存（厂家原始报价保留留痕）', 'adj_price': adj_price})


# V11.205 询价统一开标: 报价截止时间解析/锁定期判断(老数据纯日期按当日23:59计)
def _inq_deadline_dt(dl):
    try:
        dl = (dl or '').strip()
        if not dl:
            return None
        if len(dl) <= 10:
            return datetime.datetime.strptime(dl[:10], '%Y-%m-%d').replace(hour=23, minute=59, second=59)
        return datetime.datetime.strptime(dl[:16].replace('T', ' '), '%Y-%m-%d %H:%M')
    except Exception:
        return None


def _inq_locked(deadline):
    """截止前=锁定期(内部禁看报价/禁定标/禁提交审批/禁导出); 到点自动解锁"""
    _ddt = _inq_deadline_dt(deadline)
    return bool(_ddt) and datetime.datetime.now() < _ddt


@app.route('/api/inquiries/<int:iid>')
@login_required
def api_inquiry_detail(iid):
    """询价单详情: 申请信息 + 物品明细 + 供应商报价对比 + 品牌分析
    V11.205 统一开标: 截止前 locked=True → 报价字段全部脱敏(0/空), 前端显示"报价未开标,暂不可查看";
    截止后自动解锁, 按价格从低到高排序(既有 ORDER BY quote_price 即升序, 未报价排最后)"""
    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (i['req_id'],)).fetchall()
    sups = conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=? ORDER BY (quote_price=0), quote_price, id", (iid,)).fetchall()
    # V11.217: 提前开标 — 全部受邀供应商已报价(每家quote_price>0)则立即解锁(不等截止时间)
    _all_q = bool(sups) and all((s['quote_price'] or 0) > 0 for s in sups)
    conn.close()
    out = dict_row(i)
    out['request'] = dict_row(pr)
    out['items'] = [dict_row(r) for r in items]
    # V11.205: 锁定期状态(截止时间精确到分钟; 纯日期老数据按当天23:59)
    # V11.217: 全部报价完成 → 提前解锁(三家在截止前全部提前报完价 → 提前显示报价详情并可提交)
    out['locked'] = _inq_locked(i['deadline']) and not _all_q
    out['deadline_passed'] = (not out['locked']) and bool((i['deadline'] or '').strip())
    out['all_quoted'] = _all_q
    # 添加品牌分析
    supplier_list = []
    for s in sups:
        sd = dict_row(s)
        if out['locked']:
            # 开标前对内部账号脱敏报价字段(数值清0/文本清空), 不泄露任何报价信息
            for _k in ('quote_price', 'quote_details', 'quote_remark', 'quote_delivery', 'quote_warranty',
                       'quote_brand', 'brand', 'adj_price', 'adj_details', 'adj_remark'):
                if _k in ('quote_price', 'adj_price'):
                    sd[_k] = 0
                else:
                    sd[_k] = ''
            sd['brand_analysis'] = None
            sd['_locked'] = True
        else:
            brand_info = search_brand_info(sd.get('supplier_name', ''), '')
            sd['brand_analysis'] = brand_info
        supplier_list.append(sd)
    out['suppliers'] = supplier_list
    return jsonify(out)

@app.route('/inq/<token>')
def inquiry_vendor_page(token):
    """商家免登录报价页(无需登录, 链接发供应商)
    V11.164: 链接无效/询价结束/作废/截止 全部友好中文提示页, 绝不500/裸报错; 空值容错; 已报价可回显修改"""
    import datetime as _d
    _today_str = _d.date.today().strftime('%Y-%m-%d')

    def _msg_page(_icon, _title, _sub):
        return ('<div style="max-width:520px;margin:80px auto;background:#fff;border-radius:12px;padding:36px;'
                'box-shadow:0 4px 24px rgba(0,0,0,.08);font-family:-apple-system,Segoe UI,Microsoft YaHei,sans-serif;text-align:center">'
                '<div style="font-size:44px;margin-bottom:10px">%s</div>'
                '<h2 style="margin:0 0 8px;color:#333;font-size:18px">%s</h2>'
                '<p style="color:#888;font-size:13px;margin:0;line-height:1.6">%s</p></div>') % (
                    _icon, esc_html(_title), esc_html(_sub))

    conn = db()
    s = conn.execute("SELECT * FROM inquiry_suppliers WHERE token=?", (token,)).fetchone()
    if not s:
        conn.close()
        return _msg_page('🔗', '报价链接无效或已失效', '该链接可能不完整或已过期，请联系采购方获取最新报价链接。')
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (s['inquiry_id'],)).fetchone()
    if not i:
        conn.close()
        return _msg_page('📭', '询价单不存在', '该询价单已被删除，请联系采购方。')
    _st = i['status'] or '询价中'
    if _st != '询价中':
        _st_txt = {'已生成订单': '该询价已完成定标', '定标审批中': '该询价正在定标审批中', '待定标': '该询价等待定标'}.get(_st, '该询价已结束')
        conn.close()
        return _msg_page('⏳', _st_txt, '本批次询价已结束，无法继续报价。感谢参与，欢迎下次合作。')
    _deadline = (i['deadline'] or '').strip()
    # V11.205: 截止精确到分钟 — 到点(含纯日期老数据按当日23:59)后商家页关闭
    _ddt2 = _inq_deadline_dt(_deadline)
    if _ddt2 and datetime.datetime.now() >= _ddt2:
        conn.close()
        return _msg_page('⏰', '报价已截止', '该询价已于 %s 截止，无法继续报价/修改。如有疑问请联系采购方。' % _deadline)
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (i['req_id'],)).fetchall()
    conn.close()
    if not items:
        return _msg_page('📦', '询价单暂无物资明细', '该询价单尚未录入物资明细，请联系采购方确认。')
    # 已报价数据(回显/修改用)
    _prev = {}
    try:
        _qd = json.loads(s['quote_details']) if s['quote_details'] else []
        if isinstance(_qd, list):
            for _x in _qd:
                if isinstance(_x, dict):
                    _prev[len(_prev)] = _x
    except Exception:
        pass
    _already = bool(s['quote_price'] and s['quote_price'] > 0)
    # ---------- 明细行表单(含税单价/总价/交付/质保/品牌/备注, 已报价则回显) ----------
    _rows_html = []
    for idx, it in enumerate(items):
        _pv = _prev.get(idx, {}) or {}
        _v_price = esc_html(_pv.get('unit_price') or '')
        _qty = it['quantity'] or 0
        _ref = ('<span style="color:#bbb;font-size:11px">(参考¥%.0f)</span>' % ((it['total_price'] or 0) / _qty if _qty else 0))
        _rows_html.append(
            '<tr>'
            '<td style="padding:6px 8px;text-align:left;border-bottom:1px solid #eef">%s</td>'
            '<td style="padding:6px 8px;text-align:left;border-bottom:1px solid #eef;color:#888;font-size:12px">%s</td>'
            '<td style="padding:6px 8px;text-align:left;border-bottom:1px solid #eef;white-space:nowrap">%s%s</td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef"><input type="number" min="0" step="0.01" placeholder="单价" '
            'oninput="calc()" data-q="%s" id="up%d" value="%s" style="width:64px;padding:5px 6px;border:1px solid #d0d7e2;border-radius:6px;font-size:13px;text-align:right"></td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef;text-align:right;font-weight:600;color:#2e7d32;white-space:nowrap">¥<span id="ut%d">0.00</span></td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef"><input placeholder="如7天" id="dl%d" value="%s" style="width:52px;padding:5px 6px;border:1px solid #d0d7e2;border-radius:6px;font-size:12px"></td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef"><input placeholder="如3个月" id="wr%d" value="%s" style="width:56px;padding:5px 6px;border:1px solid #d0d7e2;border-radius:6px;font-size:12px"></td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef"><input placeholder="品牌" id="br%d" value="%s" style="width:80px;padding:5px 6px;border:1px solid #d0d7e2;border-radius:6px;font-size:12px"></td>'
            '<td style="padding:6px 8px;border-bottom:1px solid #eef"><input placeholder="备注" id="rm%d" value="%s" style="width:64px;padding:5px 6px;border:1px solid #d0d7e2;border-radius:6px;font-size:12px"></td>'
            '</tr>' % (
                esc_html(it['item_name']), esc_html(it['spec'] or ''),
                str(_qty) + esc_html(it['unit'] or '个'), _ref,
                str(_qty), idx, _v_price,
                idx, idx, esc_html(_pv.get('delivery') or ''),
                idx, esc_html(_pv.get('warranty') or ''),
                idx, esc_html(_pv.get('brand') or ''),
                idx, esc_html(_pv.get('remark') or '')))
    _item_rows = ''.join(_rows_html)
    # ---------- 头部提示 ----------
    _dl_txt = ''
    if _deadline:
        _dl_txt = '<div style="background:#fff3cd;border-radius:8px;padding:10px 14px;font-size:13px;margin-bottom:12px;border:1px solid #ffeeba"><b>⏰ 报价截止：%s</b></div>' % esc_html(_deadline)
    _head_note = ('<p style="color:#2e7d32;font-size:13px;margin:0 0 10px">✅ 贵司已报价，可修改后重新提交（将覆盖原报价；报价金额开标前不对外显示）</p>') if _already else ''
    _ship_val = '' if _already else esc_html(s['quote_price'])  # V11.217: 已报价回显不再带金额(不显示具体价格)
    _remark_val = esc_html(s['quote_remark'] or '')
    body = ('<div style="max-width:860px;margin:40px auto;background:#fff;border-radius:12px;padding:28px;'
            'box-shadow:0 4px 24px rgba(0,0,0,.08);font-family:-apple-system,Segoe UI,Microsoft YaHei,sans-serif">'
            '<h2 style="margin:0 0 4px;color:#1f6feb">📋 采购询价单</h2>'
            '<p style="color:#888;font-size:13px;margin:0 0 10px">尊敬的 %s，请逐项填写含税单价，总价自动计算；交付日期/质保时间按实际填写</p>%s%s'
            '<div style="background:#f5f8ff;border-radius:8px;padding:12px 16px;font-size:13px;margin-bottom:14px">'
            '<b>%s</b><br><span style="color:#888">询价编号：%s</span></div>'
            '<div style="overflow-x:auto"><table style="width:100%%;border-collapse:collapse;font-size:13px;margin-bottom:10px;min-width:700px">'
            '<tr style="background:#f5f8ff"><th style="padding:6px 8px;text-align:left">物资名称</th>'
            '<th style="padding:6px 8px;text-align:left">规格</th><th style="padding:6px 8px;text-align:left">数量</th>'
            '<th style="padding:6px 8px;text-align:left">含税单价(元)<span style="color:#e74c3c">*</span></th><th style="padding:6px 8px;text-align:left">总价（含税含运）</th>'
            '<th style="padding:6px 8px;text-align:left">交付日期</th><th style="padding:6px 8px;text-align:left">质保时间</th>'
            '<th style="padding:6px 8px;text-align:left">品牌</th><th style="padding:6px 8px;text-align:left">厂家备注</th></tr>%s</table></div>'
            '<div style="background:#f0faf0;border-radius:8px;padding:10px 14px;font-size:14px;margin-bottom:12px;display:flex;justify-content:space-between;align-items:center">'
            '<span style="color:#2e7d32"><b>总价（含税含运）合计：¥<span id="total">0.00</span></b></span>'
            '<span style="font-size:12px;color:#888">物品较多时，可<a href="javascript:void(0)" onclick="quickFill()" style="color:#1f6feb">💰 填一个总价自动分摊</a></span></div>'
            '<div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:12px">'
            '<div style="flex:1;min-width:220px;background:#fff8f0;border:1px solid #f5d9b8;border-radius:8px;padding:10px 14px">'
            '<label style="font-size:12px;color:#a05a12;display:block;margin-bottom:4px"><b>🚚 总价（含税含运）——整单含运费，必填</b></label>'
            '<input id="shipTotal" type="number" min="0" step="0.01" placeholder="含运费的总金额，如 6950" value="%s" '
            'style="width:100%%;padding:7px 8px;border:1px solid #d0d7e2;border-radius:6px;font-size:14px;box-sizing:border-box"></div>'
            '<div style="flex:1;min-width:220px;background:#f8f9fb;border:1px solid #e2e7ee;border-radius:8px;padding:10px 14px">'
            '<label style="font-size:12px;color:#555;display:block;margin-bottom:4px"><b>📝 厂家备注（整单说明，选填）</b></label>'
            '<input id="supRemark" placeholder="如：含税含运、交货条件等" value="%s" '
            'style="width:100%%;padding:7px 8px;border:1px solid #d0d7e2;border-radius:6px;font-size:13px;box-sizing:border-box"></div></div>'
            '<button onclick="sub()" style="width:100%%;padding:12px;background:#1f6feb;color:#fff;border:none;border-radius:8px;font-size:15px;cursor:pointer">提交报价</button>'
            '<div id="msg" style="margin-top:10px;font-size:13px;color:#27ae60;text-align:center"></div>'
            '<script>'
            'window.calc=function(){let t=0;document.querySelectorAll("[id^=up]").forEach((e,i)=>{const q=parseFloat(e.getAttribute("data-q"))||1;const p=parseFloat(e.value)||0;'
            'const st=p*q;t+=st;const u=document.getElementById("ut"+i);if(u)u.textContent=st.toFixed(2)});'
            'const _t=document.getElementById("total");if(_t)_t.textContent=t.toFixed(2)};'
            'window.quickFill=function(){const v=prompt("请输入报价总金额(元):");if(!v||isNaN(v))return;const n=document.querySelectorAll("[id^=up]").length;'
            'const per=parseFloat(v)/n;document.querySelectorAll("[id^=up]").forEach(e=>{e.value=per.toFixed(2)});calc();'
            'alert("已按平均分摊到每行，可再逐行微调")};'
            'window.sub=async function(){const rows=document.querySelectorAll("[id^=up]");const details=[];let emptyIdx=[];'
            'rows.forEach((e,i)=>{const p=parseFloat(e.value)||0;if(p<=0)emptyIdx.push(i+1);details.push({unit_price:p,qty:parseFloat(e.getAttribute("data-q"))||1,'
            'delivery:(document.getElementById("dl"+i)||{}).value||"",warranty:(document.getElementById("wr"+i)||{}).value||"",'
            'brand:(document.getElementById("br"+i)||{}).value||"",remark:(document.getElementById("rm"+i)||{}).value||""})});'
            'if(emptyIdx.length){alert("请填写所有物料的含税单价（第"+emptyIdx.join("、")+"行未填）");return}'
            'const shipTotal=parseFloat((document.getElementById("shipTotal")||{}).value)||0;'
            'if(shipTotal<=0){alert("请填写总价（含税含运）——整单含运费的总金额");return}'
            'const supRemark=(document.getElementById("supRemark")||{}).value||"";'
            'const btn=document.querySelector("button[onclick*=sub]");if(btn){btn.disabled=true;btn.style.opacity=.6;btn.textContent="提交中..."}'
            'try{const r=await fetch("%s",{method:"POST",headers:{"Content-Type":"application/json"},'
            'body:JSON.stringify({quote_price:shipTotal,details,quote_delivery:"",quote_warranty:"",quote_remark:supRemark})});'
            'const j=await r.json();if(j.success){document.getElementById("msg").textContent="✅ 报价提交成功";setTimeout(()=>location.reload(),800)}'
            'else{alert(j.error||"提交失败");if(btn){btn.disabled=false;btn.style.opacity=1;btn.textContent="提交报价"}}}'
            'catch(err){alert("网络异常，请重试");if(btn){btn.disabled=false;btn.style.opacity=1;btn.textContent="提交报价"}}};'
            'window.addEventListener("load",function(){try{window.calc()}catch(e){}});'
            '</script></div>') % (
                esc_html(s['supplier_name']), _head_note, _dl_txt,
                esc_html(pr['purpose'] if pr else ''), esc_html(i['inq_no']),
                _item_rows, _ship_val, _remark_val, '/api/inquiry/vendor/%s/quote' % token)
    return body

@app.route('/api/inquiry/vendor/<token>/quote', methods=['POST'])
def inquiry_vendor_quote(token):
    """商家提交报价(免登录)"""
    d = request.json
    details = d.get('details') or []
    # V11.41: 行明细报价时以明细合计为准, 总价可传0; 无明细时走旧逻辑(总价必填)
    price = float(d.get('quote_price') or 0)
    if not details and price <= 0:
        return jsonify({'error': '请填写有效报价金额'}), 400
    conn = db()
    s = conn.execute("SELECT * FROM inquiry_suppliers WHERE token=?", (token,)).fetchone()
    if not s:
        conn.close(); return jsonify({'error': '报价链接无效'}), 404
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (s['inquiry_id'],)).fetchone()
    if not i or i['status'] != '询价中':
        conn.close(); return jsonify({'error': '该询价已结束，无法报价'}), 400
    # V11.24/V11.205: 报价截止时间检查 — 到点后拒绝商家报价/修改(精确到分钟)
    if i['deadline']:
        _ddt3 = _inq_deadline_dt(i['deadline'])
        if _ddt3 and datetime.datetime.now() >= _ddt3:
            conn.close(); return jsonify({'error': '该询价已于 %s 截止，无法继续报价/修改报价' % i['deadline']}), 400
    # V11.41: 行明细报价(每行单价+备注), 合计=Σ单价×数量; 兼容旧版总价提交
    # V11.162: 含运总价=商家填的 quote_price(整单含运费), 明细合计只作参考不再覆盖
    _final_price = price
    if details:
        _sum = sum(float(x.get('unit_price') or 0) * float(x.get('qty') or 1) for x in details)
        # 旧版前端(无shipTotal)明细合计>0且总价为0时, 用明细合计兜底
        if _final_price <= 0 and _sum > 0:
            _final_price = _sum
    # V11.126: 行明细里的 交付日期/质保时间/品牌 汇总去重后存汇总字段
    # (详情页/比价表/订单备注直接显示; 旧版无details时保留原提交值)
    def _uniq_vals(vals):
        out = []
        for v in vals:
            t = str(v or '').strip()
            if t and t not in out:
                out.append(t)
        return '、'.join(out)
    if details:
        _delivery = _uniq_vals(x.get('delivery') for x in details)
        _warranty = _uniq_vals(x.get('warranty') for x in details)
        _brand = _uniq_vals(x.get('brand') for x in details)
    else:
        _delivery = d.get('quote_delivery') or ''
        _warranty = d.get('quote_warranty') or ''
        _brand = d.get('quote_brand') or ''
    conn.execute("UPDATE inquiry_suppliers SET quote_price=?, quote_remark=?, quote_details=?, quote_delivery=?, quote_warranty=?, quote_brand=?, quote_time=? WHERE id=?",
                 (_final_price, (d.get('quote_remark') or '')[:200],
                  json.dumps(details, ensure_ascii=False) if details else '',
                  _delivery[:60], _warranty[:60], _brand[:100], now(), s['id']))
    # V11.133: 取消"三家报价自动提交审批" — 改为采购员在系统里手动点"提交审批"
    # (用户要求: 商家全部报价后不自动发起审批, 由人工确认后再提交定标)
    conn.commit(); conn.close()
    return jsonify({'success': True, 'quote_price': _final_price})

@app.route('/api/inquiries/<int:iid>/submit', methods=['POST'])
@login_required
def api_inquiry_submit(iid):
    """V11.93: 手动提交询价审批（所有供应商报价后）"""
    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close()
        return jsonify({'error': '询价单不存在'}), 404
    if i['status'] != '询价中':
        conn.close()
        return jsonify({'error': '该询价已结束'}), 400
    # V11.205: 统一开标 — 截止前禁止提交定标审批(否则审批详情/钉钉会泄露报价)
    # V11.217: 例外 — 全部受邀供应商提前报完价 → 允许提前开标提交(不等截止时间)
    sups0 = conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=?", (iid,)).fetchall()
    _all_q = bool(sups0) and all((s['quote_price'] or 0) > 0 for s in sups0)
    if _inq_locked(i['deadline']) and not _all_q:
        conn.close(); return jsonify({'error': '报价未开标（截止 %s，当前尚未全部报价），请等全部供应商报完价或到截止时间后再提交定标审批' % (i['deadline'] or '')}), 400
    # V11.155f: 防重复提交 — 已有审批中记录(定标审批中)则拒绝
    _pend = conn.execute("SELECT 1 FROM inquiry_approvals WHERE inquiry_id=? AND status='审批中' LIMIT 1", (iid,)).fetchone()
    if _pend:
        conn.close()
        return jsonify({'error': '该询价已提交定标审批，请勿重复提交'}), 400
    # 检查是否所有供应商都报价
    sups = conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=?", (iid,)).fetchall()
    if not sups:
        conn.close()
        return jsonify({'error': '无供应商报价'}), 400
    quoted = [s for s in sups if s['quote_price'] and s['quote_price'] > 0]
    if len(quoted) < 2:
        conn.close()
        return jsonify({'error': f'需至少2家报价，当前{len(quoted)}家'}), 400
    # 创建采购订单草稿
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (i['req_id'],)).fetchall()
    if not pr or not items:
        conn.close()
        return jsonify({'error': '来源申请或明细缺失'}), 400
    # 取最低报价(V11.180: 按采购调整后有效价取最低, 无调整用厂家原始价)
    cheapest = conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=? AND quote_price>0 ORDER BY quote_price ASC LIMIT 1", (iid,)).fetchone()
    cheapest = dict(cheapest) if cheapest else None
    if cheapest:
        cheapest['_eff'] = inquiry_eff_price(cheapest, 'quote_price')
    total = float(cheapest['_eff'] if cheapest and cheapest.get('_eff') else (cheapest['quote_price'] if cheapest else 0))
    remark = '询价单:%s 商家已报价完成，最低报价¥%.0f(%s)，请领导定标' % (i['inq_no'], total, cheapest['supplier_name'] if cheapest else '待定')
    # V11.126: 已存在草稿订单(如驳回后重新提交)则复用更新, 不重复建单
    _draft = conn.execute("SELECT id FROM purchase_orders WHERE inquiry_id=? AND status='草稿' ORDER BY id LIMIT 1", (iid,)).fetchone()
    if _draft:
        conn.execute("UPDATE purchase_orders SET supplier=?, total_amount=?, remark=? WHERE id=?",
                     (cheapest['supplier_name'] if cheapest else '待定', total, remark, _draft['id']))
    else:
        conn.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
            supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments,status,inquiry_id) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (gen_no('CG', 'purchase_orders', 'order_no', conn), i['req_id'], i['title'][:50], '', 1, '个', 0, total, 0, 0, total,
             cheapest['supplier_name'] if cheapest else '待定', i['created_by'], '后勤类', i['created_by'], 1, (i['deadline'] or '')[:10], '货到付款',
             remark, 0, json.dumps([], ensure_ascii=False), '草稿', i['id']))
    # 创建询价审批记录 (biz_id 统一=询价单id)
    conn.execute("INSERT INTO inquiry_approvals(inquiry_id, status, created_at) VALUES(?, '审批中', ?)", (iid, now()))
    conn.execute("INSERT INTO approval_instances(biz_type, biz_id, level_no, role, approver, status) VALUES(?, ?, 1, '分管领导', 'xingguo', 'pending')", ('inquiry_approval', iid))
    conn.execute("UPDATE inquiries SET status='定标审批中', updated_at=? WHERE id=?", (now(), iid))
    conn.commit()
    conn.close()
    # V11.126: 提交后立即发起钉钉(及飞书)审批实例 — 修复"提交后钉钉没反应"
    try:
        start_instances('inquiry_approval', iid)
    except Exception as e:
        log(session['user_name'], '询价审批发起异常', str(e))
    log(session['user_name'], '提交询价审批', '%s 已提交定标审批' % i['inq_no'])
    return jsonify({'success': True, 'inq_no': i['inq_no']})

@app.route('/api/inquiries/<int:iid>/select', methods=['POST'])
@login_required
def api_inquiry_select(iid):
    """定标审批: 领导选定供应商 → 生成采购订单草稿 → 提交审批"""
    d = request.json
    sid = d.get('supplier_id')
    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    if i['status'] != '询价中':
        conn.close(); return jsonify({'error': '该询价已结束'}), 400
    # V11.24/V11.205: 统一开标 — 报价截止前禁止提前定标(防人为泄露/串通); 到点开标后才允许选中下单
    if i['deadline']:
        _ddt4 = _inq_deadline_dt(i['deadline'])
        if _ddt4 and datetime.datetime.now() < _ddt4:
            conn.close(); return jsonify({'error': '报价尚未开标（截止 %s），请等待统一开标后再定标选择供应商' % i['deadline']}), 400
    s = conn.execute("SELECT * FROM inquiry_suppliers WHERE id=? AND inquiry_id=?", (sid, iid)).fetchone()
    if not s:
        conn.close(); return jsonify({'error': '供应商不在该询价单中'}), 400
    if not s['quote_price'] or s['quote_price'] <= 0:
        conn.close(); return jsonify({'error': '该供应商尚未报价'}), 400
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (i['req_id'],)).fetchall()
    if not pr or not items:
        conn.close(); return jsonify({'error': '来源申请或明细缺失'}), 400
    # 生成采购订单: 供应商=选中家; V11.170: 明细价格优先用商家报价(quote_details按全量物资顺序存unit_price),
    # 不再按申请参考金额分摊(参考金额常为0, 导致首项吃全额/后项变0, 合同丢明细)
    # V11.180: 采购已议价(adj_*)时用调整后报价生成订单, 原始报价留痕
    no = gen_no('CG', 'purchase_orders', 'order_no', conn)
    s_dict = dict(s)
    total = float(inquiry_eff_price(s_dict, 'quote_price'))
    _qd_map = {}
    try:
        _qd_list = json.loads(inquiry_eff_price(s_dict, 'quote_details'))
        for _qi, _q in enumerate(_qd_list):
            _qd_map[_qi] = _q
    except Exception:
        _qd_map = {}
    rows = []
    grand_amt = 0.0
    for idx, it in enumerate(items):
        qty = float(it['quantity'] or 1)
        _q = _qd_map.get(idx, {})
        price = float(_q.get('unit_price') or 0) or 0
        amt = round(price * qty, 2)
        grand_amt += amt
        rows.append((it['item_name'], it['spec'] or '', it['unit'] or '个', qty, price, amt))
    # 兜底: 商家未填单价(旧数据) → 回退按申请参考金额比例分摊报价总额
    if grand_amt <= 0:
        rows = []
        base_sum = sum(float(it['total_price'] or 0) for it in items)
        grand_amt = 0.0
        for idx, it in enumerate(items):
            qty = float(it['quantity'] or 1)
            if base_sum > 0 and idx < len(items) - 1:
                amt = total * (float(it['total_price'] or 0) / base_sum)
            else:
                amt = total - grand_amt  # 最后一行吃余数, 保证合计=报价
            amt = round(amt, 2)
            price = round(amt / qty, 2) if qty else 0
            grand_amt += amt
            rows.append((it['item_name'], it['spec'] or '', it['unit'] or '个', qty, price, amt))
    first = rows[0]
    # 商家详细信息自动填入订单
    contact = (s['contact'] or '').strip()
    phone = (s['phone'] or '').strip()
    quote_remark = (s['quote_remark'] or '').strip()
    quote_delivery = (s['quote_delivery'] or '').strip()
    quote_warranty = (s['quote_warranty'] or '').strip()
    detail_parts = ['三方询价选中: %s 报价¥%.2f' % (s['supplier_name'], total)]
    if contact:
        detail_parts.append('联系人: %s' % contact)
    if phone:
        detail_parts.append('电话: %s' % phone)
    if quote_delivery:
        detail_parts.append('交付日期: %s' % quote_delivery)
    if quote_warranty:
        detail_parts.append('质保时间: %s' % quote_warranty)
    if quote_remark:
        detail_parts.append('备注: %s' % quote_remark)
    # V11.180: 采购内部备注同步订单(给领导看议价情况)
    _adj_remark = (s_dict.get('adj_remark') or '').strip()
    if _adj_remark:
        detail_parts.append('采购议价备注: %s' % _adj_remark)
    detail_parts.append('询价单号: %s' % i['inq_no'])
    remark = '; '.join(detail_parts)
    tm = '货到付款'
    if quote_remark:
        for kw in ('货到付款', '先款后货', '月结', '预付', '现款', '承兑'):
            if kw in quote_remark:
                tm = kw
                break
    # 定标审批: 领导选定后, 订单草稿 + 提交定标审批(必须领导审批通过才能下单)
    settle_type = d.get('settle_type') or '现结'
    conn.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
        supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments,status) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (no, i['req_id'], first[0], first[1], sum(r[3] for r in rows), first[2], first[4], grand_amt, 0, 0, total,
         s['supplier_name'], pr['requester'] or '', '后勤类', session['user_name'], session['user_id'],
         pr['target_date'] or '', tm, remark, 0,
         json.dumps([], ensure_ascii=False), '草稿'))
    oid = conn.execute("SELECT id FROM purchase_orders WHERE order_no=?", (no,)).fetchone()[0]
    for r in rows:
        conn.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                     (oid, r[0], r[1], r[2], r[3], r[4], r[5], 0, 0, r[5], ''))
    conn.execute("UPDATE inquiry_suppliers SET is_selected=1 WHERE id=?", (sid,))
    # V11.73: 定标审批通过后直接生效(领导已选定供应商,无需再次审批)
    conn.execute("UPDATE purchase_orders SET status='已通过', settle_type=?, updated_at=? WHERE id=?", (settle_type, now(), oid))
    conn.execute("UPDATE inquiries SET status='已生成订单', selected_supplier_id=?, updated_at=? WHERE id=?", (sid, now(), iid))
    conn.commit()
    log(session['user_name'], '询价定标', '%s → 订单%s(供应商:%s ¥%.0f,已生效)' % (i['inq_no'], no, s['supplier_name'], total))
    return jsonify({'success': True, 'order_no': no, 'id': oid, 'total_amount': total, 'status': '已通过',
                    'message': '✅ 定标通过，订单已生效'})


@app.route('/api/inquiries/<int:iid>/split-select', methods=['POST'])
@login_required
def api_inquiry_split_select(iid):
    """V11.145: 分项定标 — 每个物资可指定不同供应商, 按供应商分组生成多个订单
    请求: {items: [{item_id, supplier_id}, ...]}  item_id=request_items.id, supplier_id=inquiry_suppliers.id
    每项取该供应商对该物资的报价; 生成N个订单(按供应商分组)"""
    d = request.json or {}
    picks = d.get('items') or []
    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    if not picks:
        conn.close(); return jsonify({'error': '未选择任何物资的供应商'}), 400
    # V11.155: 分项定标仅在 询价中(采购员直接定标) 或 待定标(领导同意按最低价) 时允许
    # 已生成订单/定标审批中 禁止再分项(防重复生成订单)
    if i['status'] not in ('询价中', '待定标'):
        conn.close(); return jsonify({'error': '当前状态不可分项定标（仅询价中/待定标可操作）'}), 400
    # V11.205: 统一开标 — 截止前禁止分项定标(报价不可见阶段)
    if _inq_locked(i['deadline']):
        conn.close(); return jsonify({'error': '报价未开标（截止 %s），请等待统一开标后再分项定标' % (i['deadline'] or '')}), 400
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    if not pr:
        conn.close(); return jsonify({'error': '来源申请缺失'}), 400
    sups = {r['id']: dict(r) for r in conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=?", (iid,)).fetchall()}
    items = {r['id']: dict(r) for r in conn.execute("SELECT * FROM request_items WHERE req_id=?", (i['req_id'],)).fetchall()}
    # 校验每项选择合法
    valid = []
    for p in picks:
        it = items.get(p.get('item_id'))
        s = sups.get(p.get('supplier_id'))
        if not it or not s:
            conn.close(); return jsonify({'error': '选择项无效'}), 400
        if not s.get('quote_price') or s['quote_price'] <= 0:
            conn.close(); return jsonify({'error': '供应商%s尚未报价' % s['supplier_name']}), 400
        valid.append((it, s))
    if len(valid) != len(items):
        conn.close(); return jsonify({'error': '必须为每个物资选择供应商'}), 400
    # 按供应商分组
    groups = {}
    for it, s in valid:
        groups.setdefault(s['id'], {'sup': s, 'items': []})['items'].append(it)
    created = []
    # V11.158: 全量物资顺序映射(quote_details按全量顺序存, 取价不能按组内序号)
    all_items_list = conn.execute("SELECT * FROM request_items WHERE req_id=? ORDER BY id", (i['req_id'],)).fetchall()
    item_pos = {it['id']: pos for pos, it in enumerate(all_items_list)}
    for sid, g in groups.items():
        s = g['sup']
        its = g['items']
        no = gen_no('CG', 'purchase_orders', 'order_no', conn)
        # 该供应商的逐项报价(quote_details, 按全量物资顺序) — V11.180: 采购已议价时用调整后明细
        qd = {}
        try:
            qd_list = json.loads(inquiry_eff_price(dict(s), 'quote_details'))
            for idx, q in enumerate(qd_list):
                qd[idx] = q
        except Exception:
            qd = {}
        total = 0.0
        rows = []
        for it in its:
            qty = float(it['quantity'] or 1)
            # V11.158: 用物资在全量列表中的位置取对应报价(修复错位导致金额错误)
            pos = item_pos.get(it['id'], 0)
            q = qd.get(pos, {})
            price = float(q.get('unit_price') or 0) or 0
            amt = round(price * qty, 2)
            total += amt
            rows.append((it['item_name'], it['spec'] or '', it['unit'] or '个', qty, price, amt,
                         (q.get('brand') or ''), (q.get('delivery') or ''), (q.get('warranty') or ''), (q.get('remark') or '')))
        if total <= 0:
            total = float(inquiry_eff_price(dict(s), 'quote_price')) or float(s['quote_price'] or 0)
        first = rows[0]
        detail_parts = ['三方询价分项定标: %s' % s['supplier_name']]
        for r in rows:
            detail_parts.append('%s x%s ¥%.2f' % (r[0], r[3], r[5]))
        # V11.180: 采购内部备注同步订单
        _adj_rm = (s.get('adj_remark') or '').strip() if 'adj_remark' in s.keys() else ''
        if _adj_rm:
            detail_parts.append('采购议价备注: %s' % _adj_rm)
        detail_parts.append('询价单号: %s' % i['inq_no'])
        remark = '; '.join(detail_parts)
        conn.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
            supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments,status) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (no, i['req_id'], first[0], first[1], sum(r[3] for r in rows), first[2], first[4], total, 0, 0, total,
             s['supplier_name'], pr['requester'] or '', '后勤类', session['user_name'], session['user_id'],
             pr['target_date'] or '', '货到付款', remark, 0,
             json.dumps([], ensure_ascii=False), '已通过'))
        oid = conn.execute("SELECT id FROM purchase_orders WHERE order_no=?", (no,)).fetchone()[0]
        for r in rows:
            conn.execute("""INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark)
                VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                (oid, r[0], r[1], r[2], r[3], r[4], r[5], 0, 0, r[5],
                 '品牌:%s 交付:%s 质保:%s %s' % (r[6], r[7], r[8], r[9]) if (r[6] or r[7] or r[8]) else ''))
        conn.execute("UPDATE inquiry_suppliers SET is_selected=1 WHERE id=?", (sid,))
        # V11.158c: 分项定标生成的订单(货到付款)自动生成待入库单 → 入库验收模块同步显示
        try:
            _rno = gen_no('RK', 'receivings', 'receive_no', conn)
            _rqty = sum(r[3] for r in rows)
            _rjson = json.dumps(
                [{'item_name': r[0], 'spec': r[1] or '', 'quantity': r[3], 'unit': r[2] or '个', 'price': r[4] or 0} for r in rows],
                ensure_ascii=False)
            _rname = (rows[0][0] + ' 等%d项' % len(rows)) if len(rows) > 1 else rows[0][0]
            conn.execute("INSERT INTO receivings(receive_no,delivery_id,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,dept,items_json,is_est) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,0)",
                (_rno, None, oid, _rname, '', _rqty, rows[0][2] or '个', 0, '待入库', now(),
                 '分项定标: 订单%s自动进入入库板块(整批%d项)' % (no, len(rows)), pr['dept'] or '', _rjson))
        except Exception as _re:
            log('系统', '分项定标生成入库单失败', f'order{oid}: {str(_re)[:80]}')
        created.append({'order_no': no, 'supplier': s['supplier_name'], 'total': total})
    conn.execute("UPDATE inquiries SET status='已生成订单', updated_at=? WHERE id=?", (now(), iid))
    conn.commit()
    log(session['user_name'], '分项定标', '%s → 生成%d个订单' % (i['inq_no'], len(created)))
    conn.close()
    return jsonify({'success': True, 'orders': created, 'count': len(created),
                    'message': '✅ 分项定标完成，已按%d家供应商生成%d个订单' % (len(created), len(created))})

def gen_inquiry_xlsx_file(iid):
    """V11.135: 生成询价比价单Excel文件存uploads, 返回文件路径(导出接口/钉钉附件共用)
    复用 api_inquiry_export 的生成逻辑; 失败返回 None"""
    try:
        import io as _io
        # 在请求上下文中调导出逻辑(需登录态, 附件场景由发起审批的登录用户触发)
        with app.test_request_context():
            if 'user_name' not in session:
                session['user_name'] = '系统'
                session['user_role'] = '系统管理员'
                session['user_id'] = 1
            resp = api_inquiry_export(iid)
            data = resp.get_data()
        if not data:
            return None
        fn = '询价单_%s.xlsx' % iid
        path = os.path.join(BASE, 'uploads', fn)
        with open(path, 'wb') as f:
            f.write(data)
        return fn
    except Exception as e:
        log('系统', '询价比价单生成失败', f'inquiry#{iid}: {str(e)[:120]}')
        return None


@app.route('/api/inquiries/<int:iid>/extend', methods=['POST'])
@login_required
def api_inquiry_extend(iid):
    """V11.205: 延长报价截止时间 — 仅采购主管级(系统管理员/分管领导/总经理)可操作, 记录操作日志"""
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理'):
        return jsonify({'error': '仅采购主管及以上(分管领导/总经理/系统管理员)可延长报价截止时间'}), 403
    d = request.json or {}
    ndl = (d.get('deadline') or '').strip().replace('T', ' ')
    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    if i['status'] != '询价中':
        conn.close(); return jsonify({'error': '当前状态(%s)不可延长截止时间' % i['status']}), 400
    _ddt = _inq_deadline_dt(ndl)
    if _ddt is None:
        conn.close(); return jsonify({'error': '截止时间格式不正确（请用 YYYY-MM-DD HH:MM）'}), 400
    if _ddt < datetime.datetime.now():
        conn.close(); return jsonify({'error': '新的截止时间需晚于当前时间'}), 400
    _old = (i['deadline'] or '')
    _new = _ddt.strftime('%Y-%m-%d %H:%M')
    conn.execute("UPDATE inquiries SET deadline=?, updated_at=? WHERE id=?", (_new, now(), iid))
    conn.commit(); conn.close()
    log(session['user_name'], '延长询价截止', '%s 报价截止 %s → %s' % (i['inq_no'], _old, _new))
    return jsonify({'success': True, 'inq_no': i['inq_no'], 'deadline': _new})


@app.route('/api/inquiries/<int:iid>/export')
@login_required
def api_inquiry_export(iid):
    """V11.1 询价单导出Excel: 基础信息+物料明细+全量供应商报价+比价统计+选中高亮+决策备注"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    conn = db()
    i = conn.execute("SELECT * FROM inquiries WHERE id=?", (iid,)).fetchone()
    if not i:
        conn.close(); return jsonify({'error': '询价单不存在'}), 404
    # V11.205: 统一开标 — 截止前禁止导出比价单(报价内容不外泄)
    if _inq_locked(i['deadline']):
        conn.close(); return jsonify({'error': '报价未开标（截止 %s），开标后方可导出比价单' % (i['deadline'] or '')}), 400
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (i['req_id'],)).fetchone()
    items = conn.execute("SELECT * FROM request_items WHERE req_id=? ORDER BY id", (i['req_id'],)).fetchall()
    sups = [dict(s) for s in conn.execute("SELECT * FROM inquiry_suppliers WHERE inquiry_id=? ORDER BY id", (iid,)).fetchall()]
    conn.close()

    wb = Workbook(); ws = wb.active; ws.title = '询价单'
    # V11.132: 列数提前算(4+每家7列 — V11.180: 每家加"厂家原始报价"留痕列), 标题/章节/备注框全部按全宽合并
    n_sup = len(sups)
    col_count = 4 + n_sup * 7
    # V11.132: 横向A4+缩放, 否则16列挤在纵向A4上必乱
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9  # A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    thin = Side(style='thin', color='B0B0B0')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    title_font = Font(name='微软雅黑', size=14, bold=True)
    head_font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
    head_fill = PatternFill('solid', fgColor='2F5597')
    label_font = Font(name='微软雅黑', size=10, bold=True)
    base_font = Font(name='微软雅黑', size=10)
    note_font = Font(name='微软雅黑', size=10, italic=True, color='666666')
    sel_fill = PatternFill('solid', fgColor='FFF2CC')       # 选中行浅黄
    min_font = Font(name='微软雅黑', size=10, bold=True, color='C00000')  # 最低价标红加粗
    wrap = Alignment(vertical='center', wrap_text=True)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    ws['A1'] = '采 购 询 价 单'
    ws['A1'].font = title_font; ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30

    info = [
        ('询价单号', i['inq_no'] or '', '询价主题', (i['title'] or i['purpose'] or '')[:60]),
        ('申请单号', pr['req_no'] if pr else '', '采购事由', (pr['purpose'] if pr else '') or (i['purpose'] or '')),
        ('发起部门', (pr['dept'] if pr else '') or '', '发起人', i['created_by'] or ''),
        ('询价发起时间', i['created_at'] or '', '询价状态', i['status'] or ''),
        ('物料/服务项数', str(len(items)) + ' 项', '供应商家数', str(len(sups)) + ' 家'),
    ]
    row = 3
    # V11.131: 章节编号补全 — 信息区加"一、"标题, 与"二、比价表""三、决策备注"连续
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_count)
    ws.cell(row, 1, '一、询价基本信息').font = head_font
    for col in range(1, col_count + 1):
        ws.cell(row, col).fill = head_fill; ws.cell(row, col).border = border
    row += 1
    for left_k, left_v, right_k, right_v in info:
        ws.cell(row, 1, left_k).font = label_font
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=3)
        c = ws.cell(row, 2, left_v); c.font = base_font; c.alignment = wrap
        ws.cell(row, 4, right_k).font = label_font
        ws.merge_cells(start_row=row, start_column=5, end_row=row, end_column=col_count)
        c = ws.cell(row, 5, right_v); c.font = base_font; c.alignment = wrap
        for col in range(1, col_count + 1):
            ws.cell(row, col).border = border
        row += 1

    row += 2
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_count)
    ws.cell(row, 1, '二、供应商报价比价表（逐项对比）').font = head_font
    for col in range(1, col_count + 1):
        ws.cell(row, col).fill = head_fill; ws.cell(row, col).border = border
    row += 1
    # V11.52: 逐行三家对比 — 每家一列组(单价/总价/品牌/交付/质保/备注), 备注独立列跟商家走
    # V11.126: 增加 交付日期/质保时间 逐项对比列(商家按行报价的交付/质保直接进Excel)
    # V11.188: 每家厂家一组列用独立浅色区分(厂家A蓝/B绿/C黄/D紫/E橙...), 一眼看清各家报价; 最低价红底覆盖仍醒目
    _PALETTE = ['D6E4F0', 'D8EAD3', 'FFF2CC', 'E4DFEC', 'FDE9D9', 'D5E8D4', 'FCE4D6', 'DDEBF7']
    def _sup_fill(si):
        return PatternFill('solid', fgColor=_PALETTE[(si or 0) % len(_PALETTE)])
    quoted = [s for s in sups if s['quote_price'] and s['quote_price'] > 0]
    # 解析每家行明细(V11.180: 主明细=采购调整后; 另存厂家原始明细留痕)
    sup_details = []
    sup_ori_details = []
    for s in sups:
        try:
            sup_details.append(json.loads(inquiry_eff_price(s, 'quote_details')) or None)
        except Exception:
            sup_details.append(None)
        try:
            sup_ori_details.append(json.loads(s.get('quote_details') or '[]') or None)
        except Exception:
            sup_ori_details.append(None)
    n_sup = len(sups)
    # 表头: 序号/物料/数量/规格 | 每家7列(调整后含税单价/调整后含税总价/品牌/交付/质保/厂家备注/厂家原始报价)
    sup_head = ['序号', '物料名称', '数量', '规格型号']
    for s in sups:
        _tag_adj = '（调整后）' if inquiry_is_adjusted(s) else ''
        sup_head += [f"{s['supplier_name']} 含税单价{_tag_adj}", f"{s['supplier_name']} 总价（含税含运）{_tag_adj}", f"{s['supplier_name']} 品牌",
                     f"{s['supplier_name']} 交付", f"{s['supplier_name']} 质保", f"{s['supplier_name']} 厂家备注",
                     f"{s['supplier_name']} 厂家原始报价"]
    # col_count 已在前面按 4+每家7列 算好
    for ci, h in enumerate(sup_head, 1):
        c = ws.cell(row, ci, h); c.border = border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        # V11.188: 每家厂家的表头列组涂各自浅色(厂家名行首列), 视觉上一家一块; 文字用深色(浅底白字看不清)
        if ci > 4:
            _si = (ci - 5) // 7
            if 0 <= _si < n_sup:
                c.fill = _sup_fill(_si)
                c.font = Font(name='微软雅黑', size=10, bold=True, color='1F3864')
            else:
                c.fill = head_fill; c.font = head_font
        else:
            c.fill = head_fill; c.font = head_font
    ws.row_dimensions[row].height = 30
    row += 1
    # 每物料一行
    total_per_sup = [0.0] * n_sup
    min_font_s = Font(name='微软雅黑', size=10, bold=True, color='C00000')
    for idx, it in enumerate(items, 1):
        qty = float(it['quantity'] or 0)
        vals = [idx, it['item_name'] or '', ('%g' % qty) + (it['unit'] or '个'), it['spec'] or '']
        # 每家: 单价/总价/备注(行明细优先, 无则空)
        row_prices = []  # (unit_price, total)
        for si, s in enumerate(sups):
            # V11.77: 自动搜索品牌信息并附到备注
            brand_info = search_brand_info(s['supplier_name'], '')
            if brand_info:
                s_brand_remark = '%s 品牌分析: 优点-%s | 缺点-%s' % (s['supplier_name'], brand_info.get('优点',''), brand_info.get('缺点',''))
            else:
                s_brand_remark = ''
            det = sup_details[si]
            ori_det = sup_ori_details[si]
            unit_p = None; total_p = None; remark = ''
            ori_unit_p = None; ori_total_p = None
            if det and idx - 1 < len(det):
                d_i = det[idx - 1]
                if d_i.get('unit_price') is not None:
                    unit_p = float(d_i.get('unit_price') or 0)
                    total_p = unit_p * qty
                remark = d_i.get('remark') or ''
            if ori_det and idx - 1 < len(ori_det):
                o_i = ori_det[idx - 1]
                if o_i.get('unit_price') is not None:
                    ori_unit_p = float(o_i.get('unit_price') or 0)
                    ori_total_p = ori_unit_p * qty
            # 获取品牌/交付/质保
            brand = ''; delivery = ''; warranty = ''
            if det and idx - 1 < len(det):
                brand = det[idx - 1].get('brand') or ''
                delivery = det[idx - 1].get('delivery') or ''
                warranty = det[idx - 1].get('warranty') or ''
            if unit_p is not None:
                # V11.132: 金额千分位+两位小数, 领导看更专业
                # V11.180: 末列=厂家原始报价(留痕); 未议价时原始价=调整价且标注"—"
                _ori_txt = '{:,.2f}'.format(ori_total_p) if (ori_total_p is not None and inquiry_is_adjusted(s)) else ('{:,.2f}'.format(total_p) if ori_total_p is None else '{:,.2f}'.format(ori_total_p))
                vals += ['{:,.2f}'.format(unit_p), '{:,.2f}'.format(total_p),
                         brand, delivery, warranty, remark, _ori_txt]
                total_per_sup[si] += total_p
            else:
                vals += ['', '', brand, delivery, warranty, remark, '']
            row_prices.append((unit_p, total_p))
        # 最低单价标红
        unit_prices = [p[0] for p in row_prices if p[0] is not None]
        min_unit = min(unit_prices) if unit_prices else None
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v); c.border = border
            c.alignment = Alignment(horizontal='center' if ci <= 4 or (ci - 5) % 7 != 0 else 'left', vertical='center', wrap_text=True)
            # V11.188: 数据行厂家列组浅色底(最低价标红时红底覆盖)
            if ci > 4:
                _si = (ci - 5) // 7
                if 0 <= _si < n_sup:
                    c.fill = _sup_fill(_si)
            # 单价列: 最低标红加粗+★（V11.145: 领导一眼看到每项最便宜的厂家）— V11.180: 每组7列
            if unit_prices and min_unit is not None:
                k = (ci - 5) // 7
                if 0 <= k < n_sup and (ci - 5) % 7 == 0:
                    if row_prices[k][0] is not None and abs(row_prices[k][0] - min_unit) < 0.001:
                        c.font = Font(name='微软雅黑', size=11, bold=True, color='C00000')
                        c.fill = PatternFill('solid', fgColor='FFEB9C')  # V11.188: 最低价深黄高亮(区别于厂家浅色组)
                        if c.value:
                            c.value = '★' + str(c.value)
                        continue
            c.font = base_font
        ws.row_dimensions[row].height = 26
        row += 1
    # 合计行
    ws.cell(row, 1, '').border = border
    ws.cell(row, 2, '合计').font = label_font; ws.cell(row, 2).border = border
    for cc in (3, 4):
        ws.cell(row, cc).border = border
    total_min = None
    for si, t in enumerate(total_per_sup):
        # V11.126: 总价落位修正到 总价列(7+si*7) — V11.180: 每组7列; V11.188: 合计行厂家组同色
        ws.cell(row, 6 + si * 7, '').border = border
        c = ws.cell(row, 7 + si * 7, '{:,.2f}'.format(round(t, 2)))
        c.border = border; c.font = label_font
        c.alignment = Alignment(horizontal='center', vertical='center')
        c.fill = _sup_fill(si)
        if t > 0 and (total_min is None or t < total_min):
            total_min = t
        for cc in range(8 + si * 7, 12 + si * 7):
            ws.cell(row, cc).border = border
            ws.cell(row, cc).fill = _sup_fill(si)
    for si, t in enumerate(total_per_sup):
        if t > 0 and total_min is not None and abs(t - total_min) < 0.001:
            _min_c = ws.cell(row, 7 + si * 7)
            _min_c.font = min_font_s
            _min_c.fill = PatternFill('solid', fgColor='FFEB9C')  # V11.188: 最低总价深黄底(区别于厂家浅色)
    # V11.132: ★说明移到合计行最后列, 最低总价已标红黄底, 领导一眼看到最便宜
    ws.cell(row, col_count, '★=该项最低价').font = note_font
    ws.cell(row, col_count).border = border
    ws.row_dimensions[row].height = 22
    row += 1
    if not sups:
        for ci in range(1, col_count + 1):
            ws.cell(row, ci, '（暂无供应商）').border = border
        row += 1

    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_count)
    ws.cell(row, 1, '三、采购决策备注').font = head_font
    for col in range(1, col_count + 1):
        ws.cell(row, col).fill = head_fill; ws.cell(row, col).border = border
    row += 1
    # 备注区不预先合并, 由下方按内容行数统一合并(避免重复merge截断文本)
    # 生成品牌分析（V11.131: 拆分多品牌, 每家一行优缺点, 不依赖选中状态）
    brand_analysis_lines = []
    for s in sups:
        if s['quote_price'] and s['quote_price'] > 0:
            # 品牌可能是多个(如"得力、晨光"), 拆分逐个分析
            brand_raw = s.get('quote_brand') or s.get('brand') or ''
            brands = [b.strip() for b in re.split(r'[、,，;；/]', brand_raw) if b.strip()]
            parts = []
            for bname in brands:
                b_info = search_brand_info(bname, '')
                if b_info and (b_info.get('优点') or b_info.get('缺点')):
                    parts.append('%s(优点:%s|缺点:%s)' % (bname, b_info.get('优点', ''), b_info.get('缺点', '')))
                else:
                    parts.append('%s' % bname)
            if parts:
                brand_analysis_lines.append('%s报价¥%s：%s' % (s['supplier_name'],
                    format(float(s['quote_price'] or 0), ',.2f'), '；'.join(parts)))
            else:
                brand_analysis_lines.append('%s报价¥%s：未填写品牌' % (s['supplier_name'],
                    format(float(s['quote_price'] or 0), ',.2f')))
    
    decision = ('本批次采购已收到多家供应商报价，正在定标审批中，待领导确认后确定合作方。')
    selected = next((s for s in sups if s['is_selected']), None)
    if selected:
        decision = ('经多方询价比价，最终选择「%s」为合作方，报价¥%s，'
                    '性价比最优。%s' % (selected['supplier_name'],
                    format(float(selected['quote_price'] or 0), ',.2f'),
                    selected['quote_remark'] or '交货期及付款条件按合同约定'))
    # V11.131: 品牌对比优缺点不依赖"已选中" — 只要有报价的商家就显示在决策备注
    if brand_analysis_lines:
        brand_text = '【品牌对比】' + chr(10) + chr(10).join(brand_analysis_lines)
        decision = decision + chr(10) + chr(10) + brand_text
    # V11.162: 每家含运总价 + 厂家备注(整单) 汇总显示
    # V11.180: 已议价厂家显示调整后总价+原始价+采购内部备注
    _ship_lines = []
    for s in sups:
        if s['quote_price'] and s['quote_price'] > 0:
            _eff = inquiry_eff_price(s, 'quote_price')
            _adj_flag = '【采购调整后】' if inquiry_is_adjusted(s) else ''
            _line = '%s：总价（含税含运） ¥%s%s' % (s['supplier_name'],
                format(float(_eff), ',.2f'), _adj_flag)
            if inquiry_is_adjusted(s) and float(s['quote_price']) > 0:
                _line += '（厂家原始报价 ¥%s）' % format(float(s['quote_price']), ',.2f')
            if (s.get('quote_remark') or '').strip():
                _line += '（厂家备注：%s）' % (s['quote_remark'] or '')
            if (s.get('adj_remark') or '').strip():
                _line += '【采购议价备注】%s' % (s['adj_remark'] or '')
            _ship_lines.append(_line)
    if _ship_lines:
        decision = decision + chr(10) + chr(10) + '【总价（含税含运）】' + chr(10) + chr(10).join(_ship_lines)
    c = ws.cell(row, 1, decision)
    c.font = base_font; c.alignment = Alignment(vertical='top', wrap_text=True)
    # V11.131: 决策备注黄色高亮框(显眼), 行高按品牌分析行数自适应
    _dec_fill = PatternFill('solid', fgColor='FFF2CC')
    _dec_rows = max(4, 2 + len(decision.split(chr(10))))
    ws.merge_cells(start_row=row, start_column=1, end_row=row + _dec_rows - 1, end_column=col_count)
    for r2 in range(row, row + _dec_rows):
        for col in range(1, col_count + 1):
            ws.cell(r2, col).fill = _dec_fill
            ws.cell(r2, col).border = border
    row += _dec_rows
    ws.cell(row, 1, '编制人：').font = note_font
    ws.cell(row, 4, '审核人：').font = note_font
    ws.cell(row, 7, '日期：').font = note_font

    # 动态设置列宽，确保所有供应商列宽一致
    ws.column_dimensions['A'].width = 6   # 序号
    ws.column_dimensions['B'].width = 18  # 物料名称
    ws.column_dimensions['C'].width = 10  # 数量
    ws.column_dimensions['D'].width = 14  # 规格
    # 每个供应商4列: 单价(10) + 总价(12) + 品牌(12) + 备注(16)
    col_count = 4 + len(sups) * 4
    for ci in range(5, min(col_count + 1, 26)):
        col_letter = chr(64 + ci) if ci <= 26 else 'A' + chr(64 + ci - 26)
        if ci % 4 == 1:  # 备注列
            ws.column_dimensions[col_letter].width = 16
        elif ci % 4 == 0:  # 总价列
            ws.column_dimensions[col_letter].width = 12
        elif ci % 4 == 3:  # 品牌列
            ws.column_dimensions[col_letter].width = 12
        else:  # 单价列
            ws.column_dimensions[col_letter].width = 10

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = '询价单_%s.xlsx' % (i['inq_no'] or iid)
    resp = make_response(buf.getvalue())
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    resp.headers['Content-Disposition'] = "attachment; filename*=UTF-8''%s" % urllib.parse.quote(fname)
    log(session['user_name'], '导出询价单Excel', '%s' % (i['inq_no'] or iid))
    return resp

@app.route('/api/budgets')
@login_required
def api_budgets():
    conn = db(); rows = conn.execute("SELECT b.*,d.name as dept_name FROM budget_accounts b LEFT JOIN departments d ON b.dept_id=d.id").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/contracts')
@login_required
def api_contracts():
    # V11.159: 合同列表 — 员工仅看自己相关的(通过订单的发起人关联)
    if session.get('user_role') == '员工':
        conn = db(); rows = conn.execute("SELECT c.*,po.order_no FROM contracts c LEFT JOIN purchase_orders po ON c.order_id=po.id WHERE po.requester_id=? ORDER BY c.id DESC LIMIT 50", (session.get('user_id', 0),)).fetchall(); conn.close()
        return jsonify([dict_row(r) for r in rows])
    conn = db(); rows = conn.execute("SELECT c.*,po.order_no FROM contracts c LEFT JOIN purchase_orders po ON c.order_id=po.id ORDER BY c.id DESC LIMIT 50").fetchall(); conn.close()
    out = []
    for r in rows:
        d = dict_row(r)
        try:
            d['amount_upper'] = rmb_upper(float(d.get('amount') or 0))
        except Exception:
            d['amount_upper'] = ''
        out.append(d)
    return jsonify(out)

@app.route('/api/contracts', methods=['POST'])
@login_required
def api_create_contract():
    d = request.json; conn = db()
    no = gen_contract_no(conn)
    # V4.1: 合同创建后进入审批流(待审批→审批通过→执行中), 审批通过前不能挂账
    conn.execute("INSERT INTO contracts(contract_no,order_id,contract_name,supplier,amount,sign_date,start_date,end_date,content,status,remark,urgent,attachment) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (no,d.get('order_id'),d.get('contract_name',''),d.get('supplier',''),float(d.get('amount',0)),
         d.get('sign_date'),d.get('start_date'),d.get('end_date'),d.get('content',''),'待审批',d.get('remark',''),
         1 if d.get('urgent') else 0,(d.get('attachment') or '').strip()))
    cid = conn.execute("SELECT id FROM contracts WHERE contract_no=?", (no,)).fetchone()[0]
    if d.get('order_id'): conn.execute("UPDATE purchase_orders SET status='已签合同',updated_at=? WHERE id=?", (now(),d['order_id']))
    conn.commit()
    create_approvals('contract', cid, float(d.get('amount',0)), submitter=session['user_name'])
    start_instances('contract', cid)   # 飞书/钉钉同步发起合同审批(未配置则跳过)
    conn.close()
    log(session['user_name'],'创建合同',f'{no}')
    return jsonify({'success':True,'contract_no':no})

@app.route('/api/contracts/monthly-summary')
@login_required
def api_monthly_summary():
    """V11.68: 月结汇总 — 按厂家分组当月待月结订单(settle_type=月结且未月结)"""
    role = session.get('user_role')
    if role not in ('系统管理员', '采购员', '分管领导', '总经理'):
        return jsonify({'error': '无权限'}), 403
    conn = db()
    try: conn.execute("ALTER TABLE purchase_orders ADD COLUMN settled_at TEXT DEFAULT ''")
    except Exception: pass
    m = datetime.date.today().strftime('%Y-%m')
    # V11.200: 月结汇总条件放宽 — 月结订单只要未月结(settled_at空)且未作废/未取消就应列入,
    # 含已入库未结款(原条件限'审批通过/已通过'导致入库后状态变'已入库'的单从汇总消失)
    rows = conn.execute("""SELECT supplier, COUNT(*) cnt, COALESCE(SUM(total_amount),0) amt,
        GROUP_CONCAT(order_no || ':' || item_name || 'x' || printf('%g',quantity) || ' ¥' || printf('%.2f',total_amount), '\n') detail
        FROM purchase_orders WHERE settle_type='月结' AND status NOT IN ('已作废','已取消','草稿')
        AND (settled_at IS NULL OR settled_at='') AND created_at LIKE ?
        GROUP BY supplier ORDER BY amt DESC""", (m + '%',)).fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/contracts/monthly-generate', methods=['POST'])
@login_required
def api_monthly_generate():
    """V11.68: 生成月度合同 — 指定厂家的当月待月结订单汇总成一份合同"""
    role = session.get('user_role')
    if role not in ('系统管理员', '采购员', '分管领导', '总经理'):
        return jsonify({'error': '无权限'}), 403
    d = request.json or {}
    supplier = (d.get('supplier') or '').strip()
    if not supplier:
        return jsonify({'error': '请指定厂家'}), 400
    conn = db()
    m = datetime.date.today().strftime('%Y-%m')
    orders = conn.execute("""SELECT * FROM purchase_orders WHERE supplier=? AND settle_type='月结'
        AND status NOT IN ('已作废','已取消','草稿') AND (settled_at IS NULL OR settled_at='') AND created_at LIKE ?""",
        (supplier, m + '%')).fetchall()
    if not orders:
        conn.close(); return jsonify({'error': '该厂家本月无待月结订单'}), 400
    total = sum(float(o['total_amount'] or 0) for o in orders)
    # 明细文本
    lines = []
    for o in orders:
        lines.append(f"{o['order_no']} {o['item_name']} x{o['quantity']}{o['unit'] or '个'} ¥{float(o['total_amount'] or 0):.2f}")
    content = '\n'.join(lines)
    order_nos = '、'.join(o['order_no'] for o in orders)
    # 生成月度合同(关联第一张单, 明细在content, 关联单号在remark)
    no = gen_contract_no(conn)
    y = datetime.date.today().strftime('%Y%m')
    conn.execute("""INSERT INTO contracts(contract_no,order_id,contract_name,supplier,amount,sign_date,start_date,end_date,content,status,remark,urgent,attachment)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (no, orders[0]['id'], f'{supplier}{y}月结算合同', supplier, round(total, 2),
         datetime.date.today().strftime('%Y-%m-%d'), '', '', content, '待审批',
         f'月结汇总: {order_nos}', 0, ''))
    cid = conn.execute("SELECT id FROM contracts WHERE contract_no=?", (no,)).fetchone()[0]
    # 标记订单已月结
    oids = [o['id'] for o in orders]
    for oid in oids:
        try: conn.execute("ALTER TABLE purchase_orders ADD COLUMN settled_at TEXT DEFAULT ''")
        except Exception: pass
        conn.execute("UPDATE purchase_orders SET settled_at=?, status='已签合同' WHERE id=?", (datetime.date.today().strftime('%Y-%m-%d'), oid))
    conn.commit()
    create_approvals('contract', cid, round(total, 2), submitter=session['user_name'])
    conn.close()
    try: start_instances('contract', cid)
    except Exception: pass
    log(session['user_name'], '月结汇总', f'{supplier} {len(orders)}单 → 合同{no} ¥{total:.2f}')
    return jsonify({'success': True, 'contract_no': no, 'count': len(orders), 'amount': round(total, 2)})

@app.route('/api/deliveries')
@login_required
def api_deliveries():
    conn = db(); rows = conn.execute("SELECT * FROM deliveries ORDER BY id DESC LIMIT 50").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/deliveries', methods=['POST'])
@login_required
def api_create_delivery():
    """到货登记(模式1节点); 需求4/8: 送货单强制关联合同编号"""
    d = request.json; conn = db()
    cid = d.get('contract_id')
    cno = ''
    if cid:
        ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
        if ct: cno = ct['contract_no']
    if not cid or not cno:
        conn.close(); return jsonify({'error': '到货登记必须关联合同编号(需求硬性规则)'}), 400
    no = gen_no('DH','deliveries','delivery_no')
    conn.execute("INSERT INTO deliveries(delivery_no,order_id,contract_id,contract_no,supplier,item_name,spec,quantity,unit,driver_name,vehicle_no,delivery_date,receiver,sign_status,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (no,d.get('order_id'),cid,cno,d.get('supplier',''),d.get('item_name',''),d.get('spec',''),
         float(d.get('quantity',0)),d.get('unit','个'),d.get('driver_name',''),d.get('vehicle_no',''),
         d.get('delivery_date'),d.get('receiver',''),'待签收',d.get('remark','')))
    if d.get('order_id'): conn.execute("UPDATE purchase_orders SET status='已发货',updated_at=? WHERE id=?", (now(),d['order_id']))
    conn.commit(); conn.close()
    log(session['user_name'],'创建送货单',f'{no} 合同{cno}')
    return jsonify({'success':True,'delivery_no':no})

@app.route('/api/deliveries/<int:did>/sign', methods=['POST'])
@login_required
def api_sign_delivery(did):
    d = request.json; conn = db()
    dn = conn.execute("SELECT * FROM deliveries WHERE id=?", (did,)).fetchone()
    if not dn: return jsonify({'error':'not found'}),404
    conn.execute("UPDATE deliveries SET sign_status='已签收', receiver=?, sign_time=? WHERE id=?", (d.get('receiver','库房'),now(),did))
    # 任务4: 若该订单已自动生成入库单(货到付款下单即生成), 则复用, 不重复插入
    # V11.152e: 统一防重 — 该订单已有任意未入库状态的入库单则复用(不限于待检验)
    exist = conn.execute("SELECT id FROM receivings WHERE order_id=? AND status NOT IN ('已入库','已作废') LIMIT 1", (dn['order_id'],)).fetchone()
    if exist:
        conn.execute("UPDATE receivings SET delivery_id=?, quantity=?, updated_at=datetime('now','localtime') WHERE id=?", (did, dn['quantity'], exist['id']))
        rno = conn.execute("SELECT receive_no FROM receivings WHERE id=?", (exist['id'],)).fetchone()[0]
    else:
        # V11.31: 自动带出部门(订单→申请单链)
        _dept = ''
        try:
            _po2 = conn.execute("SELECT req_id FROM purchase_orders WHERE id=?", (dn['order_id'],)).fetchone()
            if _po2 and _po2['req_id']:
                _pr2 = conn.execute("SELECT dept FROM purchase_requests WHERE id=?", (_po2['req_id'],)).fetchone()
                if _pr2 and _pr2['dept']:
                    _dept = _pr2['dept']
        except Exception:
            pass
        rno = gen_no('RK','receivings','receive_no')
        # V11.198: 送货签收自动生成的入库单默认暂估(is_est=1, 货到票未到先暂估入账, 收票后红冲转正式)
        conn.execute("INSERT INTO receivings(receive_no,delivery_id,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,dept,is_est) VALUES(?,?,?,?,?,?,?,?,?,?,?,0)",
            (rno,did,dn['order_id'],dn['item_name'],dn['spec'],dn['quantity'],dn['unit'],dn['quantity'],'待检验',now(),_dept))
    if dn['order_id']: conn.execute("UPDATE purchase_orders SET status='已到货',updated_at=? WHERE id=?", (now(),dn['order_id']))
    conn.commit(); conn.close()
    log(session['user_name'],'签收送货单',f'#{did}')
    return jsonify({'success':True,'receive_no':rno})

@app.route('/api/receivings')
@login_required
def api_receivings():
    # V11.64: 入库单 — 库管员/采购员/财务/领导/管理员可见(采购要发票核对+跟到货, 财务对账); 员工不看
    if session.get('user_role') == '员工':
        return jsonify([])
    # V11.29: 部门/类别筛选; V11.202: 入库类型筛选(est=暂估入库/formal=正式入库)
    f_dept = (request.args.get('dept') or '').strip()
    f_cat = (request.args.get('cat') or '').strip()
    f_type = (request.args.get('type') or '').strip()
    conn = db()
    sql = "SELECT r.*, po.trade_mode, po.order_no, po.supplier FROM receivings r LEFT JOIN purchase_orders po ON r.order_id=po.id"
    where = []; args = []
    if f_dept:
        where.append("r.dept=?"); args.append(f_dept)
    if f_cat:
        where.append("(r.item_name IN (SELECT item_name FROM inventory WHERE cat_code=(SELECT code FROM categories WHERE name=?)))")
        args.append(f_cat)
    if f_type == 'est':
        where.append("r.is_est=1")
    elif f_type == 'formal':
        where.append("COALESCE(r.is_est,0)=0")
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY r.id DESC LIMIT 80"
    rows = conn.execute(sql, args).fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        items = []
        if r['order_id']:
            items = [dict_row(x) for x in conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (r['order_id'],)).fetchall()]
            cnt = conn.execute("SELECT COUNT(*), COALESCE(SUM(quantity),0) FROM order_items WHERE order_id=?", (r['order_id'],)).fetchone()
            if cnt and cnt[0]:
                d['item_count'] = cnt[0]
                # V11.198: 数量列显示本单实际验收量(qualified_qty优先, 分批入库时不再是订单整批总量)
                d['total_qty'] = r['qualified_qty'] if (r['qualified_qty'] or 0) > 0 else (r['quantity'] or cnt[1])
                d['order_total_qty'] = cnt[1]  # 订单总量(详情可对比剩余)
        if not items and r['items_json']:
            try:
                items = json.loads(r['items_json'])
            except Exception:
                items = []
        if not items:
            items = [{'item_name': r['item_name'], 'spec': r['spec'], 'quantity': r['quantity'], 'unit': r['unit']}]
        # V11.152: 有明细时列表物资名显示完整(首项+共N项)
        if len(items) > 1:
            d['item_name'] = items[0]['item_name'] + ' 等%d项' % len(items)
        elif items:
            d['item_name'] = items[0]['item_name']
        d['items'] = items
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/receivings/<int:rid>/arrived', methods=['POST'])
@login_required
def api_receiving_arrived(rid):
    """V11.37: 到货提醒单确认到货 — 货实际到了, 状态 待入库→入库中(可提交验收)
    V11.206: 到货确认时可选'需集体验收'(大型设备/关键物资/维修返回件) → 建集体验收审批实例
    """
    d = request.json or {}
    conn = db()
    rn = conn.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        conn.close(); return jsonify({'error': '入库单不存在'}), 404
    if rn['status'] != '待入库':
        conn.close(); return jsonify({'error': f'当前状态({rn["status"]})无需确认到货'}), 400
    collect_flag = 1 if d.get('collect_accept') else 0
    # 建集体验收审批(标记后): 走 approval_flow_config collect_accept 配置的角色
    if collect_flag:
        _exists = conn.execute("SELECT id FROM approval_instances WHERE biz_type='collect_accept' AND biz_id=? AND status='pending'", (rid,)).fetchone()
        if not _exists:
            try:
                create_approvals('collect_accept', rid, float(rn['quantity'] or 0), submitter=session.get('user_name', ''))
            except Exception as _e:
                conn.close(); return jsonify({'error': f'集体验收审批创建失败: {_e}'}), 500
        conn.execute("UPDATE receivings SET collect_accept=1, collect_status='待集体验收', status='入库中', received_at=? WHERE id=?",
                     (now(), rid))
    else:
        conn.execute("UPDATE receivings SET collect_accept=0, collect_status='', status='入库中', received_at=? WHERE id=?",
                     (now(), rid))
    conn.commit(); conn.close()
    if collect_flag:
        log(session['user_name'], '确认到货(集体验收)', f'#{rid} {rn["item_name"]} 需集体验收, 已推送审批')
    else:
        log(session['user_name'], '确认到货', f'#{rid} {rn["item_name"]} 合同自动生成单已确认到货')
    return jsonify({'success': True, 'collect_accept': collect_flag})

@app.route('/api/receivings/<int:rid>/complete', methods=['POST'])
@login_required
def api_complete_receiving(rid):
    """V5.0: 入库单提交审批 → 审批通过后由 finish_approvals 实际增加库存
    兼容旧调用(直接确认验收): 提交后自动走审批, 审批通过即入库
    V11.37: 到货提醒单(合同自动生成)需先点"确认到货"→状态入库中→再提交验收"""
    d = request.json; conn = db()
    rn = conn.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        conn.close(); return jsonify({'error': '入库单不存在'}), 404
    if rn['status'] in ('已入库', '已驳回', '已作废'):
        conn.close(); return jsonify({'error': f'当前状态({rn["status"]})不可提交审批'}), 400
    # V11.206 集体验收: 标记需集体验收的单必须先完成集体验收审批, 才能走常规入库审批
    _need_collect = int(rn['collect_accept'] or 0)
    if _need_collect and (rn['collect_status'] or '') != '已集体验收':
        _cpend = conn.execute("SELECT status FROM approval_instances WHERE biz_type='collect_accept' AND biz_id=? ORDER BY id DESC LIMIT 1", (rid,)).fetchone()
        _cs = (_cpend['status'] if _cpend else '') or ''
        if _cs == 'pending':
            conn.close(); return jsonify({'error': '该物资需集体验收，集体验收审批中，通过后才能提交入库审批'}), 400
        conn.close(); return jsonify({'error': '该物资需集体验收，请先完成集体验收（到货确认时勾选，审批通过后再提交入库）'}), 400
    # 防重复提交: 待审批且已有待审实例 → 提示撤回而不是重复建链
    pend = conn.execute("SELECT 1 FROM approval_instances WHERE biz_type='receiving' AND biz_id=? AND status='pending' LIMIT 1", (rid,)).fetchone()
    if rn['status'] == '待审批' and pend:
        conn.close(); return jsonify({'error': '该入库单已在审批中，请勿重复提交（如需修改请先撤回）'}), 400
    warehouse = d.get('warehouse', '主库房'); inspector = (d.get('inspector') or '').strip() or session.get('user_name', '') or '系统'
    qty_override = d.get('items') or {}
    try: conn.execute("ALTER TABLE receivings ADD COLUMN completed_at TEXT")
    except Exception: pass
    try: conn.execute("ALTER TABLE receivings ADD COLUMN warehouse TEXT DEFAULT '主库房'")
    except Exception: pass
    oi = []
    if rn['order_id']:
        oi = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (rn['order_id'],)).fetchall()
    # 更新验收信息 + 置待审批
    total_q = 0.0
    # V8.1: 批量验收明细按行匹配(支持同名不同规格), 兼容旧 dict 按品名
    qty_list = qty_override if isinstance(qty_override, list) else None
    for idx, it in enumerate(oi):
        q = float(it['quantity'] or 0)
        if qty_list is not None and idx < len(qty_list):
            try: q = float(qty_list[idx].get('quantity', q) or 0)
            except Exception: pass
        elif isinstance(qty_override, dict) and it['item_name'] in qty_override:
            try: q = float(qty_override[it['item_name']] or 0)
            except Exception: pass
        total_q += max(q, 0)
    if not oi:
        total_q = float(d.get('qualified_qty', rn['quantity'] or 0))
    # V11.198b: 入库类型(is_est)以提交时手动选择为准(0/1), 覆盖自动生成默认值
    _is_est = d.get('is_est')
    if _is_est is not None:
        _is_est = 1 if int(_is_est) == 1 else 0
    conn.execute("UPDATE receivings SET qualified_qty=?,defective_qty=?,inspector=?,warehouse=?,status='待审批',remark=?" + (",is_est=?" if _is_est is not None else "") + " WHERE id=?",
                 (total_q, float(d.get('defective_qty',0)), inspector, warehouse,
                  (rn['remark'] or '') + ' 提交审批') + ((_is_est,) if _is_est is not None else ()) + (rid,))
    conn.commit()
    # 创建审批实例(入库单审批) + 同步发起钉钉/飞书审批
    create_approvals('receiving', rid, 0, submitter=session['user_name'])
    conn.close()
    try: start_instances('receiving', rid)
    except Exception as e: print('receiving start_instances err:', e)
    log(session['user_name'], '提交入库审批', f'入库单#{rid} 待审批 {total_q}件')
    return jsonify({'success': True, 'message': f'入库单已提交审批，审批通过后自动增加库存（{total_q}件）', 'status': '待审批', 'items_in': len(oi) or 1})


def do_requisition_stock(c, rid, warehouse='主库房', operator='系统'):
    """V5.0: 出库审批通过后执行 — 扣减库存 + 写流水(幂等: 已有该单据出库流水则跳过)
    明细取 requisition_items; 扣减后允许库存为负(展示为负值, 与V5.0设计一致)"""
    rq = c.execute("SELECT * FROM requisitions WHERE id=?", (rid,)).fetchone()
    if not rq:
        return 0
    done = c.execute("SELECT 1 FROM inventory_flows WHERE doc_type='requisition' AND doc_id=? AND flow_type='出库' LIMIT 1", (rid,)).fetchone()
    if done:
        return 0
    its = c.execute("SELECT * FROM requisition_items WHERE requisition_id=? ORDER BY id", (rid,)).fetchall()
    if not its:
        return 0
    total_q = 0.0
    for it in its:
        q = float(it['quantity'] or 0)
        if q <= 0:
            continue
        total_q += q
        inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? AND (warehouse=? OR warehouse IS NULL OR warehouse='') ORDER BY quantity DESC LIMIT 1",
                        (it['item_name'], it['spec'] or '', warehouse)).fetchone()
        if inv is None:
            inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? ORDER BY quantity DESC LIMIT 1",
                            (it['item_name'], it['spec'] or '')).fetchone()
        if inv:
            new_q = (inv['quantity'] or 0) - q
            c.execute("UPDATE inventory SET quantity=?, last_move_date=?, updated_at=? WHERE id=?", (new_q, now(), now(), inv['id']))
        else:
            new_q = -q
            c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,last_move_date,updated_at) VALUES(?,?,?,?,?,?,?)",
                      (it['item_name'], it['spec'] or '', it['unit'] or '个', new_q, warehouse, now(), now()))
        c.execute("INSERT INTO inventory_flows(item_name,spec,unit,flow_type,doc_type,doc_id,doc_no,qty,balance_after,operator,remark,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                  (it['item_name'], it['spec'] or '', it['unit'] or '个', '出库', 'requisition', rid, rq['req_no'],
                   -q, new_q, operator or '系统', f'出库单{rq["req_no"]}审批通过', now()))
    return total_q


def _op_name():
    """V11.12: 取操作人 — 后台线程(审批轮询/回调)无HTTP请求上下文, flask session 会抛异常; 有请求时返回登录用户, 否则返回'系统'"""
    try:
        from flask import session as _s
        return _s.get('user_name', '系统')
    except Exception:
        return '系统'


def _rcv_doc_qty(doc):
    """单据本批实际数量 = items_json 明细合计(与 do_receiving_stock 入库存口径一致); 无明细用 quantity"""
    try:
        if doc.get('items_json'):
            its = json.loads(doc['items_json'])
            s = sum(float(x.get('quantity') or 0) for x in its if isinstance(x, dict))
            if s > 0:
                return s
    except Exception:
        pass
    return float(doc.get('quantity') or 0)


def _order_rcv_stats(c, oid):
    """V11.202 分批验收: 订单验收统计 — 总数量/已验收入库/待验收待定/暂估·正式合计/批次台账/各明细行余量。
    已验收=该订单下 status='已入库' 的入库单实际入库量(按items_json口径); 作废/待审批不计入。"""
    oi = c.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()
    if oi:
        order_total = sum(float(x['quantity'] or 0) for x in oi)
    else:
        po0 = c.execute("SELECT quantity FROM purchase_orders WHERE id=?", (oid,)).fetchone()
        order_total = float((po0['quantity'] if po0 else 0) or 0)
    docs = [dict_row(x) for x in c.execute("SELECT * FROM receivings WHERE order_id=? ORDER BY id", (oid,)).fetchall()]
    accepted = 0.0
    est_total = 0.0
    formal_total = 0.0
    per_line = []
    if oi:
        for x in oi:
            per_line.append({'item_name': x['item_name'], 'spec': x['spec'] or '',
                             'unit': x['unit'] or '个', 'order_qty': float(x['quantity'] or 0), 'accepted': 0.0})
    batches = []
    has_active_full_doc = False   # 存在未作废的老整批单(待入库/入库中/待检验/草稿/已驳回/待审批)
    for doc in docs:
        doc['_qty'] = _rcv_doc_qty(doc)
        in_stock = doc['status'] == '已入库'
        doc['_in_stock'] = in_stock
        if not doc.get('batch_no') and doc['status'] not in ('已入库', '已作废'):
            has_active_full_doc = True
        if in_stock:
            accepted += doc['_qty']
            if doc['is_est']:
                est_total += doc['_qty']
            else:
                formal_total += doc['_qty']
            # 明细行已验(按名称+规格匹配)
            if per_line:
                try:
                    dij = json.loads(doc['items_json']) if doc.get('items_json') else []
                except Exception:
                    dij = []
                if dij:
                    for it in dij:
                        for pl in per_line:
                            if pl['item_name'] == it.get('item_name') and ((pl['spec'] or '') == (it.get('spec') or '') or not (pl['spec'] or '') or not (it.get('spec') or '')):
                                pl['accepted'] += float(it.get('quantity') or 0)
                                break
                else:
                    for pl in per_line:
                        if pl['item_name'] == doc['item_name'] and (pl['spec'] or '') == (doc.get('spec') or ''):
                            pl['accepted'] += doc['_qty']
                            break
        if doc['status'] != '已作废':
            batches.append(doc)
    pending = max(0.0, order_total - accepted)
    for pl in per_line:
        pl['remaining'] = max(0.0, float(pl['order_qty']) - float(pl['accepted']))
    return {'order_id': oid, 'order_total': order_total, 'accepted': accepted, 'pending': pending,
            'est_total': est_total, 'formal_total': formal_total,
            'per_line': per_line, 'batches': batches, 'has_active_full_doc': has_active_full_doc}


def do_receiving_stock(c, rid, warehouse='主库房', inspector='管理员', qty_override=None):
    """V5.0: 入库审批通过后执行 — 增加库存 + 写流水(幂等: 已有该单据入库流水则跳过)"""
    rn = c.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        return 0
    # 幂等判断: 用流水表而非状态(父状态可能已被 finish_approvals 更新) — V11.202 兼容分批流水类型
    done = c.execute("SELECT 1 FROM inventory_flows WHERE doc_type='receiving' AND doc_id=? AND (flow_type='入库' OR flow_type LIKE '分批入库%') LIMIT 1", (rid,)).fetchone()
    if done:
        return 0
    # V11.202 分批验收: 分批单流水类型标注来源(分批入库暂估/分批入库正式), 可溯源到批次单
    _is_batch = bool(rn['batch_no'])
    _ft = ('分批入库暂估' if rn['is_est'] else '分批入库正式') if _is_batch else '入库'
    _ft_suffix = (f" 分批验收{rn['batch_no']}·{'暂估入库' if rn['is_est'] else '正式入库'}") if _is_batch else ''
    qty_override = qty_override or {}
    oi = []
    if rn['order_id']:
        oi = c.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (rn['order_id'],)).fetchall()
    # 手动入库单: 明细在 items_json
    manual_items = []
    if rn['items_json']:
        try:
            manual_items = json.loads(rn['items_json'] or '[]')
        except Exception:
            manual_items = []
    total_q = 0.0
    if manual_items:
        # ── 手动入库单: 逐条明细入库存(带含税单价) ──
        for it in manual_items:
            q = float(it.get('quantity', 0) or 0)
            if q <= 0: continue
            total_q += q
            _price = float(it.get('price', 0) or 0); _tr = float(it.get('tax_rate', 13) or 13)
            inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? AND warehouse=?",
                            (it['item_name'], it.get('spec', '') or '', warehouse)).fetchone()
            if inv:
                _up = "quantity=quantity+?, last_move_date=?, updated_at=?"
                _args = [q, now(), now()]
                if (not inv['price'] or inv['price'] == 0) and _price:
                    _up += ", price=?"; _args.append(_price)
                if _tr and (not inv['tax_rate'] or inv['tax_rate'] == 0):
                    _up += ", tax_rate=?"; _args.append(_tr)
                _args.append(inv['id'])
                c.execute("UPDATE inventory SET " + _up + " WHERE id=?", _args)
                new_bal = (inv['quantity'] or 0) + q
            else:
                c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,price,tax_rate,last_move_date,updated_at) VALUES(?,?,?,?,?,?,?,?,?)",
                          (it['item_name'], it.get('spec','') or '', it.get('unit','个') or '个', q, warehouse, _price, _tr, now(), now()))
                new_bal = q
            c.execute("INSERT INTO inventory_flows(item_name,spec,unit,flow_type,doc_type,doc_id,doc_no,qty,balance_after,operator,remark,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                      (it['item_name'], it.get('spec','') or '', it.get('unit','个') or '个', _ft, 'receiving', rid, rn['receive_no'], q, new_bal,
                       _op_name(), f'入库单{rn["receive_no"]}审批通过{_ft_suffix}', now()))
    elif oi:
        _po = c.execute("SELECT category, trade_mode, supplier FROM purchase_orders WHERE id=?", (rn['order_id'],)).fetchone()
        _po_sup = (_po['supplier'] or '') if _po else ''
        qty_list = qty_override if isinstance(qty_override, list) else None
        for idx, it in enumerate(oi):
            q = float(it['quantity'] or 0)
            if qty_list is not None and idx < len(qty_list):
                try: q = float(qty_list[idx].get('quantity', q) or 0)
                except Exception: pass
            elif isinstance(qty_override, dict) and it['item_name'] in qty_override:
                try: q = float(qty_override[it['item_name']] or 0)
                except Exception: pass
            if q <= 0: continue
            total_q += q
            _price = float(it['price'] or 0); _tr = float(it['tax_rate'] or 13)
            _cat = _po['category'] or '' if _po else ''
            inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? AND warehouse=?",
                            (it['item_name'], it['spec'] or '', warehouse)).fetchone()
            if inv:
                _up = "quantity=quantity+?, last_move_date=?, updated_at=?"
                _args = [q, now(), now()]
                if (not inv['price'] or inv['price'] == 0) and _price:
                    _up += ", price=?"; _args.append(_price)
                if _cat and not inv['cat_code']:
                    _up += ", cat_code=?"; _args.append(_cat)
                if _po_sup and not inv['supplier']:
                    _up += ", supplier=?"; _args.append(_po_sup)
                _args.append(inv['id'])
                c.execute("UPDATE inventory SET " + _up + " WHERE id=?", _args)
                new_bal = (inv['quantity'] or 0) + q
            else:
                c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,price,tax_rate,cat_code,last_move_date,updated_at,supplier) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                          (it['item_name'], it['spec'] or '', it['unit'] or '个', q, warehouse, _price, _tr, _cat, now(), now(), _po_sup))
                new_bal = q
            c.execute("INSERT INTO inventory_flows(item_name,spec,unit,flow_type,doc_type,doc_id,doc_no,qty,balance_after,operator,remark,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                      (it['item_name'], it['spec'] or '', it['unit'] or '个', _ft, 'receiving', rid, rn['receive_no'], q, new_bal,
                       _op_name(), f'入库单{rn["receive_no"]}审批通过{_ft_suffix}', now()))
    else:
        q = float(rn['quantity'] or 0)
        total_q = q
        _price = 0; _tr = 13; _cat = ''; _sup = ''
        if rn['order_id']:
            _po2 = c.execute("SELECT price, tax_rate, category, supplier FROM purchase_orders WHERE id=?", (rn['order_id'],)).fetchone()
            if _po2:
                _price = _po2['price'] or 0; _tr = _po2['tax_rate'] or 13; _cat = _po2['category'] or ''; _sup = _po2['supplier'] or ''
        inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? AND warehouse=?",
                        (rn['item_name'], rn['spec'] or '', warehouse)).fetchone()
        if inv:
            _up = "quantity=quantity+?, last_move_date=?, updated_at=?"
            _args = [q, now(), now()]
            if (not inv['price'] or inv['price'] == 0) and _price:
                _up += ", price=?"; _args.append(_price)
            if _cat and not inv['cat_code']:
                _up += ", cat_code=?"; _args.append(_cat)
            if _sup and not inv['supplier']:
                _up += ", supplier=?"; _args.append(_sup)
            _args.append(inv['id'])
            c.execute("UPDATE inventory SET " + _up + " WHERE id=?", _args)
            new_bal = (inv['quantity'] or 0) + q
        else:
            c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,price,tax_rate,cat_code,last_move_date,updated_at,supplier) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                      (rn['item_name'], rn['spec'] or '', rn['unit'] or '个', q, warehouse, _price, _tr, _cat, now(), now(), _sup))
            new_bal = q
        c.execute("INSERT INTO inventory_flows(item_name,spec,unit,flow_type,doc_type,doc_id,doc_no,qty,balance_after,operator,remark,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                  (rn['item_name'], rn['spec'] or '', rn['unit'] or '个', _ft, 'receiving', rid, rn['receive_no'], q, new_bal,
                   _op_name(), f'入库单{rn["receive_no"]}审批通过{_ft_suffix}', now()))
    c.execute("UPDATE receivings SET status='已入库',completed_at=?,warehouse=?,inspector=? WHERE id=?",
              (now(), warehouse, inspector or rn['inspector'] or '系统', rid))
    if rn['order_id']:
        po = c.execute("SELECT * FROM purchase_orders WHERE id=?", (rn['order_id'],)).fetchone()
        if _is_batch and po:
            # V11.202 分批验收: 订单状态按剩余待验收待定量流转(不是直接已入库)
            _st = _order_rcv_stats(c, rn['order_id'])
            if _st['pending'] <= 0.001:
                c.execute("UPDATE purchase_orders SET status='全部已验收',updated_at=? WHERE id=?", (now(), rn['order_id']))
            else:
                c.execute("UPDATE purchase_orders SET status='部分到货，待继续验收',updated_at=? WHERE id=?", (now(), rn['order_id']))
        elif po and po['trade_mode'] == '先款后货':
            c.execute("UPDATE purchase_orders SET status='已核销',updated_at=? WHERE id=?", (now(), rn['order_id']))
        else:
            c.execute("UPDATE purchase_orders SET status='已入库',updated_at=? WHERE id=?", (now(), rn['order_id']))
    return total_q

# ---- V6: 入库单下载(生成标准入库单 xlsx) ----
@app.route('/api/receivings/<int:rid>/download')
@login_required
def api_receiving_download(rid):
    """入库单下载(V8.4b: 紧凑布局, 无空白列)
    标题/日期+票号+供应商+品种数/仓库/表头(No|品名|规格|数量|单位|不含税单价|税率|不含税金额|含税金额|备注)/小计/合计大写/签字区"""
    conn = db()
    rn = conn.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        conn.close(); return jsonify({'error': '入库单不存在'}), 404
    po = None
    if rn['order_id']:
        po = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (rn['order_id'],)).fetchone()
    oi = []
    if rn['order_id']:
        oi = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (rn['order_id'],)).fetchall()
    conn.close()
    supplier = (po['supplier'] if po else '') or ''
    rows = []
    if oi:
        for it in oi:
            price = float(it['price'] or 0)
            tr = float(it['tax_rate'] or 13)
            rows.append({'name': it['item_name'], 'spec': it['spec'] or '', 'qty': float(it['quantity'] or 0),
                         'unit': it['unit'] or '个', 'price': price, 'tax': tr,
                         'amt_no': price * float(it['quantity'] or 0),
                         'amt_tax': price * float(it['quantity'] or 0) * (1 + tr / 100),
                         'remark': (dict(it).get('remark') or '')})
    elif rn['items_json']:
        try:
            for it in json.loads(rn['items_json']):
                price = float(it.get('price') or 0); tr = float(it.get('tax_rate') or 13)
                q = float(it.get('quantity') or 0)
                rows.append({'name': it.get('item_name', ''), 'spec': it.get('spec') or '', 'qty': q,
                             'unit': it.get('unit') or '个', 'price': price, 'tax': tr,
                             'amt_no': price * q, 'amt_tax': price * q * (1 + tr / 100),
                             'remark': it.get('remark') or ''})
        except Exception:
            rows = []
    if not rows:
        price = float((po['price'] if po else 0) or 0); tr = 13
        q = float(rn['quantity'] or 0)
        rows = [{'name': rn['item_name'], 'spec': rn['spec'] or '', 'qty': q, 'unit': rn['unit'] or '个',
                 'price': price, 'tax': tr, 'amt_no': price * q, 'amt_tax': price * q * (1 + tr / 100),
                 'remark': rn['remark'] or ''}]
    # V11.202 分批验收: 导出同步展示 批次/本批入库数量/订单待验收待定余量(老单据无批次不显示, 表头自动下移)
    _extra_lines = []
    if rn['batch_no']:
        _extra_lines.append(('批次', "%s · %s" % (rn['batch_no'], '暂估入库' if rn['is_est'] else '正式入库')))
        _extra_lines.append(('本批入库数量', "%g 件" % (_rcv_doc_qty(dict_row(rn)) or 0)))
    if rn['order_id']:
        try:
            _c2 = db()
            _ost = _order_rcv_stats(_c2, rn['order_id'])
            _c2.close()
            _extra_lines.append(('订单验收进度', '订单总数量 %g ｜ 已验收入库 %g ｜ 待验收待定 %g' % (_ost['order_total'], _ost['accepted'], _ost['pending'])))
        except Exception:
            pass
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    wb = Workbook(); ws = wb.active; ws.title = '入库单'
    thin = Side(style='thin', color='000000')
    bd = Border(left=thin, right=thin, top=thin, bottom=thin)
    CN = lambda bold=False, size=10: Font(name='宋体', bold=bold, size=size)
    # 标题
    ws.merge_cells('A1:J1')
    c = ws['A1']; c.value = '河曲县正成洗选煤有限责任公司入库单'
    c.font = CN(bold=True, size=14); c.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 26
    # 第2行: 日期/票号/供应商
    ws['A2'] = '日期：'; ws['B2'] = (rn['completed_at'] or rn['created_at'] or '')[:10]
    ws['C2'] = '票号：'; ws['D2'] = rn['receive_no']
    ws['E2'] = '供应商：'; ws.merge_cells('F2:J2'); ws['F2'] = supplier
    for cc in ('A2','B2','C2','D2','E2','F2'):
        ws[cc].font = CN()
    # 第3行: 仓库/品种数 + V11.29 归属部门
    ws['A3'] = '仓库：'; ws.merge_cells('B3:C3'); ws['B3'] = rn['warehouse'] or '生产库房'
    ws['D3'] = '品种数：'; ws['E3'] = len(rows)
    ws['F3'] = '归属部门：'; ws.merge_cells('G3:J3'); ws['G3'] = rn['dept'] if 'dept' in rn.keys() and rn['dept'] else ''
    for cc in ('A3','B3','D3','E3'):
        ws[cc].font = CN()
    # V11.202: 批次/验收进度信息行(有则显示) — 表头与明细行自动下移, 保持 10 列布局
    _info_rows = 0
    for _lab, _val in _extra_lines:
        _r0 = 4 + _info_rows
        ws.cell(row=_r0, column=1, value=_lab + '：').font = CN(bold=True)
        ws.merge_cells(start_row=_r0, start_column=2, end_row=_r0, end_column=10)
        ws.cell(row=_r0, column=2, value=_val).font = CN()
        for j in range(1, 11):
            ws.cell(row=_r0, column=j).border = bd
        _info_rows += 1
    _hr = 4 + _info_rows
    # 表头 (10列连续, 无空白列)
    headers = ['No.', '品名', '规格', '数量', '单位', '不含税单价', '税率', '不含税金额', '含税金额', '备注']
    for j, h in enumerate(headers, 1):
        cc = ws.cell(row=_hr, column=j, value=h)
        cc.font = CN(bold=True); cc.border = bd
        cc.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[_hr].height = 18
    # 明细行
    r = _hr + 1
    t_qty = 0.0; t_no = 0.0; t_tax = 0.0
    for idx, it in enumerate(rows, 1):
        t_qty += it['qty']; t_no += it['amt_no']; t_tax += it['amt_tax']
        vals = [idx, it['name'], it['spec'], it['qty'], it['unit'], it['price'],
                f"{it['tax']}%", round(it['amt_no'], 2), round(it['amt_tax'], 2), it['remark']]
        for j, v in enumerate(vals, 1):
            cc = ws.cell(row=r, column=j, value=v)
            cc.font = CN(); cc.border = bd
            cc.alignment = Alignment(vertical='center', horizontal='left' if j == 2 else 'center')
        r += 1
    # 小计行
    ws.cell(row=r, column=1, value='小计').font = CN(bold=True)
    ws.cell(row=r, column=4, value=t_qty).font = CN(bold=True)
    ws.cell(row=r, column=8, value=round(t_no, 2)).font = CN(bold=True)
    ws.cell(row=r, column=9, value=round(t_tax, 2)).font = CN(bold=True)
    for j in range(1, 11):
        ws.cell(row=r, column=j).border = bd
    r += 1
    # 合计大写
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=10)
    ws.cell(row=r, column=1, value=f'合计: 人民币大写：{rmb_upper(t_tax)}').font = CN(bold=True)
    for j in range(1, 11):
        ws.cell(row=r, column=j).border = bd
    r += 1
    # 签字区
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=10)
    ws.cell(row=r, column=1, value='库管员签字：____________    验收员签字：____________    采购员：____________').font = CN()
    for j in range(1, 11):
        ws.cell(row=r, column=j).border = bd
    ws.row_dimensions[r].height = 22
    r += 1
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=10)
    ws.cell(row=r, column=1, value='第1页/共1页').font = CN()
    ws.cell(row=r, column=1).alignment = Alignment(horizontal='center')
    for j, w in enumerate([6, 18, 18, 10, 8, 12, 8, 12, 12, 20], 1):
        ws.column_dimensions[chr(64 + j)].width = w
    # V11.27: 审批通过 → 盖章领导预录签名
    stamp_leader_sign(ws, r - 1, 'receiving', rn['id'])
    # V11.175: 底部追加驳回审批记录
    append_reject_rows(ws, r, 'receiving', rn['id'], ncols=10, CN=CN)
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    resp = send_file(bio, as_attachment=True, download_name=f'{rn["receive_no"]}入库单.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    return resp

@app.route('/api/requisitions/<int:rid>/download')
@login_required
def api_requisition_download(rid):
    """出库单下载(V8.4b: 紧凑布局, 无空白列) 数量/金额为负值"""
    conn = db()
    rq = conn.execute("SELECT * FROM requisitions WHERE id=?", (rid,)).fetchone()
    if not rq:
        conn.close(); return jsonify({'error': '出库单不存在'}), 404
    its = conn.execute("SELECT * FROM requisition_items WHERE requisition_id=? ORDER BY id", (rid,)).fetchall()
    price_map = {}
    try:
        # V9.1: 按 名称+规格 取价, 避免同名不同规格串价
        for inv in conn.execute("SELECT item_name, spec, price FROM inventory WHERE price IS NOT NULL").fetchall():
            price_map.setdefault((inv['item_name'], inv['spec'] or ''), float(inv['price'] or 0))
    except Exception:
        pass
    conn.close()
    if not its:
        its = [{'item_name': rq['item_name'], 'spec': rq['spec'] or '', 'unit': rq['unit'] or '个',
                'quantity': rq['quantity'], 'purpose': rq['purpose'] or ''}]
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    wb = Workbook(); ws = wb.active; ws.title = '出库单'
    thin = Side(style='thin', color='000000')
    bd = Border(left=thin, right=thin, top=thin, bottom=thin)
    CN = lambda bold=False, size=10: Font(name='宋体', bold=bold, size=size)
    # 标题
    ws.merge_cells('A1:G1')
    c = ws['A1']; c.value = '河曲县洗选煤有限责任公司部门出库单'
    c.font = CN(bold=True, size=14); c.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 26
    # 第2行: 日期/票号/领取人
    ws['A2'] = '日期：'; ws['B2'] = (rq['issued_at'] or rq['created_at'] or '')[:10]
    ws['C2'] = '票号：'; ws['D2'] = rq['req_no']
    ws['E2'] = '领取人：'; ws['F2'] = rq['receiver'] or rq['requester'] or ''
    for cc in ('A2','B2','C2','D2','E2','F2'):
        ws[cc].font = CN()
    # 第3行: 部门/仓库/品种数
    ws['A3'] = '部门：'; ws['B3'] = rq['dept'] or ''
    ws['C3'] = '仓库：'; ws['D3'] = '主库房'
    ws['E3'] = '领取部门：'; ws['F3'] = rq['receive_dept'] or rq['dept'] or ''
    for cc in ('A3','B3','C3','D3','E3','F3'):
        ws[cc].font = CN()
    # 第4行表头 (7列连续, 无空白列)
    headers = ['No.', '品名', '规格', '单位', '数量', '单价', '金额']
    for j, h in enumerate(headers, 1):
        cc = ws.cell(row=4, column=j, value=h)
        cc.font = CN(bold=True); cc.border = bd
        cc.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[4].height = 18
    # 明细行
    r = 5
    t_qty = 0.0; t_amt = 0.0
    for idx, it in enumerate(its, 1):
        q = float(it['quantity'] or 0); t_qty += q
        price = price_map.get((it['item_name'], (dict(it).get('spec') or '') or ''), 0)
        amt = q * price; t_amt += amt
        vals = [idx, it['item_name'], (dict(it).get('spec') or '') or '',
                (dict(it).get('unit') or '个') or '个', -q, price, -amt]
        for j, v in enumerate(vals, 1):
            cc = ws.cell(row=r, column=j, value=v)
            cc.font = CN(); cc.border = bd
            cc.alignment = Alignment(vertical='center', horizontal='left' if j == 2 else 'center')
        r += 1
    # 本页合计
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
    ws.cell(row=r, column=1, value='本页合计：').font = CN(bold=True)
    ws.cell(row=r, column=5, value=f'数量：{-t_qty}').font = CN(bold=True)
    ws.cell(row=r, column=7, value=f'金额：{-t_amt:.2f}').font = CN(bold=True)
    for j in range(1, 8):
        ws.cell(row=r, column=j).border = bd
    r += 1
    # 总金额
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    ws.cell(row=r, column=1, value=f'总金额：￥{-t_amt:.2f}元（人民币大写：{rmb_upper(t_amt)}）').font = CN(bold=True)
    for j in range(1, 8):
        ws.cell(row=r, column=j).border = bd
    r += 1
    # 签字区
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    ws.cell(row=r, column=1, value='采购员：____________    操作员：____________').font = CN()
    for j in range(1, 8):
        ws.cell(row=r, column=j).border = bd
    ws.row_dimensions[r].height = 22
    for j, w in enumerate([6, 20, 20, 8, 10, 12, 14], 1):
        ws.column_dimensions[chr(64 + j)].width = w
    # V11.27: 审批通过 → 盖章领导预录签名
    stamp_leader_sign(ws, r, 'requisition', rq['id'])
    # V11.175: 底部追加驳回审批记录
    append_reject_rows(ws, r, 'requisition', rq['id'], ncols=7, CN=CN)
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    resp = send_file(bio, as_attachment=True, download_name=f'{rq["req_no"]}出库单.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    return resp

@app.route('/api/requisitions')
@login_required
def api_requisitions():
    """出库单列表 — V11.64: 库管员/领导/部门负责人可见; 员工只看自己的; 采购员/财务不看(库房的事)"""
    role = session.get('user_role')
    if role in ('采购员', '财务'):
        return jsonify([])
    conn = db()
    if role in ('员工',):
        rows = conn.execute("SELECT * FROM requisitions WHERE requester=? ORDER BY id DESC LIMIT 100", (session.get('user_name', ''),)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM requisitions ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        cnt = conn.execute("SELECT COUNT(*), COALESCE(SUM(quantity),0) FROM requisition_items WHERE requisition_id=?", (r['id'],)).fetchone()
        d['item_count'] = cnt[0] or 1
        d['total_qty'] = cnt[1] or r['quantity']
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/requisitions/<int:rid>')
@login_required
def api_requisition_detail(rid):
    """V11.172: 出库单详情(审批页"详情"按钮) — 单据+明细行"""
    role = session.get('user_role')
    if role in ('采购员', '财务'):
        return jsonify({'error': '无权限'}), 403
    conn = db()
    r = conn.execute("SELECT * FROM requisitions WHERE id=?", (rid,)).fetchone()
    if not r:
        conn.close(); return jsonify({'error': '出库单不存在'}), 404
    items = conn.execute("SELECT * FROM requisition_items WHERE requisition_id=? ORDER BY id", (rid,)).fetchall()
    conn.close()
    d = dict_row(r)
    d['items'] = [dict_row(x) for x in items]
    return jsonify({'requisition': d, 'items': [dict_row(x) for x in items]})

@app.route('/api/requisitions', methods=['POST'])
@login_required
def api_create_requisition():
    """V5.0: 新建出库单(批量商品) — 提交走审批, 审批通过自动扣减库存(展示为负值)
    V11.159: 出库单创建限 库管员/部门负责人/领导/管理员(职能分权: 采购员只采购不管理出库)"""
    if session.get('user_role') not in ('库管员', '部门负责人', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：出库管理仅限库管员/领导使用'}), 403
    d = request.json; conn = db()
    items = d.get('items') or []
    if not items and d.get('item_name'):
        items = [{'item_name': d.get('item_name'), 'spec': d.get('spec'), 'unit': d.get('unit', '个'),
                  'quantity': d.get('quantity', 0), 'purpose': d.get('purpose', '')}]
    items = [it for it in items if it.get('item_name') and float(it.get('quantity', 0) or 0) > 0]
    if not items:
        conn.close(); return jsonify({'error': '请至少填写一个商品及数量'}), 400
    # 库存校验(拦截超量) — V9.1: 名称+规格双条件匹配独立SKU
    for it in items:
        inv = conn.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? ORDER BY quantity DESC",
                           (it['item_name'], it.get('spec', '') or '')).fetchone()
        if not inv:
            conn.close(); return jsonify({'error': '库存中无此物资: %s %s' % (it['item_name'], it.get('spec', '') or '')}), 400
        if inv['quantity'] < float(it['quantity']):
            conn.close(); return jsonify({'error': '库存不足: %s(%s) 当前%s%s' % (it['item_name'], it.get('spec', '') or '', inv['quantity'], inv['unit'] or '个')}), 400
    no = gen_no('CK', 'requisitions', 'req_no', conn)
    total_q = sum(float(it['quantity']) for it in items)
    first = items[0]
    # V11.25: 领取人/领取部门 — 出库留痕可追溯(丢了东西能找到人)
    receiver = (d.get('receiver') or '').strip()
    receive_dept = (d.get('receive_dept') or '').strip()
    if not receiver:
        receiver = session['user_name']
    conn.execute("INSERT INTO requisitions(req_no,dept,requester,item_name,spec,quantity,unit,purpose,status,receiver,receive_dept,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                 (no, d.get('dept', ''), session['user_name'], first['item_name'], first.get('spec', ''),
                  total_q, first.get('unit', '个'), d.get('purpose', first.get('purpose', '')), '待审批', receiver, receive_dept, now()))
    rid = conn.execute("SELECT id FROM requisitions WHERE req_no=?", (no,)).fetchone()[0]
    for it in items:
        conn.execute("INSERT INTO requisition_items(requisition_id,item_name,spec,unit,quantity,purpose,created_at) VALUES(?,?,?,?,?,?,?)",
                     (rid, it['item_name'], it.get('spec', ''), it.get('unit', '个'),
                      float(it['quantity']), it.get('purpose', d.get('purpose', '')), now()))
    conn.commit()
    create_approvals('requisition', rid, 0, submitter=session['user_name'])
    conn.close()
    try: start_instances('requisition', rid)
    except Exception as e: print('requisition start_instances err:', e)
    log(session['user_name'], '新建出库单', f'{no} {len(items)}项 {total_q}件 待审批')
    return jsonify({'success': True, 'req_no': no, 'id': rid, 'message': f'出库单 {no} 已提交审批，审批通过后自动扣减库存'})


@app.route('/api/receivings', methods=['POST'])
@login_required
def api_create_receiving():
    """新建入库单(不依赖订单): 多商品明细, 提交走审批, 审批通过自动加库存"""
    d = request.json; conn = db()
    items = d.get('items') or []
    if not items and d.get('item_name'):
        items = [{'item_name': d.get('item_name'), 'spec': d.get('spec'), 'unit': d.get('unit', '个'),
                  'quantity': d.get('quantity', 0), 'price': d.get('price', 0), 'tax_rate': d.get('tax_rate', 13)}]
    items = [it for it in items if it.get('item_name') and float(it.get('quantity', 0) or 0) > 0]
    if not items:
        conn.close(); return jsonify({'error': '请至少填写一个商品及数量'}), 400
    no = gen_no('RK', 'receivings', 'receive_no', conn)
    total_q = sum(float(it['quantity']) for it in items)
    first = items[0]
    try: conn.execute("ALTER TABLE receivings ADD COLUMN items_json TEXT DEFAULT ''")
    except Exception: pass
    # V11.26: 验收照片/视频附件(责任留证)
    _atts = d.get('attachments') or []
    _atts_json = json.dumps([str(a) for a in _atts if a], ensure_ascii=False) if _atts else ''
    # V11.29: 归属部门(必选, 按部门分类入库)
    _dept = (d.get('dept') or '').strip()[:30]
    # V11.54: 入库单默认"暂估"(货到发票未到), 采购收到发票后红冲转正式
    _is_est = 1 if d.get('is_est', True) else 0
    _est_amt = round(float(d.get('est_amount') or 0), 2)
    # V11.172: 经办人(inspector)必须写当前登录用户 — 否则fs_biz_info取发起人返回'系统',
    # 钉钉发起时兜底成审批人自己(发起人=审批人) → 820003审批实例参数错误
    _inspector = (d.get('inspector') or '').strip() or session.get('user_name', '') or '系统'
    conn.execute("INSERT INTO receivings(receive_no,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,items_json,attachments,dept,is_est,est_amount,inspector) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                 (no, d.get('order_id'), first['item_name'], first.get('spec', ''), total_q,
                  first.get('unit', '个'), 0, '待审批', now(), '手动入库单: %d项商品' % len(items),
                  json.dumps(items, ensure_ascii=False), _atts_json, _dept, _is_est, _est_amt, _inspector))
    rid = conn.execute("SELECT id FROM receivings WHERE receive_no=?", (no,)).fetchone()[0]
    # 手动入库单没有 order_items, 明细暂存 remark; 审批通过时按 quantity 入库
    conn.commit()
    create_approvals('receiving', rid, 0, submitter=session['user_name'])
    conn.close()
    try: start_instances('receiving', rid)
    except Exception as e: print('receiving start_instances err:', e)
    log(session['user_name'], '新建入库单', f'{no} {len(items)}项 {total_q}件 待审批')
    return jsonify({'success': True, 'receive_no': no, 'message': f'入库单 {no} 已提交审批，审批通过后自动增加库存'})

@app.route('/api/inventory')
@login_required
def api_inventory():
    """V55需求3: 库存列表含 不含税单价/税率/含税单价/合计 计算字段
    V6: 带出供应商(入库时写入, 存量空值回填自订单历史)"""
    cat = request.args.get('cat', '')
    conn = db()
    # 存量回填: supplier 为空时从订单历史取最近供应商
    conn.execute("""UPDATE inventory SET supplier=(SELECT po.supplier FROM purchase_orders po
                    WHERE po.item_name=inventory.item_name AND po.supplier!='' ORDER BY po.id DESC LIMIT 1)
                    WHERE (supplier IS NULL OR supplier='') AND item_name IN
                    (SELECT DISTINCT item_name FROM purchase_orders WHERE supplier!='')""")
    conn.commit()
    if cat:
        rows = conn.execute("SELECT i.*,c.name as cat_name FROM inventory i LEFT JOIN categories c ON i.cat_code=c.code WHERE i.cat_code=? ORDER BY i.id", (cat,)).fetchall()
    else:
        rows = conn.execute("SELECT i.*,c.name as cat_name FROM inventory i LEFT JOIN categories c ON i.cat_code=c.code ORDER BY i.id").fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict_row(r)
        # V11.64: 库管员/员工/部门负责人 看库存数量但价格脱敏(防利益)
        if not can_see_price():
            d['price'] = None
        tr = float(d.get('tax_rate') or 13)
        price0 = d.get('price') or 0
        d['price_taxed'] = round(price0 * (1 + tr / 100.0), 4)
        d['total_untaxed'] = round((d.get('quantity') or 0) * price0, 2)
        d['total_taxed'] = round((d.get('quantity') or 0) * d['price_taxed'], 2)
        out.append(d)
    return jsonify(out)

@app.route('/api/inventory/recon')
@login_required
def api_inventory_recon():
    """V11.206b 库存对账勾稽(模块四4.3) — 财务/领导/库房核对
    勾稽公式: 当前库存 = 正式入库+暂估入库(分批标注) - 出库 - 退供应商 + 领用退回 + 报溢 - 报损
    按 物资名称+规格(不同规格分开统计, 不合并) 汇总流水, 与库存表当前量比对, 差额标红提示
    """
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理', '财务', '库管员', '部门负责人'):
        return jsonify({'error': '无权限'}), 403
    conn = db()
    inv_rows = conn.execute("SELECT id,item_name,spec,unit,quantity,warehouse FROM inventory ORDER BY item_name,spec").fetchall()
    fl_rows = conn.execute("SELECT item_name,spec,flow_type,qty FROM inventory_flows").fetchall()
    conn.close()
    from collections import defaultdict
    flow = defaultdict(float)
    for f in fl_rows:
        key = (f['item_name'] or '', f['spec'] or '')
        flow[key] += float(f['qty'] or 0)  # 入库+ / 出库-(负数已存)
    out = []
    total_cur = total_flow = total_diff = 0
    for r in inv_rows:
        key = (r['item_name'] or '', r['spec'] or '')
        cur = float(r['quantity'] or 0)
        fl = flow.get(key, 0.0)
        diff = round(cur - fl, 4)
        total_cur += cur; total_flow += fl
        if abs(diff) > 0.001: total_diff += 1
        out.append({'item_name': key[0], 'spec': key[1], 'unit': r['unit'] or '个',
                    'cur_qty': cur, 'flow_qty': fl, 'diff': diff, 'ok': abs(diff) <= 0.001,
                    'warehouse': r['warehouse'] or '主库房'})
    orphan = [{'item_name': k[0], 'spec': k[1], 'flow_qty': round(v, 2)} for k, v in flow.items()
              if not any(x['item_name'] == k[0] and x['spec'] == k[1] for x in out) and abs(v) > 0.001]
    return jsonify({'items': out, 'orphans': orphan[:20], 'summary': {
        'cur_total': round(total_cur, 2), 'flow_total': round(total_flow, 2),
        'diff_items': total_diff, 'item_count': len(out), 'orphan_count': len(orphan)}})

@app.route('/api/inventory/stock-check')
@login_required
def api_inventory_stock_check():
    """V11.38: 申请单填物资时查库存 — 按物资名称(支持规格模糊)返回库存情况"""
    name = (request.args.get('name') or '').strip()
    if not name:
        return jsonify({'error': '缺少物资名称'}), 400
    conn = db()
    rows = conn.execute("SELECT * FROM inventory WHERE item_name=? ORDER BY id", (name,)).fetchall()
    if not rows:
        # 模糊匹配(名称包含)
        rows = conn.execute("SELECT * FROM inventory WHERE item_name LIKE ? ORDER BY id LIMIT 5", ('%' + name + '%',)).fetchall()
    conn.close()
    out = [{'id': r['id'], 'item_name': r['item_name'], 'spec': r['spec'] or '', 'unit': r['unit'] or '个',
            'quantity': r['quantity'] or 0, 'warehouse': r['warehouse'] or '主库房',
            'supplier': r['supplier'] or '', 'cat_name': ''} for r in rows]
    total = sum(float(x['quantity']) for x in out)
    return jsonify({'matches': out, 'total_qty': total, 'found': len(out) > 0})

@app.route('/api/logs')
@login_required
def api_logs():
    # V11.64: 操作日志仅管理员/领导可见(审计记录敏感)
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理'):
        return jsonify([])
    conn = db(); rows = conn.execute("SELECT * FROM logs ORDER BY id DESC LIMIT 100").fetchall(); conn.close()
    return jsonify([dict_row(r) for r in rows])

# ============================================================
# V55 ── 预警中心: 14类预警引擎 + 首页集中展示 + 已处理/日志/配置
# ============================================================
def build_alerts():
    """重建全部待处理预警(幂等): 先清pending再按规则扫描插入; 已处理记录保留作预警日志"""
    c = db()
    c.execute("DELETE FROM alert_items WHERE status='pending'")
    warn_arrive = int(cfg_get('warn_arrive_days', '3') or 3)
    warn_contract = int(cfg_get('warn_contract_days', '30') or 30)
    warn_pay = int(cfg_get('warn_pay_days', '3') or 3)
    warn_approve = int(cfg_get('warn_approve_hours', '24') or 24)
    warn_idle = int(cfg_get('warn_idle_days', '90') or 90)
    warns = []
    def add(at, lv, title, content, btype='', bid=0):
        warns.append((at, lv, title, content, btype, bid))
    # 1) 库存缺货
    for r in c.execute("SELECT * FROM inventory WHERE safe_stock>0 AND quantity<=safe_stock"):
        lv = 'red' if r['quantity'] <= 0 else 'orange'
        add('库存缺货', lv, '缺货: %s' % r['item_name'],
            '当前库存 %s%s，低于安全库存 %s%s，请尽快发起采购申请' % (r['quantity'], r['unit'], r['safe_stock'], r['unit']), 'inventory', r['id'])
    # 2) 库存超储
    for r in c.execute("SELECT * FROM inventory WHERE max_stock>0 AND quantity>=max_stock"):
        add('库存超储', 'orange', '超储: %s' % r['item_name'],
            '当前库存 %s%s，已达库存上限 %s%s，避免过量囤积占用资金' % (r['quantity'], r['unit'], r['max_stock'], r['unit']), 'inventory', r['id'])
    # 3) 呆滞库存
    for r in c.execute("SELECT * FROM inventory WHERE quantity>0 AND last_move_date!='' AND last_move_date < date('now',?)", ('-%d days' % warn_idle,)):
        add('呆滞库存', 'orange', '呆滞: %s' % r['item_name'],
            '库存 %s%s，已超过%d天无出入库记录，建议盘点处理' % (r['quantity'], r['unit'], warn_idle), 'inventory', r['id'])
    # 3.5) 物料临期(有保质期字段的物资, 提前30天预警优先领用)
    for r in c.execute("SELECT * FROM inventory WHERE expiry_date!='' AND expiry_date<=date('now','+30 days') AND expiry_date>=date('now')"):
        add('物料临期', 'orange', '临期: %s' % r['item_name'],
            '保质期至 %s，请优先领用或尽快处理' % r['expiry_date'], 'inventory', r['id'])
    # 4) 订单待到货(临近交货日)
    for r in c.execute("SELECT * FROM purchase_orders WHERE status NOT IN ('已完成','已关闭','已核销','已入库') AND target_date!='' AND target_date<=date('now',?) AND target_date>=date('now')", ('+%d days' % warn_arrive,)):
        add('待到货', 'orange', '待到货: %s' % r['order_no'],
            '%s x%s%s，约定交货日 %s，请跟进供应商发货' % (r['item_name'], r['quantity'], r['unit'], r['target_date']), 'order', r['id'])
    # 5) 订单超期未到货
    for r in c.execute("SELECT * FROM purchase_orders WHERE status NOT IN ('已完成','已关闭','已核销','已入库') AND target_date!='' AND target_date<date('now')"):
        add('超期未到货', 'red', '超期未到货: %s' % r['order_no'],
            '%s 约定 %s 到货，已超期，请立即联系供应商' % (r['item_name'], r['target_date']), 'order', r['id'])
    # 6) 合同到期
    for r in c.execute("SELECT * FROM contracts WHERE status IN ('执行中','待审批') AND end_date!='' AND end_date<=date('now',?) AND end_date>=date('now')", ('+%d days' % warn_contract,)):
        add('合同到期', 'orange', '合同到期: %s' % r['contract_no'],
            '%s 将于 %s 到期，请评估是否续签' % (r['contract_name'] or '', r['end_date']), 'contract', r['id'])
    # 7) 合同超期未归档
    for r in c.execute("SELECT * FROM contracts WHERE status='执行中' AND (file_path IS NULL OR file_path='')"):
        add('合同未归档', 'orange', '合同未归档: %s' % r['contract_no'],
            '交易已进行但合同文件未上传归档，请经办人补齐资料', 'contract', r['id'])
    # 8) 应付款到期
    for r in c.execute("SELECT * FROM payment_requests WHERE status IN ('待审批','已通过') AND expect_pay_date!='' AND expect_pay_date<=date('now',?) AND expect_pay_date>=date('now')", ('+%d days' % warn_pay,)):
        add('应付款到期', 'orange', '应付款到期: %s' % r['payment_no'],
            '事由: %s，金额 ¥%s，期望付款日 %s' % (r['payment_reason'] or '', '%.2f' % (r['amount'] or 0), r['expect_pay_date']), 'payment', r['id'])
    # 9) 应付款逾期
    for r in c.execute("SELECT * FROM payment_requests WHERE status IN ('待审批','已通过') AND expect_pay_date!='' AND expect_pay_date<date('now')"):
        add('应付款逾期', 'red', '应付款逾期: %s' % r['payment_no'],
            '应于 %s 付款 ¥%s，已逾期，请尽快安排资金' % (r['expect_pay_date'], '%.2f' % (r['amount'] or 0)), 'payment', r['id'])
    # 10) 预付款跟踪(先款后货已付款但订单未核销)
    for r in c.execute("SELECT * FROM payment_requests WHERE trade_mode='先款后货' AND status='已付款'"):
        if r['contract_id']:
            ct2 = c.execute("SELECT order_id FROM contracts WHERE id=?", (r['contract_id'],)).fetchone()
            if ct2 and ct2['order_id']:
                po = c.execute("SELECT * FROM purchase_orders WHERE id=? AND status NOT IN ('已核销','已完成')", (ct2['order_id'],)).fetchone()
                if po:
                    add('预付款跟踪', 'red', '预付款未核销: %s' % r['payment_no'],
                        '已付款 ¥%s，对应订单 %s 长期未入库核销，请跟进' % ('%.2f' % (r['amount'] or 0), po['order_no']), 'payment', r['id'])
    # 11) 审批超时
    for r in c.execute("SELECT * FROM approval_instances WHERE status='pending' AND created_at<=datetime('now','localtime',?)", ('-%d hours' % warn_approve,)):
        add('审批超时', 'orange', '审批超时: %s#%s' % (r['biz_type'], r['biz_id']),
            '等待 %s 审批已超过%d小时，请尽快处理' % (r['role'] or '', warn_approve), r['biz_type'], r['biz_id'])
    # 12) 加急审批(标红置顶专区)
    for t, tbl, no_col, name_col in [('purchase_request','purchase_requests','req_no','purpose'),
                                      ('purchase_order','purchase_orders','order_no','item_name'),
                                      ('contract','contracts','contract_no','contract_name'),
                                      ('payment','payment_requests','payment_no','payment_reason')]:
        for r in c.execute("SELECT * FROM %s WHERE urgent=1 AND status IN ('待审批','审批通过','已通过','执行中')" % tbl):
            add('加急审批', 'red', '加急: %s' % r[no_col], '【加急】%s 请优先审批处理' % (r[name_col] or ''), t, r['id'])
    # 13) 多次驳回汇总
    for r in c.execute("SELECT biz_type, biz_id, COUNT(*) cnt FROM approval_instances WHERE status='rejected' GROUP BY biz_type, biz_id HAVING COUNT(*)>=2"):
        add('多次驳回', 'orange', '多次驳回: %s#%s' % (r['biz_type'], r['biz_id']),
            '该单据已被驳回 %d 次，请核对资料后整改重提' % r['cnt'], r['biz_type'], r['biz_id'])
    # 14) 供应商风险
    for r in c.execute("SELECT s.id, s.name, COUNT(d.id) bad FROM suppliers s LEFT JOIN deliveries d ON d.supplier=s.name AND d.sign_status='待签收' AND d.delivery_date<date('now') GROUP BY s.id HAVING bad>=2"):
        add('供应商风险', 'orange', '供应商延迟: %s' % r['name'],
            '该供应商有 %d 笔送货超期未签收，注意供货风险' % r['bad'], 'supplier', r['id'])
    for r in c.execute("SELECT id, name FROM suppliers WHERE created_at!='' AND created_at<=datetime('now','localtime','-30 days')"):
        q = c.execute("SELECT COUNT(*) FROM price_comparisons WHERE supplier=?", (r['name'],)).fetchone()[0]
        if q == 0:
            add('供应商风险', 'green', '供应商沉默: %s' % r['name'], '已超30天无报价记录，建议评估供应商活跃度', 'supplier', r['id'])
    for w in warns:
        c.execute("INSERT INTO alert_items(alert_type,level,title,content,biz_type,biz_id) VALUES(?,?,?,?,?,?)", w)
    c.commit(); c.close()

@app.route('/api/alerts')
@login_required
def api_alerts():
    """V55: 预警中心-全部待处理预警(首页集中展示)"""
    build_alerts()
    conn = db()
    rows = conn.execute("SELECT * FROM alert_items WHERE status='pending' ORDER BY CASE level WHEN 'red' THEN 0 WHEN 'orange' THEN 1 ELSE 2 END, id DESC LIMIT 50").fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/alerts/<int:aid>/process', methods=['POST'])
@login_required
def api_process_alert(aid):
    """V55: 手动标记预警已处理(处理后从预警列表清除, 留日志)"""
    conn = db()
    r = conn.execute("SELECT * FROM alert_items WHERE id=?", (aid,)).fetchone()
    if not r: conn.close(); return jsonify({'error': '预警不存在'}), 404
    conn.execute("UPDATE alert_items SET status='processed', processed_at=?, processed_by=? WHERE id=?", (now(), session['user_name'], aid))
    conn.commit(); conn.close()
    log(session['user_name'], '处理预警', '%s' % r['title'])
    return jsonify({'success': True})

@app.route('/api/alerts/log')
@login_required
def api_alerts_log():
    """V55: 预警历史日志(含已处理)"""
    conn = db()
    rows = conn.execute("SELECT * FROM alert_items ORDER BY id DESC LIMIT 100").fetchall()
    conn.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/alerts/config', methods=['GET', 'POST'])
@login_required
def api_alerts_config():
    """V55: 预警参数后台配置(是否开启/提前天数/超时小时/呆滞天数)"""
    if request.method == 'POST':
        if not can_manage_config(): return jsonify({'error': '仅系统管理员可配置'}), 403
        d = request.json or {}
        for k in ('warn_enabled', 'warn_arrive_days', 'warn_contract_days', 'warn_pay_days', 'warn_approve_hours', 'warn_idle_days'):
            if k in d: cfg_set(k, d[k])
        changed = [k for k in ('warn_enabled','warn_arrive_days','warn_contract_days','warn_pay_days','warn_approve_hours','warn_idle_days') if k in d]
        log(session.get('user_name',''), '修改预警参数', '变更项: %s' % (','.join(changed) or '无'))
        return jsonify({'success': True})
    return jsonify({k: cfg_get(k, defv) for k, defv in [
        ('warn_enabled', '1'), ('warn_arrive_days', '3'), ('warn_contract_days', '30'),
        ('warn_pay_days', '3'), ('warn_approve_hours', '24'), ('warn_idle_days', '90')]})

# ============================================================
# V4.1 ── 飞书 API: 回调 + 管理配置
# ============================================================
@app.route('/api/feishu/callback', methods=['GET', 'POST'])
def api_feishu_callback():
    """飞书事件订阅回调: GET为URL验证, POST为事件推送(审批结果同步)"""
    if request.method == 'GET':
        token = request.args.get('token', '')
        vt = cfg_get('feishu_verification_token')
        if vt and token and token != vt: return 'invalid token', 403
        return jsonify({'challenge': request.args.get('challenge', '')})
    try:
        body = request.get_json(force=True, silent=True) or {}
    except Exception:
        body = {}
    if body.get('type') == 'url_verification':
        return jsonify({'challenge': body.get('challenge', '')})
    try:
        if body.get('encrypt'):
            body = fs_decrypt(body)
    except Exception as e:
        log('系统', '飞书回调解密失败', str(e))
        return jsonify({'code': 0})  # 返回成功, 避免飞书重试轰炸
    header = body.get('header', {}) or {}
    vt = cfg_get('feishu_verification_token')
    if vt and header.get('token') and header['token'] != vt:
        return jsonify({'code': 0})
    event = body.get('event', {}) or {}
    etype = header.get('event_type', '') or event.get('type', '')
    if 'approval' in etype:  # 审批实例状态变更: APPROVED/REJECTED/CANCELED...
        code = event.get('instance_code', '')
        status = event.get('status', '')
        if code and status:
            if status == 'APPROVED': fs_sync_result(code, 'approved')
            elif status == 'REJECTED': fs_sync_result(code, 'rejected')
            elif status in ('CANCELED', 'REVOKED', 'RECALLED', 'TERMINATED'): fs_mark_cancelled(code)
    return jsonify({'code': 0})

@app.route('/api/feishu/config', methods=['GET', 'POST'])
@login_required
def api_feishu_config():
    if request.method == 'POST':
        if not can_manage_config(): return jsonify({'error': '仅系统管理员可配置'}), 403
        d = request.json or {}
        for k in ('feishu_app_id','feishu_app_secret','feishu_verification_token','feishu_encrypt_key',
                  'feishu_report_chat','feishu_approve_hours','feishu_order_days'):
            if k in d: cfg_set(k, d[k])
        if 'feishu_enabled' in d: cfg_set('feishu_enabled', '1' if d['feishu_enabled'] else '0')
        changed = [k for k in ('feishu_app_id','feishu_app_secret','feishu_verification_token','feishu_encrypt_key',
                  'feishu_report_chat','feishu_approve_hours','feishu_order_days') if k in d]
        if 'feishu_enabled' in d: changed.append('feishu_enabled')
        log(session.get('user_name',''), '修改飞书配置', '变更项: %s' % (','.join(changed) or '无'))
        return jsonify({'success': True})
    c = db()
    rows = c.execute("SELECT key,value FROM sys_config WHERE key LIKE 'feishu_%'").fetchall()
    c.close()
    cfg = {r['key']: r['value'] for r in rows}
    cfg['feishu_approval_codes'] = fs_approval_codes()
    return jsonify(cfg)

@app.route('/api/feishu/test', methods=['POST'])
@login_required
def api_feishu_test():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    try:
        tk = fs_token()
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})
    target = (request.json or {}).get('open_id', '')
    if target:
        ok = fs_send(target, '✅ 采购系统飞书对接测试成功', 'green')
        return jsonify({'success': ok, 'token': tk[:20] + '...', 'msg': '测试消息已发送' if ok else '消息发送失败, 请检查机器人权限'})
    return jsonify({'success': True, 'token': tk[:20] + '...'})

@app.route('/api/feishu/init-definitions', methods=['POST'])
@login_required
def api_feishu_init_defs():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    return jsonify(fs_init_definitions())

@app.route('/api/feishu/bind', methods=['POST'])
@login_required
def api_feishu_bind():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    d = request.json or {}
    c = db()
    c.execute("UPDATE users SET feishu_open_id=? WHERE id=?", (str(d.get('open_id','')).strip(), int(d.get('user_id', 0))))
    c.commit(); c.close()
    return jsonify({'success': True})

@app.route('/api/feishu/lookup', methods=['POST'])
@login_required
def api_feishu_lookup():
    """按手机号查飞书open_id(需通讯录权限)"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    mobile = str((request.json or {}).get('mobile', '')).strip()
    if not mobile: return jsonify({'success': False, 'error': '请输入手机号'})
    code_r, resp = fs_post('/contact/v3/users/batch_get_id', {'mobiles': [mobile]})
    if code_r != 0: return jsonify({'success': False, 'error': resp.get('msg', '查询失败')})
    for u in resp.get('data', {}).get('user_list', []):
        if u.get('mobile') == mobile:
            return jsonify({'success': True, 'open_id': u.get('open_id', ''), 'user_id': u.get('user_id', ''), 'name': u.get('name', '')})
    return jsonify({'success': False, 'error': '未找到该手机号对应的飞书用户'})

@app.route('/api/feishu/status', methods=['GET'])
@login_required
def api_feishu_status():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    c = db()
    bound = c.execute("SELECT COUNT(*) FROM users WHERE feishu_open_id IS NOT NULL AND feishu_open_id!=''").fetchone()[0]
    insts = c.execute("SELECT COUNT(*) FROM feishu_instances").fetchone()[0]
    pend_sync = c.execute("SELECT COUNT(*) FROM feishu_instances WHERE status='pending'").fetchone()[0]
    c.close()
    return jsonify({
        'enabled': feishu_enabled(),
        'config_ok': bool(cfg_get('feishu_app_id') and cfg_get('feishu_app_secret')),
        'token_ok': bool(_TOKEN['t']),
        'bound_users': bound, 'instances': insts, 'pending_sync': pend_sync,
        'definitions': fs_approval_codes(),
    })

@app.route('/api/feishu/remind-now', methods=['POST'])
@login_required
def api_feishu_remind_now():
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    fs_send_reminders(force=True)
    return jsonify({'success': True})

@app.route('/api/feishu/sync-instances', methods=['POST'])
@login_required
def api_feishu_sync_instances():
    """把系统内所有待审批但未在飞书发起的单据, 补发到飞书"""
    if not can_manage_config(): return jsonify({'error': '仅系统管理员'}), 403
    c = db()
    rows = c.execute("SELECT DISTINCT biz_type, biz_id FROM approval_instances WHERE status='pending'").fetchall()
    c.close()
    n = 0
    for r in rows:
        if fs_start_instance(r['biz_type'], r['biz_id']): n += 1
    return jsonify({'success': True, 'pushed': n, 'total': len(rows)})

@app.route('/api/feishu/instances', methods=['GET'])
@login_required
def api_feishu_instances():
    c = db()
    rows = c.execute("SELECT * FROM feishu_instances ORDER BY id DESC LIMIT 30").fetchall()
    c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/public-url')
def api_public_url():
    """当前公网地址(由守护脚本写入 public_url.txt; 隧道重启地址变化后自动更新)"""
    f = os.path.join(BASE, 'data', 'public_url.txt')
    try:
        return jsonify({'url': open(f).read().strip()})
    except Exception:
        return jsonify({'url': ''})


@app.route('/api/notify-address-change', methods=['POST'])
def api_notify_address_change():
    """V11.146: 隧道地址变化通知 — tunnel_guard换地址后调用, 推送新地址给用户
    V11.147: 加30分钟冷却
    V11.148: 用户要求彻底停发地址变更通知 — 接口保留(守护进程仍调用), 但不再推送钉钉/飞书"""
    try:
        f = os.path.join(BASE, 'data', 'public_url.txt')
        url = ''
        try:
            url = open(f).read().strip()
        except Exception:
            pass
        if not url:
            return jsonify({'success': False, 'error': '无地址'})
        # 用户 2026-08-30 明确要求: 不再推送地址变更通知(免费隧道频繁重连, 刷屏打扰)
        # 只记日志, 不打扰用户
        log('系统', '地址变更', f'隧道地址已更新(通知已停发): {url}')
        return jsonify({'success': True, 'pushed': False, 'notify_disabled': True, 'url': url})
    except Exception as e:
        log('系统', '地址变更通知异常', str(e)[:120])
        return jsonify({'success': False, 'error': str(e)[:100]})

# ============================================================
# V4.3 ── 飞书统一认证 (OAuth 单点登录, 网页应用点开即用)
# ============================================================
@app.route('/api/feishu/oauth/url')
def api_feishu_oauth_url():
    """生成飞书授权URL(前端登录页跳转)"""
    app_id = cfg_get('feishu_app_id')
    if not app_id: return jsonify({'error': '飞书应用未配置, 请先在飞书设置填写 App ID'}), 400
    state = uuid.uuid4().hex[:16]
    session['oauth_state'] = state
    redirect_uri = request.url_root.rstrip('/') + '/api/feishu/oauth/callback'
    url = 'https://open.feishu.cn/open-apis/authen/v1/index?app_id=%s&redirect_uri=%s&state=%s' % (
        app_id, urllib.parse.quote(redirect_uri, safe=''), state)
    return jsonify({'url': url, 'redirect_uri': redirect_uri})

@app.route('/api/feishu/oauth/callback')
def api_feishu_oauth_callback():
    """飞书授权回调: code换token → 用户信息 → 匹配系统用户 → 登录"""
    code = request.args.get('code', '')
    state = request.args.get('state', '')
    if session.get('oauth_state') and state and state != session.get('oauth_state'):
        return 'state 校验失败, 请重新发起登录', 400
    if not code: return '缺少授权码', 400
    redirect_uri = request.url_root.rstrip('/') + '/api/feishu/oauth/callback'
    body = {'grant_type': 'authorization_code', 'client_id': cfg_get('feishu_app_id'),
            'client_secret': cfg_get('feishu_app_secret'), 'code': code, 'redirect_uri': redirect_uri}
    code_r, resp = fs_post('/authen/v2/oauth/token', body)
    if code_r != 0:
        return f'换取登录凭证失败: {resp.get("msg","")}', 400
    data = resp.get('data', {}) or {}
    utoken = data.get('access_token', '')
    open_id = data.get('open_id', '')
    ui = {}
    if utoken:
        try:
            req = urllib.request.Request(FS_API + '/authen/v1/user_info',
                                         headers={'Authorization': 'Bearer ' + utoken})
            with urllib.request.urlopen(req, timeout=10) as r:
                ui = json.loads(r.read().decode('utf-8')).get('data', {}) or {}
        except Exception:
            pass
    conn = db()
    u = conn.execute("SELECT * FROM users WHERE feishu_open_id=? AND is_active=1", (open_id,)).fetchone()
    if not u and ui.get('mobile'):
        u = conn.execute("SELECT * FROM users WHERE phone=? AND is_active=1", (ui['mobile'],)).fetchone()
    if not u:
        conn.close()
        return ('<html><body style="font-family:sans-serif;padding:40px;text-align:center">'
                f'<h3>⚠️ 飞书账号（{ui.get("name","")}）未绑定系统用户</h3>'
                '<p>请联系系统管理员，在「飞书设置 → 用户飞书绑定」中绑定你的账号</p>'
                '<p><a href="/">返回系统首页</a></p></body></html>'), 403
    session['user_id'] = u['id']; session['username'] = u['username']; session['user_name'] = u['name']; session['user_role'] = u['role']
    session['dept_id'] = u['dept_id']; session['dept_name'] = ''
    conn.close()
    log(u['name'], '飞书统一认证登录', f"open_id={open_id} 手机号={ui.get('mobile','')}")
    return redirect('/')

# ============================================================
# V4.2 ── 报表中心 + 库存盘点 (好生意刚需版补齐)
# ============================================================
@app.route('/api/reports')
@login_required
def api_reports():
    c = db()
    months = c.execute("""SELECT strftime('%Y-%m',created_at) m, SUM(total_amount) s
        FROM purchase_orders WHERE created_at >= datetime('now','localtime','-11 months')
        GROUP BY m ORDER BY m""").fetchall()
    sups = c.execute("""SELECT supplier, COUNT(*) cnt, SUM(total_amount) s FROM purchase_orders
        WHERE supplier!='' GROUP BY supplier ORDER BY s DESC LIMIT 10""").fetchall()
    cats = c.execute("""SELECT category, SUM(total_amount) s FROM purchase_orders
        WHERE category!='' GROUP BY category ORDER BY s DESC""").fetchall()
    todo = {
        '待审批': c.execute("SELECT COUNT(*) FROM approval_instances WHERE status='pending'").fetchone()[0],
        '待签收': c.execute("SELECT COUNT(*) FROM deliveries WHERE sign_status='待签收'").fetchone()[0],
        '待检验': c.execute("SELECT COUNT(*) FROM receivings WHERE status='待检验'").fetchone()[0],
        '待付款': c.execute("SELECT COUNT(*) FROM payment_requests WHERE status='待审批'").fetchone()[0],
        '库存预警': c.execute("SELECT COUNT(*) FROM inventory WHERE quantity<=safe_stock AND safe_stock>0").fetchone()[0],
    }
    c.close()
    return jsonify({'monthly': [dict_row(r) for r in months],
                    'suppliers': [dict_row(r) for r in sups],
                    'categories': [dict_row(r) for r in cats],
                    'todo': todo})

@app.route('/api/counts')
@login_required
def api_counts():
    """V11.159: 库存盘点仅 库管员/领导/管理员 可见(采购员不参与盘点)"""
    if session.get('user_role') not in ('库管员', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    c = db()
    rows = c.execute("SELECT * FROM inventory_counts ORDER BY id DESC LIMIT 20").fetchall()
    c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/counts', methods=['POST'])
@login_required
def api_create_count():
    """V11.159: 创建盘点仅 库管员/领导/管理员"""
    if session.get('user_role') not in ('库管员', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：库存盘点仅限库管员/领导使用'}), 403
    d = request.json
    c = db()
    no = gen_no('PD', 'inventory_counts', 'count_no')
    c.execute("INSERT INTO inventory_counts(count_no,status,remark) VALUES(?,?,?)", (no, '盘点中', d.get('remark','')))
    cid = c.execute("SELECT id FROM inventory_counts WHERE count_no=?", (no,)).fetchone()[0]
    items = c.execute("SELECT * FROM inventory").fetchall()
    for it in items:
        c.execute("INSERT INTO inventory_count_items(count_id,inventory_id,item_name,book_qty,actual_qty) VALUES(?,?,?,?,?)",
                  (cid, it['id'], it['item_name'], it['quantity'], it['quantity']))
    c.commit(); c.close()
    log(session['user_name'], '开始盘点', f'{no} 共{len(items)}项')
    return jsonify({'success': True, 'count_no': no, 'id': cid, 'items': len(items)})

@app.route('/api/counts/<int:cid>')
@login_required
def api_count(cid):
    c = db()
    ct = c.execute("SELECT * FROM inventory_counts WHERE id=?", (cid,)).fetchone()
    items = c.execute("SELECT * FROM inventory_count_items WHERE count_id=?", (cid,)).fetchall()
    c.close()
    return jsonify({'count': dict_row(ct), 'items': [dict_row(i) for i in items]})

@app.route('/api/counts/<int:cid>/finish', methods=['POST'])
@login_required
def api_finish_count(cid):
    """完成盘点: 差异项自动生成报溢/报损单(待审批), 审批通过后才调库存(V11.34 账实相符闭环)
    V11.159: 限 库管员/领导/管理员"""
    if session.get('user_role') not in ('库管员', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：库存盘点仅限库管员/领导使用'}), 403
    d = request.json
    c = db()
    diffs = 0
    for it in (d.get('items') or []):
        row = c.execute("SELECT * FROM inventory_count_items WHERE id=?", (it.get('id'),)).fetchone()
        if not row: continue
        aq = float(it.get('actual_qty', row['book_qty']))
        c.execute("UPDATE inventory_count_items SET actual_qty=? WHERE id=?", (aq, row['id']))
        if abs(aq - row['book_qty']) > 0.001:
            diffs += 1
            # V11.34: 差异 → 报溢/报损单(待审批), 不再直接改库存
            diff_qty = aq - row['book_qty']
            adj_type = '报溢' if diff_qty > 0 else '报损'
            inv = c.execute("SELECT * FROM inventory WHERE id=?", (row['inventory_id'],)).fetchone()
            adj_no = gen_no('SY' if diff_qty > 0 else 'BS', 'inventory_adjustments', 'adj_no')
            c.execute("""INSERT INTO inventory_adjustments(adj_no,adj_type,inventory_id,item_name,spec,unit,book_qty,adj_qty,reason,status,source,created_by)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""",
                (adj_no, adj_type, row['inventory_id'], row['item_name'], inv['spec'] if inv else '',
                 inv['unit'] if inv else '个', row['book_qty'], abs(diff_qty),
                 f'盘点差异自动生成: 账面{row["book_qty"]:g} 实盘{aq:g}', '待审批', '盘点', session['user_name']))
    c.execute("UPDATE inventory_counts SET status='已完成', finished_at=? WHERE id=?", (now(), cid))
    c.commit(); c.close()
    log(session['user_name'], '完成盘点', f'#{cid} 差异{diffs}项(已生成报溢/报损单待审批)')
    return jsonify({'success': True, 'diff_items': diffs})

# ============================================================
# V11.34 ── 报溢/报损单(库存账实相符)
# ============================================================
@app.route('/api/adjustments')
@login_required
def api_adjustments():
    """报溢/报损单列表 — V11.159: 仅 库管员/领导/管理员 可见"""
    if session.get('user_role') not in ('库管员', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    c = db()
    rows = c.execute("SELECT * FROM inventory_adjustments ORDER BY id DESC LIMIT 100").fetchall()
    c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/adjustments', methods=['POST'])
@login_required
def api_create_adjustment():
    """手动报溢/报损单(未盘点也可用: 到货多/物资损坏) — V11.159: 仅 库管员/领导/管理员"""
    if session.get('user_role') not in ('库管员', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：库存调整仅限库管员/领导使用'}), 403
    d = request.json
    inv_id = d.get('inventory_id')
    adj_type = d.get('adj_type', '')
    qty = float(d.get('adj_qty') or 0)
    reason = (d.get('reason') or '').strip()
    if adj_type not in ('报溢', '报损') or qty <= 0:
        return jsonify({'error': '请选择类型(报溢/报损)并填写数量'}), 400
    if not reason:
        return jsonify({'error': '请填写原因(审批依据)'}), 400
    c = db()
    inv = c.execute("SELECT * FROM inventory WHERE id=?", (inv_id,)).fetchone() if inv_id else None
    if not inv:
        c.close(); return jsonify({'error': '库存物资不存在'}), 404
    adj_no = gen_no('SY' if adj_type == '报溢' else 'BS', 'inventory_adjustments', 'adj_no')
    c.execute("""INSERT INTO inventory_adjustments(adj_no,adj_type,inventory_id,item_name,spec,unit,book_qty,adj_qty,reason,status,source,created_by)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""",
        (adj_no, adj_type, inv_id, inv['item_name'], inv['spec'] or '', inv['unit'] or '个',
         inv['quantity'] or 0, qty, reason, '待审批', '手动', session['user_name']))
    aid = c.execute("SELECT id FROM inventory_adjustments WHERE adj_no=?", (adj_no,)).fetchone()[0]
    c.commit(); c.close()
    log(session['user_name'], f'新建{adj_type}单', f'{adj_no} {inv["item_name"]} x{qty:g} {reason[:30]}')
    return jsonify({'success': True, 'adj_no': adj_no, 'id': aid})

@app.route('/api/inventory/<int:iid>/trace')
@login_required
def api_inventory_trace(iid):
    """V11.36: 库存溯源 — 该物资全部出入库/报溢报损流水"""
    c = db()
    inv = c.execute("SELECT * FROM inventory WHERE id=?", (iid,)).fetchone()
    if not inv:
        c.close(); return jsonify({'error': '库存物资不存在'}), 404
    flows = c.execute("""SELECT * FROM inventory_flows WHERE item_name=? AND (spec=? OR (?='' AND (spec='' OR spec IS NULL)))
        ORDER BY id DESC LIMIT 50""", (inv['item_name'], inv['spec'] or '', inv['spec'] or '')).fetchall()
    c.close()
    return jsonify({'inventory': dict_row(inv), 'flows': [dict_row(f) for f in flows]})

@app.route('/api/adjustments/<int:aid>/approve', methods=['POST'])
@login_required
def api_adjust_approve(aid):
    """审批报溢/报损单: 通过→调库存; 驳回→不改库存
    V11.159: 审批限 领导/管理员(库管员提交, 领导审批, 职能分离)"""
    if session.get('user_role') not in ('分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：审批仅限领导/管理员'}), 403
    d = request.json or {}
    action = d.get('action', 'approved')
    c = db()
    r = c.execute("SELECT * FROM inventory_adjustments WHERE id=?", (aid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '单据不存在'}), 404
    if r['status'] != '待审批':
        c.close(); return jsonify({'error': '该单据已处理'}), 400
    if action == 'rejected':
        c.execute("UPDATE inventory_adjustments SET status='已驳回' WHERE id=?", (aid,))
        c.commit(); c.close()
        log(session['user_name'], '驳回' + r['adj_type'], f"{r['adj_no']} {r['item_name']}")
        return jsonify({'success': True})
    # 通过 → 调库存
    if r['adj_type'] == '报溢':
        c.execute("UPDATE inventory SET quantity=quantity+? WHERE id=?", (r['adj_qty'], r['inventory_id']))
    else:
        # 报损: 库存不够则按现有量扣(不扣成负)
        inv = c.execute("SELECT quantity FROM inventory WHERE id=?", (r['inventory_id'],)).fetchone()
        cur = float(inv['quantity'] or 0) if inv else 0
        c.execute("UPDATE inventory SET quantity=? WHERE id=?", (max(0, cur - r['adj_qty']), r['inventory_id']))
    # V11.36: 报溢/报损写入库存流水(溯源完整闭环)
    _new_qty = c.execute("SELECT quantity FROM inventory WHERE id=?", (r['inventory_id'],)).fetchone()
    _bal = float(_new_qty[0] or 0) if _new_qty else 0
    c.execute("INSERT INTO inventory_flows(item_name,spec,unit,flow_type,doc_type,doc_id,doc_no,qty,balance_after,operator,remark,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
              (r['item_name'], r['spec'] or '', r['unit'] or '个', r['adj_type'], 'adjustment', aid, r['adj_no'],
               r['adj_qty'] if r['adj_type'] == '报溢' else -r['adj_qty'], _bal, session['user_name'],
               f'{r["adj_type"]}单{r["adj_no"]}审批通过', now()))
    c.execute("UPDATE inventory_adjustments SET status='已通过', approved_at=? WHERE id=?", (now(), aid))
    c.commit(); c.close()
    log(session['user_name'], f'{r["adj_type"]}审批通过', f"{r['adj_no']} {r['item_name']} x{r['adj_qty']:g}{r['unit']} 已调库存")
    return jsonify({'success': True})

# ============================================================
# V4.4 ── 需求文档补充: 加购下单/月度对账/合并开票/往来台账
# ============================================================
@app.route('/api/orders/from-requests', methods=['POST'])
@login_required
def api_orders_from_requests():
    """同批已通过申请合并为一张订单(多明细), 一次审批; 支持增补物资"""
    d = request.json
    req_ids = d.get('req_ids') or []
    tm = d.get('trade_mode', '货到付款')
    if tm not in ('货到付款', '先款后货'): tm = '货到付款'
    extras = d.get('extras') or []  # 增补物资: [{req_id, item_name, spec, quantity, unit, price}]
    if not req_ids:
        return jsonify({'error': '请选择待下单的采购申请'}), 400
    conn = db()
    rows = []
    used_reqs = []
    for rid in req_ids:
        pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (rid,)).fetchone()
        if not pr or pr['status'] != '已通过':
            continue
        if conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (rid,)).fetchone()[0] > 0:
            continue
        items = conn.execute("SELECT * FROM request_items WHERE req_id=?", (rid,)).fetchall()
        for it in items:
            rows.append((it['item_name'], it['spec'] or '', it['unit'] or '个', float(it['quantity']),
                         float(it['estimated_price'] or 0), rid))
        used_reqs.append(rid)
    for ex in extras:
        if not ex.get('req_id'): continue
        rows.append((ex.get('item_name',''), ex.get('spec','') or '', ex.get('unit','个') or '个',
                     float(ex.get('quantity',1)), float(ex.get('price',0) or 0), ex.get('req_id')))
    if not rows:
        conn.close(); return jsonify({'error': '没有可下单的申请（已通过且未下单）'}), 400
    no = gen_no('CG', 'purchase_orders', 'order_no', conn)
    total_qty = 0.0; grand_amt = 0.0; grand_tax = 0.0; grand_total = 0.0
    for it in rows:
        qty = it[3]; price = it[4]
        amt = qty*price; tax = amt*0.13
        total_qty += qty; grand_amt += amt; grand_tax += tax; grand_total += amt+tax
    first = rows[0]
    conn.execute("""INSERT INTO purchase_orders(order_no,req_id,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,
        supplier,requester,category,owner,owner_id,target_date,trade_mode,remark,urgent,attachments) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (no, used_reqs[0] if used_reqs else None, first[0], first[1], total_qty, first[2], first[4], grand_amt, 13, grand_tax, grand_total,
         d.get('supplier',''), session['user_name'], d.get('category','后勤类'), session['user_name'], session['user_id'],
         d.get('target_date',''), tm, '加购: 多申请合并下单', 1 if d.get('urgent') else 0, '[]'))
    oid = conn.execute("SELECT id FROM purchase_orders WHERE order_no=?", (no,)).fetchone()[0]
    for it in rows:
        qty = it[3]; price = it[4]; amt = qty*price; tax = amt*0.13
        conn.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
            (oid, it[0], it[1], it[2], qty, price, amt, 13, tax, amt+tax, ''))
    rno = None
    if tm == '货到付款':
        # V11.152e: 防重复 — 该订单已有任意未入库状态的入库单则不再生成
        _dup = conn.execute("SELECT id FROM receivings WHERE order_id=? AND status NOT IN ('已入库','已作废') LIMIT 1", (oid,)).fetchone()
        if not _dup:
            rno = gen_no('RK', 'receivings', 'receive_no', conn)
            # V11.152: 多物资订单自动生成入库时, 明细完整存items_json
            _items_json = json.dumps(
                [{'item_name': it[0], 'spec': it[1] or '', 'quantity': it[3], 'unit': it[2] or '个', 'price': it[4] or 0} for it in rows],
                ensure_ascii=False)
            _name = (first[0] + ' 等%d项' % len(rows)) if len(rows) > 1 else first[0]
            # V11.198: 自动生成的入库单不预设类型(is_est=0待定, 提交审批时手动选择暂估/正式)
            conn.execute("INSERT INTO receivings(receive_no,delivery_id,order_id,item_name,spec,quantity,unit,qualified_qty,status,received_at,remark,items_json,is_est) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,0)",
            (rno, None, oid, _name, '', total_qty, first[2], 0, '待入库', now(), '货到付款: 下单后自动进入入库板块(整批%d项)' % len(rows), _items_json))
    for rid in used_reqs:
        conn.execute("UPDATE purchase_requests SET status='已下单', updated_at=? WHERE id=?", (now(), rid))
    conn.commit()
    create_approvals('purchase_order', oid, grand_total, submitter=session['user_name'])   # 一张订单一次审批
    start_instances('purchase_order', oid)
    conn.close()
    log(session['user_name'], '加购下单', '%s 合并%d项商品 ¥%.0f 模式:%s' % (no, len(rows), grand_total, tm))
    return jsonify({'success': True, 'orders': [no], 'order_no': no, 'id': oid, 'receive_no': rno,
                    'item_count': len(rows), 'total_amount': grand_total})

@app.route('/api/settlements')
@login_required
def api_settlements():
    """V11.159: 月度对账仅 财务/领导/管理员 可见"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    c = db()
    rows = c.execute("SELECT * FROM settlements ORDER BY id DESC LIMIT 50").fetchall()
    c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/settlements', methods=['POST'])
@login_required
def api_create_settlement():
    """需求5-月度对账: 按合同/供应商批量生成对账单(取已入库订单金额)
    V11.159: 创建对账仅 财务/领导/管理员"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：月度对账仅限财务/领导'}), 403
    d = request.json
    period = d.get('period', '')       # 如 2026-07
    supplier = d.get('supplier', '')
    if not period:
        return jsonify({'error': '请选择对账月份'}), 400
    conn = db()
    # 该月份已入库/已挂账/已核销订单
    rows = conn.execute("""SELECT * FROM purchase_orders WHERE trade_mode='货到付款'
        AND status IN ('已入库','已挂账','已核销','已完成')
        AND strftime('%Y-%m', COALESCE(updated_at, created_at)) = ?
        AND (?='' OR supplier=?) ORDER BY id""", (period, supplier, supplier)).fetchall()
    if not rows:
        conn.close(); return jsonify({'error': '该月份没有符合条件的已入库订单'}), 400
    oids = [str(r['id']) for r in rows]
    total = sum(r['total_amount'] or 0 for r in rows)
    cnos = list({(c['contract_no'] or '') for c in [conn.execute("SELECT contract_no FROM contracts WHERE order_id=?", (r['id'],)).fetchone() for r in rows] if c})
    no = gen_no('DZ', 'settlements', 'settlement_no')
    conn.execute("INSERT INTO settlements(settlement_no,period,supplier,contract_ids,order_ids,total_amount,status,remark) VALUES(?,?,?,?,?,?,'待确认',?)",
                 (no, period, supplier, ','.join(cnos), ','.join(oids), total, d.get('remark','')))
    sid = conn.execute("SELECT id FROM settlements WHERE settlement_no=?", (no,)).fetchone()[0]
    conn.commit(); conn.close()
    log(session['user_name'], '生成对账单', f'{no} {period} ¥{total:.0f} 共{len(rows)}单')
    return jsonify({'success': True, 'settlement_no': no, 'id': sid, 'orders': len(rows), 'total': total})

@app.route('/api/settlements/<int:sid>')
@login_required
def api_settlement_detail(sid):
    """对账单详情: 对账单头 + 关联订单及明细"""
    conn = db()
    s = conn.execute("SELECT * FROM settlements WHERE id=?", (sid,)).fetchone()
    if not s:
        conn.close(); return jsonify({'error': '对账单不存在'}), 404
    oids = [int(x) for x in (s['order_ids'] or '').split(',') if x.strip().isdigit()]
    orders = []
    for oid in oids:
        o = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
        if not o: continue
        od = dict_row(o)
        od['items'] = [dict_row(i) for i in conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()]
        od['contract_no'] = ''
        ct = conn.execute("SELECT contract_no FROM contracts WHERE order_id=?", (oid,)).fetchone()
        if ct: od['contract_no'] = ct['contract_no']
        orders.append(od)
    conn.close()
    return jsonify({'settlement': dict_row(s), 'orders': orders})

@app.route('/api/settlements/<int:sid>/confirm', methods=['POST'])
@login_required
def api_confirm_settlement(sid):
    conn = db()
    st = conn.execute("SELECT * FROM settlements WHERE id=?", (sid,)).fetchone()
    if not st: conn.close(); return jsonify({'error': '对账单不存在'}), 404
    conn.execute("UPDATE settlements SET status='已确认', confirmed_at=? WHERE id=?", (now(), sid))
    conn.commit(); conn.close()
    log(session['user_name'], '确认对账单', f'#{sid}')
    return jsonify({'success': True})

# ============================================================
# V11.53 ── 月底三表(领导流程: 暂估/红冲/白入 导出Excel给财务)
# V11.56 ── 三表系统内展示 + 自动核对差异(不用人工核查)
# ============================================================
@app.route('/api/reports/est-view')
@login_required
def api_est_view():
    """三表数据+自动核对: 返回 暂估/红冲/白入 三组明细 + 差异核对结果"""
    conn = db()
    rows = conn.execute("SELECT * FROM receivings WHERE status IN ('已入库','已通过','审批通过') ORDER BY received_at DESC").fetchall()
    conn.close()
    est, hc, br = [], [], []
    for row in rows:
        is_est = bool(row['is_est']); has_inv = bool(row['invoice_no'])
        d = {
            'receive_no': row['receive_no'], 'item_name': row['item_name'], 'spec': row['spec'] or '',
            'quantity': row['quantity'], 'unit': row['unit'] or '个', 'amount': row['est_amount'] or 0,
            'invoice_no': row['invoice_no'] or '', 'date': str(row['received_at'] or '')[:10],
            'remark': row['remark'] or '', 'id': row['id'],
        }
        # V11.153: 红冲组补 发票金额/差价 (invoice_amount新列; 旧数据兼容)
        try:
            d['invoice_amount'] = float(row['invoice_amount'] or 0) if 'invoice_amount' in row.keys() else 0
        except Exception:
            d['invoice_amount'] = 0
        if has_inv and not d['invoice_amount']:
            d['invoice_amount'] = float(row['est_amount'] or 0)
        d['diff'] = round(d['invoice_amount'] - float(row['est_amount'] or 0), 2) if has_inv else 0
        if is_est and not has_inv: est.append(d)     # 暂估未红冲(发票没回)
        elif is_est and has_inv: hc.append(d)         # 已红冲
        else: br.append(d)                            # 正式入库
    # 自动核对(系统自己比对, 不用人工):
    checks = []
    # ① 暂估未红冲 = 发票没回, 要催采购
    if est:
        checks.append({'level': 'warn', 'msg': f'⚠️ {len(est)} 张暂估入库单发票未回(采购需核对红冲)', 'count': len(est)})
    else:
        checks.append({'level': 'ok', 'msg': '✅ 暂估入库单均已红冲, 无发票未回', 'count': 0})
    # ② V11.153: 红冲单(is_est=1且有发票)即已转正式, 不再要求进白入组; 检查"有发票号但is_est=0"的异常(正式单带发票=数据不一致)
    _weird = [x['receive_no'] for x in br if x['invoice_no']]
    if _weird:
        checks.append({'level': 'danger', 'msg': f'🚨 {len(_weird)} 张正式入库单带发票号(数据异常, 需检查): {", ".join(_weird[:3])}', 'count': len(_weird)})
    else:
        checks.append({'level': 'ok', 'msg': '✅ 红冲单均已转正式, 无数据异常', 'count': 0})
    # ③ 汇总
    checks.append({'level': 'info', 'msg': f'📊 暂估 {len(est)} 张 / 红冲 {len(hc)} 张 / 白入 {len(br)} 张', 'count': len(est) + len(hc) + len(br)})
    return jsonify({'est': est, 'hc': hc, 'br': br, 'checks': checks})

def _gen_est_export(kind):
    """生成 暂估/红冲/白入 明细表Excel
    kind: est=暂估入库明细(货到发票未到) / hc=红字冲销明细(发票到已冲) / br=白入明细(正式入库)"""
    import io as _io
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    conn = db()
    rows = conn.execute("SELECT * FROM receivings WHERE status IN ('已入库','已通过','审批通过') ORDER BY received_at DESC").fetchall()
    conn.close()
    wb = Workbook(); ws = wb.active
    thin = Side(style='thin', color='B0B0B0')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    tf = Font(name='微软雅黑', size=14, bold=True)
    hf = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
    hfill = PatternFill('solid', fgColor='2F5597')
    bf = Font(name='微软雅黑', size=10)
    titles = {'est': '暂估入库明细表', 'hc': '红字冲销明细表', 'br': '白入(正式入库)明细表'}
    heads = ['入库单号', '物资名称', '规格', '数量', '单位', '暂估金额(元)', '发票金额(元)', '差价(元)', '发票号', '发票类型', '入库日期', '备注']
    ws.merge_cells('A1:L1')
    ws['A1'] = titles[kind]; ws['A1'].font = tf
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30
    for ci, h in enumerate(heads, 1):
        c = ws.cell(3, ci, h); c.font = hf; c.fill = hfill; c.border = border
        c.alignment = Alignment(horizontal='center', vertical='center')
    r = 4
    cnt = 0
    for row in rows:
        is_est = bool(row['is_est'])
        has_inv = bool(row['invoice_no'])
        # 过滤: est=暂估未红冲 / hc=已红冲 / br=正式入库
        if kind == 'est' and (not is_est or has_inv): continue
        if kind == 'hc' and not (is_est and has_inv): continue
        if kind == 'br' and is_est: continue
        # V11.153: 红冲表体现 暂估价 vs 发票价 差价(invoice_amount新列; 兼容旧数据用est_amount)
        _est = float(row['est_amount'] or 0)
        _inv = float(row.get('invoice_amount') or 0) if row.keys() and 'invoice_amount' in row.keys() else 0
        if _inv == 0 and has_inv: _inv = _est  # 旧数据兼容
        _diff = round(_inv - _est, 2)
        vals = [row['receive_no'], row['item_name'], row['spec'] or '', row['quantity'],
                row['unit'] or '个', _est,
                (_inv if has_inv else 0), (_diff if has_inv else 0),
                row['invoice_no'] or '', row.get('invoice_type') or '', str(row['received_at'] or '')[:10], row['remark'] or '']
        for ci, v in enumerate(vals, 1):
            c = ws.cell(r, ci, v); c.border = border; c.font = bf
            c.alignment = Alignment(horizontal='center' if ci in (1, 3, 4, 5, 6, 7, 8, 9, 10, 11) else 'left', vertical='center', wrap_text=True)
        # 差价不为0时红色标注(差异提醒)
        if has_inv and _diff != 0:
            ws.cell(r, 8).font = Font(name='微软雅黑', size=10, bold=True, color='E74C3C')
        ws.row_dimensions[r].height = 22
        r += 1; cnt += 1
    if cnt == 0:
        ws.merge_cells(f'A{r}:L{r}')
        ws.cell(r, 1, '（本月暂无数据）').font = bf
    widths = [16, 20, 14, 10, 6, 14, 14, 12, 18, 16, 12, 20]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + j)].width = w
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9
    ws.page_setup.fitToWidth = 1
    bio = _io.BytesIO(); wb.save(bio); bio.seek(0)
    return bio, titles[kind]

@app.route('/api/reports/est-export')
@login_required
def api_est_export():
    """月底三表导出: ?kind=est|hc|br"""
    kind = request.args.get('kind', 'est')
    if kind not in ('est', 'hc', 'br'):
        return jsonify({'error': 'kind 只能是 est/hc/br'}), 400
    from flask import send_file
    bio, title = _gen_est_export(kind)
    return send_file(bio, as_attachment=True, download_name=f'{title}.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# ---- V11.203 模块一1.3: 发票台账统计(合同/供应商维度, 年/月/周/日分组, 欠票超期预警) ----
def _inv_stats_rows(args):
    """发票台账取数: 返回 {'dim','rows'(每行带已收/待收/超期标记),'summary'} 供JSON与Excel导出共用"""
    frm = (args.get('from') or '').strip()[:10]
    to = (args.get('to') or '').strip()[:10]
    sup = (args.get('supplier') or '').strip()
    status = (args.get('status') or '').strip()
    dim = args.get('dim') or 'contract'
    q = "SELECT c.*, po.order_no FROM contracts c LEFT JOIN purchase_orders po ON c.order_id=po.id WHERE c.status NOT IN ('已作废','已撤回','撤回','草稿')"
    p = []
    if frm:
        q += " AND c.created_at>=?"
        p.append(frm + ' 00:00:00')
    if to:
        q += " AND c.created_at<=?"
        p.append(to + ' 23:59:59')
    if sup:
        q += " AND c.supplier LIKE ?"
        p.append('%' + sup + '%')
    if status:
        q += " AND c.status=?"
        p.append(status)
    q += " ORDER BY c.created_at DESC, c.id DESC LIMIT 1200"
    conn = db()
    rows = conn.execute(q, p).fetchall()
    inv_map = {}
    for r in conn.execute("SELECT contract_id, COUNT(*) n, COALESCE(SUM(amount),0) amt FROM contract_invoices GROUP BY contract_id"):
        inv_map[r['contract_id']] = {'n': r['n'], 'amt': float(r['amt'] or 0)}
    conn.close()
    today = datetime.date.today().strftime('%Y-%m-%d')
    detail = []
    for r in rows:
        d = dict_row(r)
        st = inv_map.get(d['id'], {'n': 0, 'amt': 0.0})
        amt = float(d['amount'] or 0)
        pend = max(round(amt - st['amt'], 2), 0)
        over = bool(d.get('invoice_est_done')) and today > str(d['invoice_est_done'])[:10] and pend > 0.01
        d['inv_count'] = st['n']
        d['inv_amount'] = round(st['amt'], 2)
        d['pending'] = pend
        d['overdue'] = over
        detail.append(d)
    summary = {'contracts': len(detail),
               'amount': round(sum(float(x['amount'] or 0) for x in detail), 2),
               'inv_amount': round(sum(x['inv_amount'] for x in detail), 2),
               'pending': round(sum(x['pending'] for x in detail), 2),
               'overdue_cnt': sum(1 for x in detail if x['overdue'])}
    return {'dim': dim, 'rows': detail, 'summary': summary}


@app.route('/api/reports/invoice-stats')
@login_required
def api_invoice_stats():
    """发票台账: ?dim=contract|supplier & from/to/supplier/status 筛选 & group=year|month|week|day 时间分组
    金额敏感 → can_see_price 权限(库管/员工等不可见)"""
    if not can_see_price():
        return jsonify({'error': '无权限查看发票台账(金额敏感数据)'}), 403
    data = _inv_stats_rows(request.args)
    rows = data['rows']
    dim = data['dim']
    if dim == 'supplier':
        agg = {}
        for x in rows:
            s = x['supplier'] or '(未填)'
            a = agg.setdefault(s, {'supplier': s, 'contracts': 0, 'amount': 0.0, 'inv_count': 0, 'inv_amount': 0.0, 'pending': 0.0, 'overdue_cnt': 0})
            a['contracts'] += 1
            a['amount'] += float(x['amount'] or 0)
            a['inv_count'] += x['inv_count']
            a['inv_amount'] += x['inv_amount']
            a['pending'] += x['pending']
            a['overdue_cnt'] += 1 if x['overdue'] else 0
        data['rows'] = sorted((dict(v) for v in agg.values()), key=lambda z: -z['amount'])
    # 时间粒度分组(年/月/周/日) — 按合同维度计算
    gp = (request.args.get('group') or '').strip()
    if gp in ('year', 'month', 'week', 'day'):
        from collections import OrderedDict
        grp = OrderedDict()
        _base = data['rows'] if dim == 'contract' else rows
        for x in _base:
            try:
                dt0 = datetime.datetime.strptime((x.get('created_at') or '')[:10], '%Y-%m-%d').date()
            except Exception:
                continue
            if gp == 'year':
                lab = str(dt0.year)
            elif gp == 'month':
                lab = dt0.strftime('%Y-%m')
            elif gp == 'week':
                _iso = dt0.isocalendar()
                lab = '%d-W%02d' % (_iso[0], _iso[1])
            else:
                lab = dt0.strftime('%Y-%m-%d')
            g = grp.setdefault(lab, {'period': lab, 'contracts': 0, 'amount': 0.0, 'inv_amount': 0.0, 'pending': 0.0})
            g['contracts'] += 1
            g['amount'] += float(x['amount'] or 0)
            g['inv_amount'] += x['inv_amount']
            g['pending'] += x['pending']
        data['groups'] = list(grp.values())
    return jsonify(data)


@app.route('/api/reports/invoice-stats/export')
@login_required
def api_invoice_stats_export():
    """发票台账导出Excel(当前维度全明细行, 含汇总表头)"""
    if not can_see_price():
        return jsonify({'error': '无权限导出发票台账'}), 403
    data = _inv_stats_rows(request.args)
    dim = data['dim']
    rows = data['rows']
    s = data['summary']
    import io
    from openpyxl import Workbook
    from flask import send_file
    wb = Workbook()
    ws = wb.active
    ws.title = '发票台账'
    from openpyxl.styles import Font
    ws.append([('合同' if dim == 'contract' else '供应商') + '维度发票台账'])
    ws.merge_cells('A1:H1')
    ws['A1'].font = Font(bold=True, size=13)
    ws.append(['时间范围: 全部' if not request.args.get('from') else ('%s ~ %s' % (request.args.get('from'), request.args.get('to') or '至今')),
               '合同数: %d' % s['contracts'], '合同总额: ¥%.2f' % s['amount'],
               '已收发票金额: ¥%.2f' % s['inv_amount'], '待收(欠票): ¥%.2f' % s['pending'],
               '超期合同: %d' % s['overdue_cnt']])
    if dim == 'contract':
        head = ['合同编号', '供应商', '合同金额', '状态', '签订日期', '预计首次开票', '预计开票完成', '已收张数', '已收金额', '待收金额', '超期预警']
        ws.append(head)
        for x in rows:
            ws.append([x['contract_no'], x['supplier'] or '', round(float(x['amount'] or 0), 2), x['status'],
                       (x.get('created_at') or '')[:10], x.get('invoice_est_first') or '', x.get('invoice_est_done') or '',
                       x['inv_count'], round(x['inv_amount'], 2), round(x['pending'], 2), '⚠️超期' if x['overdue'] else ''])
    else:
        head = ['供应商', '合同数', '合同总金额', '已收发票张数', '已收发票金额', '待收(欠票)金额', '超期合同数']
        ws.append(head)
        for x in rows:
            ws.append([x['supplier'], x['contracts'], round(x['amount'], 2), x['inv_count'], round(x['inv_amount'], 2), round(x['pending'], 2), x['overdue_cnt']])
    from openpyxl.cell.cell import MergedCell
    for col in ws.columns:
        _vals = [c for c in col if c.value is not None and not isinstance(c, MergedCell)]
        if not _vals:
            continue
        _w = min(max(len(str(c.value)) for c in _vals) * 2 + 4, 40)
        ws.column_dimensions[_vals[0].column_letter].width = _w
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    fn = '发票台账_%s维度_%s.xlsx' % (dim, datetime.date.today().strftime('%Y%m%d'))
    return send_file(bio, as_attachment=True, download_name=fn,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/receivings/<int:rid>/invoice-match', methods=['POST'])
@login_required
def api_receiving_invoice_match(rid):
    """V11.53: 采购发票核对 — 暂估入库单收到发票后, 填发票号/金额 → 红冲转正式入库"""
    d = request.json or {}
    invoice_no = (d.get('invoice_no') or '').strip()
    amount = float(d.get('amount') or 0)
    invoice_type = d.get('invoice_type') or '增值税专用发票'
    if not invoice_no:
        return jsonify({'error': '请填写发票号'}), 400
    conn = db()
    rn = conn.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        conn.close(); return jsonify({'error': '入库单不存在'}), 404
    if not rn['is_est']:
        conn.close(); return jsonify({'error': '该入库单不是暂估单,无需红冲'}), 400
    if rn['invoice_no']:
        conn.close(); return jsonify({'error': f'该暂估单已红冲(发票{rn["invoice_no"]})'}), 400
    # 红冲: 暂估→已红冲(V11.153: 保留is_est=1表示"暂估已红冲", 月底红冲表判定 is_est=1 AND invoice_no; 正式入库才是is_est=0)
    # V11.153: 发票金额单独存invoice_amount, 暂估价est_amount保留, 差价=invoice_amount-est_amount可体现
    try: conn.execute("ALTER TABLE receivings ADD COLUMN invoice_type TEXT DEFAULT ''")
    except Exception: pass
    try: conn.execute("ALTER TABLE receivings ADD COLUMN invoice_amount REAL DEFAULT 0")
    except Exception: pass
    _inv_amt = amount if amount > 0 else rn['est_amount']
    conn.execute("UPDATE receivings SET is_est=1, invoice_no=?, est_amount=?, invoice_type=?, invoice_amount=? WHERE id=?",
                 (invoice_no, rn['est_amount'] or 0, invoice_type, _inv_amt, rid))
    conn.commit(); conn.close()
    log(session['user_name'], '发票核对红冲', f'{rn["receive_no"]} 发票{invoice_no} 暂估{rn["est_amount"]}→发票{_inv_amt} 差价{_inv_amt-(rn["est_amount"] or 0):.2f}')
    return jsonify({'success': True, 'receive_no': rn['receive_no']})

@app.route('/api/invoices')
@login_required
def api_invoices():
    """获取发票列表 — V11.159: 仅 财务/领导/管理员 可见"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    conn = db()
    rows = conn.execute("SELECT * FROM invoices ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/invoices/merge', methods=['POST'])
@login_required
def api_merge_invoice():
    """需求5-合并结算开票: 多笔合同/订单合并生成一张发票
    V11.159: 合并开票仅 财务/领导/管理员"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：合并开票仅限财务/领导'}), 403
    d = request.json
    order_ids = d.get('order_ids') or []
    if not order_ids:
        return jsonify({'error': '请选择要合并开票的订单'}), 400
    conn = db()
    rows = conn.execute(f"SELECT * FROM purchase_orders WHERE id IN ({','.join('?'*len(order_ids))})", order_ids).fetchall()
    if not rows: conn.close(); return jsonify({'error': '未找到订单'}), 400
    total = sum(r['total_amount'] or 0 for r in rows)
    amt = float(d.get('amount', total))
    supplier = d.get('supplier', '') or (rows[0]['supplier'] if rows else '')
    cno = ''
    ct = conn.execute("SELECT c.* FROM contracts c WHERE c.order_id=?", (rows[0]['id'],)).fetchone()
    if ct: cno = ct['contract_no']
    inv_no = d.get('invoice_no', '')
    if not inv_no:
        conn.close(); return jsonify({'error': '请填写发票号码'}), 400
    conn.execute("INSERT INTO invoices(invoice_no,invoice_code,order_id,order_ids,contract_id,contract_no,supplier,amount,tax_amount,total_amount,invoice_date,invoice_type,file_path,status,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (inv_no, d.get('invoice_code',''), rows[0]['id'], ','.join(str(x) for x in order_ids), ct['id'] if ct else None, cno, supplier,
         amt, float(d.get('tax_amount',0)), float(d.get('total_amount', amt)), d.get('invoice_date',''), d.get('invoice_type','增值税专用发票'),
         d.get('file_path',''), '待验证', f"合并开票{len(rows)}单"))
    conn.commit(); conn.close()
    log(session['user_name'], '合并开票', f'{inv_no} 合并{len(rows)}单 ¥{amt:.0f}')
    return jsonify({'success': True, 'invoice_no': inv_no, 'orders': len(rows)})

@app.route('/api/ledger')
@login_required
def api_ledger():
    """需求4-往来台账: 按供应商/合同/时间段筛选全流程单据
    V11.159: 往来台账仅 财务/领导/管理员 可见"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    supplier = request.args.get('supplier', '')
    contract = request.args.get('contract', '')
    start = request.args.get('start', '')
    end = request.args.get('end', '')
    c = db()
    def q(table, cols, date_col, extra=''):
        sql = f"SELECT {cols} FROM {table} WHERE 1=1"
        args = []
        if supplier: sql += ' AND supplier=?'; args.append(supplier)
        if contract and 'contract_no' in [r[1] for r in c.execute(f'PRAGMA table_info({table})').fetchall()]:
            sql += ' AND contract_no LIKE ?'; args.append(f'%{contract}%')
        if start: sql += f' AND {date_col} >= ?'; args.append(start + ' 00:00:00')
        if end: sql += f' AND {date_col} <= ?'; args.append(end + ' 23:59:59')
        sql += extra
        return c.execute(sql, args).fetchall()
    out = []
    for r in q('purchase_orders', "id,order_no,item_name,supplier,total_amount,trade_mode,status,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '采购订单', 'no': r['order_no'], 'supplier': r['supplier'],
                    'contract': '', 'item': r['item_name'], 'amount': r['total_amount'], 'status': r['status']})
    for r in q('deliveries', "id,delivery_no,item_name,supplier,contract_no,quantity,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '到货', 'no': r['delivery_no'], 'supplier': r['supplier'],
                    'contract': r['contract_no'], 'item': r['item_name'], 'amount': None, 'status': '已登记'})
    for r in q('receivings', "id,receive_no,item_name,contract_no,quantity,status,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '入库', 'no': r['receive_no'], 'supplier': '',
                    'contract': r['contract_no'], 'item': r['item_name'], 'amount': None, 'status': r['status']})
    for r in q('invoices', "id,invoice_no,supplier,contract_no,total_amount,status,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '发票', 'no': r['invoice_no'], 'supplier': r['supplier'],
                    'contract': r['contract_no'], 'item': '', 'amount': r['total_amount'], 'status': r['status']})
    for r in q('credit_notes', "id,credit_no,supplier,contract_no,item_name,amount,status,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '挂账', 'no': r['credit_no'], 'supplier': r['supplier'],
                    'contract': r['contract_no'], 'item': r['item_name'], 'amount': r['amount'], 'status': r['status']})
    for r in q('payment_requests', "id,payment_no,supplier,contract_id,amount,trade_mode,status,created_at", 'created_at'):
        out.append({'date': r['created_at'], 'type': '付款', 'no': r['payment_no'], 'supplier': r['supplier'],
                    'contract': '', 'item': '', 'amount': r['amount'], 'status': r['status']})
    c.close()
    out.sort(key=lambda x: x['date'] or '', reverse=True)
    return jsonify(out)


# ============================================================
# V5 ── 44需求: 全局搜索/合同模板自动生成/导出/作废/权限
# ============================================================
@app.route('/api/search')
@login_required
def api_search():
    """需求44-全局搜索: 菜单/单据/档案快速检索(覆盖全部业务类目)"""
    q = request.args.get('q', '').strip()
    # 兼容GBK编码客户端(飞书/钉钉内置浏览器或部分系统按ANSI编码发送URL):
    # 若UTF-8解码失败出现替换符, 用原始字节按GBK重新解码
    if q and '\ufffd' in q:
        try:
            from urllib.parse import parse_qsl
            raw_qs = request.query_string.decode('latin-1', 'replace')
            # 指定 latin-1 保留百分号解码后的原始字节, 再整体按 GBK 还原
            for k, v in parse_qsl(raw_qs, encoding='latin-1'):
                if k == 'q':
                    q2 = v.encode('latin-1', 'replace').decode('gbk', 'replace').strip()
                    if q2:
                        q = q2
                    break
        except Exception:
            pass
    if not q: return jsonify({'orders': [], 'requests': [], 'contracts': [], 'suppliers': [], 'items': [],
                              'deliveries': [], 'receivings': [], 'credits': [], 'payments': [],
                              'inventory': [], 'invoices': [], 'settlements': [], 'approvals': []})
    like = f'%{q}%'
    c = db()
    # V11.64: 搜索范围按角色域限制(员工只搜自己的申请; 库管员不搜采购/财务)
    role = session.get('user_role')
    _own = f" AND pr.requester_id={session.get('user_id', 0)}" if role in ('员工', '部门负责人') else ''
    # 主表搜索 + 明细表联查(物资名在 request_items/order_items 里, 必须覆盖否则搜不到)
    orders = c.execute("SELECT po.id,po.order_no,po.item_name,po.supplier,po.status FROM purchase_orders po WHERE po.order_no LIKE ? OR po.item_name LIKE ? OR po.supplier LIKE ? OR po.remark LIKE ? OR po.id IN (SELECT order_id FROM order_items WHERE item_name LIKE ? OR spec LIKE ? OR remark LIKE ?) LIMIT 8",
                       (like, like, like, like, like, like, like)).fetchall()
    reqs = c.execute("SELECT pr.id,pr.req_no,pr.purpose,pr.dept,pr.status FROM purchase_requests pr WHERE pr.req_no LIKE ? OR pr.purpose LIKE ? OR pr.dept LIKE ? OR pr.requester LIKE ? OR pr.id IN (SELECT req_id FROM request_items WHERE item_name LIKE ? OR spec LIKE ? OR remark LIKE ?)" + _own + " LIMIT 8",
                     (like, like, like, like, like, like, like)).fetchall()
    cons = c.execute("SELECT id,contract_no,contract_name,supplier,status FROM contracts WHERE contract_no LIKE ? OR contract_name LIKE ? OR supplier LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    sups = c.execute("SELECT id,name,contact,phone FROM suppliers WHERE name LIKE ? OR contact LIKE ? OR phone LIKE ? OR category LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    items = c.execute("SELECT id,name,spec,unit FROM items WHERE name LIKE ? OR spec LIKE ? OR cat_code LIKE ? OR supplier LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    try:
        dels = c.execute("SELECT id,delivery_no,item_name,supplier,delivery_date,sign_status FROM deliveries WHERE delivery_no LIKE ? OR item_name LIKE ? OR supplier LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    except Exception:
        dels = []
    try:
        recs = c.execute("SELECT r.id,r.receive_no,r.item_name,COALESCE(po.supplier,'') AS supplier,r.status FROM receivings r LEFT JOIN purchase_orders po ON r.order_id=po.id WHERE r.receive_no LIKE ? OR r.item_name LIKE ? OR r.warehouse LIKE ? OR r.contract_no LIKE ? OR r.remark LIKE ? OR COALESCE(po.supplier,'') LIKE ? LIMIT 8", (like, like, like, like, like, like)).fetchall()
    except Exception:
        recs = []
    try:
        crds = c.execute("SELECT id,credit_no,supplier,amount,status FROM credit_notes WHERE credit_no LIKE ? OR supplier LIKE ? OR item_name LIKE ? OR category LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like, like, like)).fetchall()
    except Exception:
        crds = []
    try:
        pays = c.execute("SELECT id,payment_no,supplier,amount,status FROM payment_requests WHERE payment_no LIKE ? OR supplier LIKE ? OR payment_reason LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    except Exception:
        pays = []
    try:
        invs = c.execute("SELECT id,item_name,warehouse,quantity,unit FROM inventory WHERE item_name LIKE ? OR warehouse LIKE ? OR spec LIKE ? OR cat_code LIKE ? LIMIT 8", (like, like, like, like)).fetchall()
    except Exception:
        invs = []
    try:
        fins = c.execute("SELECT id,invoice_no,supplier,amount,status FROM invoices WHERE invoice_no LIKE ? OR supplier LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like)).fetchall()
    except Exception:
        fins = []
    try:
        setls = c.execute("SELECT id,settlement_no,supplier,total_amount AS amount,status FROM settlements WHERE settlement_no LIKE ? OR supplier LIKE ? OR remark LIKE ? LIMIT 8", (like, like, like)).fetchall()
    except Exception:
        setls = []
    try:
        # 审批: 关联单据编号(通过 biz_type/biz_id 反查各主表单号), 搜单号也能出审批
        apps = c.execute("SELECT ai.id,ai.biz_type,ai.biz_id,ai.role,ai.status,ai.created_at FROM approval_instances ai WHERE ai.biz_type LIKE ? OR ai.role LIKE ? OR ai.status LIKE ? OR ai.id IN (SELECT a2.id FROM approval_instances a2 WHERE a2.biz_type='purchase_request' AND a2.biz_id IN (SELECT id FROM purchase_requests WHERE req_no LIKE ?)) OR ai.id IN (SELECT a3.id FROM approval_instances a3 WHERE a3.biz_type='purchase_order' AND a3.biz_id IN (SELECT id FROM purchase_orders WHERE order_no LIKE ?)) LIMIT 8",
                         (like, like, like, like, like)).fetchall()
    except Exception:
        apps = []
    try:
        usrs = c.execute("SELECT id,username,name,role FROM users WHERE is_active=1 AND (username LIKE ? OR name LIKE ? OR role LIKE ?) LIMIT 8", (like, like, like)).fetchall()
    except Exception:
        usrs = []
    c.close()
    return jsonify({'orders': [dict_row(r) for r in orders], 'requests': [dict_row(r) for r in reqs],
                    'contracts': [dict_row(r) for r in cons], 'suppliers': [dict_row(r) for r in sups],
                    'items': [dict_row(r) for r in items], 'deliveries': [dict_row(r) for r in dels],
                    'receivings': [dict_row(r) for r in recs], 'credits': [dict_row(r) for r in crds],
                    'payments': [dict_row(r) for r in pays], 'inventory': [dict_row(r) for r in invs],
                    'invoices': [dict_row(r) for r in fins], 'settlements': [dict_row(r) for r in setls],
                    'approvals': [dict_row(r) for r in apps], 'users': [dict_row(r) for r in usrs]})

# ============================================================
# V11.209 补充二: 全链路溯源档案 — 命中单据 → 聚合上下游/审批/发票/库存流水/时间轴
# 检索入口(首页全局搜索/报表中心/系统中心)共用 /api/search, 溯源共用本接口
# ============================================================
@app.route('/api/trace')
@login_required
def api_trace():
    """溯源档案: 按 biz_type+biz_id 聚合全链路。
    返回: self基础信息 / timeline时间轴 / approvals审批记录 / logs操作日志
          / upstream上游(申请/询价/合同) / downstream下游(入库/出库/退库/付款/发票)
          / flows库存流水 / attachments附件
    """
    biz_type = request.args.get('t', '')
    biz_id = request.args.get('id', '')
    kw = request.args.get('q', '').strip()  # 也支持纯关键词溯源(自动找主命中)
    c = db(); c.row_factory = sqlite3.Row
    TBL = {'order': ('purchase_orders', 'order_no'), 'request': ('purchase_requests', 'req_no'),
           'contract': ('contracts', 'contract_no'), 'receiving': ('receivings', 'receive_no'),
           'requisition': ('requisitions', 'req_no'), 'return_request': ('return_requests', 'return_no'),
           'payment': ('payment_requests', 'payment_no'), 'credit': ('credit_notes', 'credit_no'),
           'invoice': ('invoices', 'invoice_no'), 'repair': ('repair_plans', 'plan_no'),
           'inventory': ('inventory', 'item_name'), 'settlement': ('settlements', 'settlement_no')}
    out = {'self': None, 'timeline': [], 'approvals': [], 'logs': [],
           'upstream': {'requests': [], 'inquiries': [], 'contracts': [], 'orders': []},
           'downstream': {'receivings': [], 'requisitions': [], 'returns': [], 'payments': [], 'invoices': [], 'repairs': []},
           'flows': [], 'attachments': [], 'links': []}
    def _push(row): return {k: row[k] for k in row.keys()}
    if biz_type in TBL and biz_id:
        tbl, no_col = TBL[biz_type]
        try:
            r = c.execute(f"SELECT * FROM {tbl} WHERE id=?", (biz_id,)).fetchone()
        except Exception:
            r = None
        if not r: c.close(); return jsonify(out)
        row = dict(r)
        out['self'] = {'biz_type': biz_type, 'id': r['id'], 'no': row.get(no_col, ''), 'data': row}
        # 时间轴: 单据创建
        _ct = row.get('created_at') or row.get('received_at') or row.get('issued_at') or ''
        if _ct: out['timeline'].append({'t': _ct, 'event': f'单据创建({no_col}={row.get(no_col, "")})'})
        # 审批记录(该单据的全部审批实例+动作日志)
        try:
            for a in c.execute("SELECT ai.*, (SELECT comment FROM approval_action_logs al WHERE al.biz_type=ai.biz_type AND al.biz_id=ai.biz_id AND al.action IN ('approved','rejected') ORDER BY al.id DESC LIMIT 1) last_comment FROM approval_instances ai WHERE ai.biz_type=? AND ai.biz_id=? ORDER BY ai.level_no",
                               (biz_type if biz_type != 'receiving' else 'receiving', biz_id)).fetchall():
                d = dict(a)
                if biz_type == 'order': d['biz_type'] = 'purchase_order'
                out['approvals'].append(d)
                if d.get('status') == 'approved': out['timeline'].append({'t': d.get('approved_at') or d.get('updated_at') or '', 'event': f'审批通过({d.get("role") or ""})'})
                elif d.get('status') == 'rejected': out['timeline'].append({'t': d.get('updated_at') or '', 'event': f'驳回({d.get("role") or ""})'})
        except Exception:
            pass
        # 操作日志
        try:
            for l in c.execute("SELECT * FROM logs WHERE detail LIKE ? ORDER BY id DESC LIMIT 20", (f'%{row.get(no_col, "")}%',)).fetchall():
                out['logs'].append(dict(l))
        except Exception:
            pass
        # ── 上下游聚合 ──
        _item = row.get('item_name') or ''
        _sup = row.get('supplier') or ''
        _cid = row.get('contract_id') or row.get('credit_id')
        _oid = row.get('order_id')
        _rcv_no = row.get('receive_no') or ''
        # 申请(同一物料名/单号引用)
        if _item:
            for t2 in c.execute("SELECT id,req_no,purpose,status FROM purchase_requests WHERE id IN (SELECT req_id FROM request_items WHERE item_name=? OR spec=?) OR purpose LIKE ? LIMIT 5", (_item, _item, f'%{_item}%')).fetchall():
                out['upstream']['requests'].append(dict(t2))
        # 询价(物料名)
        if _item:
            try:
                for t2 in c.execute("SELECT id,inq_no,item_name,status FROM inquiries WHERE item_name=? OR purpose LIKE ? LIMIT 5", (_item, f'%{_item}%')).fetchall():
                    out['upstream']['inquiries'].append(dict(t2))
            except Exception:
                pass
        # 订单(通过合同/物料/供应商)
        if biz_type == 'request' or biz_type == 'contract':
            pass
        if _item or _sup:
            try:
                for t2 in c.execute("SELECT id,order_no,item_name,supplier,total_amount,status FROM purchase_orders WHERE (item_name=? OR supplier=?) AND id!=? LIMIT 5", (_item, _sup, biz_id if biz_type == 'order' else -1)).fetchall():
                    out['upstream']['orders'].append(dict(t2))
            except Exception:
                pass
        # 合同(供应商/物料)
        if _sup or _cid:
            try:
                for t2 in c.execute("SELECT id,contract_no,contract_name,supplier,amount,status FROM contracts WHERE supplier=? OR id=? LIMIT 5", (_sup, _cid)).fetchall():
                    out['upstream']['contracts'].append(dict(t2))
            except Exception:
                pass
        # 入库单(订单/合同/物料)
        try:
            for t2 in c.execute("SELECT id,receive_no,item_name,qualified_qty,status,received_at FROM receivings WHERE (order_id=? OR item_name=? OR id=?) LIMIT 8", (_oid or -1, _item, biz_id if biz_type == 'receiving' else -1)).fetchall():
                out['downstream']['receivings'].append(dict(t2))
                if biz_type == 'order' and t2['id'] == (biz_id if biz_type == 'receiving' else -1):
                    pass
        except Exception:
            pass
        # 出库(物料)
        if _item:
            try:
                for t2 in c.execute("SELECT id,req_no,item_name,quantity,dept,status FROM requisitions WHERE item_name=? LIMIT 5", (_item,)).fetchall():
                    out['downstream']['requisitions'].append(dict(t2))
            except Exception:
                pass
            try:
                for t2 in c.execute("SELECT id,return_no,item_name,quantity,status FROM return_requests WHERE item_name=? LIMIT 5", (_item,)).fetchall():
                    out['downstream']['returns'].append(dict(t2))
            except Exception:
                pass
        # 付款(credit关联/供应商)
        if _cid:
            try:
                for t2 in c.execute("SELECT id,payment_no,supplier,amount,status FROM payment_requests WHERE credit_id=? LIMIT 5", (_cid,)).fetchall():
                    out['downstream']['payments'].append(dict(t2))
            except Exception:
                pass
        # 发票(合同/供应商/单号)
        try:
            for t2 in c.execute("SELECT id,invoice_no,supplier,amount,status,invoice_date FROM invoices WHERE (contract_id=? OR supplier=? OR contract_no=?) LIMIT 5", (_cid or -1, _sup, row.get('contract_no') or '')).fetchall():
                out['downstream']['invoices'].append(dict(t2))
        except Exception:
            pass
        # 维修计划(设备名/物料)
        if _item:
            try:
                for t2 in c.execute("SELECT id,plan_no,device_name,status FROM repair_plans WHERE device_name=? LIMIT 5", (_item,)).fetchall():
                    out['downstream']['repairs'].append(dict(t2))
            except Exception:
                pass
        # 库存流水(物料+规格)
        _spec = row.get('spec') or ''
        if _item:
            for t2 in c.execute("SELECT item_name,spec,flow_type,doc_type,doc_no,qty,balance_after,operator,created_at FROM inventory_flows WHERE item_name=? ORDER BY id DESC LIMIT 20", (_item,)).fetchall():
                out['flows'].append(dict(t2))
        # 附件: 附件字段若有(/uploads/路径) → attachments
        _atts = row.get('attachments') or ''
        if _atts:
            import re as _re
            out['attachments'] = _re.findall(r'/uploads/[^\s",]+', str(_atts))
    elif kw and not biz_id:
        # 纯关键词溯源: 前端已先经 /api/search 命中得类型+id, 此处兜底返回空(提示前端走search)
        pass
    c.close()
    return jsonify(out)

# ---- 合同模板管理 ----
@app.route('/api/contract-templates')
@login_required
def api_ctpls():
    c = db(); rows = c.execute("SELECT * FROM contract_templates ORDER BY id").fetchall(); c.close()
    return jsonify([dict_row(r) for r in rows])

@app.route('/api/contract-templates', methods=['POST'])
@admin_required
def api_ctpl_create():
    d = request.json
    c = db()
    first = c.execute("SELECT COUNT(*) FROM contract_templates").fetchone()[0] == 0
    c.execute("INSERT INTO contract_templates(name,file_path,version,status,is_default,remark) VALUES(?,?,?,?,?,?)",
              (d.get('name', '合同模板'), d.get('file_path', ''), d.get('version', 'V1'), '启用',
               1 if first or d.get('is_default') else 0, d.get('remark', '')))
    c.commit(); c.close()
    log(session.get('user_name',''), '新增合同模板', '%s (%s)' % (d.get('name','合同模板'), d.get('version','V1')))
    return jsonify({'success': True})

@app.route('/api/contract-templates/<int:tid>', methods=['POST'])
@admin_required
def api_ctpl_update(tid):
    d = request.json; c = db()
    if 'status' in d: c.execute("UPDATE contract_templates SET status=? WHERE id=?", (d['status'], tid))
    if 'name' in d: c.execute("UPDATE contract_templates SET name=? WHERE id=?", (d['name'], tid))
    if 'file_path' in d: c.execute("UPDATE contract_templates SET file_path=? WHERE id=?", (d['file_path'], tid))
    if 'is_default' in d and d['is_default']:
        c.execute("UPDATE contract_templates SET is_default=0")
        c.execute("UPDATE contract_templates SET is_default=1 WHERE id=?", (tid,))
    c.commit(); c.close()
    log(session.get('user_name',''), '修改合同模板', '模板#%s 变更: %s' % (tid, ','.join(k for k in ('status','name','file_path','is_default') if k in d) or '无'))
    return jsonify({'success': True})

@app.route('/api/contract-templates/<int:tid>', methods=['DELETE'])
@admin_required
def api_ctpl_delete(tid):
    c = db()
    tpl = c.execute("SELECT * FROM contract_templates WHERE id=?", (tid,)).fetchone()
    c.execute("DELETE FROM contract_templates WHERE id=?", (tid,)); c.commit(); c.close()
    log(session.get('user_name',''), '删除合同模板', '模板#%s %s' % (tid, tpl['name'] if tpl else ''))
    return jsonify({'success': True})

# ---- 人民币金额大写 ----
def rmb_upper(n):
    units = ['', '拾', '佰', '仟']
    bigs = ['', '万', '亿', '万亿']
    digs = '零壹贰叁肆伍陆柒捌玖'
    n = round(float(n) + 1e-9, 2)
    int_part = int(n)
    dec = int(round((n - int_part) * 100))
    jiao, fen = dec // 10, dec % 10
    if int_part == 0:
        s = '零元'
    else:
        s = ''
        segs = []
        v = int_part
        while v > 0:
            segs.append(v % 10000); v //= 10000
        for i in range(len(segs) - 1, -1, -1):
            seg = segs[i]
            zero = False
            if seg == 0:
                if s and not s.endswith('零') and i > 0:
                    s += '零'
                continue
            if s and (s.endswith('元') or s.endswith('零')):
                pass
            if s and i < len(segs) - 1 and segs[i] < 1000 and s and not s.endswith('零'):
                s += '零'
            s += fmt_wan(seg, digs, units) + bigs[i]
            if i > 0 and seg % 10 == 0:
                s = s.rstrip('零') + bigs[i]
        s = s.replace('零零', '零') + '元'
    if jiao == 0 and fen == 0:
        s += '整'
    elif jiao == 0:
        s += '零' + digs[fen] + '分'
    elif fen == 0:
        s += digs[jiao] + '角整'
    else:
        s += digs[jiao] + '角' + digs[fen] + '分'
    return s

def fmt_wan(seg, digs, units):
    s = ''
    zero = False
    for i in range(3, -1, -1):
        d = seg // (10 ** i) % 10
        if d == 0:
            if s and not s.endswith('零'):
                zero = True
        else:
            if zero:
                s += '零'; zero = False
            s += digs[d] + units[i]
    return s

# ---- 需求44: 采购合同自动生成 ----
@app.route('/api/contracts/generate', methods=['POST'])
@login_required
def api_contract_generate():
    """订单下单后一键生成采购合同: 调默认模板→自动填充甲方/乙方/明细/结算方式→生成docx归档"""
    d = request.json
    oid = d.get('order_id')
    conn = db()
    # V11.203 模块一1.1: 每份合同独立发票条款/预计开票时间(采购员生成合同时填, 不传则空=老逻辑)
    inv_clause = (d.get('invoice_clause') or '').strip()
    inv_first = (d.get('invoice_est_first') or '').strip()[:10]
    inv_done = (d.get('invoice_est_done') or '').strip()[:10]
    o = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
    if not o:
        conn.close(); return jsonify({'error': '订单不存在'}), 400
    # V11.151: 防重复生成合同 — 订单已有有效合同(非作废/非已撤回)时禁止再次生成
    _exist = conn.execute(
        "SELECT id,contract_no,status FROM contracts WHERE order_id=? AND status NOT IN ('已作废','已撤回','撤回') ORDER BY id LIMIT 1",
        (oid,)).fetchone()
    if _exist:
        conn.close()
        return jsonify({'error': '该订单已生成合同 %s（状态:%s），如需重新生成请先撤回或作废原合同' % (_exist['contract_no'], _exist['status'])}), 400
    sup = conn.execute("SELECT * FROM suppliers WHERE name=?", (o['supplier'],)).fetchone() if o['supplier'] else None
    tpl = None
    if d.get('template_id'):
        tpl = conn.execute("SELECT * FROM contract_templates WHERE id=?", (d['template_id'],)).fetchone()
    if not tpl:
        tpl = conn.execute("SELECT * FROM contract_templates WHERE is_default=1 AND status='启用'").fetchone()
    if not tpl:
        conn.close(); return jsonify({'error': '未找到启用的合同模板, 请到 系统设置→合同模板管理 上传模板'}), 400
    # 甲方预设
    cname = cfg_get('company_name', '正成能源有限公司')
    caddr = cfg_get('company_address', '山西省')
    ccontact = cfg_get('company_contact', '采购部')
    cphone = cfg_get('company_phone', '')
    cno = gen_contract_no(conn)
    # V11.144: 结算方式由采购员在生成合同时选择(现结/月结), 覆盖订单默认值
    _settle_choice = (d.get('settle_type') or '').strip()
    if _settle_choice in ('现结', '月结'):
        conn.execute("UPDATE purchase_orders SET settle_type=?, updated_at=? WHERE id=?", (_settle_choice, now(), oid))
        conn.commit()
    tm = o['trade_mode'] or '货到付款'
    # V11.7: 结算方式跟随订单交易模式 — 自定义模式(如 预付30%)直接带入, 内置两种保留详细说明
    if _settle_choice == '现结':
        settle = '现结：一单一结，验收合格后立即付款'
    elif _settle_choice == '月结':
        settle = '月结：月底按厂家汇总对账，统一生成月度合同后付款'
    elif tm == '货到付款':
        settle = '货到付款：到货验收入库后，月度对账、合并开票、挂账后付款'
    elif tm == '先款后货':
        settle = '先款后货：合同签订后预付货款，供应商收款后发货，到货入库后挂账核销'
    else:
        settle = tm
    items_txt = ''
    _oi = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()
    if _oi:
        items_txt = '\n'.join(f"{i+1}. {r['item_name']} {r['spec'] or ''} {r['quantity']}{r['unit'] or '个'} 单价¥{r['price'] or 0}" for i, r in enumerate(_oi))
    elif o['req_id']:
        rows = conn.execute("SELECT * FROM request_items WHERE req_id=?", (o['req_id'],)).fetchall()
        if rows:
            items_txt = '\n'.join(f"{i+1}. {r['item_name']} {r['spec'] or ''} {r['quantity']}{r['unit'] or '个'}" for i, r in enumerate(rows))
    else:
        items_txt = f"1. {o['item_name']} {o['spec'] or ''} {o['quantity']}{o['unit'] or '个'} 单价¥{o['price'] or 0}"
    mapping = {
        '{合同编号}': cno, '{订单编号}': o['order_no'],
        '{甲方名称}': cname, '{甲方地址}': caddr, '{甲方联系人}': ccontact, '{甲方电话}': cphone,
        '{乙方名称}': sup['name'] if sup else (o['supplier'] or ''),
        '{乙方地址}': sup['bank'] if sup else '', '{乙方联系人}': sup['contact'] if sup else '',
        '{乙方电话}': sup['phone'] if sup else '', '{乙方开户行}': sup['bank'] if sup else '',
        '{乙方账号}': sup['account'] if sup else '',
        '{下单日期}': (o['created_at'] or '')[:10], '{预计交货日期}': o['target_date'] or '',
        '{结算方式}': settle, '{明细清单}': items_txt,
        '{合计金额}': f"¥{o['total_amount'] or 0:,.2f}（人民币大写：{rmb_upper(o['total_amount'] or 0)}）",
    }
    try:
        from docx import Document
        tpl_path = os.path.join(BASE, 'uploads', tpl['file_path'])
        if not os.path.exists(tpl_path):
            conn.close(); return jsonify({'error': '模板文件缺失, 请重新上传'}), 400
        doc = Document(tpl_path)
        # 明细行(多商品订单取 order_items 汇总)
        if _oi:
            amt = sum(float(r['amount'] or 0) for r in _oi)
            tax = sum(float(r['tax_amount'] or 0) for r in _oi)
            total = sum(float(r['total_amount'] or 0) for r in _oi)
            tax_rate = float(_oi[0]['tax_rate'] or 13) if _oi else 13
        else:
            amt = float(o['amount'] or 0)
            tax_rate = float(o['tax_rate'] or 13)
            tax = amt * tax_rate / 100.0
            total = amt + tax
        # 交付天数
        days = ''
        if o['target_date']:
            try:
                d1 = datetime.datetime.strptime(o['target_date'][:10], '%Y-%m-%d').date()
                days = str(max((d1 - datetime.date.today()).days, 1))
            except Exception:
                days = ''
        today_s = datetime.date.today().strftime('%Y年 %m月 %d日')
        # V8.4: 合同文本通用处理(段落+表格共用) — 合计金额中文大写/税率/税金/不含税/收款账户/日期
        def _apply_ct(t):
            if '合计金额：¥' in t:
                return (f"合计金额：¥{total:,.2f}元（大写金额：人民币{rmb_upper(total)}）。"
                        f"税金（税率 {tax_rate:.0f}%）为：¥{tax:,.2f}元（大写金额：人民币{rmb_upper(tax)}）；"
                        f"不含税价款为：¥{amt:,.2f}元（大写金额：人民币{rmb_upper(amt)}）。")
            reps = [
                (r'合同签订后\s+日内交付', f'合同签订后{days or "7"}日内交付'),
                (r'运抵甲方指定地点后\s+日内', '运抵甲方指定地点后1日内'),
                (r'乙方应在\s+日内更换', '乙方应在1日内更换'),
                (r'质保期为\s+年', '质保期为1年'),
                (r'签订合同后\s+日内，乙方向甲方提供全额', '签订合同后3日内，乙方向甲方提供全额'),
                (r'收到发票后\s+日内支付合同总价的\s+%', '收到发票后3日内支付合同总价的100%'),
                (r'质保期满后若无质量纠纷，\s+日内支付剩余价款', '质保期满后若无质量纠纷，30日内支付剩余价款'),
                (r'延迟交付货物超过\s+天', '延迟交付货物超过3天'),
                (r'需提前\s+天通知对方', '需提前3天通知对方'),
                (r'合同额的\s+%', '合同额的30%'),
            ]
            for pat, repx in reps:
                if re.search(pat, t):
                    t = re.sub(pat, repx, t)
            if '收款账户名称：' in t:
                t = t.replace('收款账户名称：', '收款账户名称：' + (sup['name'] if sup else (o['supplier'] or '')))
            if '收款账号：' in t:
                t = t.replace('收款账号：', '收款账号：' + (sup['account'] if sup and sup['account'] else ''))
            if '收款银行：' in t:
                t = t.replace('收款银行：', '收款银行：' + (sup['bank'] if sup and sup['bank'] else ''))
            # V11.144: 结算方式注入 — 付款条款段(甲方自收到发票后...)前插入现结/月结说明
            if ('甲方自收到发票后' in t) and _settle_choice in ('现结', '月结'):
                _sline = '现结：一单一结，验收合格后立即付款；' if _settle_choice == '现结' else '月结：月底按厂家汇总对账，统一生成月度合同后付款；'
                t = t.replace('甲方自收到发票后', _sline + '甲方自收到发票后')
            elif re.search(r'20\d\d年\s*\d+\s*月\s*\d+日', t):
                t = re.sub(r'20\d\d年\s*\d+\s*月\s*\d+日', today_s, t)
            return t
        # 1) 段落: 占位符 + 框架合同字段填充
        for para in doc.paragraphs:
            t = para.text
            for k, v in mapping.items():
                if k in t:
                    t = t.replace(k, v)
            # 框架合同: 合同编码/乙方/合计金额/交付/验收/质保/结算/解除/签署日期
            # (用正则精确匹配空白占位, 不污染行首缩进)
            if '合同编码' in t and 'HT-' not in t:
                t = '合同编码：' + cno
            elif t.strip().startswith('乙方：') and len(t.strip()) <= 5:
                t = '乙方：' + (sup['name'] if sup else (o['supplier'] or '')) + '（供应方）'
            t = _apply_ct(t)
            if t != para.text:
                para.text = t
        # 2) 表格: 先处理所有单元格段落(占位符替换 + 合计金额大写/税金/税率/收款账户等), 再填明细
        for table in doc.tables:
            for _row in table.rows:
                for _cell in _row.cells:
                    for _p in _cell.paragraphs:
                        if _p.text.strip():
                            _nt = _p.text
                            for _k, _v in mapping.items():
                                if _k in _nt:
                                    _nt = _nt.replace(_k, _v)
                            _nt = _apply_ct(_nt)
                            if _nt != _p.text:
                                _p.text = _nt
            rows = table.rows
            if len(rows) < 3:
                continue
            header = [c.text.strip() for c in rows[0].cells]
            if '标的物' in header or '标的' in header[0] or '品名' in header[0]:
                # 明细行: 多商品订单取 order_items 逐行填充, 旧单取订单单商品
                det_rows = _oi if _oi else [None]
                # 先定位合计行(含"合计"字样的行)作为明细边界
                total_row_i = None
                for i in range(1, len(rows)):
                    cells = [cc.text.strip() for cc in rows[i].cells]
                    if any('合计' in cc for cc in cells):
                        total_row_i = i
                        break
                # 再找第一个可写空行(合计行之前)
                idx = 1
                boundary = total_row_i if total_row_i is not None else len(rows)
                for i in range(1, boundary):
                    cells = [cc.text.strip() for cc in rows[i].cells]
                    if not any(cells) or all(cc.strip() == '' for cc in cells):
                        idx = i; break
                else:
                    idx = boundary  # 合计行前无空行 → 从合计行位置开始填(会先插入)
                try:
                    # 需要行数 > 可用空行 → 插入新行(在合计行之前)
                    need = len(det_rows)
                    avail = boundary - idx
                    while avail < need:
                        from docx.oxml.ns import qn as _qn
                        new_tr = rows[idx]._tr.makeelement(_qn('w:tr'), {})
                        for _ in range(len(rows[idx].cells)):
                            tc = rows[idx]._tr.makeelement(_qn('w:tc'), {})
                            new_tr.append(tc)
                        if total_row_i is not None:
                            rows[total_row_i]._tr.addprevious(new_tr)
                            total_row_i += 1
                        else:
                            table._tbl.append(new_tr)
                        avail += 1
                    rows = table.rows
                    # 逐行填充
                    for k, oi_row in enumerate(det_rows):
                        if oi_row is not None:
                            qty_v = oi_row['quantity']
                            qty_s = str(int(qty_v)) if float(qty_v).is_integer() else str(qty_v)
                            line = [oi_row['item_name'] or '', oi_row['spec'] or '', oi_row['unit'] or '',
                                    qty_s, f"{oi_row['price'] or 0:,.2f}", f"{oi_row['amount'] or 0:,.2f}", '']
                        else:
                            qty_v = o['quantity']
                            qty_s = str(int(qty_v)) if float(qty_v).is_integer() else str(qty_v)
                            line = [o['item_name'] or '', o['spec'] or '', o['unit'] or '',
                                    qty_s, f"{o['price'] or 0:,.2f}", f"{amt:,.2f}", '']
                        tr = rows[idx + k]
                        for j, val in enumerate(line):
                            if j < len(tr.cells):
                                tr.cells[j].text = str(val)
                    # 合计行: 可能有"合计"标签行 + "合计金额：¥   元"行, 两处都要填
                    for i in range(idx + len(det_rows), len(rows)):
                        cells = [cc.text.strip() for cc in rows[i].cells]
                        joined = ' | '.join(cells)
                        if any('合计' in cc for cc in cells):
                            if '合计金额' in joined:
                                # 只替换"合计金额：¥"后的数字(税金/不含税/大写已由 _apply_ct 填好, 不能整格替换)
                                for cell in rows[i].cells:
                                    if '合计金额：¥' in cell.text:
                                        cell.text = re.sub(r'(合计金额：¥)[\d,\.\s]*(元)',
                                                           lambda m: f'{m.group(1)}{total:,.2f}{m.group(2)}', cell.text)
                                        break
                            else:
                                # 照片格式: 合计行金额列写纯数字(如 2800.00)
                                rows[i].cells[-2].text = f"{total:,.2f}"
                            # 继续检查下一行是否也是"合计金额"行
                except Exception:
                    pass
            for row in rows:
                for cell in row.cells:
                    for k, v in mapping.items():
                        if k in cell.text:
                            cell.text = cell.text.replace(k, v)
        # V11.203 模块一1.1: 发票条款注入 — 在结算付款条款段(甲方自收到发票后...)后插入独立发票条款段(每份合同可编辑区域)
        if inv_clause or inv_first or inv_done:
            _inv_txt = '发票条款：' + (inv_clause or '按双方协商约定开票')
            if inv_first:
                _inv_txt += '；预计首次开票时间：' + inv_first
            if inv_done:
                _inv_txt += '；预计全部开票完成时间：' + inv_done
            _inv_txt += '。'
            _anchor = None
            for para in doc.paragraphs:
                if '甲方自收到发票后' in para.text:
                    _anchor = para
                    break
            if _anchor is not None:
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _QN
                _p = _OE('w:p')
                _r = _OE('w:r')
                _t = _OE('w:t')
                _t.text = _inv_txt
                _t.set(_QN('xml:space'), 'preserve')
                _r.append(_t)
                _p.append(_r)
                _anchor._p.addnext(_p)
        fname = f"contract_{cno}.docx"
        _fpath = os.path.join(BASE, 'uploads', fname)
        # V11.164: 防编号复用覆盖历史文件 — 目标文件已存在且无合同记录引用(孤儿残留, 如清理过contracts表)时先删除再生成
        if os.path.exists(_fpath):
            _ref = conn.execute("SELECT COUNT(*) FROM contracts WHERE file_path=?", (fname,)).fetchone()[0]
            if _ref == 0:
                os.remove(_fpath)
        doc.save(_fpath)
        # 合同全文(供在线编辑)
        full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        for t in doc.tables:
            for row in t.rows:
                full_text += '\n' + ' | '.join(c.text for c in row.cells)
    except Exception as e:
        conn.close(); return jsonify({'error': f'合同生成失败: {e}'}), 500
    conn.execute("""INSERT INTO contracts(contract_no,order_id,contract_name,supplier,amount,sign_date,start_date,end_date,content,file_path,status,remark,created_at,updated_at,invoice_clause,invoice_est_first,invoice_est_done)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (cno, oid, f"{o['item_name']}采购合同", o['supplier'] or '', o['total_amount'] or 0, (o['created_at'] or '')[:10],
         (o['created_at'] or '')[:10], o['target_date'], full_text, fname, '待审批', f"由订单{o['order_no']}自动生成", datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'), datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
         inv_clause, inv_first, inv_done))
    cid = conn.execute("SELECT id FROM contracts WHERE contract_no=?", (cno,)).fetchone()[0]
    conn.commit()
    create_approvals('contract', cid, o['total_amount'] or 0, submitter=session['user_name'])
    start_instances('contract', cid)
    conn.close()
    log(session['user_name'], '自动生成合同', f'{cno} 订单{o["order_no"]} 模式:{tm}')
    return jsonify({'success': True, 'contract_no': cno, 'file': fname})

# ---- 合同在线编辑: 读取文本(旧合同自动从Word提取) ----
@app.route('/api/contracts/<int:cid>/content')
@login_required
def api_get_contract_content(cid):
    conn = db()
    ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
    if not ct:
        conn.close(); return jsonify({'error': '合同不存在'}), 404
    text = ct['content'] or ''
    if not text and ct['file_path']:
        try:
            from docx import Document
            doc = Document(os.path.join(BASE, 'uploads', ct['file_path']))
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            for t in doc.tables:
                for row in t.rows:
                    text += '\n' + ' | '.join(cc.text for cc in row.cells)
        except Exception:
            pass
    conn.close()
    return jsonify({'content': text})

# ---- 合同在线编辑: 修改文本→重建Word文件归档 ----
@app.route('/api/contracts/<int:cid>/content', methods=['POST'])
@login_required
def api_save_contract_content(cid):
    """需求44: 允许操作人员手动修改合同文本、补充特殊约定条款;
    保存后按编辑文本重建Word合同文件, 与原订单绑定归档"""
    d = request.json
    text = (d.get('content') or '').strip()
    if not text:
        return jsonify({'error': '合同内容不能为空'}), 400
    conn = db()
    ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
    if not ct:
        conn.close(); return jsonify({'error': '合同不存在'}), 404
    conn.execute("UPDATE contracts SET content=?, updated_at=? WHERE id=?", (text, now(), cid))
    # 重建 Word 文件: 保留明细表格, 正文段落替换为编辑后文本
    try:
        from docx import Document
        from docx.oxml.ns import qn
        orig = os.path.join(BASE, 'uploads', ct['file_path']) if ct['file_path'] else ''
        fname = ct['file_path'] or f"contract_{ct['contract_no']}.docx"
        if orig and os.path.exists(orig):
            doc = Document(orig)
        else:
            tpl_def = conn.execute("SELECT file_path FROM contract_templates WHERE is_default=1 AND status='启用'").fetchone()
            doc = Document(os.path.join(BASE, 'uploads', tpl_def['file_path']))
        body = doc.element.body
        # 清除所有段落(保留表格)
        for child in list(body):
            if child.tag != qn('w:tbl'):
                body.remove(child)
        # 按编辑文本逐行重建段落, 插到表格之前
        lines = text.split('\n')
        tbl = body.find(qn('w:tbl'))
        for line in lines:
            p = doc.add_paragraph(line)
            if tbl is not None:
                tbl.addprevious(p._element)
        doc.save(os.path.join(BASE, 'uploads', fname))
    except Exception as e:
        conn.close(); return jsonify({'error': f'合同文件更新失败: {e}'}), 500
    conn.commit(); conn.close()
    log(session['user_name'], '编辑合同', f'#{cid} {ct["contract_no"]}')
    return jsonify({'success': True, 'file': fname})

# ---- V11.196 工作台公告: 发布范围/置顶/生效失效/到期自动隐藏/操作留痕 ----
def _notice_can_manage():
    """公告管理权限: 系统管理员 或 被授权发布公告的用户(可给分管领导等)"""
    me_name = session.get('user_name', '')
    if session.get('user_role') == '系统管理员':
        return True
    pub = (cfg_get('notice_publishers') or '').strip()
    return bool(pub) and me_name in [x.strip() for x in pub.split(',') if x.strip()]

def _notice_visible_sql(conn, me_role, me_name):
    """公告可见性: 状态=已发布 且 生效时间<=now 且 (失效时间为空 或 >now) 且 范围匹配当前用户"""
    now_s = now()
    rows = conn.execute("SELECT * FROM notices WHERE status='已发布' AND (effective_at='' OR effective_at<=?) AND (expire_at='' OR expire_at>?) ORDER BY pinned DESC, publish_at DESC, id DESC",
                        (now_s, now_s)).fetchall()
    out = []
    for r in rows:
        sc = r['scope'] or 'all'
        ok = False
        if sc == 'all':
            ok = True
        elif sc.startswith('roles:'):
            roles = [x.strip() for x in sc.split(':', 1)[1].split(',') if x.strip()]
            ok = me_role in roles
        elif sc.startswith('users:'):
            users = [x.strip() for x in sc.split(':', 1)[1].split(',') if x.strip()]
            ok = me_name in users
        else:
            ok = True
        if ok:
            out.append(dict_row(r))
    return out

# ---- V11.203 模块一1.1/1.2: 合同发票计划/催收状态保存(采购员登记发票条款/开票计划/标记催收, 老合同也可补录) ----
@app.route('/api/contracts/<int:cid>/invoice-plan', methods=['POST'])
@login_required
def api_contract_invoice_plan(cid):
    """字段均为可选项, 传了才更新(未传保留原值): invoice_clause/invoice_est_first/invoice_est_done/inv_collect_status"""
    d = request.json or {}
    conn = db()
    ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
    if not ct:
        conn.close(); return jsonify({'error': '合同不存在'}), 404
    sets, params = [], []
    for _k in ('invoice_clause', 'invoice_est_first', 'invoice_est_done', 'inv_collect_status'):
        if _k in d and d[_k] is not None:
            sets.append(_k + '=?')
            params.append(str(d[_k]).strip()[:500])
    if not sets:
        conn.close(); return jsonify({'error': '没有可保存的内容'}), 400
    sets.append('updated_at=?')
    params.append(now())
    params.append(cid)
    conn.execute("UPDATE contracts SET " + ', '.join(sets) + " WHERE id=?", params)
    conn.commit()
    _log_txt = '、'.join('%s=%s' % (s.split('=')[0], p) for s, p in zip(sets, params[:-2]))
    conn.close()
    log(session['user_name'], '更新合同发票计划', '%s %s' % (ct['contract_no'], _log_txt))
    return jsonify({'success': True, 'contract_no': ct['contract_no']})

# ---- V11.203 模块一1.2: 合同发票登记台账 API + 发票节点自动提醒引擎 ----
def _contract_inv_stats(c, cid):
    """某合同已收发票张数/已收金额"""
    _r = c.execute("SELECT COUNT(*) n, COALESCE(SUM(amount),0) amt FROM contract_invoices WHERE contract_id=?", (cid,)).fetchone()
    return {'received_count': _r['n'], 'received_amount': float(_r['amt'] or 0)}


@app.route('/api/contracts/<int:cid>/invoices')
@login_required
def api_contract_invoices(cid):
    conn = db()
    ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
    if not ct:
        conn.close(); return jsonify({'error': '合同不存在'}), 404
    rows = conn.execute("SELECT * FROM contract_invoices WHERE contract_id=? ORDER BY id DESC", (cid,)).fetchall()
    st = _contract_inv_stats(conn, cid)
    _plan = {'invoice_clause': ct['invoice_clause'] or '', 'invoice_est_first': ct['invoice_est_first'] or '',
             'invoice_est_done': ct['invoice_est_done'] or '', 'inv_collect_status': ct['inv_collect_status'] or ''}
    conn.close()
    return jsonify({'list': [dict_row(x) for x in rows], 'stats': st,
                    'contract_amount': float(ct['amount'] or 0), 'plan': _plan})


@app.route('/api/contracts/<int:cid>/invoices', methods=['POST'])
@login_required
def api_contract_invoice_register(cid):
    """采购专员登记已收到发票: 号码/金额/类型/收票日期 → 自动更新合同发票台账与回收状态"""
    d = request.json or {}
    no = str(d.get('invoice_no') or '').strip()
    try:
        amt = float(d.get('amount') or 0)
    except Exception:
        amt = 0
    if not no:
        return jsonify({'error': '请填写发票号码'}), 400
    if amt <= 0:
        return jsonify({'error': '请填写正确的开票金额'}), 400
    conn = db()
    ct = conn.execute("SELECT * FROM contracts WHERE id=?", (cid,)).fetchone()
    if not ct:
        conn.close(); return jsonify({'error': '合同不存在'}), 404
    conn.execute("""INSERT INTO contract_invoices(contract_id,invoice_no,amount,invoice_type,received_date,operator,remark)
                    VALUES(?,?,?,?,?,?,?)""",
                 (cid, no, amt, (d.get('invoice_type') or '').strip()[:10],
                  (d.get('received_date') or datetime.date.today().strftime('%Y-%m-%d')).strip()[:10],
                  session.get('user_name', '系统'), (d.get('remark') or '').strip()[:200]))
    # 登记发票=已收到发票, 自动更新催收状态
    conn.execute("UPDATE contracts SET inv_collect_status='已收到发票', updated_at=? WHERE id=?", (now(), cid))
    conn.commit()
    conn.close()
    log(session['user_name'], '登记发票', '%s 发票号%s ¥%.2f' % (ct['contract_no'], no, amt))
    return jsonify({'success': True})


@app.route('/api/contracts/<int:cid>/invoices/<int:iid>/delete', methods=['POST'])
@login_required
def api_contract_invoice_delete(cid, iid):
    """登记错误时删除该发票记录(留操作日志); 删除后回收状态退回'已催收待回'由采购员重标"""
    conn = db()
    row = conn.execute("SELECT * FROM contract_invoices WHERE id=? AND contract_id=?", (iid, cid)).fetchone()
    if not row:
        conn.close(); return jsonify({'error': '发票记录不存在'}), 404
    conn.execute("DELETE FROM contract_invoices WHERE id=?", (iid,))
    conn.execute("UPDATE contracts SET inv_collect_status='已催收待回', updated_at=? WHERE id=?", (now(), cid))
    conn.commit()
    conn.close()
    log(session['user_name'], '删除发票记录', '合同发票登记记录 #%s' % iid)
    return jsonify({'success': True})


# 发票节点提醒节流(最多每15分钟扫一次, 每合同每天每类由 contract_inv_reminds 唯一约束去重)
_INV_REMIND_TS = [0.0]


def check_invoice_node_reminders():
    """V11.203 模块一1.2: 按合同发票节点自动提醒对应采购专员(系统内预警+dashboard见另一函数, 此函数负责钉钉推送)
    due=到了预计首次开票日仍未收到任何发票; overdue=超过预计全部开票完成日仍有未收金额"""
    import time as _time
    if time.time() - _INV_REMIND_TS[0] < 900:
        return []
    _INV_REMIND_TS[0] = time.time()
    _today = datetime.date.today().strftime('%Y-%m-%d')
    _out = []
    try:
        c = db()
        rows = c.execute("""SELECT id, contract_no, supplier, amount, invoice_est_first, invoice_est_done
                            FROM contracts WHERE status='执行中' AND (invoice_est_first!='' OR invoice_est_done!='')""").fetchall()
        for r in rows:
            st = _contract_inv_stats(c, r['id'])
            pend = float(r['amount'] or 0) - st['received_amount']
            kinds = []
            if r['invoice_est_first'] and _today >= r['invoice_est_first'][:10] and st['received_count'] == 0:
                kinds.append(('due', '预计首次开票时间%s已到, 尚未收到任何发票' % r['invoice_est_first'][:16]))
            if r['invoice_est_done'] and _today > r['invoice_est_done'][:10] and pend > 0.01:
                kinds.append(('overdue', '超过约定开票完成时间%s, 仍未收票¥%.2f' % (r['invoice_est_done'][:16], pend)))
            for kind, _why in kinds:
                _ex = c.execute("SELECT 1 FROM contract_inv_reminds WHERE contract_id=? AND remind_date=? AND kind=?",
                                (r['id'], _today, kind)).fetchone()
                if _ex:
                    continue
                _usr = find_doc_submitter('contract', r['id'])
                _uid = ''
                if _usr:
                    _uid = str(_usr.get('dingtalk_userid') or '') if isinstance(_usr, dict) else ''
                _title = '🧾 发票催收提醒（%s）' % r['contract_no']
                _txt = ('合同编号：%s\n供应商：%s\n合同金额：¥%.2f\n已收发票：%d张 ¥%.2f\n待催收金额：¥%.2f\n%s\n'
                        '请线下联系供应商取票，收到后在系统合同详情【登记发票】。' %
                        (r['contract_no'], r['supplier'] or '-', float(r['amount'] or 0),
                         st['received_count'], st['received_amount'], max(pend, 0), _why))
                _ok = 0
                if _uid and dingtalk_enabled():
                    try:
                        _ok = 1 if dt_send_todo([_uid], _title, _txt, biz_type='contract', biz_id=r['id'], push_type='alert') else 0
                    except Exception:
                        _ok = 0
                try:
                    c.execute("INSERT INTO contract_inv_reminds(contract_id,remind_date,kind,pushed) VALUES(?,?,?,?)",
                              (r['id'], _today, kind, _ok))
                    c.commit()
                except Exception:
                    pass
                _out.append({'contract_id': r['id'], 'contract_no': r['contract_no'], 'kind': kind, 'pushed': _ok})
        c.close()
    except Exception:
        pass
    return _out


@app.route('/api/notices', methods=['GET'])
@login_required
def api_notices():
    """工作台公告列表(仅返回当前用户可见且在有效期内的已发布公告)"""
    conn = db()
    me_role = session.get('user_role', '')
    me_name = session.get('user_name', '')
    rows = _notice_visible_sql(conn, me_role, me_name)
    conn.close()
    return jsonify(rows)

@app.route('/api/notices/manage', methods=['GET'])
@login_required
def api_notices_manage():
    """公告管理列表(全部状态含草稿/已撤销; 仅管理员/授权发布人可见)"""
    if not _notice_can_manage():
        return jsonify({'error': '无权限：仅系统管理员或公告发布授权用户可管理公告'}), 403
    conn = db()
    rows = conn.execute("SELECT * FROM notices ORDER BY pinned DESC, id DESC").fetchall()
    conn.close()
    return jsonify([dict_row(x) for x in rows])

@app.route('/api/notices/manage', methods=['POST'])
@login_required
def api_notices_save():
    """新建/修改公告: draft=True存草稿 | publish=True发布 | 支持富文本HTML+图片/范围/置顶/生效失效"""
    if not _notice_can_manage():
        return jsonify({'error': '无权限：仅系统管理员或公告发布授权用户可管理公告'}), 403
    d = request.json or {}
    nid = int(d.get('id') or 0)
    title = (d.get('title') or '').strip()
    if not title:
        return jsonify({'error': '公告标题不能为空'}), 400
    content = d.get('content') or ''
    scope = d.get('scope') or 'all'
    # 校验范围格式
    if not scope.startswith(('all', 'roles:', 'users:')):
        return jsonify({'error': '发布范围格式错误'}), 400
    pinned = 1 if d.get('pinned') else 0
    eff = (d.get('effective_at') or '').strip()
    exp = (d.get('expire_at') or '').strip()
    status = '草稿'
    do_publish = bool(d.get('publish'))
    if do_publish:
        # 发布时若生效时间为空则立即生效; 校验失效>生效
        eff = eff or now()[:10]
        if exp and eff and exp < eff:
            return jsonify({'error': '失效时间不能早于生效时间'}), 400
        status = '已发布'
    conn = db()
    if nid:
        conn.execute("UPDATE notices SET title=?, content=?, scope=?, pinned=?, effective_at=?, expire_at=?, status=?, updated_at=? WHERE id=?",
                     (title, content, scope, pinned, eff, exp, status, now(), nid))
    else:
        cur = conn.execute("INSERT INTO notices(title,content,scope,pinned,status,publisher,effective_at,expire_at,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                           (title, content, scope, pinned, status, session.get('user_name', ''), eff, exp, now(), now()))
        nid = cur.lastrowid
    if do_publish:
        conn.execute("UPDATE notices SET publish_at=?, status='已发布' WHERE id=?", (now(), nid))
    conn.execute("INSERT INTO notice_action_logs(notice_id,title,action,operator,detail,created_at) VALUES(?,?,?,?,?,?)",
                 (nid, title, '发布' if do_publish else ('新建' if not d.get('is_update') else '修改'),
                  session.get('user_name', ''), f'范围:{scope} 置顶:{pinned} 生效:{eff or "立即"} 失效:{exp or "永久"}', now()))
    conn.commit(); conn.close()
    log(session['user_name'], '公告' + ('发布' if do_publish else '保存'), f'{title} ({"发布" if do_publish else "草稿"})')
    return jsonify({'success': True, 'id': nid, 'message': ('公告已发布' if do_publish else '公告已保存（草稿）')})

@app.route('/api/notices/manage/<int:nid>/<action>', methods=['POST'])
@login_required
def api_notices_action(nid, action):
    """撤销/删除/置顶切换 — 操作留痕"""
    if not _notice_can_manage():
        return jsonify({'error': '无权限'}), 403
    conn = db()
    n = conn.execute("SELECT * FROM notices WHERE id=?", (nid,)).fetchone()
    if not n:
        conn.close(); return jsonify({'error': '公告不存在'}), 404
    title = n['title']
    detail = ''
    if action == 'revoke':
        conn.execute("UPDATE notices SET status='已撤销', updated_at=? WHERE id=?", (now(), nid))
        detail = '公告已撤销（不再展示）'
    elif action == 'pin':
        conn.execute("UPDATE notices SET pinned=1-pinned, updated_at=? WHERE id=?", (now(), nid))
        detail = '置顶' if not n['pinned'] else '取消置顶'
    elif action == 'delete':
        conn.execute("DELETE FROM notices WHERE id=?", (nid,))
        conn.execute("DELETE FROM notice_action_logs WHERE notice_id=?", (nid,))
        conn.commit(); conn.close()
        log(session['user_name'], '删除公告', title)
        return jsonify({'success': True, 'message': '公告已删除'})
    else:
        conn.close(); return jsonify({'error': '未知操作'}), 400
    conn.execute("INSERT INTO notice_action_logs(notice_id,title,action,operator,detail,created_at) VALUES(?,?,?,?,?,?)",
                 (nid, title, action, session.get('user_name', ''), detail, now()))
    conn.commit(); conn.close()
    log(session['user_name'], f'公告{action}', f'{title} {detail}')
    return jsonify({'success': True, 'message': detail})

@app.route('/api/notices/logs')
@login_required
def api_notices_logs():
    """公告操作日志(管理页展示, 仅管理员/授权发布人)"""
    if not _notice_can_manage():
        return jsonify({'error': '无权限'}), 403
    conn = db()
    rows = conn.execute("SELECT * FROM notice_action_logs ORDER BY id DESC LIMIT 100").fetchall()
    conn.close()
    return jsonify([dict_row(x) for x in rows])

# ---- 文件上传: 申请附件/付款附件/合同模板等 ----
@app.route('/api/upload', methods=['POST'])
@login_required
def api_upload():
    """通用文件上传: multipart 表单 file 字段 → 存 uploads/ → 返回 {file_path}"""
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'error': '未选择文件'}), 400
    fn = os.path.basename(f.filename)  # 去路径, 防穿越
    if not fn or fn in ('.', '..'):
        return jsonify({'error': '非法文件名'}), 400
    # 限制扩展名: 文档/图片/表格/视频(V11.26: 入库验收照片视频)
    ext = os.path.splitext(fn)[1].lower()
    allow = {'.doc', '.docx', '.pdf', '.xls', '.xlsx', '.png', '.jpg', '.jpeg', '.gif', '.txt', '.zip', '.rar', '.mp4', '.mov', '.avi', '.webm'}
    if ext not in allow:
        return jsonify({'error': f'不支持的文件类型 {ext or "空"}（允许: doc/docx/pdf/xls/xlsx/图片/视频/zip/rar）'}), 400
    os.makedirs(os.path.join(BASE, 'uploads'), exist_ok=True)
    # 时间戳前缀防重名
    ts = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
    store = f"{ts}_{fn}"
    f.save(os.path.join(BASE, 'uploads', store))
    log(session['user_name'], '上传文件', store)
    return jsonify({'success': True, 'file_path': store})

# ---- 文件下载/预览: 合同Word/附件等 ----
@app.route('/uploads/<path:filename>')
def api_upload_file(filename):
    """下载/预览上传的文件(合同docx等), 安全校验防路径穿越"""
    if '..' in filename or filename.startswith('/'):
        return 'bad path', 400
    d = os.path.join(BASE, 'uploads')
    full = os.path.join(d, filename)
    if not os.path.exists(full):
        return '文件不存在', 404
    from flask import send_from_directory
    if filename.lower().endswith('.docx'):
        # V11.165: docx 强制下载(attachment) — Windows 浏览器(Chrome/Edge)无 docx 内置预览,
        # inline 响应会"打不开/下载不了"; attachment 让 Windows 直接下载文件
        return send_from_directory(d, filename, as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name=filename)
    return send_from_directory(d, filename, as_attachment=True)

# ---- V55: 采购申请单下载(单个申请生成标准xlsx含明细行) ----
def append_reject_rows(ws, start_row, biz_type, biz_id, ncols=11, CN=None):
    """V11.175/V11.184: 导出/打印单据时在底部追加审批流转日志(同意/驳回全部列出, 含附件名)。
    数据源: approval_action_logs(完整操作+附件) 为主; 无则回退驳回记录。
    返回下一可用行号; 无记录时原样返回 start_row"""
    try:
        acts = get_approval_action_logs(biz_type, biz_id)
        logs = get_reject_logs(biz_type, biz_id)
        if not acts and not logs:
            return start_row
        def _f(bold=False, size=10, color=None):
            if CN:
                return CN(bold=bold, size=size)
            from openpyxl.styles import Font
            kw = {'name': '宋体', 'bold': bold, 'size': size}
            if color:
                kw['color'] = color
            return Font(**kw)
        from openpyxl.styles import Alignment, Border, Side, PatternFill
        thin = Side(style='thin'); border = Border(left=thin, right=thin, top=thin, bottom=thin)
        log_fill = PatternFill('solid', fgColor='FFF8E1')   # 审批日志浅黄
        red_fill = PatternFill('solid', fgColor='FDECEC')   # 驳回浅红
        r = start_row + 1
        # 用 action_logs(同意+驳回+附件) 为主展示
        items = acts if acts else [dict(l, action='reject') for l in logs if isinstance(l, dict)]
        title = '📋 审批流转日志（%d条）' % len(items)
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncols)
        c = ws.cell(row=r, column=1, value=title)
        c.font = _f(bold=True, size=11, color='1F6FEB')
        c.fill = log_fill
        c.alignment = Alignment(horizontal='left', vertical='center')
        for rr in range(1, ncols + 1):
            ws.cell(row=r, column=rr).border = border
            ws.cell(row=r, column=rr).fill = log_fill
        ws.row_dimensions[r].height = 18
        r += 1
        for it in items:
            is_rej = it.get('action') == 'reject' or (not it.get('action') and it.get('source') == 'dingtalk' and '驳回' in str(it.get('comment') or ''))
            _src = '【钉钉】' if it.get('source') == 'dingtalk' else '【系统】'
            _act = '驳回' if (it.get('action') == 'reject' or (it.get('action') != 'agree' and is_rej)) else ('同意' if it.get('action') == 'agree' else '驳回')
            _att_txt = ''
            try:
                _atts = it.get('attachments') or []
                if _atts:
                    _att_txt = '  附件: ' + '、'.join(str(a.get('fileName') or '附件') for a in _atts)
            except Exception:
                pass
            txt = '%s %s 审批人：%s    时间：%s    意见：%s%s' % (_src, _act, it.get('approver') or '钉钉',
                                                          (it.get('processed_at') or '')[:16],
                                                          it.get('comment') or '-', _att_txt)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncols)
            c = ws.cell(row=r, column=1, value=txt)
            c.font = _f(size=10)
            c.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            for rr in range(1, ncols + 1):
                ws.cell(row=r, column=rr).border = border
                if _act == '驳回' or is_rej:
                    ws.cell(row=r, column=rr).fill = red_fill
            ws.row_dimensions[r].height = 20
            r += 1
        return r
    except Exception:
        return start_row


@app.route('/api/prequests/<int:rid>/download')
@login_required
def api_prequest_download(rid):
    """按桌面模板《物资采购申请单》格式生成 xlsx:
    A1:L2 合并标题行(含申请部门/申请日期/编号), 12列表头, 明细, 底部签字区"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    conn = db()
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (rid,)).fetchone()
    if not pr:
        conn.close(); return jsonify({'error': '申请单不存在'}), 404
    items = conn.execute("SELECT * FROM request_items WHERE req_id=? ORDER BY id", (rid,)).fetchall()
    # 库存量(按物品名称汇总, 供 J列展示)
    stock_map = {}
    for inv in conn.execute("SELECT item_name, quantity FROM inventory WHERE quantity>0").fetchall():
        stock_map[inv['item_name']] = stock_map.get(inv['item_name'], 0) + inv['quantity']
    conn.close()
    wb = Workbook(); ws = wb.active; ws.title = '物资采购申请单'
    thin = Side(style='thin'); border = Border(left=thin, right=thin, top=thin, bottom=thin)
    fill = PatternFill('solid', fgColor='D9E2F3')
    # V8.4: 全表统一宋体(兼容WPS/其他电脑)
    CN = lambda bold=False, size=11: Font(name='宋体', bold=bold, size=size)
    # ── 标题行 A1:L2 合并: 标题 + 部门/日期/编号 ──
    ws.merge_cells('A1:K2')
    d = pr['apply_date'] or pr['created_at'] or ''
    ds = str(d)[:10]
    try:
        if '年' in ds:
            # 中文格式 2026年8月6日
            y, m, dd = re.split(r'[年月日]', ds)[:3]
            date_cn = f'{int(y)}年{int(m)}月{int(dd)}日'
        else:
            y, m, dd = ds.split('-')
            date_cn = f'{y}年{int(m)}月{int(dd)}日'
    except Exception:
        date_cn = ds
    ws['A1'] = f'物资采购申请单\n 申请部门：{pr["dept"] or ""}        申请日期：{date_cn}        编号：{pr["req_no"]}'
    ws['A1'].font = CN(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[1].height = 42
    # ── 表头行 (第3行, 11列 — V11.47去掉到货日期列) ──
    headers = ['序号', '采购类别', '物品名称', '厂家/品牌/技术参数', '规格型号', '单位',
               '数量', '用途', '库存', '请购数量', '备注']
    for j, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=j, value=h)
        cell.font = CN(bold=True); cell.fill = fill; cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[3].height = 28
    # ── 明细行 (第4行起) ──
    r = 4
    from openpyxl.drawing.image import Image as XLImage
    import os as _os
    for i, it in enumerate(items, 1):
        use = (it['remark'] or '').strip() or (pr['purpose'] or '')
        cat = it['category'] if 'category' in it.keys() and it['category'] else ''
        brand = it['brand_param'] if 'brand_param' in it.keys() and it['brand_param'] else ''
        # V11.46: 行级附件图片(备注列插图)
        _attach = (it['attach'] if 'attach' in it.keys() and it['attach'] else '').strip()
        remark_txt = it['remark'] or ''
        # V11.47b: 填了到货日期 → 备注带"需到货:xx"(没填不显示)
        _arr = (it['arrival_date'] if 'arrival_date' in it.keys() and it['arrival_date'] else '') or ''
        if _arr:
            _arr = str(_arr)[:10]
            remark_txt = (remark_txt + ' ' if remark_txt else '') + f'【需到货:{_arr}】'
        vals = [i, cat, it['item_name'], brand, it['spec'] or '', it['unit'] or '个', it['quantity'],
                use, stock_map.get(it['item_name'], 0), it['quantity'], remark_txt]
        for j, v in enumerate(vals, 1):
            cell = ws.cell(row=r, column=j, value=v)
            cell.border = border
            cell.font = CN()
            # V11.44: 自动换行, 长文字(厂家/技术参数/规格/用途)不被遮挡
            cell.alignment = Alignment(horizontal='center' if j in (1, 2, 3, 5, 6, 7, 10, 11) else 'left',
                                       vertical='center', wrap_text=True)
        # V11.46: 备注列插入行级附件图片
        if _attach:
            _img_path = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), 'uploads', _attach.split('/')[-1])
            if _os.path.exists(_img_path):
                try:
                    img = XLImage(_img_path)
                    img.width = min(img.width, 90); img.height = min(img.height, 90)
                    _anchor = f'K{r}'
                    ws.add_image(img, _anchor)
                except Exception:
                    pass
        # V11.44: 行高按内容自动计算(参考模板: 内容多行高30-93), 至少22; 有图的行加高到90
        _maxlen = max(len(str(v or '')) for v in vals)
        _rows_h = max(22, min(90, 22 + int(_maxlen / 6) * 9))
        if _attach:
            _rows_h = max(_rows_h, 90)
        ws.row_dimensions[r].height = _rows_h
        r += 1
    # ── 底部签字区 (动态下移) ──
    sign1 = r + 1
    sign2 = sign1 + 2
    ws.merge_cells(f'A{sign1}:K{sign1 + 1}')
    ws.cell(row=sign1, column=1, value='经办人：                        部门负责人：')
    ws.cell(row=sign1, column=1).font = CN()
    ws.cell(row=sign1, column=1).alignment = Alignment(horizontal='left', vertical='center')
    for rr in range(sign1, sign1 + 2):
        for cc in range(1, 12):
            ws.cell(row=rr, column=cc).border = border
    ws.merge_cells(f'A{sign2}:K{sign2 + 2}')
    ws.cell(row=sign2, column=1, value='计划采购负责人：                         采购员：')
    ws.cell(row=sign2, column=1).font = CN()
    ws.cell(row=sign2, column=1).alignment = Alignment(horizontal='left', vertical='center')
    for rr in range(sign2, sign2 + 3):
        for cc in range(1, 12):
            ws.cell(row=rr, column=cc).border = border
    # ── 列宽: 对齐桌面模板 申请单.xlsx ──
    widths = [5.1, 10.9, 24.6, 12.1, 23.4, 5.2, 8.2, 17.2, 7.3, 9.0, 14.0]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + j)].width = w
    # V11.27: 审批通过 → 盖章领导预录签名
    stamp_leader_sign(ws, sign1, 'purchase_request', rid)
    # V11.175: 底部追加驳回审批记录(多次驳回全部列出)
    append_reject_rows(ws, sign2 + 3, 'purchase_request', rid, ncols=11, CN=CN)
    # ── 打印设置: A4 横向 ──
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9  # A4
    ws.page_setup.fitToWidth = 1
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    return send_file(bio, as_attachment=True, download_name=f'{pr["req_no"]}物资采购申请单.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# ---- V11.4: 采购订单下载(单个订单生成标准xlsx, 含商家信息/明细/金额/审批签名) ----
@app.route('/api/orders/<int:oid>/download')
@login_required
def api_order_download(oid):
    """生成标准《采购订单》xlsx: 标题行+基础信息(单号/供应商/交易模式/金额)+商品明细+合计+审批进度(含电子签名)+签字区"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.drawing.image import Image as XLImage
    import base64
    conn = db()
    o = conn.execute("SELECT * FROM purchase_orders WHERE id=?", (oid,)).fetchone()
    if not o:
        conn.close(); return jsonify({'error': '订单不存在'}), 404
    items = conn.execute("SELECT * FROM order_items WHERE order_id=? ORDER BY id", (oid,)).fetchall()
    approvals = conn.execute("SELECT * FROM approval_instances WHERE biz_type='purchase_order' AND biz_id=? ORDER BY level_no", (oid,)).fetchall()
    pr = conn.execute("SELECT * FROM purchase_requests WHERE id=?", (o['req_id'],)).fetchone() if o['req_id'] else None
    conn.close()
    wb = Workbook(); ws = wb.active; ws.title = '采购订单'
    thin = Side(style='thin'); border = Border(left=thin, right=thin, top=thin, bottom=thin)
    fill = PatternFill('solid', fgColor='D9E2F3')
    CN = lambda bold=False, size=11: Font(name='宋体', bold=bold, size=size)
    # ── 标题行 ──
    ws.merge_cells('A1:H1')
    ws['A1'] = '采 购 订 单'
    ws['A1'].font = CN(bold=True, size=18)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 36
    # ── 单据编号行 ──
    ws.merge_cells('A2:H2')
    ws['A2'] = '订单编号：%s' % o['order_no']
    ws['A2'].font = CN(size=11)
    ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[2].height = 20
    # ── 基础信息区 ──
    info = [
        ('供应商', o['supplier'] or '', '交易模式', o['trade_mode'] or ''),
        ('需求部门', o['requester'] or '', '品类', o['category'] or ''),
        ('负责人', o['owner'] or '', '下单日期', str(o['created_at'] or '')[:10]),
        ('目标到货日', o['target_date'] or '', '订单状态', o['status'] or ''),
        ('订单金额', '¥%s' % format(float(o['total_amount'] or 0), ',.2f'), '', ''),
        ('询价/备注', (o['remark'] or '')[:80], '', ''),
    ]
    r = 4
    for lk, lv, rk, rv in info:
        ws.cell(r, 1, lk).font = CN(bold=True); ws.cell(r, 1).fill = fill
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=3)
        ws.cell(r, 2, lv).font = CN()
        if rk:
            ws.cell(r, 4, rk).font = CN(bold=True); ws.cell(r, 4).fill = fill
            ws.merge_cells(start_row=r, start_column=5, end_row=r, end_column=8)
            ws.cell(r, 5, rv).font = CN()
        else:
            ws.merge_cells(start_row=r, start_column=4, end_row=r, end_column=8)
        for cc in range(1, 9):
            ws.cell(r, cc).border = border
        ws.row_dimensions[r].height = 22
        r += 1
    # ── 商品明细 ──
    r += 1
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
    ws.cell(r, 1, '商品明细').font = CN(bold=True); ws.cell(r, 1).fill = fill
    for cc in range(1, 9): ws.cell(r, cc).border = border
    r += 1
    headers = ['序号', '物资名称', '规格型号', '单位', '数量', '单价(元)', '税率%', '金额(元)']
    for j, h in enumerate(headers, 1):
        cell = ws.cell(r, j, h); cell.font = CN(bold=True); cell.fill = fill; cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center')
    r += 1
    total_qty = 0.0; total_amt = 0.0
    for i, it in enumerate(items, 1):
        vals = [i, it['item_name'], it['spec'] or '', it['unit'] or '个', it['quantity'],
                it['price'] or 0, it['tax_rate'] or 13, it['total_amount'] or it['amount'] or 0]
        total_qty += float(it['quantity'] or 0); total_amt += float(it['total_amount'] or it['amount'] or 0)
        for j, v in enumerate(vals, 1):
            cell = ws.cell(r, j, v); cell.border = border; cell.font = CN()
            if j in (1, 3, 4, 5, 6, 7): cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[r].height = 20
        r += 1
    ws.cell(r, 4, '合计').font = CN(bold=True)
    ws.cell(r, 5, total_qty).font = CN(bold=True)
    ws.cell(r, 8, round(total_amt, 2)).font = CN(bold=True)
    for cc in range(1, 9): ws.cell(r, cc).border = border
    # ── 审批进度(含电子签名) ──
    if approvals:
        r += 2
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
        ws.cell(r, 1, '审批进度').font = CN(bold=True); ws.cell(r, 1).fill = fill
        for cc in range(1, 9): ws.cell(r, cc).border = border
        r += 1
        for a in approvals:
            st = {'approved': '✅通过', 'rejected': '❌驳回', 'pending': '⏳待审批'}.get(a['status'], a['status'])
            # V11.6: 钉钉审批记录显示为"钉钉OA电子审批"(审批人实名认证, 与手写签名同效)
            apv = a['approver'] or ''
            if apv == '钉钉':
                apv = '钉钉OA电子审批'
            line = '%s ｜ %s ｜ %s %s' % (a['role'] or '', st, apv, str(a['processed_at'] or '')[:16])
            if a['comment']:
                line += ' ｜ ' + (a['comment'] or '')
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
            ws.cell(r, 1, line).font = CN(size=10)
            for cc in range(1, 9): ws.cell(r, cc).border = border
            # 电子签名图片(系统内手写签名) / 钉钉审批标记
            if a['signature'] and a['signature'].startswith('data:image/png'):
                try:
                    imgdata = base64.b64decode(a['signature'].split(',')[1])
                    img = XLImage(io.BytesIO(imgdata))
                    img.width = 70; img.height = 30
                    ws.add_image(img, 'H%d' % r)
                except Exception:
                    pass
            else:
                ws.cell(r, 8, '✔️' if (a['approver'] or '') == '钉钉' and a['status'] == 'approved' else '').font = Font(name='宋体', size=12, color='2F5597')
            ws.row_dimensions[r].height = 32
            r += 1
    # ── 签字区 ──
    r += 1
    ws.merge_cells(f'A{r}:H{r + 1}')
    ws.cell(row=r, column=1, value='采购经办人：                        部门负责人：')
    ws.cell(row=r, column=1).font = CN()
    ws.cell(row=r, column=1).alignment = Alignment(horizontal='left', vertical='center')
    for rr in range(r, r + 2):
        for cc in range(1, 9): ws.cell(row=rr, column=cc).border = border
    # ── 列宽/打印 ──
    widths = [6, 22, 18, 7, 9, 12, 8, 14]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + j)].width = w
    # V11.27: 审批通过 → 盖章领导预录签名
    stamp_leader_sign(ws, r, 'purchase_order', oid)
    # V11.175: 底部追加驳回审批记录
    append_reject_rows(ws, r + 2, 'purchase_order', oid, ncols=8, CN=CN)
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9
    ws.page_setup.fitToWidth = 1
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    return send_file(bio, as_attachment=True, download_name=f'{o["order_no"]}采购订单.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# ---- 单据导出 xlsx ----
@app.route('/api/export')
@login_required
def api_export():
    t = request.args.get('type', 'orders')
    from openpyxl import Workbook
    import io
    wb = Workbook(); ws = wb.active
    c = db()
    if t == 'orders':
        ws.append(['订单号', '物资', '规格', '数量', '单位', '单价', '含税总价', '供应商', '交易模式', '状态', '创建时间'])
        for r in c.execute("SELECT * FROM purchase_orders ORDER BY id DESC"):
            ws.append([r['order_no'], r['item_name'], r['spec'], r['quantity'], r['unit'], r['price'], r['total_amount'], r['supplier'], r['trade_mode'], r['status'], r['created_at']])
    elif t == 'requests':
        ws.append(['申请单号', '部门', '用途', '金额', '状态', '申请日期'])
        for r in c.execute("SELECT * FROM purchase_requests ORDER BY id DESC"):
            ws.append([r['req_no'], r['dept'], r['purpose'], r['total_estimated'], r['status'], r['created_at']])
    elif t == 'inventory':
        # V55需求9: 库存报表独立导出, 支持按分类展示(参数 cat=品类代码)
        # V5.0: 物料名称前新增【供应商】列
        cat = request.args.get('cat', '')
        ws.append(['序号', '供应商', '物料名称', '规格型号', '品类', '单位', '剩余库存量', '不含税单价', '税率%', '含税单价', '合计', '备注'])
        q = "SELECT * FROM inventory"
        params = ()
        if cat:
            q += " WHERE cat_code=?"
            params = (cat,)
        q += " ORDER BY cat_code, id"
        i = 0
        for r in c.execute(q, params):
            i += 1
            tr = float(r['tax_rate'] or 13)
            pt = round((r['price'] or 0) * (1 + tr / 100.0), 4)
            total = round((r['quantity'] or 0) * pt, 2)
            ws.append([i, r['supplier'] or '', r['item_name'], r['spec'], r['cat_code'], r['unit'], r['quantity'], r['price'] or 0, tr, pt, total, r['remark'] or ''])
        ws.column_dimensions['A'].width = 6; ws.column_dimensions['B'].width = 18
        ws.column_dimensions['C'].width = 16; ws.column_dimensions['D'].width = 8
        ws.column_dimensions['E'].width = 6; ws.column_dimensions['F'].width = 8
        ws.column_dimensions['G'].width = 12; ws.column_dimensions['H'].width = 8
        ws.column_dimensions['I'].width = 12; ws.column_dimensions['J'].width = 12
        ws.column_dimensions['K'].width = 20
    elif t == 'credits':
        ws.append(['挂账单号', '订单', '合同', '供应商', '金额', '状态', '创建时间'])
        for r in c.execute("SELECT * FROM credit_notes ORDER BY id DESC"):
            ws.append([r['credit_no'], r['order_id'], r['contract_no'], r['supplier'], r['amount'], r['status'], r['created_at']])
    elif t == 'payments':
        ws.append(['付款单号', '供应商', '金额', '模式', '状态', '创建时间'])
        for r in c.execute("SELECT * FROM payment_requests ORDER BY id DESC"):
            ws.append([r['payment_no'], r['supplier'], r['amount'], r['trade_mode'], r['status'], r['created_at']])
    elif t == 'contracts':
        ws.append(['合同号', '订单', '名称', '供应商', '金额', '状态', '创建时间'])
        for r in c.execute("SELECT * FROM contracts ORDER BY id DESC"):
            ws.append([r['contract_no'], r['order_id'], r['contract_name'], r['supplier'], r['amount'], r['status'], r['created_at']])
    else:
        ws.append(['名称', '联系人', '电话', '类别', '评级'])
        for r in c.execute("SELECT * FROM suppliers ORDER BY id"):
            ws.append([r['name'], r['contact'], r['phone'], r['category'], r['rating']])
    c.close()
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    return send_file(bio, as_attachment=True, download_name=f'{t}_{today()}.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# ---- 单据作废 ----
@app.route('/api/orders/<int:oid>/void', methods=['POST'])
@login_required
def api_void_order(oid):
    c = db(); c.execute("UPDATE purchase_orders SET status='已作废' WHERE id=?", (oid,)); c.commit(); c.close()
    log(session['user_name'], '作废订单', f'#{oid}')
    return jsonify({'success': True})

@app.route('/api/contracts/<int:cid>/void', methods=['POST'])
@login_required
def api_void_contract(cid):
    c = db(); c.execute("UPDATE contracts SET status='已作废' WHERE id=?", (cid,)); c.commit(); c.close()
    log(session['user_name'], '作废合同', f'#{cid}')
    return jsonify({'success': True})

@app.route('/api/credits/<int:cid>/void', methods=['POST'])
@login_required
def api_void_credit(cid):
    c = db(); c.execute("UPDATE credit_notes SET status='已作废' WHERE id=?", (cid,)); c.commit(); c.close()
    log(session['user_name'], '作废挂账', f'#{cid}')
    return jsonify({'success': True})


@app.route('/api/payment_requests')
@login_required
def api_payment_requests():
    """获取付款申请列表 — V11.159: 仅 财务/领导/管理员 可见(采购员/员工/库管员不看钱)"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    conn = db()
    rows = conn.execute("SELECT * FROM payment_requests ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        out.append(d)
    conn.close()
    return jsonify(out)
@app.route('/api/payments')
@login_required
def api_payments():
    """获取付款列表 — V11.159: 仅 财务/领导/管理员 可见"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify([])
    conn = db()
    rows = conn.execute("SELECT * FROM payment_requests ORDER BY id DESC LIMIT 100").fetchall()
    out = []
    for r in rows:
        d = dict_row(r)
        out.append(d)
    conn.close()
    return jsonify(out)

@app.route('/api/payments/<int:pid>/void', methods=['POST'])
@login_required
def api_void_payment(pid):
    """作废付款 — V11.159: 仅 财务/领导/管理员(付款是财务职能)"""
    if session.get('user_role') not in ('财务', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：付款管理仅限财务/领导'}), 403
    c = db(); c.execute("UPDATE payment_requests SET status='已作废' WHERE id=?", (pid,)); c.commit(); c.close()
    log(session['user_name'], '作废付款', f'#{pid}')
    return jsonify({'success': True})

@app.route('/api/receivings/<int:rid>/void', methods=['POST'])
@login_required
def api_receiving_void(rid):
    """V5.0: 入库单作废 — 已入库(已加库存)作废则回滚库存+写流水; 未入库直接作废"""
    c = db()
    rn = c.execute("SELECT * FROM receivings WHERE id=?", (rid,)).fetchone()
    if not rn:
        c.close(); return jsonify({'error': '入库单不存在'}), 404
    if rn['status'] == '已作废':
        c.close(); return jsonify({'error': '该入库单已作废'}), 400
    if rn['status'] == '已入库':
        # 回滚库存: 删除该单入库流水 + 扣回库存 (V11.202 兼容分批入库流水类型)
        flows = c.execute("SELECT * FROM inventory_flows WHERE doc_type='receiving' AND doc_id=? AND (flow_type='入库' OR flow_type LIKE '分批入库%')", (rid,)).fetchall()
        for f in flows:
            inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? ORDER BY quantity DESC",
                            (f['item_name'], f['spec'] or '')).fetchone()
            if inv:
                new_q = inv['quantity'] - f['qty']
                c.execute("UPDATE inventory SET quantity=?, updated_at=? WHERE id=?", (new_q, now(), inv['id']))
        c.execute("DELETE FROM inventory_flows WHERE doc_type='receiving' AND doc_id=?", (rid,))
        c.execute("UPDATE receivings SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
        c.execute("UPDATE approval_instances SET status='rejected', comment='单据作废' WHERE biz_type='receiving' AND biz_id=? AND status='pending'", (rid,))
        c.commit(); c.close()
        log(session['user_name'], '作废入库单', f'{rn["receive_no"]} (已入库, 库存已回滚)')
        return jsonify({'success': True, 'message': '入库单已作废，库存已回滚'})
    c.execute("UPDATE receivings SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
    c.execute("UPDATE approval_instances SET status='rejected', comment='单据作废' WHERE biz_type='receiving' AND biz_id=? AND status='pending'", (rid,))
    c.commit(); c.close()
    log(session['user_name'], '作废入库单', f'{rn["receive_no"]} (未入库, 库存未变动)')
    return jsonify({'success': True, 'message': '入库单已作废（未入库，库存未变动）'})

# ============================================================
# V11.193 退库模块: 已领用物资退回仓库(非退供应商) — 领用剩余/未使用/错领/质量问题
# 流程: 新建退库单(关联已出库的领用单, 带出物料与可退数量) → 草稿 → 提交钉钉审批
#       → 审批通过=待仓库清点(库存不立即加) → 仓库确认退库入库(加库存+流水+回写源单累计已退)
# ============================================================
_RETURN_REASONS = ['领用剩余', '物料未使用', '错领物料', '质量问题', '其他']

@app.route('/api/returns/source-requisitions')
@login_required
def api_return_source_requisitions():
    """可选源单: 已出库且未全部退完的领用/出库单(仅库管员/领导视角可用列表查询给所有人用于选择?)。
    约束: status='已出库' 且 累计已退<领用总量。返回单+明细(带可退数量/单价)。"""
    c = db()
    # 出库单表头: req_no/dept/receiver/quantity/unit/returned_qty; 明细: requisition_items
    rows = c.execute("""SELECT r.id, r.req_no, r.dept, r.receiver, r.quantity, r.unit, r.returned_qty,
                               r.created_at, r.purpose
                        FROM requisitions r
                        WHERE r.status='已出库'
                          AND COALESCE(r.returned_qty,0) < r.quantity
                        ORDER BY r.id DESC LIMIT 50""").fetchall()
    out = []
    for r in rows:
        its = c.execute("SELECT * FROM requisition_items WHERE requisition_id=? ORDER BY id", (r['id'],)).fetchall()
        d = dict(r)
        d['items'] = [dict_row(x) for x in its]
        # 每个明细: 该源明细累计已退(从return_items按source_item_id聚合) + 可退数量 + 单价(取库存档案价)
        for it in d['items']:
            agg = c.execute("SELECT COALESCE(SUM(return_qty),0) s FROM return_items WHERE source_item_id=?", (it['id'],)).fetchone()
            it['returned_qty'] = float(agg['s'] or 0)
            it['returnable_qty'] = float(it['quantity'] or 0) - it['returned_qty']
            _invp = c.execute("SELECT price FROM inventory WHERE item_name=? AND spec=? AND price IS NOT NULL ORDER BY id LIMIT 1",
                              (it['item_name'], it.get('spec', '') or '')).fetchone()
            it['price'] = float(_invp['price'] or 0) if _invp else 0
        out.append(d)
    c.close()
    return jsonify(out)

@app.route('/api/returns', methods=['GET'])
@login_required
def api_returns():
    """退库列表 — 支持状态/时间筛选; 权限: 库管员/部门负责人/领导/管理员/采购员(看自己提交的)"""
    c = db()
    rows = c.execute("SELECT * FROM return_requests ORDER BY id DESC").fetchall()
    c.close()
    # 行附带明细条数
    out = []
    c2 = db()
    for r in rows:
        d = dict_row(r)
        d['item_count'] = c2.execute("SELECT COUNT(*) n FROM return_items WHERE return_id=?", (r['id'],)).fetchone()['n']
        out.append(d)
    c2.close()
    return jsonify(out)

@app.route('/api/returns/<int:rid>', methods=['GET'])
@login_required
def api_return_detail(rid):
    c = db()
    r = c.execute("SELECT * FROM return_requests WHERE id=?", (rid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '退库单不存在'}), 404
    d = dict_row(r)
    d['items'] = [dict_row(x) for x in c.execute("SELECT * FROM return_items WHERE return_id=? ORDER BY id", (rid,)).fetchall()]
    # 源单信息
    src = c.execute("SELECT * FROM requisitions WHERE id=?", (d['source_req_id'],)).fetchone()
    d['source'] = dict_row(src) if src else None
    # 库存变动记录(退库确认入库流水)
    d['flows'] = [dict_row(x) for x in c.execute("SELECT * FROM inventory_flows WHERE doc_type='return_request' AND doc_id=? ORDER BY id DESC", (rid,)).fetchall()]
    c.close()
    return jsonify(d)

@app.route('/api/returns', methods=['POST'])
@login_required
def api_create_return():
    """新建退库申请单 — body: {source_req_id, reason, reason_note, items:[{source_item_id,item_name,spec,unit,issued_qty,returned_qty,return_qty,price}], attachments}
    保存状态=草稿(提交走审批); 权限: 出库相关角色(库管员/部门负责人/领导/管理员) + 源单领用人本人"""
    d = request.json or {}
    src_id = int(d.get('source_req_id') or 0)
    items = d.get('items') or []
    items = [it for it in items if it.get('item_name') and float(it.get('return_qty', 0) or 0) > 0]
    if not src_id:
        return jsonify({'error': '请选择源出库单（领用单）'}), 400
    if not items:
        return jsonify({'error': '请至少填写一项退库物资及数量'}), 400
    c = db()
    src = c.execute("SELECT * FROM requisitions WHERE id=? AND status='已出库'", (src_id,)).fetchone()
    if not src:
        c.close(); return jsonify({'error': '源出库单不存在或未完成领用出库'}), 400
    # 数量校验: 0 < 本次退库数量 <= 可退数量(可退=领用-已确认入库的已退, 草稿/待审批不占额度)
    for it in items:
        q = float(it.get('return_qty', 0) or 0)
        src_item_id = int(it.get('source_item_id') or 0)
        src_it = c.execute("SELECT * FROM requisition_items WHERE id=? AND requisition_id=?", (src_item_id, src_id)).fetchone() if src_item_id else None
        if not src_it:
            c.close(); return jsonify({'error': f"物资「{it.get('item_name','')}」未匹配到源出库单明细"}), 400
        issued = float(src_it['quantity'] or 0)
        agg = c.execute("""SELECT COALESCE(SUM(t.return_qty),0) s FROM return_items t
                           JOIN return_requests rr ON rr.id=t.return_id
                           WHERE t.source_item_id=? AND rr.status='退库已完成'""", (src_item_id,)).fetchone()
        returned = float(agg['s'] or 0)
        can = issued - returned
        if q <= 0 or q > can + 1e-9:
            c.close(); return jsonify({'error': f"「{src_it['item_name']}」本次退库数量({q:g})须大于0且不超过可退数量({can:g})"}), 400
        it['issued_qty'] = issued
        it['returned_qty'] = returned
    no = gen_no('TK', 'return_requests', 'return_no', c)
    total_amt = sum(float(it.get('price', 0) or 0) * float(it.get('return_qty', 0) or 0) for it in items)
    first = items[0]
    reason = (d.get('reason') or '').strip()
    if reason not in _RETURN_REASONS:
        reason = '其他'
    req_name = session.get('user_name', '')
    cur = c.execute("""INSERT INTO return_requests(return_no,source_req_id,source_req_no,dept,receiver,warehouse,
                       reason,reason_note,total_amount,status,requester,attachments,created_at,updated_at)
                       VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (no, src_id, src['req_no'], src['dept'] or '', src['receiver'] or '',
                     '主库房', reason, (d.get('reason_note') or '').strip()[:200], total_amt,
                     '草稿', req_name, json.dumps(d.get('attachments') or [], ensure_ascii=False), now(), now()))
    rid = cur.lastrowid
    for it in items:
        amt = float(it.get('price', 0) or 0) * float(it.get('return_qty', 0) or 0)
        c.execute("""INSERT INTO return_items(return_id,source_item_id,item_name,spec,unit,issued_qty,returned_qty,return_qty,price,amount,created_at)
                     VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                  (rid, int(it.get('source_item_id') or 0), it['item_name'], it.get('spec', ''), it.get('unit', '个'),
                   float(it.get('issued_qty', 0)), float(it.get('returned_qty', 0)), float(it.get('return_qty', 0)),
                   float(it.get('price', 0) or 0), amt, now()))
    c.commit(); c.close()
    log(req_name, '新建退库单', f'{no} 源单:{src["req_no"]} {len(items)}项 ¥{total_amt:g} 已存草稿')
    return jsonify({'success': True, 'return_no': no, 'id': rid, 'message': f'退库单 {no} 已保存（草稿），确认后请提交审批'})


@app.route('/api/returns/<int:rid>/submit', methods=['POST'])
@login_required
def api_return_submit(rid):
    """退库单提交审批 — 草稿/被驳回(已驳回回草稿) → 待审批 + 建审批实例 + 推钉钉"""
    c = db()
    r = c.execute("SELECT * FROM return_requests WHERE id=?", (rid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '退库单不存在'}), 404
    if r['status'] not in ('草稿', '已驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可提交审批'}), 400
    # 提交人/管理员 校验
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_admin = me and me['role'] == '系统管理员'
    if not is_admin and r['requester'] and me and me['name'] != r['requester']:
        c.close(); return jsonify({'error': '仅退库单提交人本人或管理员可提交审批'}), 403
    # 提交前校验明细数量仍有效(仅已完成退库占额度; 草稿/待审批不占)
    its = c.execute("SELECT * FROM return_items WHERE return_id=?", (rid,)).fetchall()
    for it in its:
        agg = c.execute("""SELECT COALESCE(SUM(t.return_qty),0) s FROM return_items t
                           JOIN return_requests rr ON rr.id=t.return_id
                           WHERE t.source_item_id=? AND rr.status='退库已完成' AND t.return_id<>?""", (it['source_item_id'], rid)).fetchone()
        src_it = c.execute("SELECT * FROM requisition_items WHERE id=?", (it['source_item_id'],)).fetchone() if it['source_item_id'] else None
        if src_it:
            can = float(src_it['quantity'] or 0) - float(agg['s'] or 0)
            if float(it['return_qty']) > can + 1e-9:
                c.close(); return jsonify({'error': f"「{it['item_name']}」可退数量已变化(现可退{can:g})，请修改后重试"}), 400
    # 已驳回重提: 清驳回标记+累计
    if r['status'] == '已驳回':
        c.execute("UPDATE return_requests SET rejected_items='', rejected_reason='', resubmit_count=resubmit_count+1 WHERE id=?", (rid,))
    else:
        c.execute("UPDATE return_requests SET resubmit_count=COALESCE(resubmit_count,0) WHERE id=?", (rid,))
    c.execute("UPDATE return_requests SET status='待审批', updated_at=? WHERE id=?", (now(), rid))
    # 清旧审批实例(重提场景), 重建
    c.execute("DELETE FROM approval_instances WHERE biz_type='return_request' AND biz_id=?", (rid,))
    c.execute("DELETE FROM dingtalk_instances WHERE biz_type='return_request' AND biz_id=?", (rid,))
    c.commit()
    amount = float(r['total_amount'] or 0)
    create_approvals('return_request', rid, amount, submitter=r['requester'] or session.get('user_name', ''))
    c.close()
    try:
        start_instances('return_request', rid)
    except Exception as e:
        print('return submit start_instances err:', e)
    log(session['user_name'], '提交退库审批', f'{r["return_no"]} 待审批 金额¥{amount:g}')
    return jsonify({'success': True, 'message': f'退库单 {r["return_no"]} 已提交审批'})


@app.route('/api/returns/<int:rid>/confirm-warehouse', methods=['POST'])
@login_required
def api_return_confirm_warehouse(rid):
    """仓库实物清点确认 → 退库入库: 库存+数量+流水+回写源单累计已退; 仅 库管员/部门负责人/领导/管理员
    状态机: 审批通过(待仓库清点) → 退库已完成。幂等: 已完成/已有流水跳过。"""
    if session.get('user_role') not in ('库管员', '部门负责人', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：仓库确认入库仅限库管员/领导'}), 403
    c = db()
    r = c.execute("SELECT * FROM return_requests WHERE id=?", (rid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '退库单不存在'}), 404
    if r['status'] == '退库已完成':
        c.close(); return jsonify({'error': '该退库单已完成入库'}), 400
    if r['status'] != '审批通过':
        c.close(); return jsonify({'error': f'仅审批通过的退库单可确认入库（当前:{r["status"]}）'}), 400
    done = c.execute("SELECT 1 FROM inventory_flows WHERE doc_type='return_request' AND doc_id=? AND flow_type='退库' LIMIT 1", (rid,)).fetchone()
    if done:
        c.close(); return jsonify({'error': '该退库单已确认入库（防重复）'}), 400
    its = c.execute("SELECT * FROM return_items WHERE return_id=?", (rid,)).fetchall()
    if not its:
        c.close(); return jsonify({'error': '退库单无明细，无法入库'}), 400
    # 1) 库存增加(按名称+规格匹配行累加; 无则新建) + 2) 流水
    its = [dict_row(x) for x in its]
    for it in its:
        q = float(it['return_qty'] or 0)
        if q <= 0:
            continue
        inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? AND (warehouse=? OR warehouse IS NULL OR warehouse='') ORDER BY quantity DESC LIMIT 1",
                        (it['item_name'], it.get('spec', '') or '', r['warehouse'] or '主库房')).fetchone()
        if inv:
            new_q = float(inv['quantity'] or 0) + q
            c.execute("UPDATE inventory SET quantity=?, updated_at=?, last_move_date=? WHERE id=?", (new_q, now(), now()[:10], inv['id']))
            inv_id = inv['id']
        else:
            cur = c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,price,updated_at) VALUES(?,?,?,?,?,?,?)",
                            (it['item_name'], it.get('spec', '') or '', it.get('unit', '个') or '个', q,
                             r['warehouse'] or '主库房', float(it.get('price', 0) or 0), now()))
            inv_id = cur.lastrowid
        c.execute("""INSERT INTO inventory_flows(flow_type,doc_type,doc_id,doc_no,item_name,spec,qty,balance_after,operator,remark,created_at)
                     VALUES('退库','return_request',?,?,?,?,?,?,?,?,?)""",
                  (rid, r['return_no'], it['item_name'], it.get('spec', '') or '', q,
                   float(c.execute("SELECT quantity FROM inventory WHERE id=?", (inv_id,)).fetchone()['quantity'] or 0),
                   session.get('user_name', ''), f'退库入库: {r["return_no"]}', now()))
        # 回写源出库单明细累计已退: 源单表头 returned_qty += q (按明细对应的源单)
        if it['source_item_id']:
            c.execute("""UPDATE requisitions SET returned_qty=COALESCE(returned_qty,0)+?
                         WHERE id=(SELECT requisition_id FROM requisition_items WHERE id=?)""", (q, it['source_item_id']))
    c.execute("UPDATE return_requests SET status='退库已完成', warehouse_confirm_by=?, warehouse_confirm_at=?, finished_at=?, updated_at=? WHERE id=?",
              (session.get('user_name', ''), now(), now(), now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '退库确认入库', f'{r["return_no"]} 库存已增加, 源单{r["source_req_no"]}累计已退回写')
    return jsonify({'success': True, 'message': f'退库单 {r["return_no"]} 确认入库完成，库存已增加'})


@app.route('/api/returns/<int:rid>/update', methods=['POST'])
@login_required
def api_return_update(rid):
    """V11.193: 修改退库单(草稿/被驳回回草稿状态) — 原因/说明/明细数量可改; 数量强校验同新建"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM return_requests WHERE id=?", (rid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '退库单不存在'}), 404
    if r['status'] not in ('草稿', '已驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可修改（已完成/审批中单据不可改）'}), 400
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_admin = me and me['role'] == '系统管理员'
    if not is_admin and r['requester'] and me and me['name'] != r['requester']:
        c.close(); return jsonify({'error': '仅退库单提交人或管理员可修改'}), 403
    items = d.get('items') or []
    items = [it for it in items if it.get('item_name') and float(it.get('return_qty', 0) or 0) > 0]
    if not items:
        c.close(); return jsonify({'error': '请至少保留一项退库物资及数量'}), 400
    # 数量校验: 已确认入库的占额度(本单自身未确认, 不算)
    for it in items:
        q = float(it.get('return_qty', 0) or 0)
        src_item_id = int(it.get('source_item_id') or 0)
        src_it = c.execute("SELECT * FROM requisition_items WHERE id=?", (src_item_id,)).fetchone() if src_item_id else None
        if not src_it:
            c.close(); return jsonify({'error': f"物资「{it.get('item_name','')}」未匹配到源出库单明细"}), 400
        issued = float(src_it['quantity'] or 0)
        agg = c.execute("""SELECT COALESCE(SUM(t.return_qty),0) s FROM return_items t
                           JOIN return_requests rr ON rr.id=t.return_id
                           WHERE t.source_item_id=? AND rr.status='退库已完成'""", (src_item_id,)).fetchone()
        can = issued - float(agg['s'] or 0)
        if q <= 0 or q > can + 1e-9:
            c.close(); return jsonify({'error': f"「{src_it['item_name']}」本次退库数量({q:g})须大于0且不超过可退数量({can:g})"}), 400
    reason = (d.get('reason') or '').strip()
    if reason not in _RETURN_REASONS:
        reason = r['reason'] or '其他'
    total_amt = sum(float(it.get('price', 0) or 0) * float(it.get('return_qty', 0) or 0) for it in items)
    c.execute("UPDATE return_requests SET reason=?, reason_note=?, total_amount=?, updated_at=? WHERE id=?",
              (reason, (d.get('reason_note') or '').strip()[:200], total_amt, now(), rid))
    c.execute("DELETE FROM return_items WHERE return_id=?", (rid,))
    for it in items:
        amt = float(it.get('price', 0) or 0) * float(it.get('return_qty', 0) or 0)
        c.execute("""INSERT INTO return_items(return_id,source_item_id,item_name,spec,unit,issued_qty,returned_qty,return_qty,price,amount,created_at)
                     VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                  (rid, int(it.get('source_item_id') or 0), it['item_name'], it.get('spec', ''), it.get('unit', '个'),
                   float(it.get('issued_qty', 0) or 0), float(it.get('returned_qty', 0) or 0), float(it.get('return_qty', 0)),
                   float(it.get('price', 0) or 0), amt, now()))
    c.commit(); c.close()
    log(session['user_name'], '修改退库单', f'{r["return_no"]} 明细已更新')
    return jsonify({'success': True, 'message': f'退库单 {r["return_no"]} 已修改保存'})


@app.route('/api/returns/<int:rid>/void', methods=['POST'])
@login_required
def api_return_void(rid):
    """退库单作废(需求: 单据永久保存不允许物理删除, 仅作废) — 未确认入库的作废不影响库存;
    已完成入库的单据不允许作废(有错误只能重新做单据冲抵)。留痕: log + 审批实例终止"""
    c = db()
    r = c.execute("SELECT * FROM return_requests WHERE id=?", (rid,)).fetchone()
    if not r:
        c.close(); return jsonify({'error': '退库单不存在'}), 404
    if r['status'] == '退库已完成':
        c.close(); return jsonify({'error': '该退库单已完成入库，不可作废（如有错误请重新做单据冲抵）'}), 400
    if r['status'] == '已作废':
        c.close(); return jsonify({'error': '该退库单已作废'}), 400
    # 提交人本人/管理员 可作废
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_admin = me and (me['role'] == '系统管理员' or me['username'] in ('xingguo', 'admin', 'mujiao'))
    if not is_admin and r['requester'] and me and me['name'] != r['requester']:
        c.close(); return jsonify({'error': '仅退库单提交人或管理员可作废'}), 403
    c.execute("UPDATE return_requests SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
    # 终止待审批实例(若有)
    c.execute("UPDATE approval_instances SET status='rejected', comment='单据作废' WHERE biz_type='return_request' AND biz_id=? AND status IN ('pending','approved')", (rid,))
    _insts = c.execute("SELECT instance_code FROM dingtalk_instances WHERE biz_type='return_request' AND biz_id=? AND status IN ('pending','synced')", (rid,)).fetchall()
    for _ins in _insts:
        try:
            if _ins['instance_code'] and not str(_ins['instance_code']).startswith('ERR-'):
                dt_terminate_instance(str(_ins['instance_code']), dt_first_bound_userid() or '')
        except Exception:
            pass
    c.execute("DELETE FROM dingtalk_instances WHERE biz_type='return_request' AND biz_id=?", (rid,))
    # 作废日志留痕(审批流转日志统一记录)
    log_approval_action('return_request', rid, 'void', session.get('user_name', ''), session.get('user_id', 0),
                        '退库单作废（不影响库存）', now(), None, 'system', '', conn=c)
    c.commit(); c.close()
    log(session['user_name'], '作废退库单', f'{r["return_no"]} 已作废(未入库, 库存未变动)')
    return jsonify({'success': True, 'message': f'退库单 {r["return_no"]} 已作废（不影响库存，记录留痕）'})

# ============================================================
# V11.210 设备维修完整工单(设备维修功能需求.docx 对标金蝶/用友):
# 报修提报→定损三选一→金额分级审批→服务商比价→维修变更四方确认→委托单/合同→发料出库委外→
# 进度跟踪→回厂性能验收→归还/回收入库→发票归档
# 审批走 approval_flow_config repair_plan 配置(金额分级, 角色在系统设置配)
# 铁律: 不动物资采购申请任何表/字段/逻辑, 全部新增
# ============================================================
def _repair_no():
    """维修工单单号 WX+年月日+流水 — 与维修车间物资申请(WX前缀)共用当天序号池防撞号"""
    c = sqlite3.connect(DB); c.row_factory = sqlite3.Row
    d = datetime.date.today().strftime('%Y%m%d')
    cur = 0
    for r in c.execute("SELECT plan_no FROM repair_plans WHERE plan_no LIKE ?", ('WX' + d + '%',)).fetchall():
        try:
            cur = max(cur, int(str(r['plan_no'])[-2:]))
        except Exception:
            pass
    # 物资申请同前缀(WX=维修车间)当天序号也计入, 防两表撞号
    for r in c.execute("SELECT req_no FROM purchase_requests WHERE req_no LIKE ?", ('WX' + d + '%',)).fetchall():
        try:
            cur = max(cur, int(str(r['req_no'])[-2:]))
        except Exception:
            pass
    c.close()
    return f'WX{d}{cur+1:02d}'

@app.route('/api/repairs')
@login_required
def api_repairs():
    """维修计划列表(采购员/员工看自己, 领导/管理员全看)"""
    conn = db()
    # V11.213 老状态批量迁移: 定损通过(老通过态)+审批approved → 审批通过(新状态机), 一次修复历史卡死单
    try:
        stale = conn.execute("SELECT rp.id FROM repair_plans rp WHERE rp.status='定损通过' AND EXISTS (SELECT 1 FROM approval_instances ai WHERE ai.biz_type='repair_plan' AND ai.biz_id=rp.id AND ai.status='approved')").fetchall()
        for (sid,) in stale:
            conn.execute("UPDATE repair_plans SET status='审批通过', updated_at=? WHERE id=? AND status='定损通过'", (now(), sid))
        if stale: conn.commit()
    except Exception:
        pass
    # V11.215: 数据权限对齐物资采购申请口径(filter_scope) — 采购员/财务/库管看全部历史, 仅员工/部门负责人看自己的
    # (修复: 原硬编码仅领导看全部→采购员邢果看不到穆娇等提交的历史维修单)
    if filter_scope(session.get('user_role')) == 'own':
        rows = conn.execute("SELECT * FROM repair_plans WHERE requester=? OR requester_id=? ORDER BY id DESC LIMIT 100",
                            (session.get('user_name', ''), session.get('user_id', 0))).fetchall()
    else:
        rows = conn.execute("SELECT * FROM repair_plans ORDER BY id DESC LIMIT 100").fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict_row(r)
        d['item_count'] = sqlite3.connect(DB).execute("SELECT COUNT(*) FROM repair_items WHERE plan_id=?", (r['id'],)).fetchone()[0]
        out.append(d)
    return jsonify(out)

@app.route('/api/repairs', methods=['POST'])
@login_required
def api_create_repair():
    """第一步: 计划提报员提交维修计划(设备/故障/更换部件+故障照片附件)"""
    d = request.json or {}
    device = str(d.get('device_name') or '').strip()
    fault = str(d.get('fault_desc') or '').strip()
    if not device: return jsonify({'error': '请填写故障设备名称'}), 400
    if not fault: return jsonify({'error': '请填写故障现象描述'}), 400
    if not d.get('dept'): return jsonify({'error': '请选择申请部门'}), 400
    no = _repair_no()
    c = db()
    c.execute("""INSERT INTO repair_plans(plan_no,device_name,device_no,fault_desc,fault_time,urgency,init_judge,est_cost,dept,requester,requester_id,status,attachments,remark)
                 VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
              (no, device, str(d.get('device_no') or '').strip(), fault, str(d.get('fault_time') or ''),
               d.get('urgency') or '普通', d.get('init_judge') or '', float(d.get('est_cost') or 0),
               d.get('dept', ''), session['user_name'], session['user_id'],
               '草稿', json.dumps(d.get('attachments') or [], ensure_ascii=False), d.get('remark', '')))
    pid = c.execute("SELECT id FROM repair_plans WHERE plan_no=?", (no,)).fetchone()[0]
    for it in (d.get('items') or []):
        if str(it.get('part_name') or '').strip():
            c.execute("INSERT INTO repair_items(plan_id,part_name,fault_note) VALUES(?,?,?)", (pid, it['part_name'], it.get('fault_note', '')))
    c.commit(); c.close()
    log(session['user_name'], '新建维修报修单', f'{no} {device}')
    return jsonify({'success': True, 'id': pid, 'plan_no': no})

@app.route('/api/repairs/<int:rid>/submit', methods=['POST'])
@login_required
def api_repair_submit(rid):
    """节点1: 报修提交 — 单据流转至定损确认环节(钉钉推定损角色, 不建审批)"""
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('草稿', '定损驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可提交报修'}), 400
    c.execute("UPDATE repair_plans SET status='待定损', updated_at=? WHERE id=?", (now(), rid))
    c.commit()
    # V11.214: 通知定损角色(分管领导/总经理/管理员) — 铃铛+钉钉; 排除提交人自己
    try:
        _leaders = c.execute("SELECT id,dingtalk_userid,name FROM users WHERE role IN ('分管领导','总经理') AND is_active=1").fetchall()
        _admin = c.execute("SELECT id,dingtalk_userid,name FROM users WHERE role='系统管理员' AND is_active=1").fetchone()
        if _admin: _leaders = list(_leaders) + [_admin]
        _notify_ids = [u['id'] for u in _leaders if u['id'] != session.get('user_id')]
        add_notif(_notify_ids, f'🔧 新设备报修待定损：{r["plan_no"]}',
                  f'{r["device_name"]}（{r["dept"] or ""}提报，紧急度:{r["urgency"] or "普通"}）请进入 采购申请-设备维修 做技术定损', 'repair_plan', rid, conn=c)
        c.commit()  # V11.214: add_notif(conn=c) 不自动commit, 须在关闭前提交
        for u in _leaders:
            if u.get('dingtalk_userid') and u['id'] != session.get('user_id'):
                try: dt_send_todo([u['dingtalk_userid']], f'🔧 新设备报修待定损 {r["plan_no"]}', f'{r["device_name"]} 请做技术定损', '', 'repair_plan', rid)
                except Exception: pass
    except Exception:
        pass
    c.close()
    log(session['user_name'], '提交维修报修', f'{r["plan_no"]} → 待定损')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/damage', methods=['POST'])
@login_required
def api_repair_damage(rid):
    """节点2: 技术定损三选一 (机电/机修/生产负责人在系统设置配, 厂长账号接入后填)
    1)内部自修: 填自修记录 → 闭环归档
    2)委外维修: 填定损清单 → 按预估金额走分级审批(节点3)
    3)更换新设备: 一键生成物资采购申请(带基础信息) → 维修单归档
    支持驳回(退回提报人修改)"""
    d = request.json or {}
    act = d.get('action')  # internal / external / replace / reject
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('待定损', '定损完成待审批', '审批驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可定损'}), 400
    # 定损人权限: 系统设置配的定损角色(默认分管领导), 管理员可代
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理'):
        c.close(); return jsonify({'error': '仅定损角色(系统设置配置)可定损'}), 403
    op_name = session['user_name']
    opinion = str(d.get('opinion') or '').strip()
    _logmsg = ''
    if act == 'internal':
        note = str(d.get('internal_note') or '').strip()
        if not note: c.close(); return jsonify({'error': '请填写自修处理记录'}), 400
        c.execute("UPDATE repair_plans SET repair_type='内部自修', internal_note=?, damage_opinion=?, damage_time=?, status='已归档', actual_finish=? WHERE id=?",
                  (note, opinion or '内部自修处理', now(), now(), rid))
        _logmsg = f'{r["plan_no"]} 自修闭环归档'
    elif act == 'external':
        # 清旧明细重录定损清单
        c.execute("DELETE FROM repair_items WHERE plan_id=?", (rid,))
        for it in (d.get('items') or []):
            if str(it.get('part_name') or '').strip():
                c.execute("INSERT INTO repair_items(plan_id,part_name,fault_note,confirm_status,price,unit) VALUES(?,?,?,'确认维修',?,?)",
                          (rid, it['part_name'], it.get('fault_note', ''), float(it.get('price') or 0), it.get('unit') or '项'))
        est = float(d.get('est_cost') or r['est_cost'] or 0)
        c.execute("UPDATE repair_plans SET repair_type='委外维修', damage_items_json=?, damage_opinion=?, damage_time=?, est_cost=?, status='定损完成待审批' WHERE id=?",
                  (json.dumps(d.get('items') or [], ensure_ascii=False), opinion, now(), est, rid))
        c.commit()
        # 节点3: 按预估金额分级审批 — 必须先提交释放写锁, create_approvals用独立连接
        create_approvals('repair_plan', rid, est, submitter=r['requester'] or '')
        try: start_instances('repair_plan', rid)
        except Exception: pass
        c.close()
        log(op_name, '定损:委外维修', f'{r["plan_no"]} 预估¥{est:.0f} 待分级审批')
        return jsonify({'success': True})
    elif act == 'replace':
        # 一键转物资采购申请(携带基础信息), 走现有物资采购流程(不触碰其逻辑, 仅INSERT)
        import time as _t
        req_no = gen_req_no(d.get('dept') or r['dept'] or '', c)
        c.execute("""INSERT INTO purchase_requests(req_no,dept,requester,requester_id,purpose,status,total_estimated,remark,req_type,created_at,apply_date)
                     VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
                  (req_no, r['dept'] or '', r['requester'] or '', r['requester_id'] or 0,
                   f"设备报废更换: {r['fault_desc'] or ''}（原维修单{r['plan_no']}转来）", '草稿', float(d.get('est_cost') or r['est_cost'] or 0),
                   f'由维修单 {r["plan_no"]} 定损"建议直接更换新设备"生成', '物资采购', now(), datetime.date.today().strftime('%Y-%m-%d')))
        c.execute("UPDATE repair_plans SET repair_type='更换新设备', convert_req_no=?, damage_opinion=?, damage_time=?, status='已归档' WHERE id=?",
                  (req_no, opinion, now(), rid))
        _logmsg = f'{r["plan_no"]} → 物资申请{req_no}'
    elif act == 'reject':
        reason = str(d.get('reason') or '').strip()
        if not reason: c.close(); return jsonify({'error': '请填写驳回原因'}), 400
        c.execute("UPDATE repair_plans SET status='定损驳回', damage_opinion=?, updated_at=? WHERE id=?", (reason, now(), rid))
        _logmsg = f'{r["plan_no"]} 退回提报人修改({reason})'
    else:
        c.close(); return jsonify({'error': '未知定损操作'}), 400
    c.commit(); c.close()
    if act == 'internal':
        log(op_name, '定损:内部自修', _logmsg)
    elif act == 'replace':
        log(op_name, '定损:更换新设备', _logmsg)
    elif act == 'reject':
        log(op_name, '定损驳回', _logmsg)
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/quote', methods=['POST'])
@login_required
def api_repair_quote(rid):
    """节点4: 采购专员录入多家服务商报价(维修项目/配件费/工时费/单价/合计/工期/质保/附件)"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('审批通过', '定损完成待审批', '已通过', '待比价', '定损通过'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可录报价'}), 400
    company = str(d.get('company') or '').strip()
    if not company: c.close(); return jsonify({'error': '请填写维修服务商'}), 400
    # 覆盖式保存该服务商报价(防重复), 保留其他家
    c.execute("DELETE FROM repair_quotes WHERE plan_id=? AND company=?", (rid, company))
    total = 0.0
    for it in (d.get('quotes') or []):
        part = float(it.get('part_cost') or 0)  # 配件费
        labor = float(it.get('labor_cost') or 0)  # 工时费
        p = float(it.get('price') or (part + labor))
        total += p
        # V11.216: 报价明细带 配件费/工时费/质保/完工期 落库
        c.execute("INSERT INTO repair_quotes(plan_id,company,item_name,price,duration,status,part_cost,labor_cost,warranty,finish_date) VALUES(?,?,?,?,?,?,?,?,?,?)",
                  (rid, company, it.get('item_name', '维修'), p, it.get('duration', ''),
                   '报价', part, labor, d.get('warranty', ''), d.get('finish_date', '')))
    # V11.216: 报价阶段不改 quote_total(那是选中服务商后的落定金额, 多家报价会互相覆盖误导);
    # 各家报价存 repair_quotes.price, 选商时才把选中家总价写入 quote_total
    c.execute("UPDATE repair_plans SET repair_company=?, finish_date=?, status='待比价', updated_at=? WHERE id=?",
              (company, d.get('finish_date', ''), now(), rid))
    # 累计所有服务商总价供选商比较
    c.commit(); c.close()
    log(session['user_name'], '录入维修服务商报价', f'{r["plan_no"]} {company} ¥{total:.2f}')
    return jsonify({'success': True, 'total': total})

@app.route('/api/repairs/<int:rid>/select-vendor', methods=['POST'])
@login_required
def api_repair_select_vendor(rid):
    """节点4b: 比价选定服务商(权限: 采购员/领导) → 可生成委托单/合同"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] != '待比价':
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可选商'}), 400
    company = str(d.get('company') or '').strip()
    if not company: c.close(); return jsonify({'error': '请选择服务商'}), 400
    q = c.execute("SELECT * FROM repair_quotes WHERE plan_id=? AND company=? ORDER BY id LIMIT 1", (rid, company)).fetchone()
    if not q: c.close(); return jsonify({'error': '该服务商尚无报价, 先录报价'}), 400
    # V11.216: 选商落定金额 = 选中服务商报价总价(该家可能多条明细)
    _tot = c.execute("SELECT SUM(price) s FROM repair_quotes WHERE plan_id=? AND company=?", (rid, company)).fetchone()[0] or 0
    c.execute("UPDATE repair_plans SET vendor_selected=?, repair_company=?, quote_total=?, status='已选服务商', updated_at=? WHERE id=?",
              (company, company, float(_tot), now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '比价选定服务商', f'{r["plan_no"]} 选定:{company} ¥{float(_tot):.0f}')
    return jsonify({'success': True, 'total': float(_tot)})

@app.route('/api/repairs/<int:rid>/entrust', methods=['POST'])
@login_required
def api_repair_entrust(rid):
    """节点6: 生成维修委托单(维修委托单号 WY+日期+流水)"""
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('已选服务商', '委外维修中'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可生成委托单'}), 400
    if not r['entrust_no']:
        d = datetime.date.today().strftime('%Y%m%d')
        cnt = c.execute("SELECT COUNT(*) FROM repair_plans WHERE entrust_no LIKE ?", ('WY' + d + '%',)).fetchone()[0]
        entrust = f'WY{d}{cnt+1:03d}'
        c.execute("UPDATE repair_plans SET entrust_no=? WHERE id=?", (entrust, rid))
    else:
        entrust = r['entrust_no']
    c.commit(); c.close()
    log(session['user_name'], '生成维修委托单', f'{r["plan_no"]} → {entrust}')
    return jsonify({'success': True, 'entrust_no': entrust})

@app.route('/api/repairs/<int:rid>/send-out', methods=['POST'])
@login_required
def api_repair_send_out(rid):
    """节点7: 委外发料出库 — 标记委外维修状态(库存台账区分在厂/委外, 不扣总库存只标状态)
    简化实现: 记录发料时间+外发对象, outer_status=委外维修中; 维修完成回厂后解除"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('已选服务商',):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可发料, 需先选定服务商'}), 400
    if not r['vendor_selected']:
        c.close(); return jsonify({'error': '请先比价选定服务商'}), 400
    c.execute("UPDATE repair_plans SET outer_status='委外维修中', start_time=?, status='委外维修中', updated_at=? WHERE id=?",
              (now(), now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '委外发料出库', f'{r["plan_no"]} 设备发往 {r["vendor_selected"]} 委外维修中')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/progress', methods=['POST'])
@login_required
def api_repair_progress(rid):
    """节点8: 进度跟踪 — 更新预计完工/实际完工/备注(委外维修中阶段可反复更新)"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    upd = []
    args = []
    if d.get('expect_finish'):
        upd.append('expect_finish=?'); args.append(str(d['expect_finish']))
    if d.get('start_time'):
        upd.append('start_time=?'); args.append(str(d['start_time']))
    if d.get('remark'):
        upd.append('remark=?'); args.append(str(d['remark']))
    if not upd:
        c.close(); return jsonify({'error': '无更新内容'}), 400
    upd.append('updated_at=?'); args.append(now()); args.append(rid)
    c.execute("UPDATE repair_plans SET " + ','.join(upd) + " WHERE id=?", args)
    c.commit(); c.close()
    log(session['user_name'], '更新维修进度', f'{r["plan_no"]}')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/done', methods=['POST'])
@login_required
def api_repair_done(rid):
    """节点8b: 厂家维修完成 → 待回厂验收"""
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('委外维修中', '报价完成'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可报完工'}), 400
    c.execute("UPDATE repair_plans SET status='待回厂验收', actual_finish=?, updated_at=? WHERE id=?", (now(), now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '维修完工', f'{r["plan_no"]} 待回厂联合验收')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/change', methods=['POST'])
@login_required
def api_repair_change(rid):
    """节点5: 维修变更单(风控) — 厂家发现新增损坏必须走变更单, 四方确认后才计入总费用
    变更记录永久留痕, 不可删除修改历史"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('委外维修中', '报价完成'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可变更, 需维修中'}), 400
    item = str(d.get('add_item') or '').strip()
    if not item: c.close(); return jsonify({'error': '请填写新增维修项目'}), 400
    c.execute("""INSERT INTO repair_changes(plan_id,add_item,add_part,add_labor,add_price,change_reason,status,created_by)
                 VALUES(?,?,?,?,?,?,'待确认',?)""",
              (rid, item, str(d.get('add_part') or ''), float(d.get('add_labor') or 0), float(d.get('add_price') or 0),
               str(d.get('change_reason') or ''), session['user_name']))
    c.execute("UPDATE repair_plans SET status='变更待四方确认', change_count=change_count+1, updated_at=? WHERE id=?", (now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '维修变更单', f'{r["plan_no"]} 新增:{item} ¥{float(d.get("add_price") or 0):.2f} 待四方确认')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/change-confirm', methods=['POST'])
@login_required
def api_repair_change_confirm(rid):
    """节点5b: 变更四方确认 — 报修提报人/采购专员/机电厂长/机修车间主任 四方全通过才生效
    任一方驳回则该变更不增加费用(记录驳回); 变更记录永久留存"""
    d = request.json or {}
    c = db()
    cid = int(d.get('change_id') or 0)
    act = d.get('action')  # confirm / reject
    c.row_factory = sqlite3.Row
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    ch = c.execute("SELECT * FROM repair_changes WHERE id=?", (cid,)).fetchone() if cid else None
    if not ch: c.close(); return jsonify({'error': '变更单不存在'}), 400
    if ch['status'] == '已驳回':
        c.close(); return jsonify({'error': '该变更已被驳回, 永久留痕不可再确认'}), 400
    me_name = session['user_name']; me_role = session.get('user_role', '')
    # 身份判定: 按当前登录人与单据关系 + 角色归属到四方槽位(同一人可占多槽)
    slots = {'requester': ch['confirm1_by'], 'buyer': ch['confirm2_by'], 'chief': ch['confirm3_by'], 'director': ch['confirm4_by']}
    is_requester = r['requester'] == me_name or me_name == '温丽'
    # 采购员/采购专员身份, 或管理员代签
    is_buyer = me_role in ('采购员', '系统管理员')
    is_chief = me_role in ('分管领导', '总经理', '系统管理员')
    is_director = me_role in ('分管领导', '总经理', '系统管理员')
    ok_any = is_requester or is_buyer or is_chief or is_director
    if not ok_any:
        c.close(); return jsonify({'error': '仅报修提报人/采购专员/机电厂长/机修车间主任可确认'}), 403
    ts = now()
    if act == 'reject':
        # 任一方驳回 → 变更不生效(费用不计), 永久留痕
        c.execute("UPDATE repair_changes SET status='已驳回' WHERE id=?", (cid,))
        # 若无其他待确认变更 → 回到维修中
        pend = c.execute("SELECT COUNT(*) FROM repair_changes WHERE plan_id=? AND status='待确认'", (rid,)).fetchone()[0]
        if pend == 0:
            c.execute("UPDATE repair_plans SET status='委外维修中', updated_at=? WHERE id=?", (now(), rid))
        log(me_name, '变更驳回', f'{r["plan_no"]} 变更#{cid} 费用不生效')
        c.commit(); c.close()
        return jsonify({'success': True, 'rejected': True})
    # confirm: 填槽(去重)
    filled = 0
    if is_requester and not ch['confirm1_by']:
        c.execute("UPDATE repair_changes SET confirm1_by=?, confirm1_at=? WHERE id=?", (me_name, ts, cid)); filled += 1
    if is_buyer and not ch['confirm2_by']:
        c.execute("UPDATE repair_changes SET confirm2_by=?, confirm2_at=? WHERE id=?", (me_name, ts, cid)); filled += 1
    if is_chief and not ch['confirm3_by']:
        c.execute("UPDATE repair_changes SET confirm3_by=?, confirm3_at=? WHERE id=?", (me_name, ts, cid)); filled += 1
    if is_director and not ch['confirm4_by']:
        c.execute("UPDATE repair_changes SET confirm4_by=?, confirm4_at=? WHERE id=?", (me_name, ts, cid)); filled += 1
    ch2 = c.execute("SELECT * FROM repair_changes WHERE id=?", (cid,)).fetchone()
    all_done = ch2['confirm1_by'] and ch2['confirm2_by'] and ch2['confirm3_by'] and ch2['confirm4_by']
    addp = float(ch2['add_price'] or 0)
    if all_done:
        c.execute("UPDATE repair_changes SET status='已确认' WHERE id=?", (cid,))
        # 变更费用计入维修总价
        c.execute("UPDATE repair_plans SET quote_total=quote_total+?, status='委外维修中', updated_at=? WHERE id=?", (addp, now(), rid))
    else:
        c.execute("UPDATE repair_plans SET status='变更待四方确认', updated_at=? WHERE id=?", (now(), rid))
    c.commit(); c.close()
    if all_done:
        log(me_name, '变更四方确认完成', f'{r["plan_no"]} 变更#{cid} +¥{addp:.2f} 生效')
    return jsonify({'success': True, 'filled': filled, 'all_done': all_done,
                    'confirm1_by': ch2['confirm1_by'], 'confirm2_by': ch2['confirm2_by'],
                    'confirm3_by': ch2['confirm3_by'], 'confirm4_by': ch2['confirm4_by']})

@app.route('/api/repairs/<int:rid>/accept', methods=['POST'])
@login_required
def api_repair_accept(rid):
    """节点9: 回厂联合性能验收(服务类验收, 非普通入库) — 结果二选一: 通过/不通过返修
    必填性能验收结果(不只外观); 上传测试报告/照片; 不通过退回服务商返修"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('待回厂验收', '验收不通过返修'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可验收'}), 400
    if session.get('user_role') not in ('系统管理员', '分管领导', '总经理', '采购员', '库管员', '部门负责人'):
        c.close(); return jsonify({'error': '无验收权限'}), 403
    result = d.get('result')  # pass / fail
    opinion = str(d.get('opinion') or '').strip()
    if result not in ('pass', 'fail'):
        c.close(); return jsonify({'error': '请选择验收结果(通过/不通过返修)'}), 400
    if result == 'pass' and not opinion:
        c.close(); return jsonify({'error': '验收通过必须填写性能验收意见(不能只填外观合格)'}), 400
    files = json.dumps(d.get('files') or [], ensure_ascii=False)
    if result == 'pass':
        c.execute("UPDATE repair_plans SET accept_result='通过', accept_opinion=?, accept_files=?, accept_time=?, outer_status='已回厂', status='验收通过', updated_at=? WHERE id=?",
                  (opinion, files, now(), now(), rid))
    else:
        c.execute("UPDATE repair_plans SET accept_result='不通过返修', accept_opinion=?, accept_files=?, status='验收不通过返修', updated_at=? WHERE id=?",
                  (opinion or '退回服务商返修', files, now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '验收' + ('通过' if result == 'pass' else '不通过'), f'{r["plan_no"]} {"性能验收通过" if result=="pass" else "退回服务商返修"}')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/handle', methods=['POST'])
@login_required
def api_repair_handle(rid):
    """节点10: 验收通过后实物处理 — ①归还部门 ②回收入库(复用入库流程/登记)
    同时解委外状态, 流转发票登记"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] != '验收通过':
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可处理'}), 400
    ht = d.get('handle_type')  # return_dept / warehouse
    if ht not in ('return_dept', 'warehouse'):
        c.close(); return jsonify({'error': '请选择处理方式(归还部门/回收入库)'}), 400
    c.execute("UPDATE repair_plans SET handle_type=?, outer_status='已回厂', status='待发票登记', updated_at=? WHERE id=?",
              ('归还部门' if ht == 'return_dept' else '回收入库', now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '验收后处理', f'{r["plan_no"]} {"归还使用部门" if ht=="return_dept" else "回收入库"} → 待发票登记')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/invoice', methods=['POST'])
@login_required
def api_repair_invoice(rid):
    """节点10b: 发票登记 — 登记维修发票(号码/金额) 归集到对应设备/部门, 进入付款流程"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] != '待发票登记':
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可登记发票'}), 400
    inv = str(d.get('invoice_no') or '').strip()
    amt = float(d.get('invoice_amount') or r['quote_total'] or 0)
    if not inv: c.close(); return jsonify({'error': '请填写发票号码'}), 400
    c.execute("UPDATE repair_plans SET invoice_no=?, invoice_amount=?, status='已归档', updated_at=? WHERE id=?", (inv, amt, now(), rid))
    c.commit(); c.close()
    log(session['user_name'], '维修发票登记归档', f'{r["plan_no"]} 发票{inv} ¥{amt:.2f} → 已归档')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/detail')
@login_required
def api_repair_detail(rid):
    c = db(); c.row_factory = sqlite3.Row
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    # V11.213 老状态迁移: V11.208时代的"定损通过"(当时审批通过态), 若其审批已approved且无报价 → 升为"审批通过"
    # (V11.210新状态机通过态='审批通过', 老单不迁移则前端无按钮可点=卡死); 幂等只升不降, 不影响业务
    if r['status'] == '定损通过':
        ok_cnt = c.execute("SELECT COUNT(*) FROM approval_instances WHERE biz_type='repair_plan' AND biz_id=? AND status='approved'", (rid,)).fetchone()[0]
        if ok_cnt > 0:
            c.execute("UPDATE repair_plans SET status='审批通过', updated_at=? WHERE id=? AND status='定损通过'", (now(), rid))
            c.commit()
            r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    items = c.execute("SELECT * FROM repair_items WHERE plan_id=? ORDER BY id", (rid,)).fetchall()
    quotes = c.execute("SELECT * FROM repair_quotes WHERE plan_id=? ORDER BY id", (rid,)).fetchall()
    changes = c.execute("SELECT * FROM repair_changes WHERE plan_id=? ORDER BY id", (rid,)).fetchall()
    c.close()
    return jsonify({'plan': dict_row(r), 'items': [dict_row(i) for i in items],
                    'quotes': [dict_row(q) for q in quotes], 'changes': [dict_row(ch) for ch in changes]})

@app.route('/api/repairs/<int:rid>/void', methods=['POST'])
@login_required
def api_repair_void(rid):
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修计划不存在'}), 404
    if r['status'] in ('已完成', '待返库'):
        c.close(); return jsonify({'error': '该计划已进入返库/完成, 不可作废'}), 400
    # V11.218: 删除按钮全状态显示(与物资采购同构), 但仅 草稿/定损驳回/终态前 可作废;
    # 审批中/流程中(待定损/审批中/委外/验收等)必须先撤回或走完流程 — 防误删真实在办单&产生幽灵审批
    if r['status'] not in ('草稿', '定损驳回', '已归档', '验收通过', '待发票登记', '已选服务商'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可删除。请在「详情」操作或先撤回后再删除'}), 400
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_admin = me and (me['role'] == '系统管理员' or session.get('user_role') == '分管领导')
    if not is_admin and r['requester'] and me and me['name'] != r['requester']:
        c.close(); return jsonify({'error': '仅提交人或领导可作废'}), 403
    c.execute("UPDATE repair_plans SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
    c.execute("UPDATE approval_instances SET status='rejected', comment='计划作废' WHERE biz_type='repair_plan' AND biz_id=? AND status IN ('pending','approved')", (rid,))
    c.commit(); c.close()
    log(session['user_name'], '作废维修计划', f'{r["plan_no"]}')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>', methods=['PUT'])
@login_required
def api_repair_update(rid):
    """V11.214: 编辑维修报修单(限 草稿/定损驳回) — 对齐采购申请'修改'能力; 保存后保持原状态可重新提交"""
    d = request.json or {}
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('草稿', '定损驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可修改(仅草稿/被驳回可编辑)'}), 400
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_leader = session.get('user_role') in ('系统管理员', '分管领导', '总经理')
    if not is_leader and r['requester_id'] != session.get('user_id'):
        c.close(); return jsonify({'error': '仅提交人本人可修改'}), 403
    device = str(d.get('device_name') or '').strip()
    if not device: c.close(); return jsonify({'error': '请填写故障设备名称'}), 400
    # 更新主表(保留 plan_no/requester/status)
    c.execute("""UPDATE repair_plans SET device_name=?, device_no=?, fault_desc=?, fault_time=?, urgency=?,
                 init_judge=?, est_cost=?, dept=?, attachments=?, remark=?, updated_at=? WHERE id=?""",
              (device, str(d.get('device_no') or '').strip(), str(d.get('fault_desc') or '').strip(),
               str(d.get('fault_time') or ''), d.get('urgency') or '普通', d.get('init_judge') or '',
               float(d.get('est_cost') or 0), d.get('dept') or r['dept'], json.dumps(d.get('attachments') or [], ensure_ascii=False),
               str(d.get('remark') or ''), now(), rid))
    # 重建部件明细
    c.execute("DELETE FROM repair_items WHERE plan_id=?", (rid,))
    for it in (d.get('items') or []):
        if str(it.get('part_name') or '').strip():
            c.execute("INSERT INTO repair_items(plan_id,part_name,fault_note) VALUES(?,?,?)", (rid, it['part_name'], it.get('fault_note', '')))
    c.commit(); c.close()
    log(session['user_name'], '修改维修报修单', f'{r["plan_no"]} {device}')
    return jsonify({'success': True})

@app.route('/api/repairs/<int:rid>/download')
@login_required
def api_repair_download(rid):
    """V11.216: 设备维修单导出 xlsx(报修信息/定损清单/报价/变更/验收/处理全记录) — 对齐采购申请下载能力"""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    conn = db()
    rp = conn.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not rp:
        conn.close(); return jsonify({'error': '维修单不存在'}), 404
    rp = dict(rp)  # V11.216: Row无.get, 转dict
    items = [dict(x) for x in conn.execute("SELECT * FROM repair_items WHERE plan_id=? ORDER BY id", (rid,)).fetchall()]
    quotes = [dict(x) for x in conn.execute("SELECT * FROM repair_quotes WHERE plan_id=? ORDER BY id", (rid,)).fetchall()]
    changes = [dict(x) for x in conn.execute("SELECT * FROM repair_changes WHERE plan_id=? ORDER BY id", (rid,)).fetchall()]
    conn.close()
    wb = Workbook(); ws = wb.active; ws.title = '设备维修单'
    thin = Side(style='thin', color='999999')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    title_f = Font(bold=True, size=14)
    head_f = Font(bold=True, size=10)
    wrap = Alignment(vertical='center', wrap_text=True)
    ws.merge_cells('A1:F1')
    ws['A1'] = '设备维修单'
    ws['A1'].font = title_f; ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 28
    ws.merge_cells('A2:F2')
    ws['A2'] = f'编号: {rp["plan_no"]}    申请部门: {rp["dept"] or ""}    提报人: {rp["requester"] or ""}    提报时间: {str(rp["created_at"] or "")[:16]}    状态: {rp["status"]}'
    ws['A2'].font = Font(size=9)
    rows = []
    rows.append(['故障设备名称', rp['device_name'] or '', '设备编号', rp['device_no'] or ''])
    rows.append(['故障发生时间', str(rp['fault_time'] or '')[:16], '紧急等级', rp['urgency'] or ''])
    rows.append(['初步故障判断', rp['init_judge'] or '', '预估维修费用', f"¥{float(rp['est_cost'] or 0):.0f}"])
    rows.append(['故障现象描述', rp['fault_desc'] or '', '', ''])
    rows.append(['定损意见', rp['damage_opinion'] or '', '定损类型', rp['repair_type'] or ''])
    rows.append(['验收结果', f"{rp['accept_result'] or ''} {rp['accept_opinion'] or ''}", '验收后处理', rp['handle_type'] or ''])
    rows.append(['发票号码', rp['invoice_no'] or '', '发票金额', f"¥{float(rp['invoice_amount'] or 0):.2f}" if 'invoice_amount' in rp.keys() and rp['invoice_amount'] else ''])
    if rp.get('entrust_no'): rows.append(['维修委托单', rp['entrust_no'], '维修服务商', rp['vendor_selected'] or rp['repair_company'] or ''])
    row_i = 4
    for row in rows:
        for c_i, v in enumerate(row):
            cell = ws.cell(row=row_i, column=c_i + 1, value=v)
            cell.border = border; cell.alignment = wrap
            if c_i % 2 == 0: cell.font = head_f
        ws.row_dimensions[row_i].height = 20
        row_i += 1
    # 定损清单表
    row_i += 1
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
    ws.cell(row=row_i, column=1, value='定损清单').font = head_f; row_i += 1
    for c_i, h in enumerate(['部件/部位', '损坏说明', '单位', '单价', '', '']):
        cell = ws.cell(row=row_i, column=c_i + 1, value=h); cell.font = head_f; cell.border = border
    row_i += 1
    if items:
        for it in items:
            vals = [it['part_name'], it.get('fault_note') or '', it.get('unit') or '', it.get('price') or '', '', '']
            for c_i, v in enumerate(vals):
                ws.cell(row=row_i, column=c_i + 1, value=v).border = border
            row_i += 1
    else:
        ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
        ws.cell(row=row_i, column=1, value='（无）'); row_i += 1
    # 服务商报价表
    row_i += 1
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
    ws.cell(row=row_i, column=1, value='服务商报价').font = head_f; row_i += 1
    for c_i, h in enumerate(['服务商', '维修项目', '配件费', '工时费', '工期', '质保']):
        cell = ws.cell(row=row_i, column=c_i + 1, value=h); cell.font = head_f; cell.border = border
    row_i += 1
    if quotes:
        for q in quotes:
            vals = [q['company'], q.get('item_name') or '维修', q.get('part_cost') or '', q.get('labor_cost') or '', q.get('duration') or '', q.get('warranty') or '']
            for c_i, v in enumerate(vals):
                ws.cell(row=row_i, column=c_i + 1, value=v).border = border
            row_i += 1
    else:
        ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
        ws.cell(row=row_i, column=1, value='（暂无报价）'); row_i += 1
    # 变更记录
    if changes:
        row_i += 1
        ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
        ws.cell(row=row_i, column=1, value='维修变更记录').font = head_f; row_i += 1
        for c_i, h in enumerate(['变更项目', '配件费', '工时费', '合计', '原因', '状态']):
            cell = ws.cell(row=row_i, column=c_i + 1, value=h); cell.font = head_f; cell.border = border
        row_i += 1
        for ch in changes:
            vals = [ch['add_item'], ch.get('add_part') or '', ch.get('add_labor') or '', ch.get('add_price') or '', ch.get('change_reason') or '', ch['status']]
            for c_i, v in enumerate(vals):
                ws.cell(row=row_i, column=c_i + 1, value=v).border = border
            row_i += 1
    row_i += 1
    ws.merge_cells(start_row=row_i, start_column=1, end_row=row_i, end_column=6)
    ws.cell(row=row_i, column=1, value=f'维修费合计: ¥{float(rp["quote_total"] or 0):.0f}    填报人: {session.get("user_name", "")}    打印时间: {now()[:16]}')
    for col, w in zip('ABCDEF', [18, 26, 14, 14, 14, 14]):
        ws.column_dimensions[col].width = w
    bio = io.BytesIO(); wb.save(bio); bio.seek(0)
    from flask import send_file
    fname = f"设备维修单_{rp['plan_no']}.xlsx"
    from urllib.parse import quote
    return send_file(bio, as_attachment=True, download_name=fname, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@app.route('/api/repairs/<int:rid>/withdraw', methods=['POST'])
@login_required
def api_repair_withdraw(rid):
    """V11.214: 撤回已提交报修(待定损/定损完成待审批/审批驳回时) → 回草稿可修改; 对齐采购申请'撤回'能力"""
    c = db()
    r = c.execute("SELECT * FROM repair_plans WHERE id=?", (rid,)).fetchone()
    if not r: c.close(); return jsonify({'error': '维修单不存在'}), 404
    if r['status'] not in ('待定损', '定损完成待审批', '审批驳回'):
        c.close(); return jsonify({'error': f'当前状态({r["status"]})不可撤回'}), 400
    me = c.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
    is_leader = session.get('user_role') in ('系统管理员', '分管领导', '总经理')
    if not is_leader and r['requester_id'] != session.get('user_id'):
        c.close(); return jsonify({'error': '仅提交人本人可撤回'}), 403
    c.execute("UPDATE repair_plans SET status='草稿', updated_at=? WHERE id=?", (now(), rid))
    c.execute("UPDATE approval_instances SET status='rejected', comment='提交人撤回' WHERE biz_type='repair_plan' AND biz_id=? AND status IN ('pending','approved')", (rid,))
    c.commit(); c.close()
    log(session['user_name'], '撤回维修报修', f'{r["plan_no"]} → 草稿')
    return jsonify({'success': True})

@app.route('/api/requisitions/<int:rid>/void', methods=['POST'])
@login_required
def api_requisition_void(rid):
    """V5.0: 出库单作废 — 已出库(已扣库存)作废则回滚库存+写流水; 未出库直接作废
    V11.159: 限 库管员/部门负责人/领导/管理员"""
    if session.get('user_role') not in ('库管员', '部门负责人', '分管领导', '总经理', '系统管理员'):
        return jsonify({'error': '无权限：出库管理仅限库管员/领导使用'}), 403
    c = db()
    rq = c.execute("SELECT * FROM requisitions WHERE id=?", (rid,)).fetchone()
    if not rq:
        c.close(); return jsonify({'error': '出库单不存在'}), 404
    if rq['status'] == '已作废':
        c.close(); return jsonify({'error': '该出库单已作废'}), 400
    if rq['status'] == '已出库':
        # 回滚库存: 删除该单出库流水 + 加回库存
        flows = c.execute("SELECT * FROM inventory_flows WHERE doc_type='requisition' AND doc_id=? AND flow_type='出库'", (rid,)).fetchall()
        for f in flows:
            inv = c.execute("SELECT * FROM inventory WHERE item_name=? AND spec=? ORDER BY quantity DESC",
                            (f['item_name'], f['spec'] or '')).fetchone()
            if inv:
                new_q = inv['quantity'] - f['qty']  # f['qty'] 为负值, 减负=加回
                c.execute("UPDATE inventory SET quantity=?, updated_at=? WHERE id=?", (new_q, now(), inv['id']))
            else:
                c.execute("INSERT INTO inventory(item_name,spec,unit,quantity,warehouse,created_at) VALUES(?,?,?,?,?,?)",
                          (f['item_name'], f['spec'] or '', f['unit'] or '个', -f['qty'], '主库房', now()))
        c.execute("DELETE FROM inventory_flows WHERE doc_type='requisition' AND doc_id=?", (rid,))
        c.execute("UPDATE requisitions SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
        c.execute("UPDATE approval_instances SET status='rejected', comment='单据作废' WHERE biz_type='requisition' AND biz_id=? AND status='pending'", (rid,))
        c.commit(); c.close()
        log(session['user_name'], '作废出库单', f'{rq["req_no"]} (已出库, 库存已回滚)')
        return jsonify({'success': True, 'message': '出库单已作废，库存已回滚'})
    c.execute("UPDATE requisitions SET status='已作废', updated_at=? WHERE id=?", (now(), rid))
    c.execute("UPDATE approval_instances SET status='rejected', comment='单据作废' WHERE biz_type='requisition' AND biz_id=? AND status='pending'", (rid,))
    c.commit(); c.close()
    log(session['user_name'], '作废出库单', f'{rq["req_no"]} (未出库, 库存未变动)')
    return jsonify({'success': True, 'message': '出库单已作废（未出库，库存未变动）'})

@app.route('/api/docs/<biz_type>/<int:bid>/withdraw', methods=['POST'])
@login_required
def api_doc_withdraw(biz_type, bid):
    """V8.0: 撤回审批 — 每个环节即使已审批也可撤回
    规则: 审批中(待审批/部分通过)或已通过 → 撤回 → 回到'待审批'重新走完整流程
    - 撤回后原审批实例作废, 重新生成审批链(按当前配置)
    - 钉钉实例终止(若有), 重新发起
    - 已执行库存操作(入库已入库/出库已出库)的: 需先作废回滚, 不可直接撤回
    - 已生成下游单据(申请→订单/订单→入库/合同)的: 禁止撤回"""
    conn = db()
    if not can_manage_config():
        # V11.192: 单据提交人本人也可撤回自己的单(管理员不受限); 其他角色 403
        _me = conn.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
        if not _me:
            conn.close(); return jsonify({'error': '未登录'}), 401
        _who = find_doc_submitter(biz_type, bid)
        if not _who or _who.get('name') != _me['name']:
            conn.close(); return jsonify({'error': '仅单据提交人或系统管理员可撤回审批'}), 403
    if biz_type not in _DELETE_TABLE:
        return jsonify({'error': f'不支持的撤回类型: {biz_type}'}), 400
    table = _DELETE_TABLE[biz_type]
    no_col = _DELETE_NO_COL[biz_type]
    biz = _DELETE_BIZTYPE[biz_type]
    row = conn.execute(f"SELECT * FROM {table} WHERE id=?", (bid,)).fetchone()
    if not row:
        conn.close(); return jsonify({'error': '单据不存在'}), 404
    status = row['status'] if 'status' in row.keys() else ''
    no = row[no_col] if no_col in row.keys() else str(bid)
    # 状态校验: 仅 待审批/审批中/已通过 可撤回
    if status in ('已驳回', '已作废', '草稿'):
        conn.close(); return jsonify({'error': f'当前状态({status})无需撤回'}), 400
    # 下游保护
    if biz_type == 'purchase_request':
        if conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该申请已被订单引用，无法撤回（可先删除订单）'}), 400
    if biz_type == 'purchase_order':
        if conn.execute("SELECT COUNT(*) FROM receivings WHERE order_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该订单已有入库记录，无法撤回（可先作废入库）'}), 400
        if conn.execute("SELECT COUNT(*) FROM contracts WHERE order_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该订单已生成合同，无法撤回（可先作废合同）'}), 400
    # 库存操作保护
    if biz_type == 'receiving' and status == '已入库':
        conn.close(); return jsonify({'error': '该入库单已完成入库(已加库存)，请先作废回滚后再撤回'}), 400
    if biz_type == 'requisition' and status == '已出库':
        conn.close(); return jsonify({'error': '该出库单已完成出库(已扣库存)，请先作废回滚后再撤回'}), 400
    # V11.192: 撤回 = 单据回到"草稿"(提交人可修改), 审批实例留痕作废, 钉钉实例终止; 不自动重新审批
    # 提交人改完点「再次提交审批」(resubmit) 才重新进入审批流+推钉钉 — 修复: 撤回后一直审批中/无按钮/推不到钉钉
    conn.execute(f"UPDATE {table} SET status='草稿', updated_at=? WHERE id=?", (now(), bid))
    # 审批实例留痕: 待审/已过节点置 withdrawn(撤回), 保留历史可见(审批流转日志仍显示原流程)
    conn.execute("UPDATE approval_instances SET status='rejected', comment='发起人撤回' WHERE biz_type=? AND biz_id=? AND status IN ('pending','approved')", (biz, bid))
    # 钉钉实例终止(若钉钉侧还有 RUNNING) — 避免撤回后钉钉审批人仍能批
    _insts = conn.execute("SELECT instance_code FROM dingtalk_instances WHERE biz_type=? AND biz_id=? AND status IN ('pending','synced')", (biz, bid)).fetchall()
    for _ins in _insts:
        try:
            if _ins['instance_code'] and not str(_ins['instance_code']).startswith('ERR-'):
                dt_terminate_instance(str(_ins['instance_code']), dt_first_bound_userid() or '')
        except Exception:
            pass
    conn.execute("DELETE FROM dingtalk_instances WHERE biz_type=? AND biz_id=?", (biz, bid))
    # 撤回操作留痕(审批流转日志统一记录, 申请人/审批人可见)
    log_approval_action(biz_type, bid, 'withdraw', session.get('user_name',''), session.get('user_id',0),
                        '发起人撤回，单据退回草稿，修改后可再次提交审批', now(), None, 'system', '', conn=conn)
    # 撤回次数累计(留痕用, 语义=被打回修改过几次)
    try:
        conn.execute(f"UPDATE {table} SET reject_count=COALESCE(reject_count,0)+1 WHERE id=?", (bid,))
    except Exception:
        pass
    conn.commit()
    conn.close()
    log(session['user_name'], '撤回审批', f'{biz_type}#{bid} {no} 撤回退回草稿(未自动重审)')
    return jsonify({'success': True, 'message': f'单据 {no} 已撤回退回草稿。修改确认后请点「🔄 再次提交审批」重新进入审批流'})


@app.route('/api/docs/<biz_type>/<int:bid>/update', methods=['POST'])
@login_required
def api_doc_update(biz_type, bid):
    """V8.0: 单据通用编辑 — 全字段可修改, 保存实时生效
    支持: purchase_request/purchase_order/contract/receiving/requisition/inventory
    - 主表全字段更新(白名单过滤, 防注入)
    - 明细子表重建(request_items/order_items/requisition_items/items_json)
    - 金额字段自动重算(单价x数量x税率)
    - 审批中/已通过单据: 编辑后回到待审批重新走流程(撤回原审批)"""
    if not can_manage_config():
        return jsonify({'error': '仅系统管理员可编辑单据'}), 403
    # V11.201 权限收紧: 库存修改属敏感业务写操作 — 仅 系统管理员/分管领导/总经理 可改, 采购员/财务/员工只读
    # (即使采购员被加进 config_users 也在此拦截, 库存只读是硬边界)
    if biz_type == 'inventory' and session.get('user_role') not in ('系统管理员', '分管领导', '总经理'):
        return jsonify({'error': '无权限：库存数据仅管理员/分管领导可修改，采购员只读'}), 403
    d = request.json or {}
    if biz_type not in _DELETE_TABLE:
        return jsonify({'error': f'不支持的编辑类型: {biz_type}'}), 400
    table = _DELETE_TABLE[biz_type]
    no_col = _DELETE_NO_COL[biz_type]
    biz = _DELETE_BIZTYPE[biz_type]
    conn = db()
    row = conn.execute(f"SELECT * FROM {table} WHERE id=?", (bid,)).fetchone()
    if not row:
        conn.close(); return jsonify({'error': '单据不存在'}), 404
    old_status = row['status'] if 'status' in row.keys() else ''
    no = row[no_col] if no_col in row.keys() else str(bid)

    # ---- 主表字段白名单 + 更新 ----
    EDITABLE = {
        'purchase_request': ['dept', 'requester', 'budget_code', 'purpose', 'target_date', 'remark', 'urgent', 'apply_date'],
        'purchase_order': ['supplier', 'requester', 'category', 'target_date', 'remark', 'trade_mode', 'urgent'],
        'contract': ['contract_name', 'supplier', 'amount', 'sign_date', 'start_date', 'end_date', 'content', 'remark', 'urgent'],
        'receiving': ['warehouse', 'inspector', 'remark', 'urgent'],
        'requisition': ['dept', 'requester', 'purpose', 'issued_at'],
        'inventory': ['item_name', 'spec', 'cat_code', 'unit', 'quantity', 'safe_stock', 'warehouse', 'price', 'tax_rate', 'remark', 'max_stock', 'expiry_date', 'supplier'],
    }
    fields = EDITABLE.get(biz_type, [])
    updates = []
    vals = []
    for f in fields:
        if f in d and d[f] is not None:
            updates.append(f"{f}=?")
            vals.append(d[f])
    if updates:
        updates.append("updated_at=?")
        vals.append(now())
        vals.append(bid)
        conn.execute(f"UPDATE {table} SET {', '.join(updates)} WHERE id=?", vals)
        conn.commit()

    # ---- 明细子表重建 ----
    items = d.get('items')
    if items is not None and isinstance(items, list):
        items = [it for it in items if it.get('item_name') and float(it.get('quantity', 0) or 0) > 0]
        if biz_type == 'purchase_request':
            total = sum(float(it.get('quantity', 1)) * float(it.get('estimated_price', 0)) for it in items)
            conn.execute("UPDATE purchase_requests SET total_estimated=? WHERE id=?", (total, bid))
            conn.execute("DELETE FROM request_items WHERE req_id=?", (bid,))
            for it in items:
                tp = float(it.get('quantity', 1)) * float(it.get('estimated_price', 0))
                conn.execute("INSERT INTO request_items(req_id,item_name,spec,unit,quantity,estimated_price,total_price,remark,category,brand_param,arrival_date) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                             (bid, it.get('item_name', ''), it.get('spec', ''), it.get('unit', '个'), float(it.get('quantity', 1)),
                              float(it.get('estimated_price', 0)), tp, it.get('remark', ''),
                              it.get('category', ''), it.get('brand_param', ''), it.get('arrival_date', '')))
        elif biz_type == 'purchase_order':
            grand_amt = 0.0; grand_tax = 0.0; grand_total = 0.0
            conn.execute("DELETE FROM order_items WHERE order_id=?", (bid,))
            for it in items:
                qty = float(it.get('quantity', 1)); price = float(it.get('price', 0))
                tr = float(it.get('tax_rate', 13))
                amt = qty * price; tax = amt * tr / 100; tot = amt + tax
                grand_amt += amt; grand_tax += tax; grand_total += tot
                conn.execute("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount,remark) VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                             (bid, it.get('item_name', ''), it.get('spec', ''), it.get('unit', '个'),
                              qty, price, amt, tr, tax, tot, it.get('remark', '')))
            first = items[0]
            conn.execute("UPDATE purchase_orders SET item_name=?, spec=?, quantity=?, unit=?, price=?, amount=?, tax_amount=?, total_amount=?, tax_rate=? WHERE id=?",
                         (first.get('item_name', ''), first.get('spec', ''), sum(float(it.get('quantity', 1)) for it in items),
                          first.get('unit', '个'), first.get('price', 0), grand_amt, grand_tax, grand_total, first.get('tax_rate', 13), bid))
        elif biz_type == 'requisition':
            total_q = sum(float(it.get('quantity', 0)) for it in items)
            conn.execute("DELETE FROM requisition_items WHERE requisition_id=?", (bid,))
            for it in items:
                conn.execute("INSERT INTO requisition_items(requisition_id,item_name,spec,unit,quantity,purpose,created_at) VALUES(?,?,?,?,?,?,?)",
                             (bid, it.get('item_name', ''), it.get('spec', ''), it.get('unit', '个'),
                              float(it.get('quantity', 0)), it.get('purpose', d.get('purpose', '')), now()))
            first = items[0]
            conn.execute("UPDATE requisitions SET item_name=?, spec=?, quantity=?, unit=? WHERE id=?", (first.get('item_name', ''), first.get('spec', ''), total_q, first.get('unit', '个'), bid))
        elif biz_type == 'receiving':
            total_q = sum(float(it.get('quantity', 0)) for it in items)
            conn.execute("UPDATE receivings SET item_name=?, spec=?, quantity=?, unit=?, items_json=? WHERE id=?",
                         (items[0].get('item_name', ''), items[0].get('spec', ''), total_q, items[0].get('unit', '个'),
                          json.dumps(items, ensure_ascii=False), bid))

    # ---- 审批状态处理: 已完成的单据编辑后重新走审批 ----
    if biz_type in ('purchase_request', 'purchase_order', 'contract', 'receiving', 'requisition'):
        if old_status in ('已通过', '审批通过', '已入库', '已出库'):
            conn.execute(f"UPDATE {table} SET status='待审批', updated_at=? WHERE id=?", (now(), bid))
            conn.execute("UPDATE approval_instances SET status='rejected', comment='编辑后重新审批' WHERE biz_type=? AND biz_id=? AND status IN ('pending','approved')", (biz, bid))
            conn.execute("DELETE FROM dingtalk_instances WHERE biz_type=? AND biz_id=?", (biz, bid))
            conn.commit()
            amount = 0
            try:
                if 'total_amount' in row.keys(): amount = float(d.get('total_amount', row['total_amount']) or 0)
                elif 'total_estimated' in row.keys(): amount = float(d.get('total_estimated', row['total_estimated']) or 0)
                elif 'amount' in row.keys(): amount = float(d.get('amount', row['amount']) or 0)
            except Exception:
                amount = 0
            create_approvals(biz, bid, amount, submitter=session.get('user_name',''))
            start_instances(biz, bid)
        else:
            conn.commit()

    conn.close()
    log(session['user_name'], '编辑单据', f'{biz_type}#{bid} {no}')
    return jsonify({'success': True, 'message': f'单据 {no} 已更新（数据实时生效）'})



_DELETE_TABLE = {
    'purchase_request': 'purchase_requests', 'purchase_order': 'purchase_orders',
    'contract': 'contracts', 'receiving': 'receivings',
    'requisition': 'requisitions', 'payment': 'payment_requests',
    'inventory': 'inventory', 'return_request': 'return_requests',
}
_DELETE_NO_COL = {
    'purchase_request': 'req_no', 'purchase_order': 'order_no', 'contract': 'contract_no',
    'receiving': 'receive_no', 'requisition': 'req_no', 'payment': 'pay_no',
    'inventory': 'item_name', 'return_request': 'return_no',
}
_DELETE_BIZTYPE = {
    'purchase_request': 'purchase_request', 'purchase_order': 'purchase_order', 'contract': 'contract',
    'receiving': 'receiving', 'requisition': 'requisition', 'payment': 'payment',
    'inventory': 'inventory', 'return_request': 'return_request',
}

@app.route('/api/docs/<biz_type>/<int:bid>/delete', methods=['POST'])
@login_required
def api_doc_delete(biz_type, bid):
    """V8.0: 单据删除 — 支持 申请/订单/合同/入库/出库/付款
    安全规则:
    - 仅系统管理员可删(需传 confirm=1)
    - 已审批通过且已影响库存(入库已入库/出库已出库)的单据: 需先作废回滚库存, 不可直接删
    - 有下游单据引用(订单引用申请/入库引用订单/合同引用订单)的: 禁止删除
    - 级联清理: 明细行/审批实例/钉钉实例/库存流水"""
    conn = db()
    if not can_manage_config():
        # V11.201: 单据提交人本人可删除自己的 草稿/已驳回 单据(未进审批流/无下游), 其余仅管理员
        _me = conn.execute("SELECT * FROM users WHERE id=?", (session.get('user_id', 0),)).fetchone()
        if _me:
            _who = find_doc_submitter(biz_type, bid)
            _tbl = _DELETE_TABLE.get(biz_type, '')
            _st = ''
            if _tbl:
                _rw = conn.execute(f"SELECT status FROM {_tbl} WHERE id=?", (bid,)).fetchone()
                if _rw: _st = _rw['status'] or ''
            if not (_who and _who.get('name') == _me['name'] and _st in ('草稿', '已驳回')):
                return jsonify({'error': '仅单据提交人本人（草稿/已驳回）或管理员可删除'}), 403
        else:
            return jsonify({'error': '仅系统管理员可删除单据'}), 403
    d = request.json or {}
    if not d.get('confirm'):
        conn.close(); return jsonify({'error': '请确认删除(confirm=1)'}), 400
    if biz_type not in _DELETE_TABLE:
        conn.close(); return jsonify({'error': f'不支持的删除类型: {biz_type}'}), 400
    table = _DELETE_TABLE[biz_type]
    no_col = _DELETE_NO_COL[biz_type]
    biz = _DELETE_BIZTYPE[biz_type]
    row = conn.execute(f"SELECT * FROM {table} WHERE id=?", (bid,)).fetchone()
    if not row:
        conn.close(); return jsonify({'error': '单据不存在'}), 404
    status = row['status'] if 'status' in row.keys() else ''
    no = row[no_col] if no_col in row.keys() else str(bid)
    # 1. 下游引用保护
    if biz_type == 'purchase_request':
        if conn.execute("SELECT COUNT(*) FROM purchase_orders WHERE req_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该申请已被订单引用，无法删除（可作废）'}), 400
    if biz_type == 'purchase_order':
        if conn.execute("SELECT COUNT(*) FROM receivings WHERE order_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该订单已有入库记录，无法删除（可作废）'}), 400
        if conn.execute("SELECT COUNT(*) FROM contracts WHERE order_id=?", (bid,)).fetchone()[0] > 0:
            conn.close(); return jsonify({'error': '该订单已生成合同，无法删除（可作废）'}), 400
    if biz_type == 'receiving' and status == '已入库':
        conn.close(); return jsonify({'error': '该入库单已完成入库(已增加库存)，请先作废回滚库存后再删除'}), 400
    if biz_type == 'requisition' and status == '已出库':
        conn.close(); return jsonify({'error': '该出库单已完成出库(已扣减库存)，请先作废回滚库存后再删除'}), 400
    # 2. 级联删除
    rel_map = {
        'purchase_request': [('request_items', 'req_id')],
        'purchase_order': [('order_items', 'order_id')],
        'receiving': [],
        'requisition': [('requisition_items', 'requisition_id')],
        'contract': [], 'payment': [],
    }
    for rel, fk in rel_map.get(biz_type, []):
        conn.execute(f"DELETE FROM {rel} WHERE {fk}=?", (bid,))
    conn.execute("DELETE FROM approval_instances WHERE biz_type=? AND biz_id=?", (biz, bid))
    conn.execute("DELETE FROM dingtalk_instances WHERE biz_type=? AND biz_id=?", (biz, bid))
    conn.execute("DELETE FROM inventory_flows WHERE doc_type=? AND doc_id=?", (biz, bid))
    conn.execute(f"DELETE FROM {table} WHERE id=?", (bid,))
    conn.commit(); conn.close()
    log(session['user_name'], '删除单据', f'{biz_type}#{bid} {no}')
    return jsonify({'success': True, 'message': f'单据 {no} 已删除'})

# ---- 报表中心扩充数据 ----
@app.route('/api/reports2')
@login_required
def api_reports2():
    """需求44-报表中心: 采购/库存/往来报表聚合"""
    c = db()
    # 采购订单执行表
    exec_rows = c.execute("""SELECT order_no,item_name,supplier,total_amount,trade_mode,status,target_date,created_at
        FROM purchase_orders ORDER BY id DESC LIMIT 200""").fetchall()
    # 进货统计(按商品)
    by_item = c.execute("""SELECT item_name, COUNT(*) cnt, SUM(quantity) qty, SUM(total_amount) s FROM purchase_orders
        GROUP BY item_name ORDER BY s DESC LIMIT 20""").fetchall()
    # 进货统计(按供应商)
    by_sup = c.execute("""SELECT supplier, COUNT(*) cnt, SUM(total_amount) s FROM purchase_orders
        WHERE supplier!='' GROUP BY supplier ORDER BY s DESC LIMIT 20""").fetchall()
    # 采购历史价格跟踪
    price_track = c.execute("""SELECT item_name,spec,price,supplier,created_at FROM purchase_orders
        ORDER BY created_at DESC LIMIT 50""").fetchall()
    # 库存: 缺货预警
    stock_low = c.execute("SELECT * FROM inventory WHERE quantity<=safe_stock AND safe_stock>0").fetchall()
    # 出入库流水
    inout = c.execute("""SELECT '入库' type, receive_no no, item_name, qualified_qty qty, unit, received_at t FROM receivings
        UNION ALL SELECT '出库', req_no, item_name, quantity, unit, created_at FROM requisitions
        ORDER BY t DESC LIMIT 100""").fetchall()
    # 出入库汇总(近30天)
    in_sum = c.execute("SELECT COALESCE(SUM(qualified_qty),0) FROM receivings WHERE received_at >= datetime('now','localtime','-30 days')").fetchone()[0]
    out_sum = c.execute("SELECT COALESCE(SUM(quantity),0) FROM requisitions WHERE created_at >= datetime('now','localtime','-30 days')").fetchone()[0]
    # 付款统计
    pay_sum = c.execute("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status='已付款'").fetchone()[0]
    # 超期应付款: 已挂账但未付款超过30天
    overdue_pay = c.execute("""SELECT cn.credit_no, cn.supplier, cn.amount, cn.created_at FROM credit_notes cn
        WHERE cn.status='已通过' AND cn.created_at <= datetime('now','localtime','-30 days')
        AND NOT EXISTS (SELECT 1 FROM payment_requests pr WHERE pr.credit_id=cn.id AND pr.status IN ('已付款','已通过'))""").fetchall()
    # 供应商账本: 每家供应商往来汇总
    sup_ledger = c.execute("""SELECT supplier, COUNT(*) cnt, SUM(total_amount) s FROM purchase_orders
        WHERE supplier!='' GROUP BY supplier""").fetchall()
    # V11.206c 模块七: 入库统计(按类型: 正式/暂估/分批) + 出库统计(按部门)
    rcv_stats = c.execute("""SELECT
            CASE WHEN batch_no!='' AND is_est=1 THEN '分批入库-暂估'
                 WHEN batch_no!='' THEN '分批入库-正式'
                 WHEN is_est=1 AND invoice_no!='' THEN '暂估已红冲'
                 WHEN is_est=1 THEN '暂估入库'
                 ELSE '正式入库' END typ,
            COUNT(*) cnt, COALESCE(SUM(qualified_qty),0) qty
        FROM receivings WHERE status='已入库' GROUP BY typ ORDER BY qty DESC""").fetchall()
    req_stats = c.execute("""SELECT dept, COUNT(*) cnt, COALESCE(SUM(quantity),0) qty
        FROM requisitions WHERE status IN ('已出库','已通过') AND dept!='' GROUP BY dept ORDER BY qty DESC LIMIT 15""").fetchall()
    c.close()
    return jsonify({
        'exec': [dict_row(r) for r in exec_rows],
        'by_item': [dict_row(r) for r in by_item],
        'by_sup': [dict_row(r) for r in by_sup],
        'price_track': [dict_row(r) for r in price_track],
        'stock_low': [dict_row(r) for r in stock_low],
        'inout': [dict_row(r) for r in inout],
        'in_sum': in_sum, 'out_sum': out_sum,
        'rcv_stats': [dict_row(r) for r in rcv_stats],
        'req_stats': [dict_row(r) for r in req_stats],
        'pay_sum': pay_sum,
        'overdue_pay': [dict_row(r) for r in overdue_pay],
        'sup_ledger': [dict_row(r) for r in sup_ledger],
    })

# ---- 系统设置: 选项/企业信息 ----
@app.route('/api/health')
def api_health():
    db_ok = True
    try:
        c = db(); c.execute("SELECT 1"); c.close()
    except Exception:
        db_ok = False
    return jsonify({
        'status': 'ok' if db_ok else 'degraded',
        'system': '正成能源智慧采购系统',
        'version': 'V5.1 专业加固版',
        'time': now(),
        'database': 'ok' if db_ok else 'error',
        'last_backup': last_backup_name(),
        'login_user': session.get('user_name', ''),
    })

@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    if request.method == 'POST':
        if not can_manage_config(): return jsonify({'error': '仅系统管理员可操作'}), 403
        d = request.json or {}
        for k in ('company_name', 'company_address', 'company_contact', 'company_phone', 'write_lock'):
            if k in d: cfg_set(k, d[k])
        changed = [k for k in ('company_name','company_address','company_contact','company_phone','write_lock') if k in d]
        log(session.get('user_name',''), '修改系统设置', '变更项: %s' % (','.join(changed) or '无'))
        return jsonify({'success': True})
    info = {k: cfg_get(k) for k in ('company_name', 'company_address', 'company_contact', 'company_phone')}
    info.update({
        'write_lock': cfg_get('write_lock', '1'),
        'version': 'V5.1 专业加固版',
        'last_backup': last_backup_name(),
        'db_status': 'ok',
        'notice_publishers': cfg_get('notice_publishers', ''),  # V11.196 公告发布授权
    })
    return jsonify(info)

@app.route('/api/settings/notice-publishers', methods=['POST'])
@login_required
def api_settings_notice_publishers():
    """公告发布授权(V11.196): 仅系统管理员可设置; 授权后该用户登录系统设置可见公告管理并可发布"""
    if session.get('user_role') != '系统管理员':
        return jsonify({'error': '仅系统管理员可设置公告发布授权'}), 403
    d = request.json or {}
    pubs = (d.get('publishers') or '').strip()
    # 校验用户都存在
    if pubs:
        names = [x.strip() for x in pubs.split(',') if x.strip()]
        conn = db()
        for nm in names:
            if not conn.execute("SELECT 1 FROM users WHERE name=? AND is_active=1", (nm,)).fetchone():
                conn.close(); return jsonify({'error': f'用户「{nm}」不存在或未启用'}), 400
        conn.close()
    cfg_set('notice_publishers', pubs)
    log(session.get('user_name', ''), '修改公告发布授权', f'授权用户: {pubs or "(仅管理员)"}')
    return jsonify({'success': True, 'message': '公告发布授权已保存'})

# ============================================================
# ── PAGES ──
# ============================================================
@app.route('/test-login')
def test_login():
    return render_template('test_login.html')

@app.route('/')
def login_page():
    # V8.0: 禁用页面缓存 — 每次访问强制最新版本, 避免旧JS导致按钮失灵
    resp = make_response(render_template('index.html'))
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

# ============================================================
# V11.220 数据看板(分析统计大屏) — 只读聚合, 不动物资采购/维修/库存等任何业务逻辑
# 数据源=现有业务表实时读取; 新增无业务表(仅复用已有); 权限按角色分级(全量/域/仅自己)
# ============================================================
_DASH_LIVE = {'草稿', '已驳回', '已作废', '已撤销', '作废'}          # 统计排除的无效态
_DASH_FULL_ROLES = ('系统管理员', '分管领导', '总经理', '财务', '采购员')

def _dash_scope():
    """看板数据域: full=全量(管理/领导/财务/采购) / own=仅自己发起的(普通员工)"""
    r = session.get('user_role', '员工')
    if r in _DASH_FULL_ROLES:
        return 'full'
    return 'own'

def _dash_filters(c, alias, year=None, supplier=None, dept=None, time_range=None):
    """按看板顶部筛选拼 WHERE; 返回 (sql片段, 参数list). alias=单据时间列所在表别名, 时间字段固定 created_at"""
    conds, ps = [], []
    if year:
        conds.append(f"substr({alias}.created_at,1,4)=?")
        ps.append(str(year))
    if supplier:
        conds.append(f"{alias}.supplier=?")
        ps.append(supplier)
    if dept:
        conds.append(f"{alias}.dept=?")
        ps.append(dept)
    if time_range:
        conds.append(f"{alias}.created_at>=?")
        ps.append(time_range)
    return (' AND '.join(conds), ps) if conds else ('1=1', [])

def _dash_month_keys(year):
    return ['%s-%02d' % (year, m) for m in range(1, 13)]

def _dash_scope_own(c, alias, user_name, user_id):
    return f"({alias}.requester=? OR {alias}.requester_id=?)", [user_name, user_id]

@app.route('/api/dashboard/meta')
@login_required
def api_dashboard_meta():
    """筛选下拉: 年份/供应商/物料类别/部门/仓库 + 权限提示"""
    c = db()
    yrs = [r[0] for r in c.execute("SELECT DISTINCT substr(created_at,1,4) y FROM purchase_orders WHERE created_at!='' ORDER BY y DESC")]
    if not yrs:
        yrs = [str(datetime.date.today().year)]
    sups = [r[0] for r in c.execute("SELECT DISTINCT supplier FROM purchase_orders WHERE supplier!='' ORDER BY supplier")]
    cats = [r[0] for r in c.execute("SELECT DISTINCT category FROM purchase_orders WHERE category!='' ORDER BY category")]
    deps = [r[0] for r in c.execute("SELECT DISTINCT dept FROM purchase_requests WHERE dept!='' ORDER BY dept")]
    whs = [r[0] for r in c.execute("SELECT DISTINCT warehouse FROM inventory WHERE warehouse!='' ORDER BY warehouse")]
    c.close()
    return jsonify({'years': yrs, 'suppliers': sups, 'categories': cats, 'departments': deps,
                    'warehouses': whs, 'scope': _dash_scope(),
                    'role': session.get('user_role', '')})

@app.route('/api/dashboard/overview')
@login_required
def api_dashboard_overview():
    """Tab1 经营总览: 6指标卡 + 近5年柱线对比 + 本年月度面积趋势 + 4组明细"""
    y = int(request.args.get('year') or datetime.date.today().year)
    sup = request.args.get('supplier') or ''
    c = db()
    own = _dash_scope() == 'own'
    scope_sql = scope_ord = scope_con = ''
    if own:
        # 各表人员列不同: 申请/维修单有 requester_id; 订单表只有 owner_id; 合同表无人员列(经订单EXISTS)
        scope_sql = " AND (requester=? OR requester_id=?)"
        scope_ord = " AND (requester=? OR owner_id=?)"
        scope_con = " AND EXISTS(SELECT 1 FROM purchase_orders po WHERE po.id=c.order_id AND (po.requester=? OR po.owner_id=?))"
        own_ps = [session.get('user_name', ''), session.get('user_id', 0)]
    else:
        own_ps = []

    def q(sql, ps=()):
        return c.execute(sql, ps).fetchall()

    # --- 核心指标 ---
    # 本年采购总额/订单数(生效订单)
    yyyy = str(y)
    pre = str(y - 1)
    base = "SELECT COALESCE(SUM(total_amount),0) a, COUNT(*) n FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + (scope_sql if False else '')
    # 订单金额/数量: 本年+去年(同比)
    this_o = q("SELECT COALESCE(SUM(total_amount),0), COUNT(*) FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=? AND (?='' OR supplier=?)" + (scope_ord if own else ''), (yyyy, sup, sup) + (tuple(own_ps) if own else ()))[0]
    last_o = q("SELECT COALESCE(SUM(total_amount),0), COUNT(*) FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=? AND (?='' OR supplier=?)" + (scope_ord if own else ''), (pre, sup, sup) + (tuple(own_ps) if own else ()))[0]
    # 平均单价(本年)
    avg_price = (this_o[0] / this_o[1]) if this_o[1] else 0
    avg_last = (last_o[0] / last_o[1]) if last_o[1] else 0
    # 应付余额: 合同金额-已付款(全部在账)
    pay_total = q("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status NOT IN ('草稿','已驳回','已作废')")[0][0]
    con_total = q("SELECT COALESCE(SUM(amount),0) FROM contracts WHERE status NOT IN ('已驳回','已作废')")[0][0]
    ap_bal = con_total - pay_total
    # 库存金额(现库存)
    inv_val = q("SELECT COALESCE(SUM(quantity*price),0) FROM inventory")[0][0]
    # 维修费用(本年): 外委用 invoice_amount||quote_total||est_cost
    repair_amt = q("SELECT COALESCE(SUM(CASE WHEN invoice_amount>0 THEN invoice_amount WHEN quote_total>0 THEN quote_total ELSE est_cost END),0) FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + (scope_sql if own else ''), (yyyy,) + (tuple(own_ps) if own else ()))[0][0]
    repair_amt_last = q("SELECT COALESCE(SUM(CASE WHEN invoice_amount>0 THEN invoice_amount WHEN quote_total>0 THEN quote_total ELSE est_cost END),0) FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + (scope_sql if own else ''), (pre,) + (tuple(own_ps) if own else ()))[0][0]

    def yoy(cur, last):
        if last and last > 0:
            return round((cur - last) / last * 100, 1)
        return None

    cards = [
        {'k': '本年采购总额', 'v': round(this_o[0] / 10000, 2), 'unit': '万元', 'cmp': yoy(this_o[0], last_o[0]), 'cmp_label': '同比'},
        {'k': '采购订单总数', 'v': this_o[1], 'unit': '笔', 'cmp': yoy(this_o[1], last_o[1]), 'cmp_label': '同比'},
        {'k': '平均采购单价', 'v': round(avg_price, 2), 'unit': '元/笔', 'cmp': yoy(avg_last - avg_price, avg_last) if avg_last else None, 'cmp_label': '降本'},
        {'k': '应付账款余额', 'v': round(ap_bal / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '在账'},
        {'k': '库存总金额', 'v': round(inv_val / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '现值'},
        {'k': '维修费用总额', 'v': round(repair_amt / 10000, 2), 'unit': '万元', 'cmp': yoy(repair_amt, repair_amt_last), 'cmp_label': '同比'},
    ]
    # --- 近5年 采购金额+降本 ---
    years5 = []
    for yy in range(y - 4, y + 1):
        o = q("SELECT COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + (scope_ord if own else ''), (str(yy),) + (tuple(own_ps) if own else ()))[0][0]
        rq = q("SELECT COALESCE(SUM(total_estimated),0) FROM purchase_requests WHERE status NOT IN ('草稿','已驳回','已作废','已撤销') AND substr(created_at,1,4)=?" + (scope_sql if own else ''), (str(yy),) + (tuple(own_ps) if own else ()))[0][0]
        years5.append({'year': yy, 'amt': round(o / 10000, 2), 'save': round(max(rq - o, 0) / 10000, 2)})
    # --- 本年月度采购趋势(金额) ---
    mrows = q("SELECT substr(created_at,6,2) m, COALESCE(SUM(total_amount),0) a FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=? AND (?='' OR supplier=?)" + (scope_ord if own else '') + " GROUP BY m", (yyyy, sup, sup) + (tuple(own_ps) if own else ()))
    mdict = {r[0]: round(r[1] / 10000, 2) for r in mrows}
    months = [mdict.get(m, 0) for m in ['%02d' % i for i in range(1, 13)]]
    # --- 明细 ---
    scope_d = scope_sql if own else ''     # 申请/维修表域过滤
    scope_do = scope_ord if own else ''    # 订单表域过滤
    scope_dc = scope_con if own else ''    # 合同表域过滤(经订单EXISTS)
    scope_ps = own_ps if own else []
    lim = ' LIMIT 200'
    mat_drop = q("SELECT pr.id, pr.req_no, ri.item_name, ri.spec, pr.dept, pr.total_estimated FROM purchase_requests pr LEFT JOIN request_items ri ON ri.req_id=pr.id WHERE pr.status NOT IN ('草稿','已驳回','已作废') AND ri.id IS NOT NULL" + scope_d + lim, tuple(scope_ps))
    sup_rank = q("SELECT supplier, COALESCE(SUM(total_amount),0) amt, COUNT(*) n FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废')" + scope_do + " GROUP BY supplier ORDER BY amt DESC LIMIT 10", tuple(scope_ps))
    ap_rows = q("SELECT contract_no, supplier, amount, COALESCE((SELECT SUM(amount) FROM payment_requests p WHERE p.contract_id=c.id AND p.status NOT IN ('草稿','已驳回','已作废')),0) paid FROM contracts c WHERE status NOT IN ('已驳回','已作废')" + scope_dc + " ORDER BY amount DESC" + lim, tuple(scope_ps))
    c.close()
    return jsonify({'cards': cards, 'years5': years5, 'months': months, 'labels_m': ['%d月' % i for i in range(1, 13)],
                    'tables': {
                        'material_drop': [dict_row({'req_no': r[1], 'item_name': r[2], 'spec': r[3], 'dept': r[4], 'est': r[5]}) for r in mat_drop],
                        'supplier_rank': [dict_row({'supplier': r[0], 'amt': r[1], 'orders': r[2]}) for r in sup_rank],
                        'ap_detail': [dict_row({'contract_no': r[0], 'supplier': r[1], 'amount': r[2], 'paid': r[3]}) for r in ap_rows],
                    }})


@app.route('/api/dashboard/purchase')
@login_required
def api_dashboard_purchase():
    """Tab2 采购分析: 6指标 + 申请vs订单双柱 + 类别占比环 + 供应商TOP10横柱 + 3明细"""
    y = int(request.args.get('year') or datetime.date.today().year)
    sup = request.args.get('supplier') or ''
    cat = request.args.get('category') or ''
    c = db()
    own = _dash_scope() == 'own'
    scope_d = scope_ord = scope_inq = scope_inqa = ''
    scope_ps = []
    if own:
        scope_d = " AND (requester=? OR requester_id=?)"   # 申请单表
        scope_ord = " AND (requester=? OR owner_id=?)"     # 订单表(无requester_id)
        scope_inq = " AND created_by=?"                    # 询价表(created_by存姓名)
        scope_inqa = " AND i.created_by=?"                 # 询价表带别名
        scope_ps = [session.get('user_name', ''), session.get('user_id', 0)]
    yyyy = str(y)
    supf = " AND (?='' OR supplier=?)"
    sups_ps = [sup, sup]
    catf = " AND (?='' OR category=?)"
    cats_ps = [cat, cat]
    iq_name = session.get('user_name', '') if own else ''

    # 指标
    rq_n = c.execute("SELECT COUNT(*), COALESCE(SUM(total_estimated),0) FROM purchase_requests WHERE status NOT IN ('草稿','已驳回','已作废','已撤销') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps)).fetchone()
    od_n = c.execute("SELECT COUNT(*), COALESCE(SUM(total_amount),0) FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + supf + scope_ord, (yyyy,) + tuple(sups_ps) + tuple(scope_ps)).fetchone()
    inq_n = c.execute("SELECT COUNT(*) FROM inquiries WHERE substr(created_at,1,4)=?" + scope_inq, (yyyy,) + ((iq_name,) if own else ())).fetchone()[0]
    # 询价节约: Σ(选中供应商对应申请预算-成交价) 简化=申请预算总额-订单总额(该年, 有询价来源)
    save = max(rq_n[1] - od_n[1], 0)
    cards = [
        {'k': '采购申请总数', 'v': rq_n[0], 'unit': '笔', 'cmp': None, 'cmp_label': '本年'},
        {'k': '采购申请总金额', 'v': round(rq_n[1] / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年'},
        {'k': '采购订单总数', 'v': od_n[0], 'unit': '笔', 'cmp': None, 'cmp_label': '本年'},
        {'k': '采购订单总金额', 'v': round(od_n[1] / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年'},
        {'k': '询价比价次数', 'v': inq_n, 'unit': '次', 'cmp': None, 'cmp_label': '本年'},
        {'k': '采购节约金额', 'v': round(save / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '预算-成交'},
    ]
    # 申请vs订单 月度双柱(金额万元)
    def monthly(sql, ps_extra=()):
        rows = c.execute(sql, ps_extra).fetchall()
        d = {r[0]: round(r[1] / 10000, 2) for r in rows}
        return [d.get('%02d' % i, 0) for i in range(1, 13)]
    req_m = monthly("SELECT substr(created_at,6,2) m, COALESCE(SUM(total_estimated),0) a FROM purchase_requests WHERE status NOT IN ('草稿','已驳回','已作废','已撤销') AND substr(created_at,1,4)=?" + scope_d + " GROUP BY m", (yyyy,) + tuple(scope_ps))
    ord_m = monthly("SELECT substr(created_at,6,2) m, COALESCE(SUM(total_amount),0) a FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + supf + scope_ord + " GROUP BY m", (yyyy,) + tuple(sups_ps) + tuple(scope_ps))
    # 类别占比(物资采购 vs 设备维修采购)按 req_type + order.category 近似
    cat_rows = c.execute("SELECT COALESCE(NULLIF(req_type,''),'物资采购') t, COALESCE(SUM(total_estimated),0) a FROM purchase_requests WHERE status NOT IN ('草稿','已驳回','已作废','已撤销') AND substr(created_at,1,4)=?" + scope_d + " GROUP BY t", (yyyy,) + tuple(scope_ps)).fetchall()
    cat_pie = [{'name': '物资采购' if '物资' in (r[0] or '') else '设备维修采购', 'value': round(r[1], 0)} for r in cat_rows] or [{'name': '物资采购', 'value': 0}]
    # 供应商TOP10(横柱)
    sup_top = c.execute("SELECT supplier, COALESCE(SUM(total_amount),0) a, COUNT(*) n FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND supplier!=''" + scope_ord + " GROUP BY supplier ORDER BY a DESC LIMIT 10", tuple(scope_ps)).fetchall()
    # 明细
    def rows(sql, ps=()):
        return [dict_row(dict(zip([d[0] for d in c.description], r))) for r in c.execute(sql, ps).fetchall()] if sql else []
    # 用统一 dict 转换
    def rows2(sql, ps=()):
        cur = c.execute(sql, ps)
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    rq_det = rows2("SELECT req_no, dept, requester, created_at, purpose, total_estimated, status FROM purchase_requests WHERE status NOT IN ('草稿','已驳回','已作废','已撤销') AND substr(created_at,1,4)=?" + scope_d + " ORDER BY created_at DESC LIMIT 200", (yyyy,) + tuple(scope_ps))
    od_det = rows2("SELECT order_no, supplier, total_amount, created_at, status, trade_mode FROM purchase_orders WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + supf + scope_ord + " ORDER BY created_at DESC LIMIT 200", (yyyy,) + tuple(sups_ps) + tuple(scope_ps))
    iq_det = rows2("SELECT i.inq_no, i.title, i.status, i.created_at, (SELECT COUNT(*) FROM inquiry_suppliers s WHERE s.inquiry_id=i.id) sup_cnt FROM inquiries i WHERE substr(i.created_at,1,4)=?" + scope_inqa + " ORDER BY i.created_at DESC LIMIT 200", (yyyy,) + ((iq_name,) if own else ()))
    c.close()
    return jsonify({'cards': cards,
                    'req_vs_ord': {'months': ['%d月' % i for i in range(1, 13)], 'req': req_m, 'ord': ord_m},
                    'cat_pie': cat_pie,
                    'sup_top': [dict_row({'supplier': r[0], 'amt': r[1], 'orders': r[2]}) for r in sup_top],
                    'tables': {'req': rq_det, 'ord': od_det, 'inq': iq_det}})

@app.route('/api/dashboard/inventory')
@login_required
def api_dashboard_inventory():
    """Tab3 库存分析: 6指标 + 入出库双折线 + 类别占比环 + 季度周转柱 + 4明细"""
    y = int(request.args.get('year') or datetime.date.today().year)
    wh = request.args.get('warehouse') or ''
    c = db()
    own = _dash_scope() == 'own'
    scope_d = " AND warehouse=?" if wh else ''
    scope_ps = [wh] if wh else []
    yyyy = str(y)

    def ex(sql, ps=()):
        return c.execute(sql, ps).fetchall()
    def ex1(sql, ps=()):
        r = c.execute(sql, ps).fetchone()
        return r[0] if r else 0
    inv = ex("SELECT COALESCE(SUM(quantity),0), COALESCE(SUM(quantity*price),0) FROM inventory" + (" WHERE " + scope_d[4:] if scope_d else ''), scope_ps)
    if not inv:
        inv = [(0, 0)]
    inv_qty, inv_val = inv[0][0], inv[0][1]
    low_n = ex1("SELECT COUNT(*) FROM inventory WHERE quantity < safe_stock AND safe_stock > 0" + scope_d, scope_ps)
    # 入/出库数量(本年)
    in_q = ex1("SELECT COALESCE(SUM(quantity),0) FROM receivings WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废')" + scope_d, (yyyy,) + tuple(scope_ps))
    out_q = ex1("SELECT COALESCE(SUM(quantity),0) FROM requisitions WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废')", (yyyy,))
    # 金额口径: 入库金额 receivings.est_amount 或 items 价估算; 出库金额= requisitions量×inventory.price 估算
    in_amt = ex1("SELECT COALESCE(SUM(CASE WHEN est_amount>0 THEN est_amount ELSE quantity*(SELECT COALESCE(price,0) FROM inventory i WHERE i.item_name=r.item_name LIMIT 1) END),0) FROM receivings r WHERE substr(r.created_at,1,4)=? AND r.status NOT IN ('草稿','已驳回','已作废')" + scope_d.replace('warehouse','r.warehouse'), (yyyy,) + tuple(scope_ps))
    avg_inv_val = inv_val or 1
    turn = round(out_q * (inv_val / inv_qty if inv_qty else 0) / avg_inv_val, 2) if inv_qty else 0  # 出库成本/平均库存
    cards = [
        {'k': '当前库存总数量', 'v': inv_qty, 'unit': '', 'cmp': None, 'cmp_label': '现量'},
        {'k': '当前库存总金额', 'v': round(inv_val / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '现值'},
        {'k': '本年入库总数量', 'v': in_q, 'unit': '', 'cmp': None, 'cmp_label': '本年'},
        {'k': '本年出库总数量', 'v': out_q, 'unit': '', 'cmp': None, 'cmp_label': '本年'},
        {'k': '库存周转率', 'v': turn, 'unit': '次', 'cmp': None, 'cmp_label': '估算'},
        {'k': '库存预警物料数', 'v': low_n, 'unit': '种', 'cmp': None, 'cmp_label': '低于安全库存'},
    ]
    # 月度 入出库数量
    def mrows(sql, ps=()):
        d = {r[0]: r[1] for r in ex(sql, ps)}
        return [d.get('%02d' % i, 0) for i in range(1, 13)]
    in_m = mrows("SELECT substr(created_at,6,2) m, COALESCE(SUM(quantity),0) FROM receivings WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废')" + scope_d + " GROUP BY m", (yyyy,) + tuple(scope_ps))
    out_m = mrows("SELECT substr(created_at,6,2) m, COALESCE(SUM(quantity),0) FROM requisitions WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废') GROUP BY m", (yyyy,))
    # 类别占比(库存金额按cat_code)
    cat_pie = ex("SELECT COALESCE(cat_code,'其他') t, COALESCE(SUM(quantity*price),0) a FROM inventory" + (" WHERE " + scope_d[4:] if scope_d else '') + " GROUP BY t", scope_ps)
    cats = {}
    for code, a in cat_pie:
        nm = '未分类'
        try:
            row = c.execute('SELECT name FROM categories WHERE code=?', (code,)).fetchone()
            nm = row[0] if row else (code or nm)
        except Exception:
            pass
        cats[nm] = cats.get(nm, 0) + a
    cat_list = [{'name': k, 'value': round(v, 0)} for k, v in cats.items()] or [{'name': '暂无', 'value': 0}]
    # 季度周转柱
    q_turn = []
    for qq in range(1, 5):
        qin = ex1("SELECT COALESCE(SUM(quantity),0) FROM receivings WHERE substr(created_at,1,4)=? AND CAST(substr(created_at,6,2) AS INTEGER) BETWEEN ? AND ?" + scope_d, (yyyy, (qq - 1) * 3 + 1, qq * 3) + tuple(scope_ps))
        qout = ex1("SELECT COALESCE(SUM(quantity),0) FROM requisitions WHERE substr(created_at,1,4)=? AND CAST(substr(created_at,6,2) AS INTEGER) BETWEEN ? AND ?", (yyyy, (qq - 1) * 3 + 1, qq * 3))
        q_turn.append(round(qout / (avg_inv_val + 1) * 100, 2) if avg_inv_val else 0)
    # 明细
    def rows2(sql, ps=()):
        cur = c.execute(sql, ps)
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    inv_det = rows2("SELECT item_name, spec, unit, quantity, price, (quantity*price) amt, warehouse FROM inventory" + (" WHERE " + scope_d[4:] if scope_d else '') + " ORDER BY amt DESC LIMIT 200", scope_ps)
    in_det = rows2("SELECT receive_no, item_name, quantity, est_amount, is_est, received_at, warehouse FROM receivings WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废')" + scope_d + " ORDER BY created_at DESC LIMIT 200", (yyyy,) + tuple(scope_ps))
    out_det = rows2("SELECT req_no, item_name, quantity, dept, receiver, created_at FROM requisitions WHERE substr(created_at,1,4)=? AND status NOT IN ('草稿','已驳回','已作废') ORDER BY created_at DESC LIMIT 200", (yyyy,))
    low_det = rows2("SELECT item_name, quantity, safe_stock, (safe_stock-quantity) gap FROM inventory WHERE quantity < safe_stock AND safe_stock > 0" + scope_d + " ORDER BY gap DESC LIMIT 100", scope_ps)
    c.close()
    return jsonify({'cards': cards,
                    'io': {'months': ['%d月' % i for i in range(1, 13)], 'in': in_m, 'out': out_m},
                    'cat_pie': cat_list,
                    'quarter_turn': {'quarters': ['Q1', 'Q2', 'Q3', 'Q4'], 'turn': q_turn},
                    'tables': {'inv': inv_det, 'in': in_det, 'out': out_det, 'low': low_det}})


@app.route('/api/dashboard/finance')
@login_required
def api_dashboard_finance():
    """Tab4 财务分析: 6指标 + 合同月度柱 + 开票vs付款双折线 + 账龄饼 + 4明细(财务角色可见)"""
    if session.get('user_role') not in _DASH_FULL_ROLES:
        return jsonify({'error': '无财务分析查看权限'}), 403
    y = int(request.args.get('year') or datetime.date.today().year)
    c = db()
    yyyy = str(y)

    def ex1(sql, ps=()):
        r = c.execute(sql, ps).fetchone()
        return r[0] if r else 0
    def rows2(sql, ps=()):
        cur = c.execute(sql, ps)
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in cur.fetchall()]

    con_amt = ex1("SELECT COALESCE(SUM(amount),0) FROM contracts WHERE status NOT IN ('已驳回','已作废') AND substr(sign_date,1,4)=?", (yyyy,))
    inv_amt = ex1("SELECT COALESCE(SUM(amount),0) FROM invoices WHERE status NOT IN ('作废') AND substr(created_at,1,4)=?", (yyyy,)) + ex1("SELECT COALESCE(SUM(amount),0) FROM contract_invoices WHERE substr(created_at,1,4)=?", (yyyy,))
    pay_amt = ex1("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(paid_at,1,4)=?", (yyyy,))
    ap_bal = ex1("SELECT COALESCE(SUM(amount),0) FROM contracts WHERE status NOT IN ('已驳回','已作废')") - ex1("SELECT COALESCE(SUM(amount),0) FROM payment_requests WHERE status NOT IN ('草稿','已驳回','已作废')")
    due_n = ex1("SELECT COUNT(*) FROM contracts WHERE inv_collect_status IN ('已催收待回','待收票','待催收') AND status NOT IN ('已驳回','已作废')")
    cards = [
        {'k': '合同总金额', 'v': round(con_amt / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年签订'},
        {'k': '已开票金额', 'v': round(inv_amt / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年'},
        {'k': '未开票金额', 'v': round(max(con_amt - inv_amt, 0) / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '差额'},
        {'k': '已付款金额', 'v': round(pay_amt / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年'},
        {'k': '应付账款余额', 'v': round(ap_bal / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '在账'},
        {'k': '发票催收预警数', 'v': due_n, 'unit': '个', 'cmp': None, 'cmp_label': '合同'},
    ]
    def mrows(sql, ps=()):
        d = {r[0]: r[1] for r in c.execute(sql, ps).fetchall()}
        return [d.get('%02d' % i, 0) for i in range(1, 13)]
    con_m = [round(x / 10000, 2) for x in mrows("SELECT substr(sign_date,6,2) m, COALESCE(SUM(amount),0) FROM contracts WHERE status NOT IN ('已驳回','已作废') AND sign_date LIKE ? GROUP BY m", (yyyy + '%',))]
    inv_m = [round(x / 10000, 2) for x in mrows("SELECT substr(created_at,6,2) m, COALESCE(SUM(amount),0) FROM invoices WHERE status NOT IN ('作废') AND created_at LIKE ? GROUP BY m", (yyyy + '%',))]
    pay_m = [round(x / 10000, 2) for x in mrows("SELECT substr(paid_at,6,2) m, COALESCE(SUM(amount),0) FROM payment_requests WHERE status NOT IN ('草稿','已驳回','已作废') AND paid_at LIKE ? GROUP BY m", (yyyy + '%',))]
    # 账龄饼: 应付余额按合同签订时间分桶(30/60/90/90+)
    aging = {'30天内': 0, '30-60天': 0, '60-90天': 0, '90天以上': 0}
    import datetime as _dt
    today = _dt.date.today()
    for r in c.execute("SELECT sign_date, amount FROM contracts WHERE status NOT IN ('已驳回','已作废')"):
        amt = r[1] or 0
        try:
            sd = _dt.datetime.strptime((r[0] or '')[:10], '%Y-%m-%d').date()
            days = (today - sd).days
        except Exception:
            days = 999
        if days <= 30:
            aging['30天内'] += amt
        elif days <= 60:
            aging['30-60天'] += amt
        elif days <= 90:
            aging['60-90天'] += amt
        else:
            aging['90天以上'] += amt
    aging_pie = [{'name': k, 'value': round(v, 0)} for k, v in aging.items()]
    # 明细
    con_det = rows2("SELECT contract_no, supplier, amount, sign_date, status, inv_collect_status FROM contracts WHERE status NOT IN ('已驳回','已作废') AND substr(sign_date,1,4)=? ORDER BY amount DESC LIMIT 200", (yyyy,))
    inv_det = rows2("SELECT invoice_no, supplier, amount, invoice_date, status FROM invoices WHERE status NOT IN ('作废') ORDER BY invoice_date DESC LIMIT 200")
    pay_det = rows2("SELECT payment_no, supplier, amount, paid_at, payment_type FROM payment_requests WHERE status NOT IN ('草稿','已驳回','已作废') ORDER BY paid_at DESC LIMIT 200")
    due_det = rows2("SELECT contract_no, supplier, amount, invoice_clause, inv_collect_status FROM contracts WHERE inv_collect_status IN ('已催收待回','待收票','待催收') AND status NOT IN ('已驳回','已作废') LIMIT 100")
    c.close()
    return jsonify({'cards': cards, 'con_m': con_m, 'inv_m': inv_m, 'pay_m': pay_m,
                    'months': ['%d月' % i for i in range(1, 13)], 'aging_pie': aging_pie,
                    'tables': {'con': con_det, 'inv': inv_det, 'pay': pay_det, 'due': due_det}})

@app.route('/api/dashboard/repair')
@login_required
def api_dashboard_repair():
    """Tab5 维修分析: 6指标 + 费用面积趋势 + 类型占比环 + 设备TOP10 + 3明细"""
    y = int(request.args.get('year') or datetime.date.today().year)
    c = db()
    own = _dash_scope() == 'own'
    scope_d = ''
    scope_ps = []
    if own:
        scope_d = " AND (requester=? OR requester_id=?)"
        scope_ps = [session.get('user_name', ''), session.get('user_id', 0)]
    yyyy = str(y)

    def ex1(sql, ps=()):
        r = c.execute(sql, ps).fetchone()
        return r[0] if r else 0
    def rows2(sql, ps=()):
        cur = c.execute(sql, ps)
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in cur.fetchall()]
    cost_sql = "CASE WHEN invoice_amount>0 THEN invoice_amount WHEN quote_total>0 THEN quote_total ELSE est_cost END"
    n_all = ex1("SELECT COUNT(*) FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    cost_all = ex1("SELECT COALESCE(SUM(" + cost_sql + "),0) FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    cost_ext = ex1("SELECT COALESCE(SUM(" + cost_sql + "),0) FROM repair_plans WHERE repair_type='委外维修' AND status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    cost_int = ex1("SELECT COALESCE(SUM(est_cost),0) FROM repair_plans WHERE repair_type='内部自修' AND status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    chg_n = ex1("SELECT COUNT(*) FROM repair_changes rc JOIN repair_plans rp ON rp.id=rc.plan_id WHERE rp.status NOT IN ('草稿','已驳回','已作废') AND substr(rp.created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    # 平均维修周期(报修→验收完成天数)
    cyc = ex1("SELECT COALESCE(AVG(julianday(actual_finish)-julianday(created_at)),0) FROM repair_plans WHERE actual_finish!='' AND status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d, (yyyy,) + tuple(scope_ps))
    cards = [
        {'k': '维修工单总数', 'v': n_all, 'unit': '单', 'cmp': None, 'cmp_label': '本年'},
        {'k': '维修费用总额', 'v': round(cost_all / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '本年'},
        {'k': '外委维修费用', 'v': round(cost_ext / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '委外'},
        {'k': '内部自修费用', 'v': round(cost_int / 10000, 2), 'unit': '万元', 'cmp': None, 'cmp_label': '自修'},
        {'k': '维修变更次数', 'v': chg_n, 'unit': '次', 'cmp': None, 'cmp_label': '本年'},
        {'k': '平均维修周期', 'v': round(cyc, 1), 'unit': '天', 'cmp': None, 'cmp_label': '报修→完工'},
    ]
    # 月度费用
    def mrows(sql, ps=()):
        d = {r[0]: r[1] for r in c.execute(sql, ps).fetchall()}
        return [round(d.get('%02d' % i, 0) / 10000, 2) for i in range(1, 13)]
    cost_m = mrows("SELECT substr(created_at,6,2) m, COALESCE(SUM(" + cost_sql + "),0) FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d + " GROUP BY m", (yyyy,) + tuple(scope_ps))
    # 类型占比(数量+金额)
    t_rows = c.execute("SELECT COALESCE(repair_type,'未定损') t, COUNT(*) n, COALESCE(SUM(" + cost_sql + "),0) a FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d + " GROUP BY t", (yyyy,) + tuple(scope_ps)).fetchall()
    type_pie = [{'name': r[0], 'n': r[1], 'value': round(r[2], 0)} for r in t_rows] or [{'name': '暂无', 'value': 0}]
    # 设备TOP10
    dev_top = c.execute("SELECT device_name, COALESCE(SUM(" + cost_sql + "),0) a, COUNT(*) n FROM repair_plans WHERE device_name!='' AND status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d + " GROUP BY device_name ORDER BY a DESC LIMIT 10", (yyyy,) + tuple(scope_ps)).fetchall()
    # 明细
    wd_det = rows2("SELECT plan_no, device_name, fault_desc, dept, requester, created_at, repair_type, est_cost, quote_total, invoice_amount, status FROM repair_plans WHERE status NOT IN ('草稿','已驳回','已作废') AND substr(created_at,1,4)=?" + scope_d + " ORDER BY created_at DESC LIMIT 200", (yyyy,) + tuple(scope_ps))
    chg_det = rows2("SELECT rp.plan_no, rc.add_part, rc.add_labor, rc.add_price, rc.status, rc.confirm1_by, rc.created_at FROM repair_changes rc JOIN repair_plans rp ON rp.id=rc.plan_id WHERE rp.status NOT IN ('草稿','已驳回','已作废')" + scope_d + " ORDER BY rc.created_at DESC LIMIT 200", tuple(scope_ps))
    sup_stat = c.execute("SELECT repair_company, COUNT(*) n, COALESCE(SUM(" + cost_sql + "),0) a FROM repair_plans WHERE repair_company!='' AND status NOT IN ('草稿','已驳回','已作废')" + scope_d + " GROUP BY repair_company ORDER BY a DESC LIMIT 50", tuple(scope_ps)).fetchall()
    c.close()
    return jsonify({'cards': cards, 'cost_m': cost_m, 'months': ['%d月' % i for i in range(1, 13)],
                    'type_pie': type_pie,
                    'dev_top': [dict_row({'device_name': r[0], 'amt': r[1], 'n': r[2]}) for r in dev_top],
                    'tables': {'work': wd_det, 'change': chg_det,
                               'vendor': [dict_row({'company': r[0], 'n': r[1], 'amt': r[2]}) for r in sup_stat]}})

@app.route('/dashboard')
@login_required
def page_dashboard():
    return render_template('dashboard.html')


@app.route('/api/dashboard/export')
@login_required
def api_dashboard_export():
    """看板数据导出 Excel(openpyxl): 指标+图表数据+明细多sheet"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except Exception:
        return jsonify({'error': '导出组件缺失'}), 500
    tab = request.args.get('tab') or 'overview'
    y = request.args.get('year') or str(datetime.date.today().year)
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = '看板指标'
    ws.append(['数据看板导出', 'Tab=' + tab, '年度=' + y, '导出时间=' + now()])
    hdr_fill = PatternFill('solid', fgColor='1a2a4a'); hdr_font = Font(color='FFFFFF', bold=True)
    def sheet_from(name, headers, rows):
        s = wb.create_sheet(re.sub(r'[\\/?*\[\]:]', '_', name)[:31])
        s.append(headers)
        for c in s[1]:
            c.fill = hdr_fill; c.font = hdr_font
        for r in rows:
            s.append([r.get(h) if isinstance(r, dict) else r for h in headers])
        return s
    def gen(tabname, endpoint, extra_cols=None):
        d = {}
        try:
            c = db()
            cur = c.execute('SELECT 1')
            # 直接内联各接口聚合(轻量复用): 调用对应函数再取json
        finally:
            pass
    # ---- 复用各看板接口的JSON, 转sheet ----
    import json as _json
    data_map = {}
    for t, ep in [('overview', 'overview'), ('purchase', 'purchase'), ('inventory', 'inventory'),
                  ('finance', 'finance'), ('repair', 'repair')]:
        try:
            # 模拟请求参数取数(直接调用聚合函数; 注入当前用户会话以通过login_required)
            with app.test_request_context('/api/dashboard/' + ep + '?year=' + y + '&supplier=&category=&dept=&warehouse='):
                session['user_id'] = session.get('user_id') or 1
                session['user_name'] = session.get('user_name') or '导出'
                session['user_role'] = session.get('user_role') or '系统管理员'
                fn = {'overview': api_dashboard_overview, 'purchase': api_dashboard_purchase,
                      'inventory': api_dashboard_inventory, 'finance': api_dashboard_finance,
                      'repair': api_dashboard_repair}[ep]
                resp = fn()
                data_map[t] = _json.loads(resp.get_data(as_text=True))
        except Exception as e:
            data_map[t] = {'error': str(e)}
    d = data_map.get(tab, {})
    if 'error' in d and tab != 'finance':
        return jsonify({'error': '导出失败: ' + d['error']}), 500
    # 指标卡 sheet
    for c_ in (d.get('cards') or []):
        ws.append([c_.get('k'), c_.get('v'), c_.get('unit', ''), c_.get('cmp_label', ''), c_.get('cmp', '')])
    # 明细 sheets
    tbl_defs = {
        'overview': [('物料降本/申请明细', 'material_drop', ['req_no', 'item_name', 'spec', 'dept', 'est']),
                     ('前十大供应商', 'supplier_rank', ['supplier', 'orders', 'amt']),
                     ('应付账款明细', 'ap_detail', ['contract_no', 'supplier', 'amount', 'paid'])],
        'purchase': [('采购申请明细', 'req', ['req_no', 'dept', 'requester', 'created_at', 'purpose', 'total_estimated', 'status']),
                     ('采购订单明细', 'ord', ['order_no', 'supplier', 'total_amount', 'created_at', 'status']),
                     ('询价比价明细', 'inq', ['inq_no', 'title', 'sup_cnt', 'created_at', 'status'])],
        'inventory': [('库存台账', 'inv', ['item_name', 'spec', 'unit', 'quantity', 'amt', 'warehouse']),
                      ('入库明细', 'in', ['receive_no', 'item_name', 'quantity', 'est_amount', 'received_at']),
                      ('出库明细', 'out', ['req_no', 'item_name', 'quantity', 'dept', 'receiver', 'created_at']),
                      ('库存预警', 'low', ['item_name', 'quantity', 'safe_stock', 'gap'])],
        'finance': [('合同明细', 'con', ['contract_no', 'supplier', 'amount', 'sign_date', 'status']),
                    ('发票台账', 'inv', ['invoice_no', 'supplier', 'amount', 'invoice_date', 'status']),
                    ('付款明细', 'pay', ['payment_no', 'supplier', 'amount', 'paid_at', 'payment_type']),
                    ('发票催收预警', 'due', ['contract_no', 'supplier', 'amount', 'invoice_clause', 'inv_collect_status'])],
        'repair': [('维修工单明细', 'work', ['plan_no', 'device_name', 'fault_desc', 'dept', 'created_at', 'repair_type', 'status']),
                   ('维修变更明细', 'change', ['plan_no', 'add_part', 'add_price', 'status', 'created_at']),
                   ('维修服务商统计', 'vendor', ['company', 'n', 'amt'])],
    }
    for label, key, cols in tbl_defs.get(tab, []):
        rows = (d.get('tables') or {}).get(key) or []
        sheet_from(label, cols, rows)
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    from flask import send_file as _sf
    fname = '数据看板_%s_%s.xlsx' % (tab, y)
    resp = _sf(bio, as_attachment=True, download_name=fname,
               mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp.headers['Cache-Control'] = 'no-store'
    return resp


# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    init_db()
    # V5.1 安全加固: 启动自动备份(每天首次启动备份一次)
    _startup_backup()
    # V4.1 预警调度(仅reloader子进程启动, 避免debug模式双线程重复推送)
    if os.environ.get('WERKZEUG_RUN_MAIN') == 'true' or not app.debug:
        threading.Thread(target=scheduler_loop, daemon=True).start()
        print('  [预警调度] 已启动 (每60秒扫描, 飞书机器人推送)')
        threading.Thread(target=dt_poll_loop, daemon=True).start()
        print('  [钉钉审批同步] 已启动 (每15秒轮询即时同步)')
        threading.Thread(target=_daily_backup_loop, daemon=True).start()
        print('  [自动备份] 每日03:00自动备份已启动')
    port = int(os.environ.get('PORT', 5899))
    # V11.1 多人协作: 页面/模板改动即时生效, 无需重启(后端app.py改动由 app_watchdog.py 自动重启)
    app.config['TEMPLATES_AUTO_RELOAD'] = True
    print(f"""
╔══════════════════════════════════════════════════╗
║  正成能源采购系统 v9.0 (UI美化版)      ║
║  安全: PBKDF2密码加密 | 登录失败锁定 | CSRF防护  ║
║  启动: http://127.0.0.1:{port}                     ║
║  默认账号: admin / admin123                       ║
║  飞书回调: http://<公网地址>:{port}/api/feishu/callback  ║
╚══════════════════════════════════════════════════╝
""")
    app.run(host='0.0.0.0', port=port, debug=False)  # 0.0.0.0=局域网可直连(李总等同事可用 http://本机IP:5899)


@app.route('/api/_debug')
@login_required
def api_debug():
    """临时调试: 查看 query 解析"""
    from urllib.parse import parse_qsl
    qs_bytes = request.query_string
    q_arg = request.args.get('q', '')
    out = {
        'query_string_bytes': repr(qs_bytes),
        'args_q_repr': repr(q_arg),
        'args_q_has_ufffd': '\ufffd' in q_arg,
        'raw_qs_latin1': repr(qs_bytes.decode('latin-1', 'replace')),
    }
    try:
        parsed = []
        for k, v in parse_qsl(qs_bytes.decode('latin-1', 'replace'), encoding='latin-1'):
            parsed.append({'k': k, 'v_repr': repr(v), 'gbk': v.encode('latin-1', 'replace').decode('gbk', 'replace')})
        out['parsed_latin1'] = parsed
    except Exception as e:
        out['parse_err'] = str(e)
    return jsonify(out)

@app.route('/simple-test')
def simple_test():
    return render_template('simple_test.html')

@app.route('/test')
def test_page():
    return render_template('test.html')


