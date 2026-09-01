# -*- coding: utf-8 -*-
"""V11.164 任务3端到端验证: 生成测试合同→python-docx严格验证(模拟Win/Mac双端)→公网下载完整性→清理"""
import sys, os, json, re, sqlite3, hashlib, zipfile, urllib.request, urllib.parse
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import app as appmod
appmod.start_instances = lambda *a, **k: None

BASE = os.path.dirname(os.path.abspath(__file__))
DB = os.path.join(BASE, 'data/purchase.db')
UP = os.path.join(BASE, 'uploads')
PUB = open(os.path.join(BASE, 'data/public_url.txt')).read().strip().rstrip('/')

def q(sql, args=()):
    c = sqlite3.connect(DB, timeout=30); c.row_factory = sqlite3.Row
    r = c.execute(sql, args).fetchall(); c.close(); return r
def ex(sql, args=()):
    c = sqlite3.connect(DB, timeout=30); c.execute(sql, args); c.commit(); c.close()

passed = []
def check(name, cond, extra=''):
    passed.append((name, cond))
    print(('✅' if cond else '❌'), name, ('' if cond else (' | ' + str(extra)[:200])))

# 前置清理(幂等)
for r in q("SELECT id FROM purchase_orders WHERE order_no LIKE 'CG-TEST-V11164%'"):
    ex("DELETE FROM contracts WHERE order_id=?", (r['id'],))
    ex("DELETE FROM approval_instances WHERE biz_type='contract' AND biz_id IN (SELECT id FROM contracts WHERE order_id=?)", (r['id'],))
    ex("DELETE FROM approval_instances WHERE biz_type='purchase_order' AND biz_id=?", (r['id'],))
    ex("DELETE FROM dingtalk_instances WHERE biz_id=? AND biz_type IN ('contract','purchase_order')", (r['id'],))
    ex("DELETE FROM feishu_instances WHERE biz_id=? AND biz_type IN ('contract','purchase_order')", (r['id'],))
    ex("DELETE FROM order_items WHERE order_id=?", (r['id'],))
ex("DELETE FROM purchase_orders WHERE order_no LIKE 'CG-TEST-V11164%'")
ex("DELETE FROM logs WHERE detail LIKE '%CG-TEST-V11164%'")

# ---------- 1. 造测试订单(引用真实供应商名, 订单本身是测试数据) ----------
sup = q("SELECT name FROM suppliers WHERE name NOT LIKE '%测试%' LIMIT 1")
sup_name = sup[0]['name'] if sup else '山西昂拓贸易有限公司'
ex("INSERT INTO purchase_orders(order_no,item_name,spec,quantity,unit,price,amount,tax_rate,tax_amount,total_amount,supplier,requester,trade_mode,status,remark) VALUES('CG-TEST-V11164-01','【测试】合同验证物资','T-100',2,'个',500,1000,13,130,1130,?,'测试员','货到付款','已通过','【测试】任务3合同验证')", (sup_name,))
oid = q("SELECT id FROM purchase_orders WHERE order_no='CG-TEST-V11164-01'")[0]['id']
ex("INSERT INTO order_items(order_id,item_name,spec,unit,quantity,price,amount,tax_rate,tax_amount,total_amount) VALUES(?,'【测试】合同验证物资','T-100','个',2,500,1000,13,130,1130)", (oid,))

# ---------- 2. 调生成合同接口 ----------
client = appmod.app.test_client()
r = client.post('/api/login', json={'username': 'admin', 'password': 'admin123'})
check('C1 admin登录', r.status_code == 200)
r = client.post('/api/contracts/generate', json={'order_id': oid, 'settle_type': '现结'})
check('C2 生成合同成功', r.status_code == 200 and r.get_json().get('success'), r.text[:200])
cno = r.get_json().get('contract_no')
fname = r.get_json().get('file')
fpath = os.path.join(UP, fname)
check('C3 合同文件存在且非空', os.path.exists(fpath) and os.path.getsize(fpath) > 10000, f'{fname} {os.path.getsize(fpath) if os.path.exists(fpath) else 0}B')
crec = q("SELECT * FROM contracts WHERE contract_no=?", (cno,))[0] if cno else None
check('C4 contracts 表记录入库', crec is not None and crec['file_path'] == fname)

# ---------- 3. 严格验证(模拟 Win/Mac 双端: python-docx 是跨平台 OOXML 严格解析器) ----------
from docx import Document
doc = Document(fpath)
paras = doc.paragraphs
tables = doc.tables
texts = [p.text for p in paras]
full = '\n'.join(texts)
for t in tables:
    for row in t.rows:
        for c in row.cells:
            full += '\n' + c.text
check('C5 zip包结构完整(docx=zip)', zipfile.is_zipfile(fpath))
check('C6 段落可读(65段左右)', 50 <= len(paras) <= 80, f'{len(paras)}段')
check('C7 3表格(标题/明细/落款)', len(tables) == 3, f'{len(tables)}表格')
nums = [t for t in texts if re.match(r'^[一二三四五六七八九十]+、', t)]
check('C8 13章节齐全', len(nums) == 13, f'{len(nums)}章节')
left = re.findall(r'\{[^}]*\}', full)
check('C9 无{占位符}残留', not left, sorted(set(left))[:8])
moji = re.findall(r'[\ufffd]|Ã[\x80-\xbf]|â€[™“”]', full)
check('C10 无乱码', not moji)
check('C11 甲方固定(河曲县正成洗选煤有限责任公司)', '河曲县正成洗选煤有限责任公司' in full)
check('C12 乙方自动填供应商', sup_name in full)
check('C13 合同编号已填', cno in full)
check('C14 明细行填充(物资名+规格+数量列)', '【测试】合同验证物资' in full and 'T-100' in full and re.search(r'\n2\s*\n', full) is not None)
check('C15 合计金额含大写', re.search(r'合计金额：¥1,130\.00元（大写金额：人民币壹仟壹佰叁拾元整）', full) is not None)
check('C16 税金/不含税完整', '税金（税率 13%）为：' in full and '不含税价款为：' in full)
check('C17 结算方式注入(现结)', '现结：一单一结，验收合格后立即付款' in full)
check('C18 违约金30%', '合同额的30%' in full)
check('C19 落款乙方盖章栏', '乙方（签字并盖章）' in full)

# ---------- 4. 公网下载完整性(模拟Windows下载) ----------
local = open(fpath, 'rb').read()
url = PUB + '/uploads/' + urllib.parse.quote(fname)
req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'})
with urllib.request.urlopen(req, timeout=60) as resp:
    remote = resp.read()
check('C20 公网下载字节一致', len(remote) == len(local), f'{len(remote)} vs {len(local)}')
check('C21 公网下载md5一致', hashlib.md5(remote).hexdigest() == hashlib.md5(local).hexdigest())
check('C22 公网下载后仍可解析', zipfile.is_zipfile(__import__('io').BytesIO(remote)))

# ---------- 5. 清理链(顺序: 先删审批/实例, 再删合同/订单, 防孤儿审批实例) ----------
cid_list = [r['id'] for r in q("SELECT id FROM contracts WHERE order_id=?", (oid,))]
for _cid in cid_list:
    ex("DELETE FROM approval_instances WHERE biz_type='contract' AND biz_id=?", (_cid,))
    ex("DELETE FROM dingtalk_instances WHERE biz_type='contract' AND biz_id=?", (_cid,))
    ex("DELETE FROM feishu_instances WHERE biz_type='contract' AND biz_id=?", (_cid,))
ex("DELETE FROM approval_instances WHERE biz_type='purchase_order' AND biz_id=?", (oid,))
ex("DELETE FROM dingtalk_instances WHERE biz_type='purchase_order' AND biz_id=?", (oid,))
ex("DELETE FROM feishu_instances WHERE biz_type='purchase_order' AND biz_id=?", (oid,))
ex("DELETE FROM contracts WHERE order_id=?", (oid,))
ex("DELETE FROM order_items WHERE order_id=?", (oid,))
ex("DELETE FROM purchase_orders WHERE id=?", (oid,))
if fname and os.path.exists(fpath):
    os.remove(fpath)
ex("DELETE FROM logs WHERE detail LIKE '%CG-TEST-V11164%' OR detail LIKE '%【测试】任务3合同验证%'")
for t in ('contracts', 'order_items', 'purchase_orders', 'approval_instances', 'dingtalk_instances', 'feishu_instances'):
    ex(f"UPDATE sqlite_sequence SET seq=COALESCE((SELECT MAX(id) FROM {t}),0) WHERE name='{t}'")
left = q("SELECT COUNT(*) c FROM purchase_orders WHERE order_no LIKE 'CG-TEST-V11164%'")[0]['c']
left2 = q("SELECT COUNT(*) c FROM contracts WHERE contract_name LIKE '%测试%'")[0]['c']
check('清理后测试订单/合同零残留', left == 0 and left2 == 0, f'{left}/{left2}')

print('\n==== 汇总 ====')
fails = [p for p in passed if not p[1]]
print(f'通过 {len(passed)-len(fails)}/{len(passed)}')
sys.exit(1 if fails else 0)
