#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成正式买卖合同模板 tpl_maimai.docx（参照已签合同照片结构）
甲方固定: 河曲县正成洗选煤有限责任公司
乙方: {乙方名称}等占位符, 生成时按订单供应商自动填入
明细: 表格自动填充(order_items 逐行)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = '/Users/a0/Desktop/正成能源/01_系统程序/采购系统程序/purchase_system'
OUT = os.path.join(BASE, 'uploads', 'tpl_maimai.docx')

doc = Document()

# 默认字体: 仿宋/宋体, 小四
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def para(text='', bold=False, align=None, size=12, space_after=6, font='仿宋'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return p

def heading(text):
    return para(text, bold=True, size=12, space_after=4)

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)

def set_cell_text(cell, text, bold=False, size=10.5, center=True):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = '仿宋'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# ============ 标题 ============
para('买卖合同', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, space_after=10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('合同编号：{合同编号}')
r.font.size = Pt(11)
r.font.name = '仿宋'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
doc.add_paragraph()

# ============ 双方信息 ============
para('甲方（买方）：河曲县正成洗选煤有限责任公司', bold=False, size=12)
para('地址：山西省忻州市河曲县    联系人：采购部    电话：', size=12)
para('乙方（卖方）：{乙方名称}', size=12)
para('地址：{乙方地址}    联系人：{乙方联系人}    电话：{乙方电话}', size=12)
para('开户行：{乙方开户行}    账号：{乙方账号}', size=12)
doc.add_paragraph()

# 前言
para('根据《中华人民共和国民法典》及相关法律法规的规定，甲乙双方本着平等自愿、协商一致的原则，'
     '就甲方向乙方采购标的物事宜达成如下协议，以资共同遵守。', size=12)
doc.add_paragraph()

# ============ 一、标的、规格、数量、价款 ============
heading('一、标的、规格、数量、价款')
table = doc.add_table(rows=3, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['标的物', '规格型号', '计量单位', '数量', '单价（元）', '金额（元）', '备注']
for j, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[j], h, bold=True)
# 第二行: 空明细行(代码自动填)
for j in range(7):
    set_cell_text(table.rows[1].cells[j], '')
# 第三行: 合计行
set_cell_text(table.rows[2].cells[0], '合计', bold=True)
for j in range(1, 5):
    set_cell_text(table.rows[2].cells[j], '')
set_cell_text(table.rows[2].cells[5], '合计金额：¥    元（人民币大写金额：     ）', bold=True)
set_cell_text(table.rows[2].cells[6], '')
# 列宽
widths = [Cm(3.2), Cm(2.8), Cm(1.8), Cm(1.4), Cm(2.2), Cm(2.6), Cm(2.0)]
for row in table.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w
for row in table.rows:
    for cell in row.cells:
        set_cell_border(cell)
doc.add_paragraph()

# 合计金额说明(段落, _apply_ct 自动填税金/不含税/大写)
para('合计金额：¥    元（人民币大写金额：     ）。', size=12)
doc.add_paragraph()

# ============ 二、质量标准 ============
heading('二、质量标准')
para('标的物必须符合国家现行标准及行业标准、企业标准（按国标执行）。乙方提供的标的物应符合双方确认的规格、技术参数及质量要求。')
doc.add_paragraph()

# ============ 三、包装 ============
heading('三、包装')
para('乙方应采用国家或行业标准包装，适应长途运输，防潮、防震、防碰撞，确保标的物完好运抵甲方指定地点。包装费用由乙方承担，包装物不回收。')
doc.add_paragraph()

# ============ 四、交付 ============
heading('四、交付')
para('1. 乙方应在合同签订后  日内交付，将标的物运送至甲方指定地点，并负责安装调试至正常运行。')
para('2. 乙方应随货提供质量检验合格证书、出厂检验报告、使用说明书及技术资料等。')
para('3. 合同价款包含设计、制造、运输、装卸、安装、调试及售后服务等全部费用。')
doc.add_paragraph()

# ============ 五、运输 ============
heading('五、运输')
para('1. 乙方负责将标的物运输至甲方指定地点。')
para('2. 运输方式：汽运。')
para('3. 交付前标的物的费用与风险由乙方承担。')
doc.add_paragraph()

# ============ 六、验收 ============
heading('六、验收')
para('1. 标的物运抵甲方指定地点后  日内由甲方组织验收。')
para('2. 验收合格的，甲方签署验收单；验收不合格的，乙方应在  日内更换或补足。')
doc.add_paragraph()

# ============ 七、质量保证 ============
heading('七、质量保证')
para('标的物质保期为  年，自验收合格之日起计算。质保期内出现质量问题，乙方负责免费更换、维修；因乙方原因造成甲方损失的，乙方应承担赔偿责任。')
doc.add_paragraph()

# ============ 八、结算 ============
heading('八、结算')
para('1. 乙方应于  日内向甲方提供全额增值税专用发票（税率13%）。')
para('2. 甲方自收到发票后  日内支付合同总价的  %，乙方安排发货。')
para('3. 收款账户信息：')
para('    收款账户名称：')
para('    收款账号：')
para('    收款银行：')
para('    银行行号：')
doc.add_paragraph()

# ============ 九、违约责任 ============
heading('九、违约责任')
para('1. 乙方迟延交货应承担合同总额的  %的违约金。')
para('2. 标的物质量、规格、数量等不符合约定要求的，乙方除应按约定期限及时更换、补发外，还应承担上述迟延交货的违约责任。')
para('3. 因乙方违约造成甲方损失的，违约金不足以弥补甲方全部损失时，乙方应承担继续赔偿责任。')
para('4. 乙方在货物装卸、运输、安装、调试等过程中发生的一切安全事故及第三人人身或财产损失的，均由其自行负责，与甲方无关。')
para('5. 甲方因维护自身权利支出的包括但不限于诉讼费、律师费、评估费、鉴定费、保全费等由乙方承担。')
doc.add_paragraph()

# ============ 十、合同解除 ============
heading('十、合同解除')
para('1. 乙方延迟交货超过  天，甲方有权单方面解除合同。')
para('2. 乙方在质量保证期内，拒不履行"三包"义务，甲方有权单方面解除合同。')
para('3. 乙方更换标的物后仍有瑕疵，甲方有权单方面解除合同。')
para('4. 其他致使合同目的不能实现的情况，任意一方可单方面解除合同，但需提前  天通知对方。')
doc.add_paragraph()

# ============ 十一、知识产权 ============
heading('十一、知识产权')
para('乙方提供的货物须保证不侵犯任何第三方的知识产权，如因乙方提供的货物存在侵权行为导致甲方承担赔偿责任的，乙方应全额赔偿并承担由此导致的甲方全部损失。')
doc.add_paragraph()

# ============ 十二、争议解决方式 ============
heading('十二、争议解决方式')
para('本合同在履行过程中发生争议的，由双方当事人协商解决；协商不成的，双方均可向甲方所在地人民法院提起诉讼。')
doc.add_paragraph()

# ============ 十三、其他 ============
heading('十三、其他')
para('1. 本合同自双方法定代表人或委托代理人签字并加盖单位印章之日起生效，若非真实签字、签章，本合同无效。')
para('2. 本合同未尽事宜，可以另行签订补充协议，补充协议与本合同具有同等法律效力。')
para('3. 本合同一式肆份，甲方执叁份，乙方执壹份，各份具有同等法律效力。')
doc.add_paragraph()

# ============ 落款 ============
# 双栏: 用表格 2 列
sign = doc.add_table(rows=3, cols=2)
for row in sign.rows:
    for cell in row.cells:
        cell.text = ''
set_cell_text(sign.rows[0].cells[0], '甲方（签字并盖章）：', bold=False, center=False)
set_cell_text(sign.rows[0].cells[1], '乙方（签字并盖章）：', bold=False, center=False)
set_cell_text(sign.rows[1].cells[0], '（河曲县正成洗选煤有限责任公司）', center=False, size=10)
set_cell_text(sign.rows[1].cells[1], '（{乙方名称}）', center=False, size=10)
set_cell_text(sign.rows[2].cells[0], '日期：    年    月    日', center=False)
set_cell_text(sign.rows[2].cells[1], '日期：    年    月    日', center=False)
for row in sign.rows:
    for cell in row.cells:
        cell.width = Cm(8)

doc.save(OUT)
print('模板已生成:', OUT)
print('大小:', os.path.getsize(OUT), 'bytes')
